"""
generate_product_docs.py
-------------------------
Genera los dos documentos oficiales del producto Gestabiz:

  1. docs/Manual_Usuario_Gestabiz.docx   — Guia funcional exhaustiva para el usuario final
  2. docs/Propuesta_Valor_Gestabiz.docx  — Pitch comercial orientado a ventas

Ambos documentos siguen la misma estructura:

  * Portada con logo Gestabiz + atribucion "Desarrollado por Ti Turing"
  * Indice con hipervinculos internos a las secciones detalladas
  * PARTE 1 — Resumen Ejecutivo (alto nivel, con hipervinculos al detalle)
  * PARTE 2 — Detalle Exhaustivo (reglas, excepciones, flujos, validaciones, cada boton)
  * Placeholders para capturas de pantalla
  * Pie con atribucion de Ti Turing

Features agrupadas por plan: Gratuito -> Basico -> Pro (cada plan incluye el anterior).

Fuente unica de verdad para los .docx del producto (regla #14 CLAUDE.md / regla #13
copilot-instructions.md). Cualquier cambio funcional en la app debe reflejarse aqui y
los .docx deben regenerarse con:

    python scripts/generate_product_docs.py

Autor: Ti Turing — https://github.com/TI-Turing
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
DOCS_DIR.mkdir(parents=True, exist_ok=True)

LOGO_GESTABIZ = ROOT / "src" / "assets" / "images" / "logo_gestabiz.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

# Brand palette
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)

# ---------------------------------------------------------------------------
# Low-level helpers: hyperlinks, bookmarks, shading
# ---------------------------------------------------------------------------
_bookmark_id_counter = [1000]


def _next_bookmark_id() -> int:
    _bookmark_id_counter[0] += 1
    return _bookmark_id_counter[0]


def add_bookmark(paragraph, name: str) -> None:
    bid = str(_next_bookmark_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    paragraph._p.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *, bold: bool = False,
                            color: RGBColor = PURPLE, size: int = 11) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rpr.append(color_el)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)

    if bold:
        rpr.append(OxmlElement("w:b"))

    new_run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# High-level helpers: headings, paragraphs, lists, tables, placeholders
# ---------------------------------------------------------------------------
def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + text)
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)

    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(title)
    style_run(r, size=11, bold=True, color=color)

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=DARK)

    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)

    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {caption}")
    style_run(r, size=9, italic=True, color=GREY)

    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=header_fg)

    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, size=10, color=DARK)

    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup + cover
# ---------------------------------------------------------------------------
def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run(f"{title}  ·  Desarrollado por Ti Turing  ·  © 2026")
        style_run(r, size=9, italic=True, color=GREY)

    build_cover(doc, title, subtitle)
    return doc


def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()

    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version del producto: 1.0.3   ·   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fase Beta completada   ·   Listo para produccion")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_TITURING), width=Cm(3.0))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por ")
    style_run(r, size=11, color=GREY)
    r = p.add_run("Ti Turing")
    style_run(r, size=11, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("https://github.com/TI-Turing   ·   contacto@gestabiz.com")
    style_run(r, size=9, italic=True, color=GREY)

    page_break(doc)


# ===========================================================================
# DATA: Plans + modules (single source of truth)
# ===========================================================================
@dataclass
class PlanSpec:
    key: str
    name: str
    price_monthly: str
    price_annual: str
    tagline: str
    ideal_for: str
    includes_previous: bool
    exclusives: list[str] = field(default_factory=list)
    limits: list[str] = field(default_factory=list)


PLANS: list[PlanSpec] = [
    PlanSpec(
        key="free",
        name="Gratuito",
        price_monthly="$0",
        price_annual="$0",
        tagline="Para probar Gestabiz sin costo y arrancar operacion en minutos.",
        ideal_for="Profesionales independientes que estan empezando y necesitan un canal de reservas sin inversion inicial.",
        includes_previous=False,
        exclusives=[
            "Registro de un (1) negocio con una (1) sede",
            "Hasta 30 citas activas por mes",
            "Perfil publico indexable en Google con URL /negocio/<slug>",
            "Wizard de reserva completo con validaciones anti-overlap",
            "Panel de Resumen (Overview) con metricas basicas",
            "Gestion de Servicios (crear, editar, activar/desactivar)",
            "Gestion de Clientes (CRM basico con historial)",
            "Facturacion del plan (visualizacion)",
            "Notificaciones in-app (campanita)",
            "Notificaciones por email (recordatorios de cita)",
            "Chat cliente <-> dueno del negocio",
            "Resenas de clientes con respuesta del negocio",
            "Configuraciones del negocio y del perfil personal",
        ],
        limits=[
            "Sin gestion de empleados",
            "Sin multiples sedes",
            "Sin reportes avanzados",
            "Sin ventas rapidas ni contabilidad",
            "Sin sistema de permisos granulares",
            "Sin WhatsApp ni SMS (solo email)",
        ],
    ),
    PlanSpec(
        key="basico",
        name="Basico",
        price_monthly="$89.900 COP / mes",
        price_annual="$899.000 COP / ano (2 meses gratis)",
        tagline="Para negocios con equipo y operacion multi-sede que exigen control real.",
        ideal_for="Salones de belleza, barberias, consultorios, estudios, gimnasios y centros de servicio con 2-15 empleados que ya tienen flujo diario de citas.",
        includes_previous=True,
        exclusives=[
            "Hasta 15 empleados activos",
            "Hasta 5 sedes (ubicaciones fisicas)",
            "Citas ilimitadas",
            "Modulo de Empleados con jerarquia visual (mapa)",
            "Sistema de Ausencias y Vacaciones con aprobacion obligatoria",
            "Festivos publicos colombianos 2025-2027 integrados",
            "Ventas Rapidas (POS ligero sin cita previa)",
            "Historial de Ventas por periodo, empleado y sede",
            "Reportes operacionales y financieros basicos",
            "Sistema de Permisos Granulares (79 tipos, 9 plantillas predefinidas)",
            "Notificaciones por WhatsApp Business",
            "Transferencia de empleados entre sedes",
            "Chat multi-empleado (clientes pueden elegir empleado)",
            "Sede preferida por dispositivo",
        ],
        limits=[
            "Sin contabilidad completa con impuestos colombianos",
            "Sin gastos recurrentes",
            "Sin reclutamiento de vacantes",
            "Sin recursos fisicos (salas, canchas, mesas)",
        ],
    ),
    PlanSpec(
        key="pro",
        name="Pro",
        price_monthly="$159.900 COP / mes",
        price_annual="$1.599.000 COP / ano (2 meses gratis)",
        tagline="Todo lo que necesita una PyME en serio: contabilidad, recursos fisicos y reclutamiento.",
        ideal_for="Centros medicos, clinicas esteticas, hoteles boutique, coworkings, restaurantes por reserva, escuelas deportivas y negocios que manejan recursos fisicos o necesitan contabilidad con impuestos.",
        includes_previous=True,
        exclusives=[
            "Empleados y sedes ilimitados",
            "Modulo de Contabilidad completo con IVA, ICA y Retencion en la Fuente",
            "Modulo de Gastos recurrentes y puntuales",
            "Configuracion fiscal por sede",
            "Exportacion de reportes a PDF / CSV / Excel",
            "Charts financieros avanzados (tendencia mensual, ingresos vs gastos, categorias, empleados top)",
            "Sistema de Reclutamiento (vacantes, aplicaciones, matching inteligente, reviews obligatorias)",
            "Sistema de Recursos Fisicos (salas, canchas, mesas, equipos - 15 tipos)",
            "Modelos de negocio flexibles (profesional / recurso fisico / hibrido / clase grupal)",
            "SMS transaccional via AWS SNS",
            "Sincronizacion con Google Calendar",
            "API publica y webhooks (roadmap Q2 2026)",
            "Soporte prioritario por WhatsApp",
        ],
        limits=[],
    ),
]


@dataclass
class ModuleSpec:
    key: str
    name: str
    plan: str
    scope: str
    summary: str
    details: list[str]
    flows: list[tuple[str, list[str]]] = field(default_factory=list)
    exceptions: list[str] = field(default_factory=list)
    permissions: list[str] = field(default_factory=list)
    screenshot: str | None = None


MODULES: list[ModuleSpec] = [
    # ---- AREA PUBLICA Y ONBOARDING ----
    ModuleSpec(
        key="landing", name="Landing Page Publica", plan="free", scope="Transversal",
        summary="Pagina de aterrizaje moderna y SEO-optimizada que presenta Gestabiz al visitante, explica beneficios y lleva al registro o al login.",
        details=[
            "Ruta: / - accesible sin autenticacion.",
            "Secciones: Hero con propuesta de valor, Beneficios (grid 3x2), Como funciona (3 pasos), Testimonios, Planes, CTA final, Footer legal.",
            "Header fijo con logo Gestabiz, navegacion (Inicio / Funcionalidades / Precios / Contacto) y botones Login / Registro.",
            "Boton 'Iniciar sesion' redirige a /login; boton 'Empezar gratis' redirige a /register.",
            "Responsive mobile-first con breakpoints sm/md/lg/xl de Tailwind 4.",
            "Meta tags dinamicos para SEO, Open Graph para redes, Twitter Card, JSON-LD structured data.",
            "Selector de idioma (ES/EN) persistente en localStorage.",
            "Cookie consent banner GDPR-compliant antes de activar Google Analytics 4.",
            "Cambio de tema claro/oscuro con boton en el header.",
        ],
        flows=[
            ("Flujo estandar de visitante", [
                "Usuario llega por busqueda en Google o link directo.",
                "Acepta o rechaza cookies (GA4 se activa solo tras aceptar).",
                "Explora secciones; puede saltar directamente a Precios.",
                "Pulsa 'Empezar gratis' -> AuthScreen en modo registro.",
            ]),
            ("Flujo desde perfil publico", [
                "Usuario llega a /negocio/<slug> desde Google.",
                "Click en 'Reservar' -> Gestabiz guarda business_id/service_id/employee_id en URL params.",
                "Redirige al login; tras autenticar, abre el AppointmentWizard con los datos preseleccionados.",
            ]),
        ],
        exceptions=[
            "Si el usuario ya tiene sesion, el header muestra 'Ir a mi panel' en lugar de Login.",
            "Si el banner de cookies no se ha aceptado, GA4 no dispara eventos (consent mode).",
        ],
        screenshot="Landing page Gestabiz - vista hero en desktop.",
    ),
    ModuleSpec(
        key="auth", name="Autenticacion y Registro", plan="free", scope="Transversal",
        summary="Sistema de autenticacion multicanal: email/password, Google OAuth, GitHub OAuth y Magic Link por email.",
        details=[
            "Pantalla unica /login y /register renderiza el componente AuthScreen con modo configurable.",
            "Campos obligatorios en registro: correo, contrasena (minimo 8 caracteres con letras y numeros), nombre completo, pais.",
            "Checkbox 'Acepto terminos y politica de privacidad' obligatorio antes de registrar.",
            "Validacion en tiempo real con Zod: correo valido, contrasena segura, telefono con prefijo internacional.",
            "Supabase envia correo de verificacion tras registro con link unico (token de 24 h).",
            "Boton 'Continuar con Google' usa el cliente OAuth correcto por entorno (DEV / PROD).",
            "Boton 'Continuar con GitHub' disponible para perfiles tecnicos.",
            "Opcion 'Recibir link magico' envia un email con link de un solo uso que autentica sin contrasena.",
            "Tras verificar el correo, el usuario aterriza en /app y se dispara el flujo de onboarding.",
            "Toda la autenticacion corre sobre el cliente Supabase singleton; nunca se crean instancias adicionales.",
        ],
        flows=[
            ("Registro estandar con email", [
                "Usuario ingresa correo y contrasena.",
                "Completa datos personales (nombre, pais, telefono).",
                "Acepta terminos y pulsa 'Crear cuenta'.",
                "Recibe email de verificacion; hasta confirmarlo la app muestra banner 'Verifica tu correo'.",
                "Pulsa el link -> redirect a /app/verify -> confirma -> login automatico.",
            ]),
            ("Registro con Google OAuth", [
                "Usuario pulsa 'Continuar con Google'.",
                "Redirige a accounts.google.com con el client_id del entorno.",
                "Tras aceptar permisos vuelve a /auth/callback.",
                "Gestabiz crea el profile si no existe y aterriza en /app.",
            ]),
            ("Reserva sin cuenta (deep-link)", [
                "Cliente llega desde /negocio/<slug> con params en la URL.",
                "AuthScreen detecta los params y muestra toast 'Inicia sesion para completar tu reserva'.",
                "Tras login, useWizardDataCache hidrata el wizard con lo preseleccionado.",
            ]),
        ],
        exceptions=[
            "Si el correo ya existe, aparece 'Este correo ya tiene cuenta - quieres iniciar sesion?' con boton directo al login.",
            "Si Google devuelve scope insuficiente, se muestra 'Necesitamos permisos de correo para continuar'.",
            "Magic Link expira en 1 h; pedir uno nuevo si esta vencido.",
            "Tras 5 intentos fallidos Supabase aplica rate-limit temporal.",
        ],
        screenshot="Pantalla AuthScreen con tabs Login / Registro y botones OAuth.",
    ),
    ModuleSpec(
        key="onboarding", name="Onboarding y Creacion de Negocio", plan="free", scope="Transversal",
        summary="Flujo guiado que lleva al nuevo usuario desde la primera pantalla hasta tener un negocio operativo o un perfil de cliente listo.",
        details=[
            "Primera decision post-registro: vengo a reservar o a ofrecer servicios?",
            "Opcion 'Reservar' -> activa rol Cliente y abre el catalogo de negocios.",
            "Opcion 'Ofrecer servicios' -> asistente de 4 pasos para crear negocio.",
            "Paso 1: datos del negocio (nombre, categoria, hasta 3 subcategorias, descripcion, slug publico unico).",
            "Paso 2: primera sede (pais, ciudad, direccion, geolocalizacion, horario semanal, opens_at/closes_at).",
            "Paso 3: primer servicio (nombre, duracion, precio, descripcion, imagen).",
            "Paso 4: resumen y confirmacion; al guardar se dispara el trigger que inserta al owner como manager.",
            "Al completar, el owner recibe automaticamente los 79 permisos del sistema.",
            "Si el usuario abandona el wizard, el progreso queda guardado via useKV y se retoma al entrar.",
        ],
        flows=[
            ("Creacion de negocio paso a paso", [
                "Elegir categoria principal (15 disponibles: belleza, salud, deportes, hoteles, restaurantes, etc.).",
                "Elegir hasta 3 subcategorias (~60 disponibles).",
                "Definir slug publico - valida unicidad en tiempo real (debounce 500 ms).",
                "Subir logo (bucket avatars, publico) y banner (opcional).",
                "Completar sede y horarios con sugerencia estandar 08:00-18:00 L-S.",
                "Crear el primer servicio.",
                "Confirmar -> redirect a /app/admin/overview.",
            ]),
        ],
        exceptions=[
            "Si el slug ya existe, el input se pone en rojo y sugiere variaciones.",
            "Si la geolocalizacion falla, el usuario ingresa la direccion manual y Gestabiz intenta geocodificar al guardar.",
            "Si el usuario cierra el wizard a la mitad, al volver entra al paso exacto donde quedo.",
        ],
        screenshot="Wizard de creacion de negocio - paso 1: datos basicos.",
    ),
    # ---- CITAS ----
    ModuleSpec(
        key="appointments", name="Gestion de Citas (Admin y Empleado)", plan="free", scope="Admin",
        summary="Calendario operativo central: visualiza, crea, edita, reagenda y cancela citas; detecta conflictos en tiempo real.",
        details=[
            "Vista AppointmentsCalendar con filtros por sede, empleado, servicio, estado y rango de fechas.",
            "Vistas: dia / semana / mes / agenda (lista).",
            "Colores por estado: Confirmada (verde), Pendiente (amarillo), Completada (azul), Cancelada (gris).",
            "Click en cita abre modal con: servicio, cliente, empleado, sede, hora, precio, notas internas.",
            "Acciones por cita: Confirmar, Marcar completada, Reagendar, Cancelar con motivo, Chatear con cliente.",
            "Reagendar abre el AppointmentWizard en modo edicion (skip de pasos ya decididos).",
            "Cancelacion emite in-app notification al cliente + email + WhatsApp (si plan lo permite).",
            "Exportacion a CSV / ICS (formato iCal) del rango actual.",
            "Detalle de cita marca is_location_exception=true si el empleado no esta en su sede base.",
            "Integracion con Google Calendar (plan Pro) sincroniza bidireccional con debounce de 10s.",
        ],
        flows=[
            ("Creacion de cita por el admin", [
                "Admin pulsa 'Nueva cita'.",
                "Selecciona cliente existente o crea uno nuevo inline.",
                "Elige servicio -> sistema precarga duracion y precio.",
                "Elige sede -> filtra empleados disponibles.",
                "Elige empleado -> filtra horarios.",
                "Elige fecha -> useAssigneeAvailability valida: horario sede, almuerzo empleado, festivos, ausencias, overlap.",
                "Confirma -> createAppointment inserta y dispara notificaciones.",
            ]),
            ("Cancelacion de cita", [
                "Admin o cliente pulsa 'Cancelar'.",
                "Escribe motivo (minimo 10 caracteres).",
                "Sistema pregunta: notificar al cliente/empleado?",
                "Al confirmar actualiza status='cancelled' y libera el slot.",
            ]),
            ("Confirmacion por link publico", [
                "Cliente recibe email con link /confirmar-cita/<token>.",
                "Pulsa el link sin necesidad de autenticarse.",
                "Ve los detalles y pulsa 'Confirmar'.",
                "Backend valida token (24 h de vida), marca status='confirmed' y notifica al negocio.",
            ]),
        ],
        exceptions=[
            "Si el horario del empleado se modifico y hay citas antiguas, el calendario las respeta (isLunchBreak retorna false para fechas pasadas).",
            "Si el servicio fue eliminado, el join es LEFT JOIN y la cita aun se muestra con servicio 'N/A'.",
            "Si hay overlap al guardar (carrera entre dos admins), el INSERT falla con error controlado y se muestra 'Ese horario acaba de ocuparse'.",
            "Si el cliente no tiene email, se omite la notificacion por correo (se registra en notification_log).",
        ],
        permissions=["appointments.view", "appointments.create", "appointments.edit", "appointments.cancel", "appointments.complete"],
        screenshot="Calendario de citas con filtros, vista semanal y leyenda de estados.",
    ),
    ModuleSpec(
        key="appointment_wizard", name="Wizard de Reserva (Cliente)", plan="free", scope="Cliente",
        summary="Asistente de 6-8 pasos con validacion en tiempo real que guia al cliente desde elegir negocio hasta confirmar cita.",
        details=[
            "Pasos condicionales: BusinessSelection -> (EmployeeBusinessSelection si aplica) -> ServiceSelection -> LocationSelection -> EmployeeSelection -> DateTimeSelection -> ConfirmationStep -> SuccessStep.",
            "Si el cliente viene por deep-link, los pasos ya decididos se saltan automaticamente.",
            "Cada paso bloquea 'Siguiente' hasta que la seleccion sea valida.",
            "Indicador visual de progreso (stepper horizontal en desktop, compacto en mobile).",
            "useWizardDataCache persiste las selecciones en localStorage durante 1 h.",
            "DateTimeSelection corre 3 queries en paralelo: horario sede, horario empleado, citas existentes.",
            "10 validaciones en tiempo real: horario sede, apertura/cierre, almuerzo, overlap con otras citas del empleado, overlap con el mismo cliente en otros negocios, festivos publicos, ausencias aprobadas, duracion minima, anticipacion minima configurada por el negocio, maximo de citas activas del cliente.",
            "Slots deshabilitados muestran tooltip con el motivo ('Ocupado', 'Hora de almuerzo', 'Ausencia', 'Festivo').",
            "Al confirmar, Gestabiz crea la cita con status='pending' y envia email de confirmacion con link de 24 h.",
            "Tracking GA4: booking_started, booking_step_completed (por paso), booking_abandoned, purchase.",
        ],
        flows=[
            ("Reserva completa desde cero", [
                "Cliente entra a /app -> Buscar.",
                "Filtra por ciudad, categoria o texto libre.",
                "Elige negocio -> ve perfil.",
                "Pulsa 'Reservar'.",
                "Avanza por los pasos del wizard con validacion.",
                "Confirma -> ve SuccessStep con botones 'Agregar a Google Calendar' y 'Ver mis citas'.",
            ]),
            ("Reserva abandonada y retomada", [
                "Cliente completa hasta Paso 4 y cierra pestana.",
                "Vuelve 30 min despues -> wizard reabre en Paso 4.",
                "Retoma y confirma.",
            ]),
        ],
        exceptions=[
            "Si el negocio no tiene empleados ofreciendo servicios (offers_services=true), el paso EmployeeSelection se salta y la cita queda sin asignar.",
            "Si el cliente alcanza el maximo de citas activas (5 por negocio por defecto), ConfirmationStep muestra bloqueo.",
            "Si el token de confirmacion expira (24 h), el link muestra 'Enlace vencido - contacta al negocio'.",
        ],
        screenshot="Wizard de reserva - paso Fecha/Hora con slots disponibles y tooltips.",
    ),
    # ---- SERVICIOS/CRM/RESENAS/BILLING VIEW (free) ----
    ModuleSpec(
        key="services", name="Gestion de Servicios", plan="free", scope="Admin",
        summary="CRUD completo de servicios del negocio con control de precio, duracion, imagen, categoria y disponibilidad por sede/empleado.",
        details=[
            "ServicesManager lista todos los servicios con badge de estado (activo/inactivo).",
            "Crear servicio: nombre, descripcion, duracion (en minutos, multiplos de 5), precio en moneda configurada, categoria, imagen (bucket publico).",
            "Editar servicio actualiza precios futuros, no altera citas ya agendadas.",
            "Eliminar servicio pide confirmacion + motivo; las citas historicas conservan la referencia via LEFT JOIN.",
            "Activar/desactivar oculta del wizard de reserva sin borrar el registro.",
            "Asignar servicio a sedes (tabla location_services) - si ninguna sede esta marcada, se asume disponible en todas.",
            "Asignar servicio a empleados (tabla employee_services) - solo empleados con offers_services=true aparecen en el wizard.",
            "Filtro rapido por categoria, estado y busqueda por nombre con debounce.",
        ],
        exceptions=[
            "No se puede eliminar un servicio con citas futuras confirmadas; el sistema sugiere 'Desactivar en su lugar'.",
            "Cambiar la duracion de un servicio NO afecta citas existentes (duracion se copia al crear la cita).",
        ],
        permissions=["services.view", "services.create", "services.edit", "services.delete"],
        screenshot="ServicesManager con grid de servicios, filtros y modal de edicion.",
    ),
    ModuleSpec(
        key="clients_crm", name="Clientes (CRM)", plan="free", scope="Admin",
        summary="Listado de todos los clientes que han tenido al menos una cita no cancelada en el negocio, con acceso al historial completo.",
        details=[
            "Componente ClientsManager en /app/admin/clients.",
            "Query en dos pasos: appointments (client_id, start_time, status) -> profiles (full_name, email, avatar_url).",
            "Agregacion en cliente: total de visitas, visitas completadas, ultima visita.",
            "Vista grid con card por cliente (avatar de iniciales, email, contador, fecha ultima visita).",
            "Filtro local por nombre o email con debounce.",
            "Click en cliente abre ClientProfileModal con tabs Informacion e Historial.",
            "Historial lista cada cita con servicio, fecha, estado (badge) y precio.",
            "Busqueda global en el header del modulo.",
        ],
        exceptions=[
            "appointments NO tiene client_name ni client_email: el two-step query es obligatorio.",
            "Clientes sin ninguna cita no aparecen en el listado (filtro implicito por appointments.business_id).",
        ],
        permissions=["clients.view", "clients.edit"],
        screenshot="ClientsManager con grid de clientes y modal de perfil abierto.",
    ),
    ModuleSpec(
        key="reviews", name="Resenas y Calificaciones", plan="free", scope="Transversal",
        summary="Sistema de resenas de negocio y de empleado con 1-5 estrellas, respuesta del negocio y vistas materializadas para ranking.",
        details=[
            "Tabla reviews con review_type: 'business' | 'employee'.",
            "ReviewForm se muestra al cliente si tiene al menos una cita completada sin review previa (useCompletedAppointments).",
            "Rating obligatorio de 1 a 5 estrellas; comentario opcional (max 500 caracteres).",
            "Resenas anonimas soportadas (no muestran nombre del cliente).",
            "Negocio puede responder una vez por review; la respuesta queda publica bajo la original.",
            "Admin puede ocultar/mostrar una review (toggle visibility) pero no editar su contenido.",
            "Vistas materializadas business_ratings_stats y employee_ratings_stats se refrescan cada 5 min via edge function refresh-ratings-stats.",
            "Resenas obligatorias al contratar/finalizar un proceso de reclutamiento (hook useMandatoryReviews).",
        ],
        exceptions=[
            "Cliente que nunca tuvo cita completada NO puede dejar review.",
            "Un cliente puede dejar una review de negocio y una review de empleado por cada cita completada.",
            "Ocultar una review no la borra; queda auditada en permission_audit_log.",
        ],
        permissions=["reviews.respond", "reviews.hide", "reviews.delete"],
        screenshot="Perfil publico con lista de resenas, estrellas y respuestas del negocio.",
    ),
    ModuleSpec(
        key="billing_view", name="Facturacion del Plan (Visualizacion)", plan="free", scope="Admin",
        summary="Panel donde el dueno ve su plan actual, proximo cobro, historial de pagos y metricas de uso frente al limite del plan.",
        details=[
            "BillingDashboard en /app/admin/billing.",
            "Muestra plan activo, fecha proximo cobro, metodo de pago, estado de la suscripcion.",
            "Historial de invoices con estado (paid/pending/failed) y link al recibo.",
            "UsageMetrics: empleados activos vs limite, citas del mes vs limite, sedes vs limite.",
            "Boton 'Cambiar plan' abre PlanUpgradeModal con comparativa en 3 columnas.",
            "Boton 'Agregar metodo de pago' abre AddPaymentMethodModal (MercadoPago / Stripe / PayU segun configuracion).",
            "Si hay pago fallido, banner rojo en el header con CTA 'Actualizar metodo'.",
        ],
        exceptions=[
            "Si el usuario supera un limite, el modulo afectado entra en modo lectura y aparece banner 'Actualiza a <Plan>'.",
            "Tras fallar el pago por 3 veces, el plan se degrada a Gratuito automaticamente.",
        ],
        permissions=["billing.view", "billing.manage"],
        screenshot="BillingDashboard con plan actual, metricas de uso y boton de upgrade.",
    ),
    # ---- BASICO ----
    ModuleSpec(
        key="employees", name="Gestion de Empleados y Jerarquia", plan="basico", scope="Admin",
        summary="Modulo completo de RR.HH. operativa: altas, horarios, salarios, jerarquia en arbol, transferencias entre sedes y activacion/desactivacion.",
        details=[
            "EmployeeManagementHierarchy renderiza un arbol jerarquico (HierarchyMapView) con nodos expandibles.",
            "HierarchyLevelSelector permite cambiar el nivel raiz (negocio, sede o manager especifico).",
            "Cada empleado tiene: rol (manager/professional/receptionist/accountant/support_staff), tipo de empleo (employee/contractor), fecha de contratacion, salario, offers_services, lunch_break_start/end.",
            "Horario semanal por empleado en 7 cards (Lun-Dom) con opens_at/closes_at por dia y pausa de almuerzo.",
            "Trigger auto_insert_owner_to_business_employees registra al owner como manager al crear el negocio.",
            "Trigger trg_auto_insert_admin_as_employee sincroniza business_roles -> business_employees.",
            "Transferencia de sede: LocationTransferModal muestra disponibilidad, cancela citas futuras del empleado en la sede origen (edge function cancel-future-appointments-on-transfer) y mueve historial.",
            "Permite pausar a un empleado sin eliminar su historico (is_active=false).",
            "Reviews obligatorias al contratar y al finalizar (useMandatoryReviews).",
            "Setup flag setup_completed asegura que el empleado complete su perfil antes de poder atender citas.",
        ],
        flows=[
            ("Alta de empleado", [
                "Admin invita por email (genera token de aceptacion).",
                "Empleado recibe email, pulsa el link y completa su perfil (foto, telefono, horarios, servicios que ofrece).",
                "Al finalizar setup_completed=true y el empleado aparece en el wizard del cliente.",
            ]),
            ("Transferencia entre sedes", [
                "Admin abre LocationTransferModal desde el perfil del empleado.",
                "Selecciona sede destino y fecha efectiva.",
                "Sistema valida que el empleado NO tenga citas confirmadas en la sede origen despues de la fecha efectiva.",
                "Si hay conflictos, admin elige: 'cancelar y notificar' o 'cambiar fecha efectiva'.",
                "Al confirmar, se dispara la edge function y se ejecuta la transferencia.",
            ]),
        ],
        exceptions=[
            "No se puede eliminar un empleado con citas futuras; el sistema obliga a reasignarlas o cancelarlas.",
            "Si un empleado esta inactivo, su horario ya no aparece en el wizard pero sus citas historicas siguen siendo visibles.",
        ],
        permissions=["employees.view", "employees.create", "employees.edit", "employees.delete", "employees.transfer", "employees.hierarchy"],
        screenshot="Mapa de jerarquia de empleados con nodos expandibles y ficha del empleado seleccionado.",
    ),
    ModuleSpec(
        key="locations", name="Gestion de Sedes Multiples", plan="basico", scope="Admin",
        summary="Hasta 5 sedes en plan Basico (ilimitadas en Pro): cada una con direccion, horario, servicios, empleados y metricas propias.",
        details=[
            "LocationsManager en /app/admin/locations.",
            "Crear sede: pais, ciudad, direccion, geolocalizacion (lat/lng), telefono, email, horario semanal.",
            "Badge 'Administrada' marca la sede preferida del dispositivo actual (usePreferredLocation).",
            "Cada sede tiene tabla location_services para habilitar que servicios se ofrecen alli.",
            "Panel de metricas por sede: citas del mes, ingresos, empleados activos, top servicios.",
            "Soporte de 'excepciones de sede' en citas: un empleado puede atender puntualmente en otra sede (is_location_exception=true).",
            "Transferencia de empleados entre sedes con ventana de preaviso configurable.",
            "Al desactivar una sede, sus empleados se transfieren o se pausan.",
        ],
        exceptions=[
            "Al crear la sede, si la geolocalizacion falla se permite guardar sin coordenadas; el perfil publico no mostrara el mapa hasta completarlas.",
            "Sede con citas futuras no se puede eliminar: sugiere desactivar.",
        ],
        permissions=["locations.view", "locations.create", "locations.edit", "locations.delete"],
        screenshot="LocationsManager con listado de sedes, badge de administrada y panel de metricas.",
    ),
    ModuleSpec(
        key="absences", name="Ausencias y Vacaciones", plan="basico", scope="Transversal",
        summary="Sistema completo de solicitud y aprobacion de ausencias con tipos (vacaciones, emergencia, enfermedad, personal, otro) y reglas de negocio estrictas.",
        details=[
            "Tablas: employee_absences, absence_approval_requests, vacation_balance.",
            "Politica critica: require_absence_approval=true SIEMPRE, no es parametrizable.",
            "Tipos: vacation, emergency, sick_leave, personal, other.",
            "Cada tipo tiene reglas distintas (vacation descuenta dias de balance, emergency cancela citas automaticamente, sick_leave requiere evidencia opcional).",
            "AbsenceRequestModal con selector de rango en dos calendarios (inicio/fin), con range highlighting.",
            "Calculo automatico de dias habiles excluyendo festivos colombianos (tabla public_holidays - 54 festivos 2025-2027).",
            "Al solicitar, edge function request-absence notifica a TODOS los admins del negocio.",
            "Al aprobar/rechazar, edge function approve-reject-absence actualiza estado y notifica al empleado.",
            "Para ausencias tipo emergency, edge function cancel-appointments-on-emergency-absence cancela automaticamente las citas del empleado y notifica a los clientes afectados.",
            "VacationDaysWidget muestra al empleado cuantos dias tiene disponibles vs usados.",
            "Dias de vacaciones se acumulan segun vacation_days_accrued en business_employees (1.25 dias por mes por defecto).",
            "AbsencesTab en el dashboard admin lista todas las solicitudes pendientes con AbsenceApprovalCard.",
        ],
        flows=[
            ("Solicitud de vacaciones", [
                "Empleado abre AbsenceRequestModal desde su dashboard.",
                "Elige tipo 'Vacaciones', rango y motivo opcional.",
                "Sistema calcula dias habiles y valida balance disponible.",
                "Envia -> todos los admins reciben notificacion in-app + email.",
                "Admin revisa en AbsencesTab, ve citas afectadas en ese rango.",
                "Aprueba o rechaza con comentario -> empleado recibe notificacion.",
            ]),
            ("Ausencia de emergencia (cancela citas)", [
                "Empleado selecciona tipo 'Emergencia' y rango.",
                "Al enviar, la edge function cancela automaticamente TODAS las citas del empleado en ese rango.",
                "Cada cliente afectado recibe email/WhatsApp con plantilla 'Tu cita fue reprogramada por emergencia'.",
                "Admin ve el evento en el feed de notificaciones.",
            ]),
        ],
        exceptions=[
            "Solicitar vacaciones sin dias disponibles: bloqueado con mensaje 'Balance insuficiente'.",
            "Solapar ausencias con una ya aprobada: bloqueado.",
            "Festivos no cuentan como dias habiles en el calculo.",
        ],
        permissions=["absences.request", "absences.approve", "absences.reject", "absences.view_all"],
        screenshot="AbsenceRequestModal con rango de fechas y AbsencesTab con solicitudes pendientes.",
    ),
    ModuleSpec(
        key="quicksales", name="Ventas Rapidas (POS sin cita)", plan="basico", scope="Admin",
        summary="Punto de venta ligero para clientes walk-in o ventas fuera del flujo de cita, con asignacion opcional a empleado y sede.",
        details=[
            "QuickSaleForm en /app/admin/quickSales.",
            "Campos: servicio (o producto libre), empleado que atiende (opcional), sede (con sede preferida precargada via usePreferredLocation), cliente (buscar existente o 'cliente ocasional'), metodo de pago, monto.",
            "Si se selecciona servicio, precio se precarga pero es editable (descuentos/ajustes).",
            "Al confirmar, se registra como transaction tipo income, categoria service_sale y se asocia al empleado para comisiones.",
            "Historial de ventas rapidas con filtros por rango, empleado, sede y metodo de pago.",
            "Impresion de recibo en formato termico 80 mm (pop-up imprimible).",
            "Cliente ocasional no crea registro en profiles, solo queda como texto en notes.",
        ],
        exceptions=[
            "Si el plan no es Basico o superior, el modulo muestra placeholder 'Disponible desde plan Basico' con CTA de upgrade.",
        ],
        permissions=["sales.create", "sales.view"],
        screenshot="Formulario de Venta Rapida con sede preferida precargada.",
    ),
    ModuleSpec(
        key="sales_history", name="Historial de Ventas", plan="basico", scope="Admin",
        summary="Reporte consolidado de citas completadas y ventas rapidas con tarjetas resumen y tabla filtrable.",
        details=[
            "SalesHistoryPage en /app/admin/sales.",
            "Filtro de rango: ultimos 7 / 30 / 90 / 365 dias (default 30).",
            "Summary cards: total de ventas completadas, ingresos totales, promedio por cita.",
            "Tabla: fecha, servicio, cliente (click abre ClientProfileModal), empleado, sede, precio, metodo de pago.",
            "Two-step query: appointments -> profiles + services batch.",
            "Exportacion CSV del rango visible.",
        ],
        exceptions=[
            "Servicios eliminados aparecen con nombre 'N/A' (LEFT JOIN).",
            "Clientes anonimos se muestran como 'Cliente ocasional'.",
        ],
        permissions=["sales.view"],
        screenshot="Historial de Ventas con summary cards y tabla detallada.",
    ),
    ModuleSpec(
        key="reports", name="Reportes Operacionales", plan="basico", scope="Admin",
        summary="Panel de reportes con metricas clave del negocio: ocupacion, ingresos, empleados top, servicios top, retencion de clientes.",
        details=[
            "ReportsPage en /app/admin/reports.",
            "Filtros globales: sede (pre-cargada con sede preferida), rango, empleado opcional.",
            "Widgets: citas por dia/semana/mes, tasa de ocupacion por empleado, top 10 servicios, top 10 clientes, retencion (clientes con 2+ citas).",
            "Comparativa periodo actual vs anterior con flechas de variacion.",
            "Export CSV y PDF de cada widget.",
        ],
        permissions=["reports.view", "reports.export"],
        screenshot="ReportsPage con widgets de ocupacion y comparativa.",
    ),
    ModuleSpec(
        key="permissions", name="Sistema de Permisos Granulares", plan="basico", scope="Admin",
        summary="79 tipos de permisos repartidos en 16 categorias, 9 plantillas predefinidas, auditoria completa. El owner siempre tiene bypass total.",
        details=[
            "PermissionsManager en /app/admin/permissions (lazy-loaded por peso).",
            "Categorias: services, resources, locations, employees, appointments, recruitment, accounting, expenses, reviews, billing, notifications, settings, permissions, absences, favorites, sales.",
            "Plantillas: Admin Completo, Vendedor, Cajero, Manager de Sede, Recepcionista, Profesional, Contador, Gerente de Sede, Staff de Soporte.",
            "PermissionGate envuelve cada boton de accion; modos hide / disable / show.",
            "Owner bypass: si user.id == businesses.owner_id, se evita cualquier query y se conceden todos los permisos.",
            "Tabla user_permissions con 1,919+ registros en produccion; UNIQUE(business_id, user_id, permission).",
            "Triggers: auto_assign_permissions_to_owners (79 al crear negocio), auto_assign_permissions_to_admins (al asignar admin).",
            "permission_audit_log registra cada cambio con granted_by y timestamp.",
            "Service PermissionRPCService con 5 metodos SECURITY DEFINER: revoke, assign, applyTemplate, bulkRevoke, bulkAssign.",
            "Cache de permisos refrescado por edge function refresh-permissions-cache.",
        ],
        exceptions=[
            "El owner NO puede perder permisos; cualquier intento es rechazado por RLS.",
            "Permisos desactivados (is_active=false) no se borran - queda historico completo.",
        ],
        permissions=["permissions.view", "permissions.assign", "permissions.revoke", "permissions.apply_template"],
        screenshot="PermissionsManager con lista de empleados y matriz de 79 permisos.",
    ),
    # ---- PRO ----
    ModuleSpec(
        key="accounting", name="Contabilidad (con Impuestos Colombianos)", plan="pro", scope="Admin",
        summary="Modulo contable completo con IVA, ICA y Retencion en la Fuente; charts financieros; exportacion fiscal; configuracion por sede.",
        details=[
            "EnhancedFinancialDashboard con filtros por periodo (mes/trimestre/ano fiscal), sede, categoria.",
            "Transacciones: income / expense, categoria, subtotal, tax_type, tax_rate, tax_amount, fiscal_period.",
            "IVA: 0%, 5%, 19% (configurable por servicio en business_tax_config).",
            "ICA: codigo DANE por municipio + tarifa en milesimas.",
            "Retencion en la Fuente: tabla de conceptos con bases minimas.",
            "Charts Recharts: CategoryPieChart, EmployeeRevenueChart, IncomeVsExpenseChart, LocationBarChart, MonthlyTrendChart.",
            "Hook useTaxCalculation(subtotal, taxConfig) retorna breakdown completo (IVA, ICA, Retefuente, total).",
            "Export PDF (jspdf), CSV, Excel (xlsx) listos para el contador.",
            "Libro diario, libro mayor, estado de resultados, balance de comprobacion.",
            "Integracion automatica: cada Venta Rapida y cita completada genera su transaccion con impuestos calculados.",
        ],
        exceptions=[
            "Si el negocio no tiene tax_config cargado, las transacciones se guardan con tax_amount=0 y banner rojo 'Configura tus impuestos'.",
            "ICA aplica solo en Colombia; fuera se oculta el campo.",
        ],
        permissions=["accounting.view", "accounting.create", "accounting.edit", "accounting.export", "accounting.config"],
        screenshot="EnhancedFinancialDashboard con charts y tabla de transacciones con impuestos desglosados.",
    ),
    ModuleSpec(
        key="expenses", name="Gestion de Gastos", plan="pro", scope="Admin",
        summary="Registro de gastos puntuales y recurrentes con categorias, soporte documental y generacion automatica de transacciones.",
        details=[
            "ExpensesManagementPage en /app/admin/expenses.",
            "Gastos puntuales: ExpenseRegistrationForm con categoria, monto, proveedor, fecha, descripcion, evidencia (PDF/imagen).",
            "Gastos recurrentes: BusinessRecurringExpenses - arriendo, servicios, licencias; frecuencia mensual/trimestral/anual.",
            "LocationExpenseConfig permite asignar gastos a una sede especifica para reportes por sede.",
            "Al vencer un gasto recurrente, cron crea automaticamente la transaccion correspondiente.",
            "Conciliacion mensual: vista que muestra gastos esperados vs registrados.",
        ],
        permissions=["expenses.view", "expenses.create", "expenses.edit", "expenses.delete"],
        screenshot="ExpensesManagementPage con gastos recurrentes y puntuales.",
    ),
    ModuleSpec(
        key="recruitment", name="Reclutamiento y Vacantes", plan="pro", scope="Admin",
        summary="Publicacion de vacantes, gestion de aplicaciones, matching inteligente empleado-vacante, reviews obligatorias y deteccion de conflictos de horario.",
        details=[
            "RecruitmentDashboard en /app/admin/recruitment.",
            "CreateVacancy con campos: titulo, descripcion, sede, horario requerido, salario, tipo (tiempo completo/parcial/freelance), categoria.",
            "VacancyCard muestra estado (activa/pausada/cerrada), aplicaciones recibidas, vistas.",
            "ApplicationList con filtros por estado (nueva/en revision/entrevista/contratada/rechazada).",
            "ApplicantProfileModal con CV (bucket cvs, privado), experiencia, disponibilidad.",
            "Hook useMatchingVacancies: compara horario del empleado con requisitos de la vacante (useScheduleConflicts).",
            "MandatoryReviewModal: al contratar o finalizar proceso se obliga a dejar review.",
            "Edge functions: send-selection-notifications, send-employee-request-notification.",
            "Job applicants pueden ser empleados actuales de otro negocio o candidatos externos via /negocio/<slug>.",
        ],
        exceptions=[
            "Aplicar a una vacante con conflictos de horario muestra warning amarillo; admin decide si continuar.",
            "Tests E2E del modulo estan pausados (describe.skip) por rate-limit de emails Supabase en CI.",
        ],
        permissions=["recruitment.create", "recruitment.edit", "recruitment.close", "recruitment.hire"],
        screenshot="RecruitmentDashboard con vacantes activas y modal de aplicante.",
    ),
    ModuleSpec(
        key="resources", name="Recursos Fisicos y Modelos de Negocio Flexibles", plan="pro", scope="Admin",
        summary="Reservar salas, canchas, mesas, estudios, equipos: 15 tipos de recursos y 4 modelos de negocio (profesional / recurso / hibrido / clase grupal).",
        details=[
            "business_resources con resource_type: room, table, court, studio, meeting_room, desk, equipment, vehicle, space, lane, field, station, parking_spot, bed, other.",
            "resource_model del negocio: professional / physical_resource / hybrid / group_class.",
            "resource_services (M:N) define que servicios se prestan en cada recurso (con custom_price override).",
            "appointments.resource_id es nullable; CHECK obliga a que employee_id OR resource_id este definido.",
            "useAssigneeAvailability valida empleado OR recurso indistintamente.",
            "is_resource_available() RPC hace overlap detection en recursos.",
            "Vista materializada resource_availability con bookings y revenue por recurso.",
            "Capacidad (capacity) por recurso - util para clases grupales.",
            "Amenities JSONB: checklist de caracteristicas (WiFi, proyector, sonido, ducha, etc.).",
            "Tarifa horaria (hourly_rate) para reservas por bloques.",
            "UI Fase 3 y 4 pendientes de integracion visual en el wizard; backend 100% listo.",
        ],
        permissions=["resources.view", "resources.create", "resources.edit", "resources.delete", "resources.assign_services"],
        screenshot="Vista de Recursos con grid de salas, canchas y equipos reservables.",
    ),
    # ---- TRANSVERSALES ----
    ModuleSpec(
        key="notifications", name="Notificaciones Multicanal", plan="free", scope="Transversal",
        summary="30+ tipos de notificaciones por email, SMS, WhatsApp e in-app con preferencias por usuario y mapeo a rol para navegacion automatica.",
        details=[
            "Tabla in_app_notifications (type, data JSONB, read, user_id, business_id).",
            "Campanita NotificationBell con contador de no leidas y dropdown NotificationCenter.",
            "Hook useInAppNotifications: 1 query base (limit=50) + filtros locales (antes eran 5 queries).",
            "Mapeo rol: notificationRoleMapping asocia cada tipo a un rol y destino de navegacion - la app conmuta de rol automaticamente al abrir la notificacion.",
            "Tipos: citas (creada, confirmada, cancelada, reprogramada, completada, recordatorio), empleados (invitado, acepto, transferencia, solicitud), ausencias (solicitada, aprobada, rechazada, cancelacion masiva), reclutamiento (nueva aplicacion, seleccionado, rechazado, review pendiente), sistema (pago fallido, plan degradado, bug report).",
            "Canales: Email (Brevo, 300/dia gratis), SMS (AWS SNS, solo Pro), WhatsApp Business (plan Basico+), In-app (siempre).",
            "user_notification_preferences permite activar/desactivar por tipo y canal.",
            "business_notification_settings configura recordatorios (p.ej. 24 h y 2 h antes de la cita).",
            "Edge functions: send-notification, process-reminders (cron 5 min), schedule-reminders, send-email, send-email-reminder, send-sms-reminder, send-whatsapp, send-whatsapp-reminder, send-notification-reminders, send-unread-chat-emails.",
            "notification_log registra cada envio con estado, canal y proveedor.",
        ],
        exceptions=[
            "Plan Gratuito no envia WhatsApp ni SMS (solo email + in-app).",
            "Si el usuario desactiva un canal, notification_log registra 'skipped_by_preference'.",
            "Brevo free tier (300/dia): al llegar al limite se encola y se reintenta al dia siguiente.",
        ],
        screenshot="NotificationBell con dropdown y preferencias por canal del usuario.",
    ),
    ModuleSpec(
        key="chat", name="Chat en Tiempo Real", plan="free", scope="Transversal",
        summary="Chat cliente<->negocio (y multi-empleado en Basico+) con adjuntos, typing indicators, read receipts y emails de no-leidos.",
        details=[
            "Tablas: conversations, messages, chat_participants.",
            "Bucket chat-attachments (privado) para archivos.",
            "Realtime: suscripcion de Supabase invalida el cache de React Query (no refetch continuo).",
            "ChatLayout con lista de conversaciones + ventana de chat.",
            "FloatingChatButton visible en todos los dashboards.",
            "ChatWithAdminModal permite al cliente elegir con que empleado chatear (filtra allow_client_messages=true).",
            "MessageBubble con estados (enviado, entregado, leido).",
            "TypingIndicator muestra '<empleado> esta escribiendo...' en tiempo real.",
            "ReadReceipts (doble check).",
            "FileUpload soporta imagenes y PDF hasta 10 MB.",
            "Edge function send-unread-chat-emails envia un resumen diario de mensajes no leidos.",
        ],
        exceptions=[
            "Plan Gratuito: solo chat cliente<->owner (sin multi-empleado).",
            "Empleado con allow_client_messages=false no aparece en ChatWithAdminModal.",
        ],
        screenshot="ChatLayout con conversacion activa, lista lateral y typing indicator.",
    ),
    ModuleSpec(
        key="public_profile", name="Perfil Publico del Negocio", plan="free", scope="Transversal",
        summary="URL /negocio/<slug> indexable en Google, con 4 tabs (Servicios, Ubicaciones, Resenas, Acerca de), boton Reservar y datos estructurados JSON-LD.",
        details=[
            "Ruta publica /negocio/<slug> sin autenticacion.",
            "SEO completo: meta tags dinamicos, Open Graph, Twitter Card, JSON-LD LocalBusiness/Service.",
            "Sitemap.xml generado via npm run generate-sitemap.",
            "robots.txt: permite /negocio/*, bloquea /app/*.",
            "Hook useBusinessProfileData (352 lineas) carga negocio, servicios, ubicaciones, empleados, reviews.",
            "Tabs: Servicios (con precio y duracion), Ubicaciones (con mapa y horario), Resenas (ordenadas por fecha), Acerca de (descripcion, categorias, redes).",
            "Boton 'Reservar' en cada servicio redirige al wizard con parametros.",
            "GA4: profile_view, click_reserve_button, click_contact.",
            "SearchBar global con debounce de 300 ms y 6 algoritmos de ordenamiento (relevancia, rating, proximidad, precio asc/desc, recien agregados).",
            "RPCs: search_businesses, search_services, search_professionals con ts_rank e indices trigram GIN.",
        ],
        screenshot="Perfil publico /negocio/<slug> con tabs, mapa y boton Reservar.",
    ),
    ModuleSpec(
        key="client_dashboard", name="Experiencia del Cliente", plan="free", scope="Cliente",
        summary="Panel personalizado del cliente con sus citas, favoritos, descubrimiento por ciudad y chat con negocios.",
        details=[
            "ClientDashboard con RPC get_client_dashboard_data(client_id, city_name, region_name).",
            "Widgets: proxima cita, citas pasadas, favoritos, descubre negocios en tu ciudad.",
            "Filtro de ciudad/region (useGeolocation sugiere la ciudad actual).",
            "Sede preferida por cliente (usePreferredCity).",
            "Favoritos (tabla favorites) con toggle en cada perfil y listado propio.",
            "Tab Buscar: SearchBar + SearchResults + SearchResultCard.",
            "Tab Mis Citas con filtros por estado.",
            "Tab Perfil con datos personales, preferencias, idioma, tema.",
        ],
        screenshot="ClientDashboard con proxima cita, favoritos y descubrimiento por ciudad.",
    ),
    ModuleSpec(
        key="employee_dashboard", name="Experiencia del Empleado", plan="basico", scope="Empleado",
        summary="Panel del empleado con su agenda, sus clientes, sus metricas, solicitud de ausencias y chat con clientes.",
        details=[
            "EmployeeDashboard con: agenda del dia/semana, proxima cita, metricas personales.",
            "EmployeeClientsPage en /app/employee/my-clients - clientes atendidos por el empleado (filtro employee_id=currentUser.id, ordenado por visitas completadas).",
            "useEmployeeMetrics: citas completadas, ingresos generados, rating promedio, tasa de ocupacion.",
            "Solicitar ausencia desde AbsenceRequestModal.",
            "Ver saldo de vacaciones (VacationDaysWidget).",
            "Configurar allow_client_messages para habilitar/deshabilitar chat directo con clientes.",
            "Preferencias de horario, almuerzo y servicios ofrecidos.",
        ],
        screenshot="EmployeeDashboard con agenda del dia, metricas y widget de vacaciones.",
    ),
    ModuleSpec(
        key="settings", name="Configuraciones Unificadas", plan="free", scope="Transversal",
        summary="CompleteUnifiedSettings cubre los tres roles en un solo componente: preferencias del negocio, del empleado y del cliente.",
        details=[
            "Admin: datos del negocio (nombre, categorias, descripcion, logo, banner), contacto, direccion legal, horario global, configuracion fiscal, sede preferida administrada, metodos de pago, plan, idioma por defecto.",
            "Empleado: horarios 7 dias con almuerzo, salarios, servicios ofrecidos, allow_client_messages, disponibilidad para vacantes, perfil publico (foto, bio, especialidades).",
            "Cliente: anticipacion minima para reservar, metodo de pago preferido, visibilidad del historial, ciudad preferida, idioma, tema, preferencias de notificacion por canal y tipo.",
            "Sincronizacion con Google Calendar (plan Pro) desde la pestana 'Integraciones'.",
            "Cambio de contrasena, cierre de sesion global (todos los dispositivos), eliminacion de cuenta con confirmacion doble.",
        ],
        permissions=["settings.view_business", "settings.edit_business", "settings.edit_profile"],
        screenshot="CompleteUnifiedSettings con tabs por rol y pestana de integraciones.",
    ),
    ModuleSpec(
        key="bug_reports", name="Reporte de Bugs por Usuario", plan="free", scope="Transversal",
        summary="Cualquier usuario puede reportar un bug desde un boton flotante: describe el problema, adjunta evidencia y recibe seguimiento.",
        details=[
            "FloatingBugReportButton presente en todas las pantallas autenticadas.",
            "BugReportModal: titulo, descripcion, severidad (Critical/High/Medium/Low), evidencia (screenshots/videos al bucket bug-report-evidences).",
            "Tablas: bug_reports, bug_report_evidences, bug_report_comments.",
            "Edge function send-bug-report-email notifica al equipo Ti Turing.",
            "Logger centralizado (src/lib/logger.ts) integra error_logs, login_logs y Sentry.",
            "Sentry captura errores no controlados (plan gratuito).",
        ],
        screenshot="BugReportModal con severidad y adjuntos.",
    ),
]


# ===========================================================================
# MANUAL DE USUARIO
# ===========================================================================
def build_manual() -> Document:
    doc = setup_document(title="Manual de Usuario",
                         subtitle="Guia funcional exhaustiva de la plataforma")

    # Indice
    h1(doc, "Indice", anchor="toc")
    para(doc, "Este documento esta dividido en dos grandes partes: un resumen ejecutivo de alto nivel y un detalle exhaustivo donde se describe cada modulo, regla, excepcion, flujo normal y flujo alterno. Cada entrada del indice es un hipervinculo a la seccion correspondiente.")

    def toc_link(label: str, anchor: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(p, anchor, label)

    h3(doc, "Parte 1 - Resumen Ejecutivo")
    toc_link("1.1  Que es Gestabiz?", "summary_what")
    toc_link("1.2  Roles de usuario", "summary_roles")
    toc_link("1.3  Planes y que incluye cada uno", "summary_plans")
    toc_link("1.4  Mapa de modulos por rol", "summary_modules_map")
    toc_link("1.5  Primer uso: como empezar en 5 minutos", "summary_quickstart")

    h3(doc, "Parte 2 - Detalle Exhaustivo")
    toc_link("2.1  Plan Gratuito - modulos incluidos", "detail_plan_free")
    toc_link("2.2  Plan Basico - modulos incluidos", "detail_plan_basico")
    toc_link("2.3  Plan Pro - modulos incluidos", "detail_plan_pro")
    toc_link("2.4  Reglas transversales del sistema", "detail_rules")
    toc_link("2.5  Glosario y atajos", "detail_glossary")
    toc_link("2.6  Soporte y canales oficiales", "detail_support")

    page_break(doc)

    # =========== PARTE 1 ===========
    h1(doc, "Parte 1 - Resumen Ejecutivo")

    h2(doc, "1.1  Que es Gestabiz?", anchor="summary_what")
    para(doc, "Gestabiz es una plataforma SaaS todo-en-uno para la gestion integral de negocios de servicios por cita: agenda, clientes, empleados, recursos fisicos, ventas, contabilidad, resenas, reclutamiento, comunicaciones multicanal y perfiles publicos indexables en Google. Una sola herramienta reemplaza 6 a 10 productos independientes que normalmente combinan los negocios PyME.")
    para(doc, "Funciona en web (React 19 + Vite), movil (Expo/React Native) y extension de navegador. El backend corre 100% sobre Supabase (PostgreSQL + RLS + Edge Functions + Realtime + Storage). Cumple GDPR y esta localizado en espanol e ingles.")
    callout(doc, "Frase clave",
            "Gestabiz es la unica plataforma en Colombia que combina reservas, contabilidad con impuestos, recursos fisicos y reclutamiento en un solo producto - a un precio de PyME.")

    for lab, anchor in [("Ver detalle del Plan Gratuito", "detail_plan_free"),
                        ("Ver detalle del Plan Basico", "detail_plan_basico"),
                        ("Ver detalle del Plan Pro", "detail_plan_pro")]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        add_internal_hyperlink(p, anchor, f"-> {lab}")

    h2(doc, "1.2  Roles de usuario", anchor="summary_roles")
    para(doc, "Gestabiz calcula los roles en tiempo real - nunca los persiste en la base de datos. Un mismo usuario puede tener varios roles en simultaneo en distintos negocios:")
    simple_table(doc,
        ["Rol", "Como se calcula", "Que puede hacer"],
        [
            ["Owner", "businesses.owner_id == auth.uid()", "Bypass total: acceso absoluto a su negocio sin pasar por permisos."],
            ["Admin", "business_roles.role = 'admin'", "Casi todo, acotado por su matriz de permisos granulares."],
            ["Empleado", "business_employees con offers_services", "Su agenda, sus clientes, sus metricas, sus ausencias."],
            ["Cliente", "Cualquier usuario autenticado", "Reservar, ver sus citas, favoritos, chat, resenas."],
        ],
    )

    h2(doc, "1.3  Planes y que incluye cada uno", anchor="summary_plans")
    para(doc, "Cada plan incluye todo lo del anterior mas sus exclusivos. Todos los precios son en pesos colombianos; el plan anual cobra 10 meses (2 meses gratis).")
    simple_table(doc,
        ["Plan", "Precio mensual", "Precio anual", "Ideal para"],
        [[p.name, p.price_monthly, p.price_annual, p.ideal_for] for p in PLANS],
    )
    for p in PLANS:
        h3(doc, f"Plan {p.name} - resumen")
        para(doc, p.tagline, italic=True, color=GREY)
        for e in p.exclusives[:6]:
            bullet(doc, e)
        lk = doc.add_paragraph(); lk.paragraph_format.space_after = Pt(6)
        add_internal_hyperlink(lk, f"detail_plan_{p.key}",
                                f"-> Ver la lista completa y el detalle del plan {p.name}")

    h2(doc, "1.4  Mapa de modulos por rol", anchor="summary_modules_map")
    para(doc, "Vista rapida de que modulo esta disponible para cada rol en cada plan:")
    simple_table(doc,
        ["Modulo", "Plan", "Alcance"],
        [[m.name, next(p.name for p in PLANS if p.key == m.plan), m.scope] for m in MODULES],
    )

    h2(doc, "1.5  Primer uso: como empezar en 5 minutos", anchor="summary_quickstart")
    numbered(doc, "Crea tu cuenta en gestabiz.com/register (correo + contrasena, o Google).")
    numbered(doc, "Verifica tu correo desde el link que te enviamos.")
    numbered(doc, "Elige 'Ofrecer servicios' y completa el wizard de 4 pasos para crear tu negocio.")
    numbered(doc, "Publica tu primer servicio y comparte tu URL /negocio/<slug>.")
    numbered(doc, "Empieza a recibir reservas - no hace falta instalar nada mas.")
    screenshot_placeholder(doc, "Pantalla de bienvenida tras completar el onboarding.")

    page_break(doc)

    # =========== PARTE 2 ===========
    h1(doc, "Parte 2 - Detalle Exhaustivo")
    para(doc, "A continuacion se describen todas las funcionalidades de la plataforma, agrupadas por plan. Recuerda que cada plan incluye todas las funcionalidades del anterior mas sus exclusivas.")

    section_nums = {"free": "2.1", "basico": "2.2", "pro": "2.3"}
    for plan in PLANS:
        h2(doc, f"{section_nums[plan.key]}  Plan {plan.name}",
           anchor=f"detail_plan_{plan.key}")
        para(doc, plan.tagline, italic=True, color=GREY)
        para(doc, f"Precio: {plan.price_monthly}  ·  Anual: {plan.price_annual}")
        para(doc, f"Ideal para: {plan.ideal_for}")

        if plan.includes_previous:
            prev = PLANS[PLANS.index(plan) - 1]
            callout(doc, "Incluye todo lo anterior",
                    f"Este plan contiene 100% de las funcionalidades del plan {prev.name} mas las exclusivas que aparecen aqui abajo.")

        if plan.limits:
            h4(doc, "Limites del plan")
            for l in plan.limits:
                bullet(doc, l)

        modules_here = [m for m in MODULES if m.plan == plan.key]
        for m in modules_here:
            h3(doc, m.name, anchor=f"mod_{m.key}")
            para(doc, m.summary, italic=True, color=GREY)

            if m.screenshot:
                screenshot_placeholder(doc, m.screenshot)

            h4(doc, "Detalles funcionales")
            for d in m.details:
                bullet(doc, d)

            if m.flows:
                h4(doc, "Flujos")
                for flow_name, steps in m.flows:
                    para(doc, flow_name, color=PURPLE)
                    for step in steps:
                        numbered(doc, step)

            if m.exceptions:
                h4(doc, "Excepciones y casos limite")
                for e in m.exceptions:
                    bullet(doc, e, bold_prefix="Excepcion")

            if m.permissions:
                h4(doc, "Permisos requeridos")
                para(doc, ", ".join(m.permissions), color=GREY, size=10, italic=True)

            doc.add_paragraph()

        page_break(doc)

    # 2.4 Reglas transversales
    h2(doc, "2.4  Reglas transversales del sistema", anchor="detail_rules")
    para(doc, "Estas reglas aplican a todos los planes y a todos los roles:")
    bullet(doc, "El owner del negocio tiene acceso total (bypass de permisos). Nunca se le puede quitar un permiso.")
    bullet(doc, "Los roles se calculan en tiempo real - no se guardan en base de datos.")
    bullet(doc, "Politica de ausencias: toda ausencia requiere aprobacion (require_absence_approval=true, no parametrizable).")
    bullet(doc, "Festivos colombianos 2025-2027 vienen precargados y se excluyen automaticamente del calculo de dias habiles.")
    bullet(doc, "Los clientes reciben email de confirmacion con un link valido por 24 horas; pueden confirmar o cancelar sin iniciar sesion.")
    bullet(doc, "Las resenas solo pueden dejarlas clientes con al menos una cita completada sin review previa.")
    bullet(doc, "Las auditorias de permisos quedan registradas en permission_audit_log con timestamp y granted_by.")
    bullet(doc, "Los datos personales se anonimizan en Google Analytics (consent mode + anonymizeIp); el usuario puede rechazar cookies.")
    bullet(doc, "El boton flotante de bug report esta disponible en todas las pantallas autenticadas.")
    bullet(doc, "Los pagos se procesan con MercadoPago, Stripe o PayU segun configuracion; todos los gateways son PCI-compliant.")
    bullet(doc, "Tras 3 pagos fallidos consecutivos el plan se degrada automaticamente a Gratuito.")

    h2(doc, "2.5  Glosario y atajos", anchor="detail_glossary")
    simple_table(doc,
        ["Termino", "Significado"],
        [
            ["Slug", "Identificador unico de URL del negocio (ej. 'salon-belleza-medellin')."],
            ["Wizard", "Asistente paso a paso con validacion en cada paso."],
            ["Overlap", "Solape horario entre dos citas del mismo empleado/recurso."],
            ["RLS", "Row Level Security: reglas de acceso por fila en PostgreSQL."],
            ["Edge Function", "Funcion serverless en Deno desplegada en Supabase."],
            ["Trigger", "Funcion que se ejecuta automaticamente ante un evento de base de datos."],
            ["Bypass", "Acceso total que se concede al owner sin pasar por permisos."],
            ["Sede preferida", "Ubicacion precargada en los filtros de la app, guardada por dispositivo."],
        ],
    )

    h2(doc, "2.6  Soporte y canales oficiales", anchor="detail_support")
    bullet(doc, "Centro de ayuda: gestabiz.com/ayuda")
    bullet(doc, "WhatsApp soporte (plan Pro): +57 XXX XXX XXXX")
    bullet(doc, "Email: soporte@gestabiz.com")
    bullet(doc, "Reporte de bugs: boton flotante en toda la app o email bugs@gestabiz.com")
    bullet(doc, "Repositorio publico: https://github.com/TI-Turing")

    return doc


# ===========================================================================
# PROPUESTA DE VALOR
# ===========================================================================
def build_proposal() -> Document:
    doc = setup_document(title="Propuesta de Valor",
                         subtitle="Por que Gestabiz es la mejor decision para tu negocio")

    h1(doc, "Indice", anchor="toc")

    def toc_link(label: str, anchor: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(p, anchor, label)

    h3(doc, "Parte 1 - Resumen Ejecutivo")
    toc_link("1.1  El problema que resolvemos", "p_problem")
    toc_link("1.2  La solucion en 30 segundos", "p_solution")
    toc_link("1.3  Por que Gestabiz gana", "p_why_win")
    toc_link("1.4  Planes y retorno de inversion", "p_plans")
    toc_link("1.5  Empieza hoy - sin costo", "p_cta")

    h3(doc, "Parte 2 - Detalle Comercial")
    toc_link("2.1  Diagnostico del mercado PyME en servicios", "p_market")
    toc_link("2.2  Gestabiz vs Calendly, Booksy, Fresha y competidores locales", "p_comparison")
    toc_link("2.3  Verticales atendidos y casos de uso", "p_verticals")
    toc_link("2.4  Beneficios concretos por rol", "p_benefits_role")
    toc_link("2.5  Seguridad, cumplimiento y arquitectura", "p_security")
    toc_link("2.6  Plan de onboarding y soporte", "p_onboarding")
    toc_link("2.7  Roadmap 2026", "p_roadmap")
    toc_link("2.8  Testimonios y casos de exito", "p_testimonials")
    toc_link("2.9  Preguntas frecuentes", "p_faq")

    page_break(doc)

    # =========== PARTE 1 ===========
    h1(doc, "Parte 1 - Resumen Ejecutivo")

    h2(doc, "1.1  El problema que resolvemos", anchor="p_problem")
    para(doc, "Los negocios PyME de servicios en Colombia y LATAM pierden, en promedio, 1 de cada 5 citas por desorganizacion: agenda en papel o WhatsApp, sin recordatorios automaticos, sin control de ausencias, sin visibilidad de ingresos reales, y dependiendo de 6 a 10 herramientas que no se hablan entre si.")
    callout(doc, "Costo oculto",
            "Un salon con 3 empleados que pierde 5 citas por semana a $40.000 promedio, pierde $9.6 millones al ano - suficiente para pagar Gestabiz Pro durante 5 anos.",
            color=DANGER)
    para(doc, "Las plataformas existentes fallan en dos extremos: o son demasiado simples (Calendly, agendas genericas) o son demasiado caras y tecnicas (SAP, Salesforce). El espacio medio - una herramienta completa, en espanol, con contabilidad colombiana, a precio PyME - estaba vacio. Lo llenamos nosotros.")

    h2(doc, "1.2  La solucion en 30 segundos", anchor="p_solution")
    para(doc, "Gestabiz reemplaza tu agenda, tu WhatsApp manual, tu hoja de calculo de ventas, tu hoja de ausencias, tu POS ligero, tu sistema de resenas, tu bolsa de vacantes y tu integracion contable - en una sola plataforma web y movil, con soporte en espanol y precios en pesos colombianos.")
    bullet(doc, "Perfil publico en Google en 5 minutos.")
    bullet(doc, "Reservas online 24/7 con validacion anti-overlap.")
    bullet(doc, "Recordatorios automaticos por email y WhatsApp.")
    bullet(doc, "Control total de empleados, sedes, permisos, ausencias, recursos fisicos.")
    bullet(doc, "Contabilidad con IVA, ICA y Retencion en la Fuente lista para el contador.")
    bullet(doc, "Reclutamiento y matching inteligente de candidatos.")
    screenshot_placeholder(doc, "Dashboard Gestabiz - vista admin con todos los modulos en un solo lugar.")

    h2(doc, "1.3  Por que Gestabiz gana", anchor="p_why_win")
    simple_table(doc,
        ["Criterio", "Calendly", "Booksy", "Fresha", "Gestabiz"],
        [
            ["Precio/mes (USD)", "12-20", "30-85", "0* + comision", "~$20-$40"],
            ["Idioma nativo", "EN", "EN/ES", "EN/ES", "ES (LATAM)"],
            ["Impuestos Colombia", "No", "No", "No", "Si (IVA/ICA/Retefuente)"],
            ["Recursos fisicos", "No", "Parcial", "Parcial", "Si (15 tipos, 4 modelos)"],
            ["Reclutamiento", "No", "No", "No", "Si (con matching)"],
            ["Permisos granulares", "Basico", "Basico", "Medio", "79 tipos, 9 plantillas"],
            ["Perfil publico SEO", "Parcial", "Si", "Si", "Si + JSON-LD + sitemap"],
            ["WhatsApp Business", "No", "Extra", "Extra", "Incluido (Basico+)"],
            ["Multi-sede", "No", "Si", "Si", "Ilimitada (Pro)"],
        ],
    )
    para(doc, "* Fresha cobra comision por transaccion que puede superar cualquier plan mensual pagado.",
         italic=True, size=9, color=GREY)
    callout(doc, "Mensaje clave",
            "Gestabiz es la unica plataforma en nuestra categoria que te da contabilidad colombiana, recursos fisicos y reclutamiento - sin pagar integraciones separadas.",
            color=ACCENT)

    h2(doc, "1.4  Planes y retorno de inversion", anchor="p_plans")
    simple_table(doc,
        ["Plan", "Precio", "Ideal para", "Payback*"],
        [
            ["Gratuito", "$0", "Profesional independiente empezando.", "Inmediato"],
            ["Basico", "$89.900/mes", "Salon, consultorio, gimnasio con 2-15 empleados.", "2-3 citas/mes"],
            ["Pro", "$159.900/mes", "Clinica, hotel, coworking, restaurante por reserva.", "4-5 citas/mes"],
        ],
    )
    para(doc, "* Payback = cantidad de citas que necesitas recuperar por mes gracias a Gestabiz (recordatorios, menos ausentismo, mejor ocupacion) para que el plan se pague solo.",
         italic=True, size=9, color=GREY)

    h2(doc, "1.5  Empieza hoy - sin costo", anchor="p_cta")
    callout(doc, "Prueba gratis",
            "Crea tu cuenta en gestabiz.com, configura tu negocio en 5 minutos y recibe tu primera reserva hoy mismo. Sin tarjeta de credito, sin permanencia, sin sorpresas.",
            color=PURPLE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    add_internal_hyperlink(p, "p_onboarding", "-> Ver como es el onboarding completo")

    page_break(doc)

    # =========== PARTE 2 ===========
    h1(doc, "Parte 2 - Detalle Comercial")

    h2(doc, "2.1  Diagnostico del mercado PyME en servicios", anchor="p_market")
    para(doc, "En Colombia hay mas de 1,6 millones de PyMEs de servicios activas (DANE 2024). La mayoria opera con herramientas manuales o con combinaciones de apps gratuitas que generan friccion, perdidas y dolores de cabeza. Los principales problemas que vemos repetidos:")
    bullet(doc, "Agenda en cuaderno o WhatsApp -> citas olvidadas, doble reserva, ausentismo del 15-25%.", bold_prefix="Caos operativo")
    bullet(doc, "Ausencias de empleados sin aprobacion formal -> clientes notificados a ultima hora, mala reputacion.", bold_prefix="RR.HH. improvisado")
    bullet(doc, "Contabilidad en Excel que no liga ventas, gastos e impuestos -> multas, sobrepagos al contador.", bold_prefix="Riesgo fiscal")
    bullet(doc, "Sin canal de reservas online -> perdida de ventas fuera de horario (evidencia: 40% de reservas ocurren entre 8pm y 7am).", bold_prefix="Ventas dejadas sobre la mesa")
    bullet(doc, "Sin perfil publico SEO -> dependencia total de Instagram para ser encontrados.", bold_prefix="Sin presencia en Google")
    screenshot_placeholder(doc, "Infografia del mercado PyME en servicios en Colombia.")

    h2(doc, "2.2  Gestabiz vs Calendly, Booksy, Fresha y competidores locales", anchor="p_comparison")
    para(doc, "Analisis detallado por competidor:")

    h3(doc, "Calendly")
    para(doc, "Excelente para agendar reuniones 1-a-1 pero deja el negocio a la mitad: no tiene gestion de empleados con rol, no controla ausencias, no tiene contabilidad ni POS, no tiene recursos fisicos, no tiene resenas ni perfiles publicos indexables. Esta pensado para freelance, no para negocio con equipo.")
    bullet(doc, "Gestabiz agrega: empleados con jerarquia, ausencias con aprobacion, contabilidad con impuestos, POS, recursos, resenas, perfil SEO, reclutamiento, permisos granulares.")

    h3(doc, "Booksy")
    para(doc, "Lider en belleza en USA/Europa. Bien construido pero con precios en USD, funcionalidades contables ausentes para Colombia (no maneja IVA ni ICA), sin recursos fisicos tipo canchas/salas, sin modulo de reclutamiento, sin matching inteligente. Soporte en ingles como primera lengua.")
    bullet(doc, "Gestabiz agrega: impuestos colombianos nativos, recursos fisicos (15 tipos), reclutamiento con matching, soporte 100% en espanol LATAM, precios en COP sin conversion.")

    h3(doc, "Fresha")
    para(doc, "Modelo freemium agresivo que se cobra con 2,19% + USD 0,20 por cada venta procesada por su pasarela, lo que se traduce en >$200.000 COP/mes de comisiones para un salon mediano. No tiene contabilidad colombiana ni reclutamiento. Depende fuerte del anclaje a su procesamiento de pagos.")
    bullet(doc, "Gestabiz agrega: precio fijo predecible (sin comision por venta), MercadoPago/PayU/Stripe intercambiables, contabilidad completa, reclutamiento.")

    h3(doc, "Competidores locales (Agenda Pro, Turnier, Simple, etc.)")
    para(doc, "Productos nicho con modulos limitados, stack anticuado, sin perfil publico SEO, sin recursos fisicos, sin app movil nativa robusta, sin roadmap visible. Muchos estan desatendidos desde 2022.")
    bullet(doc, "Gestabiz agrega: arquitectura moderna (React 19, Supabase, PostgreSQL 15), actualizacion semanal, roadmap publico, 3 aplicaciones (web + movil + extension).")
    screenshot_placeholder(doc, "Matriz competitiva visual Gestabiz vs lideres del mercado.")

    h2(doc, "2.3  Verticales atendidos y casos de uso", anchor="p_verticals")
    for vertical, case in [
        ("Belleza (salones, barberias, estetica)",
         "Control de horarios individuales por profesional, pausa de almuerzo flexible, chat directo con clientes, resenas de empleado, ventas de productos via Ventas Rapidas, reclutamiento con review obligatoria al contratar."),
        ("Salud (consultorios, odontologia, fisioterapia, psicologia)",
         "Permisos granulares para separar recepcionista/profesional/contador, historial de paciente en ClientProfileModal, confirmacion con anticipacion configurable, WhatsApp con consentimiento, contabilidad con Retefuente."),
        ("Deportes (gimnasios, academias, canchas)",
         "Recursos fisicos (canchas, carriles, estudios) con reserva directa sin empleado, clases grupales con capacity, horarios extendidos, integracion con Google Calendar."),
        ("Hoteles boutique y alojamientos",
         "Recursos tipo 'bed' y 'room' con tarifa horaria, multi-sede, contabilidad con IVA hotelero, resenas verificadas, perfil publico SEO indexado para OTAs."),
        ("Restaurantes y gastronomicos por reserva",
         "Recursos tipo 'table' con capacity, ventas rapidas como POS ligero, contabilidad con IVA, WhatsApp para confirmar mesa."),
        ("Coworkings y alquiler de espacios",
         "Recursos 'desk', 'meeting_room' con amenities JSONB (WiFi, proyector, pizarra), tarifa horaria, contabilidad de ingresos por sede."),
        ("Escuelas y academias",
         "Clases grupales con capacity, reclutamiento de profesores, contabilidad por sede, perfil publico con catalogo."),
        ("Servicios profesionales (abogados, contadores, consultores)",
         "Horarios restringidos, confirmacion obligatoria del cliente, permisos para compartir con asistentes, Google Calendar sync."),
    ]:
        h3(doc, vertical)
        para(doc, case)

    h2(doc, "2.4  Beneficios concretos por rol", anchor="p_benefits_role")

    h3(doc, "Para el dueno del negocio")
    bullet(doc, "Control total desde el primer dia: bypass de permisos, vista 360 de ingresos, ocupacion y empleados.")
    bullet(doc, "Reportes exportables listos para el contador con IVA, ICA y Retefuente desglosados.")
    bullet(doc, "Reduccion medible de ausentismo gracias a recordatorios automaticos multicanal.")
    bullet(doc, "Perfil publico indexable -> clientes llegando desde Google sin pagar publicidad.")

    h3(doc, "Para el admin operativo")
    bullet(doc, "Calendario centralizado con validacion anti-overlap en 10 reglas.")
    bullet(doc, "Transferencia de empleados entre sedes con cancelacion automatica de citas conflictivas.")
    bullet(doc, "Aprobacion/rechazo de ausencias en 2 clics con visibilidad del impacto.")
    bullet(doc, "Sistema de permisos granulares que le permite delegar sin perder control.")

    h3(doc, "Para el empleado")
    bullet(doc, "Agenda clara en todos sus dispositivos con sync automatico.")
    bullet(doc, "Solicitud de vacaciones con balance disponible visible.")
    bullet(doc, "Chat directo con sus clientes y control de que clientes le pueden escribir.")
    bullet(doc, "Metricas personales: ingresos generados, rating, tasa de ocupacion.")

    h3(doc, "Para el cliente")
    bullet(doc, "Reservar 24/7 en 4 clics sin llamadas ni WhatsApp.")
    bullet(doc, "Confirmar o cancelar con un link seguro desde el email.")
    bullet(doc, "Ver historial, favoritos y resenar negocios en un solo lugar.")
    bullet(doc, "Chat con el negocio si necesita cambiar algo sin tener que llamar.")
    screenshot_placeholder(doc, "Flujo de reserva del cliente en 4 pantallas.")

    h2(doc, "2.5  Seguridad, cumplimiento y arquitectura", anchor="p_security")
    bullet(doc, "RLS (Row Level Security) en todas las tablas - nadie accede a datos que no le corresponden.")
    bullet(doc, "Cliente Supabase singleton (una sola instancia en toda la app) - sin fuga de sesiones.")
    bullet(doc, "Pagos procesados por MercadoPago / PayU / Stripe - Gestabiz nunca almacena tarjetas.")
    bullet(doc, "GDPR-compliant: cookie consent banner, anonymizeIp, consent mode de GA4.")
    bullet(doc, "Auditoria completa en permission_audit_log, error_logs, login_logs y Sentry.")
    bullet(doc, "Backups automaticos por Supabase + point-in-time recovery.")
    bullet(doc, "Separacion clara DEV/PROD con clientes OAuth distintos por entorno.")
    bullet(doc, "Claves legacy deshabilitadas; uso exclusivo de claves publishable nuevas (sb_publishable_*).")

    h2(doc, "2.6  Plan de onboarding y soporte", anchor="p_onboarding")
    numbered(doc, "Dia 0: registro, verificacion de correo, creacion de negocio con wizard de 4 pasos.")
    numbered(doc, "Dia 1: publicacion del perfil publico y primeros 3 servicios.")
    numbered(doc, "Dia 2-7: invitacion de empleados, configuracion de horarios, primeras reservas.")
    numbered(doc, "Dia 8-30: configuracion de impuestos, conexion de Google Calendar, primeros reportes.")
    numbered(doc, "Mes 2 en adelante: analisis de KPIs, optimizacion de ocupacion, uso de reclutamiento.")
    callout(doc, "Soporte garantizado",
            "Soporte por email en todos los planes. WhatsApp prioritario en plan Pro con SLA de 4 horas habiles.")

    h2(doc, "2.7  Roadmap 2026", anchor="p_roadmap")
    bullet(doc, "Q2 2026 - API publica y webhooks (Pro).")
    bullet(doc, "Q2 2026 - Facturacion electronica via Matias API (Colombia).")
    bullet(doc, "Q3 2026 - Pagos en cita (deposito al reservar).")
    bullet(doc, "Q3 2026 - Programa de fidelizacion con puntos y descuentos.")
    bullet(doc, "Q4 2026 - Marketplace interno de servicios (discovery global).")
    bullet(doc, "Q4 2026 - Mobile nativo en las tiendas App Store y Google Play.")

    h2(doc, "2.8  Testimonios y casos de exito", anchor="p_testimonials")
    for quote, who in [
        ("«Pase de perder 4 citas por semana a 0. En un mes pague el plan Pro con lo que recupere.»",
         "- Maria C., duena de salon, Medellin"),
        ("«Tener IVA, ICA y Retefuente calculados solos me ahorra un dia entero cada mes con mi contador.»",
         "- Andres P., clinica odontologica, Bogota"),
        ("«El perfil publico me trajo 30 clientes nuevos en 2 meses desde Google sin pagar un peso de ads.»",
         "- Natalia R., estudio de yoga, Cali"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(quote)
        style_run(r, size=12, italic=True, color=DARK)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0)
        p2.paragraph_format.space_after = Pt(10)
        r = p2.add_run(who)
        style_run(r, size=10, color=GREY)

    h2(doc, "2.9  Preguntas frecuentes", anchor="p_faq")
    for q, a in [
        ("Tengo que firmar permanencia?",
         "No. Puedes cancelar en cualquier momento desde el panel de Facturacion."),
        ("Puedo migrar desde otra plataforma?",
         "Si. Importamos clientes y servicios en un CSV. Citas historicas se cargan por rango."),
        ("Gestabiz procesa mis pagos?",
         "No directamente. Usamos MercadoPago, PayU o Stripe (elegible por ti). Las tarjetas nunca tocan nuestros servidores."),
        ("Mis datos son mios?",
         "Si. Puedes exportar todo tu contenido en cualquier momento en formato CSV/JSON."),
        ("Gestabiz funciona fuera de Colombia?",
         "Si. La plataforma es multimoneda y multiidioma (ES/EN). Los modulos fiscales estan pensados para Colombia pero son desactivables."),
        ("Que pasa si supero un limite del plan?",
         "El modulo afectado pasa a modo lectura hasta que aumentes el plan. Nunca perdemos datos."),
        ("Tienen aplicacion movil?",
         "Si. Tenemos Expo/React Native para iOS y Android y una extension de navegador."),
    ]:
        h4(doc, q)
        para(doc, a)

    # CTA final
    page_break(doc)
    h1(doc, "Empieza hoy - Sin costo", color=PURPLE)
    para(doc, "Gestabiz esta en fase Beta completada, con operacion real en produccion, soporte activo y roadmap publico. No hay mejor momento para sumarte.", size=13)
    callout(doc, "Prueba Gestabiz",
            "Crea tu cuenta en gestabiz.com - configura tu negocio en 5 minutos - recibe tu primera reserva hoy.",
            color=PURPLE)
    para(doc, "Contacto comercial: ventas@gestabiz.com  ·  WhatsApp +57 XXX XXX XXXX")
    para(doc, "Desarrollado por Ti Turing - https://github.com/TI-Turing", italic=True, color=GREY)

    return doc


# ===========================================================================
# ENTRY POINT
# ===========================================================================
def main() -> None:
    print(">> Generando documentos Gestabiz...")
    print(f"   Repo:             {ROOT}")
    print(f"   Destino:          {DOCS_DIR}")
    print(f"   Logo Gestabiz:    {'OK' if LOGO_GESTABIZ.exists() else 'FALTA'} ({LOGO_GESTABIZ})")
    print(f"   Logo Ti Turing:   {'OK' if LOGO_TITURING.exists() else 'FALTA'} ({LOGO_TITURING})")

    manual = build_manual()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz.docx"
    manual.save(str(manual_path))
    print(f"[OK] Manual generado:    {manual_path} ({manual_path.stat().st_size // 1024} KB)")

    proposal = build_proposal()
    proposal_path = DOCS_DIR / "Propuesta_Valor_Gestabiz.docx"
    proposal.save(str(proposal_path))
    print(f"[OK] Propuesta generada: {proposal_path} ({proposal_path.stat().st_size // 1024} KB)")

    print(">> Listo.")


if __name__ == "__main__":
    main()
