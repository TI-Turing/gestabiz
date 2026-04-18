#!/usr/bin/env python3
"""
Generador de documentos Gestabiz — Version Copilot — Parte 2: Rol Empleado
==========================================================================
Genera:
  - docs/Manual_Usuario_Gestabiz - copilot - parte2.docx
  - docs/Propuesta_Valor_Gestabiz - copilot - parte2.docx

Requisitos: python-docx >= 1.0.0
Ejecucion:  python scripts/generate_copilot_docs_part2.py
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths & Brand Colors
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
LOGO_GESTABIZ = ROOT / "src" / "assets" / "images" / "logo_gestabiz.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def add_bookmark(paragraph, name: str) -> None:
    tag = OxmlElement("w:bookmarkStart")
    tag.set(qn("w:id"), str(id(name) % 2**31))
    tag.set(qn("w:name"), name)
    paragraph._p.append(tag)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(id(name) % 2**31))
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *,
                           bold: bool = False, color: str = "7C3AED") -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# High-level helpers
# ---------------------------------------------------------------------------

def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + text)
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(title)
    style_run(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {caption}")
    style_run(r, size=9, italic=True, color=GREY)
    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=header_fg)
    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, size=10, color=DARK)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup + cover
# ---------------------------------------------------------------------------

def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run(f"{title}  |  Desarrollado por Ti Turing  |  (c) 2026")
        style_run(r, size=9, italic=True, color=GREY)
    build_cover(doc, title, subtitle)
    return doc


def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Parte 2 de 5: Experiencia del Empleado")
    style_run(r, size=14, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version del producto: 1.0.3   |   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fase Beta completada   |   Listo para produccion")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_TITURING), width=Cm(3.0))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por ")
    style_run(r, size=11, color=GREY)
    r = p.add_run("Ti Turing")
    style_run(r, size=11, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("https://github.com/TI-Turing   |   contacto@gestabiz.com")
    style_run(r, size=9, italic=True, color=GREY)

    page_break(doc)


# ===========================================================================
# MANUAL DE USUARIO — PARTE 2: ROL EMPLEADO
# ===========================================================================

def build_manual_part2() -> Document:
    doc = setup_document(
        title="Manual de Usuario",
        subtitle="Guia funcional exhaustiva de la plataforma"
    )

    # =====================================================================
    # INDICE
    # =====================================================================
    h1(doc, "Indice — Parte 2: Rol Empleado", anchor="toc")
    para(doc, "Este documento cubre de manera exhaustiva la experiencia del rol Empleado en Gestabiz. Cada entrada es un hipervinculo a la seccion correspondiente. Este es el segundo de cinco volumenes que componen el Manual completo.")

    toc_items = [
        ("sec_resumen",        "Parte 1 — Resumen Ejecutivo del Rol Empleado"),
        ("sec_acceso",         "1. Acceso y Primer Inicio de Sesion"),
        ("sec_phone_gate",     "2. Verificacion de Telefono Obligatoria"),
        ("sec_dashboard",      "3. Dashboard del Empleado: Vision General"),
        ("sec_sidebar",        "4. Barra Lateral y Navegacion"),
        ("sec_selector_neg",   "5. Selector de Negocio"),
        ("sec_mis_empleos",    "6. Mis Empleos — Lista de Negocios"),
        ("sec_tarjeta_empleo", "7. Tarjeta de Empleo y Detalle"),
        ("sec_unirse",         "8. Unirse a un Negocio"),
        ("sec_onboarding",     "9. Onboarding con Codigo de Invitacion"),
        ("sec_citas",          "10. Mis Citas — Vista Lista"),
        ("sec_calendario",     "11. Mis Citas — Vista Calendario"),
        ("sec_detalle_cita",   "12. Detalle de una Cita"),
        ("sec_clientes",       "13. Mis Clientes"),
        ("sec_ausencias",      "14. Ausencias y Vacaciones"),
        ("sec_solicitar_aus",  "15. Solicitar una Ausencia"),
        ("sec_time_off",       "16. Solicitud Rapida de Ausencia desde Tarjeta de Empleo"),
        ("sec_balance_vac",    "17. Widget de Balance de Vacaciones"),
        ("sec_historial_aus",  "18. Historial de Ausencias (Tabs)"),
        ("sec_festivos",       "19. Festivos Publicos y Validacion"),
        ("sec_horario",        "20. Configuracion de Horario de Trabajo"),
        ("sec_servicios",      "21. Gestion de Servicios Ofrecidos"),
        ("sec_sedes",          "22. Sedes y Selector de Ubicacion"),
        ("sec_traslado",       "23. Traslado de Sede"),
        ("sec_vacantes",       "24. Marketplace de Vacantes"),
        ("sec_aplicar",        "25. Aplicar a una Vacante"),
        ("sec_mis_apps",       "26. Mis Aplicaciones"),
        ("sec_conflictos",     "27. Deteccion de Conflictos de Horario"),
        ("sec_chat",           "28. Chat en Tiempo Real"),
        ("sec_notificaciones", "29. Notificaciones"),
        ("sec_settings",       "30. Configuraciones del Empleado"),
        ("sec_perfil_prof",    "31. Perfil Profesional Publico"),
        ("sec_cert",           "32. Certificaciones, Idiomas y Especializaciones"),
        ("sec_mensajes_pref",  "33. Preferencias de Mensajes de Clientes"),
        ("sec_fin_empleo",     "34. Finalizar Empleo"),
        ("sec_permisos",       "35. Permisos del Empleado"),
        ("sec_modales",        "36. Modales y Dialogos del Empleado"),
        ("sec_estados",        "37. Estados de Carga, Error y Vacio"),
        ("sec_plan_limits",    "38. Limitaciones por Plan"),
        ("sec_glosario",       "39. Glosario de Terminos del Empleado"),
    ]
    for anchor, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor, label)

    page_break(doc)

    # =====================================================================
    # PARTE 1 — RESUMEN EJECUTIVO
    # =====================================================================
    h1(doc, "Parte 1 — Resumen Ejecutivo del Rol Empleado", anchor="sec_resumen")

    para(doc, "El rol Empleado en Gestabiz esta disenado para profesionales de servicios — peluqueros, esteticistas, medicos, terapeutas, entrenadores y cualquier persona que ofrece servicios personales a traves de un negocio. Este rol proporciona un espacio de trabajo completo donde el empleado puede gestionar su agenda, sus clientes, sus ausencias y su perfil profesional de forma autonoma, sin depender constantemente del administrador.")

    h2(doc, "Que puede hacer un Empleado en Gestabiz?")
    bullet(doc, "Ver y gestionar sus citas asignadas en vista lista o calendario.", bold_prefix="Agenda")
    bullet(doc, "Consultar y gestionar la lista de clientes que ha atendido.", bold_prefix="Mis Clientes")
    bullet(doc, "Solicitar vacaciones, licencias de enfermedad, emergencias y ausencias personales con aprobacion obligatoria.", bold_prefix="Ausencias")
    bullet(doc, "Configurar su horario semanal de trabajo y hora de almuerzo.", bold_prefix="Horario")
    bullet(doc, "Seleccionar que servicios del negocio ofrece, con nivel de experiencia y comision.", bold_prefix="Servicios")
    bullet(doc, "Explorar vacantes laborales de otros negocios y aplicar con CV y carta de presentacion.", bold_prefix="Vacantes")
    bullet(doc, "Chatear en tiempo real con clientes y colegas.", bold_prefix="Chat")
    bullet(doc, "Recibir notificaciones in-app y por email sobre citas, ausencias y solicitudes.", bold_prefix="Notificaciones")
    bullet(doc, "Mantener un perfil profesional publico con certificaciones, idiomas y especializaciones.", bold_prefix="Perfil Profesional")
    bullet(doc, "Trabajar en multiples negocios simultaneamente con selector de negocio.", bold_prefix="Multi-negocio")
    bullet(doc, "Solicitar traslado de sede dentro del mismo negocio con gestion de citas afectadas.", bold_prefix="Traslado de Sede")

    screenshot_placeholder(doc, "Dashboard del Empleado — Vista general")

    h2(doc, "Pantallas y Secciones del Empleado")
    simple_table(doc,
        ["Seccion", "Descripcion", "Acceso"],
        [
            ["Mis Empleos", "Lista de negocios donde trabaja con tarjetas enriquecidas", "Sidebar: Mis Empleos"],
            ["Buscar Vacantes", "Marketplace de vacantes con filtros avanzados", "Sidebar: Buscar Vacantes"],
            ["Mis Ausencias", "Balance de vacaciones, historial y formulario de solicitud", "Sidebar: Mis Ausencias"],
            ["Mis Citas", "Lista y calendario de citas asignadas con filtros", "Sidebar: Mis Citas"],
            ["Horario", "Editor visual de horario semanal + almuerzo", "Sidebar: Horario"],
            ["Mis Clientes", "Clientes atendidos con historial y perfil", "Sidebar: Mis Clientes"],
            ["Unirse a Negocio", "Busqueda y solicitud de vinculacion", "Boton en Mis Empleos"],
            ["Onboarding", "Ingreso con codigo de invitacion o QR", "Boton en Mis Empleos"],
            ["Configuraciones", "Perfil profesional, mensajes, horario, certificaciones", "Menu de usuario"],
            ["Chat", "Conversaciones en tiempo real", "Boton flotante / Notificaciones"],
        ]
    )

    h2(doc, "Flujos principales del Empleado")
    simple_table(doc,
        ["Flujo", "Pasos", "Resultado"],
        [
            ["Unirse a negocio", "Buscar negocio o ingresar codigo -> Enviar solicitud -> Admin aprueba", "Vinculado como empleado"],
            ["Ver citas del dia", "Abrir Mis Citas -> Filtrar por hoy -> Ver lista o calendario", "Agenda actualizada"],
            ["Solicitar ausencia", "Abrir Mis Ausencias -> Llenar formulario -> Enviar -> Admin aprueba", "Ausencia registrada, citas afectadas canceladas"],
            ["Aplicar a vacante", "Buscar Vacantes -> Ver detalle -> Llenar formulario -> Enviar", "Aplicacion registrada"],
            ["Configurar horario", "Abrir Horario -> Activar/desactivar dias -> Guardar", "Horario actualizado"],
            ["Seleccionar servicios", "Abrir Servicios -> Marcar servicios -> Nivel + comision -> Guardar", "Servicios actualizados"],
            ["Solicitar traslado", "Abrir Sedes -> Elegir sede -> Preaviso -> Confirmar", "Traslado programado"],
            ["Finalizar empleo", "Tarjeta de empleo -> Finalizar -> Confirmar", "Empleo desactivado"],
        ]
    )

    callout(doc, "Disponible en todos los planes",
            "Las funcionalidades del rol Empleado estan disponibles en los planes Gratuito, Basico y Pro. No existen limitaciones por plan para el empleado individual; las restricciones de cantidad de empleados aplican al administrador del negocio.",
            color=ACCENT)

    page_break(doc)

    # =====================================================================
    # PARTE 2 — DETALLE EXHAUSTIVO
    # =====================================================================
    h1(doc, "Parte 2 — Detalle Exhaustivo", anchor="sec_detalle")
    para(doc, "A continuacion se documenta cada seccion, boton, validacion, flujo normal y flujo alterno de la experiencia del Empleado en Gestabiz.")

    # ------------------------------------------------------------------
    # 1. Acceso
    # ------------------------------------------------------------------
    h2(doc, "1. Acceso y Primer Inicio de Sesion", anchor="sec_acceso")
    para(doc, "Para acceder como empleado, el usuario debe tener al menos un registro activo en la tabla business_employees. Este registro se crea cuando un administrador acepta la solicitud de vinculacion del empleado o cuando el empleado reclama un codigo de invitacion. Si el usuario no tiene ningun registro de empleado, el rol Empleado no aparecera en el selector de roles.")
    h3(doc, "Flujo de primer acceso")
    numbered(doc, "El usuario se registra en Gestabiz (email + contrasena o Google/GitHub).")
    numbered(doc, "Navega al Dashboard y ve el selector de roles. Si aun no esta vinculado a ningun negocio, solo vera los roles Cliente y Admin.")
    numbered(doc, "Hace clic en el rol Empleado (si esta disponible) o navega a Mis Empleos y hace clic en Unirse a un nuevo negocio.")
    numbered(doc, "Completa el flujo de vinculacion (ver seccion 8).")
    numbered(doc, "Al ser aprobado, el rol Empleado se activa y aparece en el selector.")
    callout(doc, "Importante: Roles dinamicos",
            "Los roles en Gestabiz se calculan en tiempo real. No se almacenan en la base de datos. Si un empleado es desvinculado (is_active = false), el rol Empleado desaparece automaticamente del selector de roles.",
            color=PURPLE)
    screenshot_placeholder(doc, "Selector de roles mostrando Empleado activo")

    # ------------------------------------------------------------------
    # 2. Phone Gate
    # ------------------------------------------------------------------
    h2(doc, "2. Verificacion de Telefono Obligatoria", anchor="sec_phone_gate")
    para(doc, "Al ingresar al Dashboard del Empleado, el sistema verifica si el usuario tiene un numero de telefono registrado en su perfil. Si no lo tiene, se muestra un modal bloqueante (PhoneRequiredModal) que impide el acceso hasta verificar el telefono via SMS OTP.")
    h3(doc, "Flujo de verificacion")
    numbered(doc, "Se muestra el modal con titulo 'Telefono requerido' y descripcion explicativa.")
    numbered(doc, "El usuario ingresa su numero de telefono con prefijo internacional.")
    numbered(doc, "El sistema envia un codigo OTP via SMS.")
    numbered(doc, "El usuario ingresa el codigo de 6 digitos.")
    numbered(doc, "Si el codigo es correcto, el telefono se guarda en el perfil y la pagina se recarga.")
    h3(doc, "Flujos alternos")
    bullet(doc, "Si el codigo OTP es incorrecto, se muestra un error y el usuario puede reintentar.")
    bullet(doc, "Si el SMS no llega, el usuario puede solicitar un reenvio despues de 60 segundos.")
    bullet(doc, "El modal NO se puede cerrar sin verificar (es bloqueante).")
    screenshot_placeholder(doc, "Modal de verificacion de telefono")

    # ------------------------------------------------------------------
    # 3. Dashboard
    # ------------------------------------------------------------------
    h2(doc, "3. Dashboard del Empleado: Vision General", anchor="sec_dashboard")
    para(doc, "El EmployeeDashboard es el componente principal del rol Empleado. Es un layout con barra lateral izquierda y area de contenido principal. El contenido cambia segun la seccion seleccionada en la barra lateral.")
    h3(doc, "Componentes del Dashboard")
    bullet(doc, "Header con nombre del empleado, boton de notificaciones (NotificationBell), menu de usuario.", bold_prefix="Header")
    bullet(doc, "Barra lateral con 6 items de navegacion + icono, ademas de links programaticos a Unirse, Onboarding, Perfil y Configuraciones.", bold_prefix="Sidebar")
    bullet(doc, "Selector de negocio en la parte superior (visible si el empleado trabaja en multiples negocios).", bold_prefix="Business Selector")
    bullet(doc, "Area de contenido principal que renderiza el componente correspondiente a la seccion activa.", bold_prefix="Content Area")
    bullet(doc, "Boton flotante de chat (FloatingChatButton) en la esquina inferior derecha.", bold_prefix="Chat")
    para(doc, "La URL sigue el patron /app/employee/{seccion}. Si el usuario navega a /app sin seccion, se redirige automaticamente a /app/employee/employments (Mis Empleos).")
    screenshot_placeholder(doc, "Dashboard del Empleado completo")

    # ------------------------------------------------------------------
    # 4. Sidebar
    # ------------------------------------------------------------------
    h2(doc, "4. Barra Lateral y Navegacion", anchor="sec_sidebar")
    simple_table(doc,
        ["Icono", "Etiqueta", "ID de ruta", "Descripcion"],
        [
            ["Briefcase", "Mis Empleos", "employments", "Lista de negocios donde trabaja"],
            ["Search", "Buscar Vacantes", "vacancies", "Marketplace de oportunidades laborales"],
            ["CalendarOff", "Mis Ausencias", "absences", "Balance de vacaciones, solicitudes e historial"],
            ["Calendar", "Mis Citas", "appointments", "Lista y calendario de citas asignadas"],
            ["Clock", "Horario", "schedule", "Editor visual de horario semanal"],
            ["Users", "Mis Clientes", "my-clients", "Clientes atendidos por el empleado"],
        ]
    )
    para(doc, "Ademas de estos 6 items del sidebar, existen paginas accesibles programaticamente: Unirse a Negocio (join-business), Onboarding (onboarding), Perfil (profile) y Configuraciones (settings).")
    callout(doc, "Navegacion responsiva",
            "En pantallas moviles, la barra lateral se colapsa en un menu hamburguesa. El contenido se muestra a pantalla completa con navegacion inferior simplificada.",
            color=ACCENT)

    # ------------------------------------------------------------------
    # 5. Selector de Negocio
    # ------------------------------------------------------------------
    h2(doc, "5. Selector de Negocio", anchor="sec_selector_neg")
    para(doc, "En la parte superior del dashboard hay un dropdown que muestra todos los negocios donde el empleado esta vinculado activamente. Al cambiar de negocio, todo el contexto del dashboard cambia: citas, ausencias, clientes, horario y servicios se filtran por el negocio seleccionado.")
    h3(doc, "Comportamiento")
    bullet(doc, "Si el empleado trabaja en un solo negocio, el selector se muestra pero con un unico item.")
    bullet(doc, "La seleccion se persiste en localStorage con la clave gestabiz-employee-business-{userId}.")
    bullet(doc, "Al cambiar de negocio, se establece un effectiveBusinessId que se usa como scope para todas las queries.")
    bullet(doc, "Si el negocio guardado en localStorage ya no existe o el empleado fue desvinculado, se selecciona automaticamente el primer negocio disponible.")
    screenshot_placeholder(doc, "Selector de negocio con multiples negocios")

    # ------------------------------------------------------------------
    # 6. Mis Empleos
    # ------------------------------------------------------------------
    h2(doc, "6. Mis Empleos — Lista de Negocios", anchor="sec_mis_empleos")
    para(doc, "La seccion Mis Empleos muestra todos los negocios donde el usuario esta registrado como empleado. Utiliza el componente MyEmploymentsEnhanced que ejecuta 4 queries paralelas por negocio para enriquecer la tarjeta.")
    h3(doc, "Informacion mostrada por negocio")
    bullet(doc, "Logo y nombre del negocio.", bold_prefix="Identidad")
    bullet(doc, "Rating promedio del negocio (estrellas).", bold_prefix="Rating")
    bullet(doc, "Rol del empleado en ese negocio (Propietario, Manager, Profesional, Recepcionista, etc.).", bold_prefix="Rol")
    bullet(doc, "Cantidad de servicios que ofrece en ese negocio.", bold_prefix="Servicios")
    bullet(doc, "Sede asignada con nombre y direccion.", bold_prefix="Ubicacion")
    bullet(doc, "Salario base si esta configurado.", bold_prefix="Salario")
    h3(doc, "Estadisticas resumen")
    bullet(doc, "Total de negocios activos.")
    bullet(doc, "Cantidad como Propietario.")
    bullet(doc, "Cantidad como Empleado.")
    h3(doc, "Acciones disponibles")
    bullet(doc, "Click en tarjeta: abre EmploymentDetailModal con 6 tabs de informacion detallada.")
    bullet(doc, "Menu contextual (tres puntos): 5 opciones de solicitar ausencia (vacaciones, enfermedad, personal, emergencia, otros) + Marcar como Finalizado.")
    bullet(doc, "Boton 'Unirse a un nuevo negocio': navega al flujo de vinculacion.")
    callout(doc, "Nota sobre Propietarios",
            "Si el usuario es el owner (dueno) de un negocio, aparece como Propietario en la tarjeta. La opcion 'Marcar como Finalizado' esta deshabilitada para propietarios — no puede autodesvincularse de su propio negocio.",
            color=PURPLE)
    screenshot_placeholder(doc, "Mis Empleos con multiples negocios")

    # ------------------------------------------------------------------
    # 7. Tarjeta de Empleo y Detalle
    # ------------------------------------------------------------------
    h2(doc, "7. Tarjeta de Empleo y Detalle", anchor="sec_tarjeta_empleo")
    para(doc, "Al hacer clic en cualquier tarjeta de empleo se abre el EmploymentDetailModal, un dialogo completo con 6 pestanas que muestran toda la informacion del empleo del profesional en ese negocio.")
    h3(doc, "Pestanas del Detalle de Empleo")
    simple_table(doc,
        ["Pestana", "Contenido"],
        [
            ["Info", "Datos generales: nombre del negocio, rol asignado, tipo de empleado, fecha de contratacion, estado activo, negocio activo o inactivo."],
            ["Sedes", "Sedes del negocio con horarios de apertura/cierre por dia, direccion, telefono, email, hasta 6 fotos por sede."],
            ["Servicios", "Lista de servicios que el empleado ofrece en ese negocio con nombre, precio, duracion, nivel de experiencia (estrellas)."],
            ["Horario", "Horario semanal del empleado en ese negocio: 7 dias con hora de inicio y fin, indicador de dia activo/inactivo."],
            ["Salario", "Salario base mensual con desglose de beneficios colombianos: Seguridad Social (10%), Salud (5%), Pension (5%), Total neto = salario x 1.20. Formato COP colombiano."],
            ["Stats", "Metricas de rendimiento: ocupacion, rating promedio, ingresos generados, citas completadas."],
        ]
    )
    screenshot_placeholder(doc, "EmploymentDetailModal — Pestana Info")
    screenshot_placeholder(doc, "EmploymentDetailModal — Pestana Salario con desglose")

    # ------------------------------------------------------------------
    # 8. Unirse a un Negocio
    # ------------------------------------------------------------------
    h2(doc, "8. Unirse a un Negocio", anchor="sec_unirse")
    para(doc, "El componente JoinBusiness permite al empleado vincularse a un nuevo negocio. Tiene dos mecanismos de vinculacion en pestanas separadas.")
    h3(doc, "Pestana 1: Buscar Negocio")
    numbered(doc, "El empleado escribe el nombre del negocio en el campo de busqueda.")
    numbered(doc, "El sistema busca con ILIKE (case-insensitive) en negocios activos, limitado a 12 resultados.")
    numbered(doc, "Se muestran tarjetas con nombre, logo y categoria del negocio.")
    numbered(doc, "El empleado hace clic en 'Solicitar Acceso' en la tarjeta deseada.")
    numbered(doc, "Se crea un registro en employee_join_requests con status 'pending'.")
    numbered(doc, "El administrador del negocio recibe una notificacion y puede aprobar o rechazar.")
    h3(doc, "Pestana 2: Codigo de Invitacion")
    numbered(doc, "El empleado ingresa un codigo alfanumerico de 6 caracteres proporcionado por el administrador.")
    numbered(doc, "El sistema valida el codigo via RPC claim_invite_code (SECURITY DEFINER).")
    numbered(doc, "Si el codigo es valido y no ha expirado (24 horas), se vincula al empleado automaticamente.")
    numbered(doc, "Se notifica al propietario del negocio via Edge Function send-notification.")
    h3(doc, "Solicitudes existentes")
    para(doc, "En la parte superior de JoinBusiness se muestra una seccion con las solicitudes existentes del empleado, indicando el estado de cada una: pendiente (amarillo), aprobada (verde) o rechazada (rojo).")
    h3(doc, "Validaciones")
    bullet(doc, "No se puede enviar una solicitud duplicada si ya existe una pendiente para el mismo negocio.")
    bullet(doc, "El codigo de invitacion expira despues de 24 horas.")
    bullet(doc, "Si el empleado ya esta vinculado al negocio, no puede solicitar nuevamente.")
    screenshot_placeholder(doc, "JoinBusiness — Busqueda de negocios")
    screenshot_placeholder(doc, "JoinBusiness — Codigo de invitacion")

    # ------------------------------------------------------------------
    # 9. Onboarding
    # ------------------------------------------------------------------
    h2(doc, "9. Onboarding con Codigo de Invitacion", anchor="sec_onboarding")
    para(doc, "El EmployeeOnboarding es un flujo simplificado que permite al empleado unirse directamente ingresando un codigo o escaneando un QR.")
    h3(doc, "Flujo")
    numbered(doc, "Se muestra un campo de texto para ingresar el codigo de 6 caracteres.")
    numbered(doc, "Alternativa: boton para escanear QR que abre la camara del dispositivo.")
    numbered(doc, "Al enviar, se llama al RPC claim_invite_code.")
    numbered(doc, "Si es exitoso, se notifica al propietario y se redirige al dashboard.")
    h3(doc, "Diferencia con JoinBusiness")
    para(doc, "EmployeeOnboarding es mas simple y directo — solo maneja codigos, no tiene busqueda de negocios. Se usa tipicamente cuando el administrador comparte un enlace o QR con el codigo pre-cargado.")

    # ------------------------------------------------------------------
    # 10. Mis Citas — Vista Lista
    # ------------------------------------------------------------------
    h2(doc, "10. Mis Citas — Vista Lista", anchor="sec_citas")
    para(doc, "La pagina de citas del empleado muestra todas las citas asignadas a el en el negocio seleccionado. Ofrece dos vistas: lista y calendario, con un toggle para alternar entre ellas.")
    h3(doc, "Estadisticas superiores")
    simple_table(doc,
        ["Metrica", "Descripcion"],
        [
            ["Citas Hoy", "Cantidad de citas para el dia actual"],
            ["Programadas", "Citas con status 'scheduled' pendientes"],
            ["Confirmadas", "Citas con status 'confirmed'"],
            ["Completadas", "Citas con status 'completed' (historicas)"],
        ]
    )
    h3(doc, "Filtros disponibles")
    bullet(doc, "Busqueda por nombre del cliente (filtrado local).", bold_prefix="Busqueda")
    bullet(doc, "Filtro por estado: Todos, Programada, Confirmada, Completada, Cancelada, No asistio.", bold_prefix="Estado")
    bullet(doc, "Filtro por servicio (dropdown con servicios del negocio).", bold_prefix="Servicio")
    h3(doc, "Agrupacion")
    para(doc, "Las citas se agrupan por fecha. Cada grupo muestra una etiqueta: 'Hoy', 'Manana', o la fecha en formato largo en espanol (ej: 'Lunes 21 Abril 2026'). Dentro de cada grupo, las citas se ordenan por hora de inicio ascendente. La zona horaria es America/Bogota.")
    h3(doc, "Componente de cada cita")
    para(doc, "Cada cita se renderiza como un AppointmentCard (componente reutilizable de src/components/cards/) con initialData del padre. Muestra: hora, nombre del cliente, servicio, estado (badge de color), sede. Al hacer clic, se abre un dialogo con el detalle completo.")
    h3(doc, "Tiempo real")
    para(doc, "Las citas se actualizan en tiempo real via suscripcion a Supabase Realtime (postgres_changes en tabla appointments, filtrando por employee_id). Cualquier INSERT, UPDATE o DELETE dispara un refetch automatico sin necesidad de recargar la pagina.")
    screenshot_placeholder(doc, "Mis Citas — Vista lista con filtros")

    # ------------------------------------------------------------------
    # 11. Vista Calendario
    # ------------------------------------------------------------------
    h2(doc, "11. Mis Citas — Vista Calendario", anchor="sec_calendario")
    para(doc, "La vista calendario del empleado permite visualizar las citas de forma grafica con 3 modos de vista.")
    h3(doc, "Modos de vista")
    simple_table(doc,
        ["Vista", "Descripcion"],
        [
            ["Dia", "Bloques horarios de una sola jornada. Cada cita ocupa su slot de tiempo."],
            ["Semana", "Grilla de 7 columnas (Lunes-Domingo). Cada celda muestra las citas del dia."],
            ["Mes", "Calendario mensual clasico. Cada dia muestra el numero de citas."],
        ]
    )
    h3(doc, "Codificacion por colores")
    simple_table(doc,
        ["Color", "Estado"],
        [
            ["Amarillo", "Programada / Reprogramada"],
            ["Verde", "Confirmada"],
            ["Azul", "Completada"],
            ["Rojo", "Cancelada"],
            ["Morado", "En progreso"],
        ]
    )
    h3(doc, "Navegacion")
    bullet(doc, "Botones Anterior / Siguiente para navegar entre periodos.")
    bullet(doc, "Boton 'Hoy' para volver rapidamente al periodo actual.")
    bullet(doc, "El dia actual tiene un resaltado visual (ring-2).")
    bullet(doc, "Selector de vista: Dia / Semana / Mes.")
    screenshot_placeholder(doc, "Vista calendario — Modo semana")

    # ------------------------------------------------------------------
    # 12. Detalle de una Cita
    # ------------------------------------------------------------------
    h2(doc, "12. Detalle de una Cita", anchor="sec_detalle_cita")
    para(doc, "Al hacer clic en una cita (en lista o calendario), se abre EmployeeAppointmentModal, un dialogo de solo lectura que muestra toda la informacion de la cita.")
    h3(doc, "Informacion mostrada")
    bullet(doc, "Badge de estado con color correspondiente (scheduled, confirmed, completed, cancelled, no_show, rescheduled, in_progress).")
    bullet(doc, "Nombre completo del cliente.")
    bullet(doc, "Servicio reservado.")
    bullet(doc, "Fecha y hora en zona horaria Colombia (America/Bogota).")
    bullet(doc, "Ubicacion (sede) donde se realizara la cita.")
    bullet(doc, "Notas adicionales del cliente (si las hay).")
    h3(doc, "Limitaciones")
    para(doc, "El modal de cita del empleado es de solo lectura. El empleado NO puede modificar, cancelar ni reprogramar citas desde este modal. Esas acciones son exclusivas del administrador o del cliente.")
    callout(doc, "Solo lectura para empleados",
            "A diferencia del cliente (que puede cancelar o reprogramar) y del administrador (que tiene control total), el empleado solo puede visualizar los detalles de sus citas. Esta restriccion es una decision de diseno para evitar conflictos de agenda.",
            color=PURPLE)

    # ------------------------------------------------------------------
    # 13. Mis Clientes
    # ------------------------------------------------------------------
    h2(doc, "13. Mis Clientes", anchor="sec_clientes")
    para(doc, "La pagina Mis Clientes muestra todos los clientes que el empleado ha atendido en el negocio seleccionado. Se construye mediante un two-step query: primero se consultan las citas del empleado, luego se extraen los IDs unicos de clientes y se consultan sus perfiles.")
    h3(doc, "Datos por cliente")
    bullet(doc, "Avatar con iniciales del nombre (color generado a partir del nombre).")
    bullet(doc, "Nombre completo y email.")
    bullet(doc, "Total de visitas (todas las citas).")
    bullet(doc, "Visitas completadas (solo citas con status 'completed').")
    bullet(doc, "Fecha de ultima visita.")
    h3(doc, "Busqueda")
    para(doc, "Campo de busqueda que filtra localmente por nombre o email del cliente.")
    h3(doc, "Ordenamiento")
    para(doc, "Los clientes se ordenan por cantidad de visitas completadas, mostrando primero los clientes mas frecuentes.")
    h3(doc, "Accion: Ver Perfil del Cliente")
    para(doc, "Al hacer clic en una tarjeta de cliente se abre ClientProfileModal (componente compartido con el administrador). El modal tiene 2 tabs: 'Informacion' (estadisticas, fecha de primera y ultima visita) e 'Historial' (lista de citas con servicio, fecha, estado y precio). Las citas mostradas se filtran al contexto del negocio y empleado seleccionado.")
    screenshot_placeholder(doc, "Mis Clientes — Grid de tarjetas")
    screenshot_placeholder(doc, "ClientProfileModal — Tab Historial")

    # ------------------------------------------------------------------
    # 14. Ausencias y Vacaciones
    # ------------------------------------------------------------------
    h2(doc, "14. Ausencias y Vacaciones", anchor="sec_ausencias")
    para(doc, "El sistema de ausencias permite al empleado solicitar tiempo libre que debe ser aprobado obligatoriamente por un administrador. No existe la opcion de auto-aprobacion; esta politica es una regla de negocio no negociable.")
    h3(doc, "Tipos de ausencia")
    simple_table(doc,
        ["Tipo", "Etiqueta", "Observacion"],
        [
            ["vacation", "Vacaciones", "Descuenta del balance de vacaciones. 15 dias por ano."],
            ["sick_leave", "Licencia por Enfermedad", "Si excede 3 dias puede requerir certificado medico."],
            ["personal", "Personal", "Ausencia personal generica."],
            ["emergency", "Emergencia", "Requiere justificacion. Cancela citas automaticamente."],
            ["other", "Otros", "Cualquier otro tipo de ausencia no categorizada."],
        ]
    )
    callout(doc, "Aprobacion SIEMPRE obligatoria",
            "La configuracion require_absence_approval esta forzada a TRUE en todos los negocios, siempre. Ningun empleado puede tomar ausencias sin autorizacion previa de un administrador o manager.",
            color=DANGER)

    # ------------------------------------------------------------------
    # 15. Solicitar una Ausencia
    # ------------------------------------------------------------------
    h2(doc, "15. Solicitar una Ausencia", anchor="sec_solicitar_aus")
    para(doc, "El AbsenceRequestModal es el formulario principal para solicitar ausencias. Es un modal avanzado con calendarios visuales y multiples validaciones.")
    h3(doc, "Campos del formulario")
    bullet(doc, "Tipo de ausencia (selector con 5 opciones).", bold_prefix="Tipo")
    bullet(doc, "Fecha de inicio (calendario visual, no input de texto).", bold_prefix="Desde")
    bullet(doc, "Fecha de fin (calendario visual, no input de texto).", bold_prefix="Hasta")
    bullet(doc, "Razon (campo de texto libre, requerido).", bold_prefix="Razon")
    bullet(doc, "Notas adicionales (opcional).", bold_prefix="Notas")
    h3(doc, "Validaciones en tiempo real")
    numbered(doc, "Dias laborales: se identifican automaticamente los dias de fin de semana en el rango seleccionado y se muestran con alerta roja.")
    numbered(doc, "Festivos publicos: se cruza el rango con la tabla public_holidays (54 festivos colombianos 2025-2027) y se muestran con alerta naranja.")
    numbered(doc, "Citas afectadas: se cuenta cuantas citas tiene el empleado en el rango seleccionado y se muestran con alerta amarilla.")
    numbered(doc, "Balance de vacaciones: se verifica que el empleado tenga dias disponibles suficientes (solo para tipo 'vacation').")
    numbered(doc, "Los calendarios muestran range highlighting: el dia de inicio y fin con marcador solido, los dias intermedios con marcador suave (20% opacidad).")
    h3(doc, "Flujo de envio")
    numbered(doc, "El empleado llena todos los campos requeridos.")
    numbered(doc, "Hace clic en 'Enviar Solicitud'.")
    numbered(doc, "Se crea un registro en employee_absences con status 'pending'.")
    numbered(doc, "TODOS los administradores y managers del negocio reciben notificacion in-app y email.")
    numbered(doc, "El administrador aprueba o rechaza desde su dashboard (ver Parte 3: Rol Admin).")
    numbered(doc, "El empleado recibe notificacion del resultado.")
    h3(doc, "Flujo alterno: Ausencia de Emergencia")
    para(doc, "Cuando una ausencia de tipo 'emergency' es aprobada, el sistema automaticamente cancela todas las citas del empleado en el rango de la ausencia via Edge Function cancel-appointments-on-emergency-absence. Los clientes afectados reciben notificacion de la cancelacion.")
    screenshot_placeholder(doc, "AbsenceRequestModal con calendarios y alertas")

    # ------------------------------------------------------------------
    # 16. Solicitud rapida desde tarjeta
    # ------------------------------------------------------------------
    h2(doc, "16. Solicitud Rapida de Ausencia desde Tarjeta de Empleo", anchor="sec_time_off")
    para(doc, "Desde la tarjeta de empleo (BusinessEmploymentCard), el menu contextual permite solicitar una ausencia directamente sin navegar a la seccion de Ausencias. Al seleccionar un tipo de ausencia, se abre el TimeOffRequestModal.")
    h3(doc, "Campos del formulario rapido")
    bullet(doc, "Tipo de ausencia (preseleccionado segun la opcion elegida del menu).")
    bullet(doc, "Fecha de inicio (CustomDateInput, minimo = hoy).")
    bullet(doc, "Calculo automatico de dias.")
    bullet(doc, "Notas (textarea).")
    h3(doc, "Alertas contextuales")
    bullet(doc, "Licencia por enfermedad > 3 dias: 'Se puede requerir un certificado medico para licencias superiores a 3 dias.'")
    bullet(doc, "Emergencia: 'Las ausencias por emergencia requieren justificacion y pueden necesitar aprobacion inmediata.'")
    h3(doc, "Permiso requerido")
    para(doc, "Este formulario esta protegido con PermissionGate usando el permiso employees.request_time_off. Si el empleado no tiene este permiso asignado, el boton no aparece.")

    # ------------------------------------------------------------------
    # 17. Widget de Balance de Vacaciones
    # ------------------------------------------------------------------
    h2(doc, "17. Widget de Balance de Vacaciones", anchor="sec_balance_vac")
    para(doc, "El VacationDaysWidget es un componente visual que muestra el estado actual del balance de vacaciones del empleado en el negocio seleccionado.")
    h3(doc, "Informacion mostrada")
    bullet(doc, "Numero grande con los dias restantes disponibles.", bold_prefix="Dias Restantes")
    bullet(doc, "Barra de progreso tricolor: verde (dias usados), amarillo (dias pendientes de aprobacion), azul (dias restantes).", bold_prefix="Barra de Progreso")
    bullet(doc, "Tres cajas de estadisticas: Total asignados, Usados, Pendientes.", bold_prefix="Stats")
    h3(doc, "Calculo del balance")
    para(doc, "El balance se obtiene de la tabla vacation_balance que tiene campos: year, total_days, used_days, pending_days, remaining_days. El calculo es: remaining = total - used - pending. Los defaults del negocio son 15 dias de vacaciones por ano.")
    screenshot_placeholder(doc, "VacationDaysWidget con barra de progreso")

    # ------------------------------------------------------------------
    # 18. Historial de ausencias
    # ------------------------------------------------------------------
    h2(doc, "18. Historial de Ausencias (Tabs)", anchor="sec_historial_aus")
    para(doc, "El componente EmployeeAbsencesTab organiza el historial de ausencias en 3 pestanas con contadores.")
    h3(doc, "Pestanas")
    simple_table(doc,
        ["Pestana", "Contenido", "Acciones"],
        [
            ["Pendientes", "Ausencias con status 'pending' esperando aprobacion", "Boton 'Cancelar' (con dialogo de confirmacion)"],
            ["Aprobadas", "Ausencias con status 'approved'", "Solo lectura"],
            ["Rechazadas", "Ausencias con status 'rejected'", "Solo lectura, muestra notas del admin"],
        ]
    )
    h3(doc, "Informacion por ausencia")
    bullet(doc, "Tipo de ausencia con badge de color (verde=vacaciones, rojo=emergencia, azul=enfermedad, gris=personal).")
    bullet(doc, "Rango de fechas (inicio - fin).")
    bullet(doc, "Cantidad de dias.")
    bullet(doc, "Razon proporcionada por el empleado.")
    bullet(doc, "Notas del administrador (solo en rechazadas o aprobadas con comentarios).")
    h3(doc, "Cancelar una solicitud pendiente")
    para(doc, "El empleado puede cancelar una solicitud que aun esta en estado 'pending'. Al hacer clic en Cancelar, se muestra un AlertDialog de confirmacion. Si acepta, el status cambia a 'cancelled' en la tabla employee_absences.")

    # ------------------------------------------------------------------
    # 19. Festivos
    # ------------------------------------------------------------------
    h2(doc, "19. Festivos Publicos y Validacion", anchor="sec_festivos")
    para(doc, "El sistema mantiene una tabla public_holidays con 54 festivos colombianos precargados para los anos 2025 a 2027. Estos festivos se validan automaticamente en los formularios de ausencia y en la seleccion de citas.")
    h3(doc, "Tipos de festivos")
    bullet(doc, "13 festivos fijos por ano: Ano Nuevo, Dia del Trabajo, Independencia, Navidad, etc.")
    bullet(doc, "5 festivos moviles por ano: basados en la fecha de Pascua (Carnaval, Semana Santa, Corpus Christi, etc.).")
    h3(doc, "Hook usePublicHolidays")
    para(doc, "El hook consulta la tabla public_holidays filtrando por pais y ano actual. Los datos se cachean durante 24 horas (staleTime) y 7 dias (gcTime). Expone dos helpers: isHoliday(date) que retorna true/false, y getHolidayName(date) que retorna el nombre del festivo.")

    # ------------------------------------------------------------------
    # 20. Horario de trabajo
    # ------------------------------------------------------------------
    h2(doc, "20. Configuracion de Horario de Trabajo", anchor="sec_horario")
    para(doc, "El WorkScheduleEditor permite al empleado configurar su horario semanal de trabajo y su hora de almuerzo. Cada dia de la semana tiene un control independiente.")
    h3(doc, "Elementos por dia")
    bullet(doc, "Switch de activacion: activo/inactivo (verde/gris).")
    bullet(doc, "Input de hora inicio (ej: 09:00).")
    bullet(doc, "Input de hora fin (ej: 18:00).")
    h3(doc, "Valores por defecto")
    simple_table(doc,
        ["Dia", "Activo", "Inicio", "Fin"],
        [
            ["Lunes", "Si", "09:00", "18:00"],
            ["Martes", "Si", "09:00", "18:00"],
            ["Miercoles", "Si", "09:00", "18:00"],
            ["Jueves", "Si", "09:00", "18:00"],
            ["Viernes", "Si", "09:00", "18:00"],
            ["Sabado", "No", "09:00", "14:00"],
            ["Domingo", "No", "09:00", "14:00"],
        ]
    )
    h3(doc, "Seccion de almuerzo")
    bullet(doc, "Toggle global para activar/desactivar hora de almuerzo.")
    bullet(doc, "Si esta activo: inputs de hora inicio y fin del almuerzo (default 12:00 - 13:00).")
    bullet(doc, "La hora de almuerzo se aplica a todos los dias activos y se guarda en business_employees (campos lunch_break_start y lunch_break_end).")
    h3(doc, "Persistencia")
    para(doc, "El horario semanal se guarda en la tabla work_schedules con upsert por employee_id + day_of_week. Los datos de almuerzo se guardan directamente en business_employees.")
    h3(doc, "Validaciones")
    bullet(doc, "La hora de inicio debe ser menor que la hora de fin para cada dia activo.")
    bullet(doc, "La hora de inicio del almuerzo debe ser menor que la hora de fin.")
    bullet(doc, "Permiso employees.edit_own_schedule requerido (pero se bypasea automaticamente si el empleado edita su propio horario).")
    screenshot_placeholder(doc, "WorkScheduleEditor — 7 dias + almuerzo")

    # ------------------------------------------------------------------
    # 21. Servicios
    # ------------------------------------------------------------------
    h2(doc, "21. Gestion de Servicios Ofrecidos", anchor="sec_servicios")
    para(doc, "El ServiceSelector permite al empleado seleccionar cuales de los servicios del negocio ofrece, con nivel de experiencia y porcentaje de comision por cada servicio.")
    h3(doc, "Elementos por servicio")
    bullet(doc, "Checkbox: toggle para activar/desactivar el servicio.")
    bullet(doc, "Nivel de experiencia: estrellas 1-5 (expertise_level).")
    bullet(doc, "Porcentaje de comision: slider 0-100% (commission_percentage).")
    h3(doc, "Flujo de guardado")
    numbered(doc, "El empleado marca/desmarca servicios y ajusta niveles y comisiones.")
    numbered(doc, "Hace clic en Guardar.")
    numbered(doc, "El sistema calcula el diff entre el estado original y el actual: servicios a agregar, a eliminar y a actualizar.")
    numbered(doc, "Se insertan nuevos registros en employee_services, se desactivan los eliminados (soft delete: is_active = false), y se actualizan los modificados.")
    h3(doc, "Requisitos")
    para(doc, "Se requiere que el empleado tenga una sede asignada (currentLocationId) para guardar los servicios.")
    h3(doc, "Deteccion de cambios")
    para(doc, "El boton Guardar solo se habilita si hay cambios (hasChanges se computa comparando el estado original con el estado actual). Si no hay cambios, el boton permanece deshabilitado.")
    screenshot_placeholder(doc, "ServiceSelector — Servicios con estrellas y comision")

    # ------------------------------------------------------------------
    # 22. Sedes
    # ------------------------------------------------------------------
    h2(doc, "22. Sedes y Selector de Ubicacion", anchor="sec_sedes")
    para(doc, "El LocationSelector muestra todas las sedes activas del negocio y permite al empleado seleccionar o cambiar su sede de trabajo.")
    h3(doc, "Informacion por sede")
    bullet(doc, "Nombre de la sede.")
    bullet(doc, "Badge 'Tu Sede' si es la sede actual del empleado.")
    bullet(doc, "Badge 'Primaria' si es la sede principal del negocio.")
    bullet(doc, "Direccion completa.")
    bullet(doc, "Contacto: email y telefono.")
    bullet(doc, "Horarios de apertura y cierre por dia de la semana (7 dias, formato en espanol).")
    bullet(doc, "Hasta 6 fotos de la sede.")
    h3(doc, "Acciones")
    bullet(doc, "Si el empleado NO tiene sede asignada: boton 'Seleccionar esta sede' que actualiza directamente business_employees.location_id.")
    bullet(doc, "Si el empleado YA tiene sede: boton 'Programar traslado' que abre LocationTransferModal.")
    h3(doc, "Estado de traslado")
    para(doc, "Si el empleado tiene un traslado pendiente, se muestra el TransferStatusBadge con: icono MapPin, texto 'Traslado programado', nombre de la sede destino y fecha efectiva del traslado.")

    # ------------------------------------------------------------------
    # 23. Traslado de Sede
    # ------------------------------------------------------------------
    h2(doc, "23. Traslado de Sede", anchor="sec_traslado")
    para(doc, "El LocationTransferModal permite programar un traslado de una sede a otra dentro del mismo negocio, con manejo de citas afectadas.")
    h3(doc, "Flujo completo")
    numbered(doc, "El empleado selecciona la sede destino.")
    numbered(doc, "Configura el periodo de preaviso con un slider de 1 a 30 dias.")
    numbered(doc, "El sistema calcula la fecha efectiva: hoy + dias de preaviso.")
    numbered(doc, "Se ejecuta getTransferImpact que calcula: citas que se mantienen (antes de la fecha efectiva) vs citas que se cancelaran (despues).")
    numbered(doc, "Si hay citas que se cancelaran, se muestra un warning con checkbox de confirmacion obligatorio.")
    numbered(doc, "Al confirmar, se actualizan los campos transfer_* en business_employees.")
    numbered(doc, "Se invoca la Edge Function cancel-future-appointments-on-transfer que cancela las citas afectadas y notifica a los clientes.")
    h3(doc, "Cancelar un traslado pendiente")
    para(doc, "Si el empleado tiene un traslado con status 'pending', puede cancelarlo. La cancelacion limpia todos los campos transfer_* y cambia el status a 'cancelled'.")
    h3(doc, "Campos almacenados")
    simple_table(doc,
        ["Campo", "Descripcion"],
        [
            ["transfer_from_location_id", "ID de la sede de origen"],
            ["transfer_to_location_id", "ID de la sede destino"],
            ["transfer_effective_date", "Fecha en que se hace efectivo el traslado"],
            ["transfer_notice_period_days", "Dias de preaviso configurados (1-30)"],
            ["transfer_scheduled_at", "Fecha/hora en que se programo el traslado"],
            ["transfer_scheduled_by", "ID del usuario que programo el traslado"],
            ["transfer_status", "Estado: pending, completed, cancelled"],
        ]
    )
    screenshot_placeholder(doc, "LocationTransferModal — Slider de preaviso y citas afectadas")

    # ------------------------------------------------------------------
    # 24. Marketplace de Vacantes
    # ------------------------------------------------------------------
    h2(doc, "24. Marketplace de Vacantes", anchor="sec_vacantes")
    para(doc, "El AvailableVacanciesMarketplace es la pagina donde los empleados exploran oportunidades laborales publicadas por otros negocios en la plataforma. Es un marketplace completo con filtros avanzados y sistema de matching.")
    h3(doc, "Barra de busqueda")
    para(doc, "Campo de texto que filtra en tiempo real (client-side) por titulo del cargo, nombre de la empresa o descripcion de la vacante.")
    h3(doc, "Panel de filtros (colapsable)")
    simple_table(doc,
        ["Filtro", "Tipo", "Descripcion"],
        [
            ["Pais", "Fijo", "Colombia (preseleccionado, no modificable)"],
            ["Departamento", "Select", "RegionSelect — lista de departamentos colombianos"],
            ["Ciudad", "Select", "CitySelect — cascada del departamento seleccionado"],
            ["Tipo de Posicion", "Select", "4 opciones: tiempo completo, medio tiempo, freelance, otro"],
            ["Nivel de Experiencia", "Select", "3 niveles: junior, intermedio, senior"],
            ["Salario Minimo", "Input", "Valor en COP"],
            ["Salario Maximo", "Input", "Valor en COP"],
            ["Solo Remoto", "Switch", "Toggle para filtrar solo vacantes remotas"],
        ]
    )
    h3(doc, "Ordenamiento")
    simple_table(doc,
        ["Criterio", "Descripcion"],
        [
            ["match_score", "Puntaje de matching con el perfil del empleado (default)"],
            ["salary", "Por salario (mayor a menor)"],
            ["published_at", "Por fecha de publicacion (mas recientes primero)"],
            ["applications_count", "Por cantidad de aplicaciones recibidas"],
        ]
    )
    h3(doc, "Tarjeta de vacante (VacancyCard)")
    para(doc, "Cada vacante se muestra como una tarjeta con: titulo del cargo, nombre del negocio, ubicacion, rango salarial (formato COP), tipo de posicion, nivel de experiencia, fecha de publicacion, numero de aplicaciones. Dos botones: 'Ver Detalles' (abre modal) y 'Aplicar' (abre formulario de aplicacion).")
    h3(doc, "Boton Mis Aplicaciones")
    para(doc, "En la parte superior de la pagina hay un boton que abre MyApplicationsModal con el historial de todas las aplicaciones enviadas por el empleado.")
    screenshot_placeholder(doc, "Marketplace de Vacantes con filtros expandidos")

    # ------------------------------------------------------------------
    # 25. Aplicar a una Vacante
    # ------------------------------------------------------------------
    h2(doc, "25. Aplicar a una Vacante", anchor="sec_aplicar")
    para(doc, "El ApplicationFormModal es el formulario para aplicar a una vacante seleccionada.")
    h3(doc, "Campos del formulario")
    bullet(doc, "Carta de presentacion (textarea, minimo 50 caracteres).", bold_prefix="Carta")
    bullet(doc, "Salario esperado (input numerico, en COP, opcional pero validado).", bold_prefix="Salario")
    bullet(doc, "Fecha de disponibilidad (CustomDateInput, minimo = hoy, requerido).", bold_prefix="Disponibilidad")
    bullet(doc, "CV (archivo adjunto, se sube al bucket 'cvs' de Supabase Storage).", bold_prefix="CV")
    h3(doc, "Validaciones")
    numbered(doc, "La carta de presentacion debe tener al menos 50 caracteres.")
    numbered(doc, "El salario esperado debe ser positivo.")
    numbered(doc, "El salario esperado no puede exceder el maximo de la vacante (si esta definido).")
    numbered(doc, "La fecha de disponibilidad es obligatoria y no puede ser en el pasado.")
    h3(doc, "Verificacion de conflictos de horario")
    para(doc, "Al abrir el modal, si la vacante tiene un work_schedule definido, el sistema automaticamente verifica conflictos con los empleos actuales del empleado. Si se detectan solapamientos, se muestra ScheduleConflictAlert con los detalles del conflicto. Esto es informativo, no bloquea la aplicacion.")
    h3(doc, "Flujo de envio")
    numbered(doc, "El empleado completa todos los campos requeridos.")
    numbered(doc, "Hace clic en Enviar.")
    numbered(doc, "Se sube el CV al storage (si existe).")
    numbered(doc, "Se crea un registro en job_applications con status 'pending'.")
    numbered(doc, "El administrador del negocio recibe notificacion.")
    numbered(doc, "La pagina se recarga para actualizar la lista de vacantes.")
    screenshot_placeholder(doc, "ApplicationFormModal con alerta de conflicto de horario")

    # ------------------------------------------------------------------
    # 26. Mis Aplicaciones
    # ------------------------------------------------------------------
    h2(doc, "26. Mis Aplicaciones", anchor="sec_mis_apps")
    para(doc, "El MyApplicationsModal muestra el historial de todas las aplicaciones enviadas por el empleado.")
    h3(doc, "Pestanas por estado")
    simple_table(doc,
        ["Estado", "Etiqueta", "Color", "Icono"],
        [
            ["all", "Todas", "—", "—"],
            ["pending", "Pendiente", "Amarillo", "Clock"],
            ["reviewing", "En revision", "Azul", "AlertCircle"],
            ["in_selection_process", "En Proceso de Seleccion", "Morado", "AlertCircle"],
            ["accepted", "Aceptada", "Verde", "CheckCircle"],
            ["rejected", "Rechazada", "Rojo", "XCircle"],
            ["withdrawn", "Retirada", "Gris", "XCircle"],
            ["completed", "Completada", "Esmeralda", "CheckCircle"],
        ]
    )
    h3(doc, "Informacion por aplicacion")
    bullet(doc, "Titulo de la vacante y nombre del negocio.")
    bullet(doc, "Estado con badge de color y icono.")
    bullet(doc, "Fecha de aplicacion.")
    bullet(doc, "Salario esperado (formato COP).")
    bullet(doc, "Fecha de disponibilidad.")
    bullet(doc, "Boton para descargar el CV adjunto (desde Supabase Storage).")
    h3(doc, "Contadores")
    para(doc, "Cada pestana muestra un contador con la cantidad de aplicaciones en ese estado. La pestana 'Todas' muestra el total.")

    # ------------------------------------------------------------------
    # 27. Conflictos de horario
    # ------------------------------------------------------------------
    h2(doc, "27. Deteccion de Conflictos de Horario", anchor="sec_conflictos")
    para(doc, "Cuando un empleado aplica a una vacante que tiene horario definido, el sistema verifica automaticamente si hay solapamientos con sus empleos actuales.")
    h3(doc, "Como funciona")
    numbered(doc, "Al abrir ApplicationFormModal, se obtiene el work_schedule de la vacante.")
    numbered(doc, "Se convierte al formato WorkSchedule: objeto con dias como claves, cada uno con enabled, start_time, end_time.")
    numbered(doc, "Se llama a checkConflict() del hook useScheduleConflicts.")
    numbered(doc, "El hook compara el horario de la vacante con los horarios de los empleos actuales del empleado.")
    numbered(doc, "Si hay solapamiento, retorna un array de ScheduleConflict con detalles: dia, rango en conflicto, nombre del negocio conflictivo.")
    h3(doc, "ScheduleConflictAlert")
    para(doc, "Si existen conflictos, se renderiza ScheduleConflictAlert que muestra un Alert de tipo warning con la lista de conflictos. El empleado puede proceder con la aplicacion a pesar de los conflictos (es informativo, no bloqueante).")
    callout(doc, "Los conflictos no bloquean la aplicacion",
            "El sistema informa al empleado sobre posibles conflictos de horario, pero no impide que aplique. Es responsabilidad del empleado y del administrador resolver los conflictos si se concreta la contratacion.",
            color=ACCENT)

    # ------------------------------------------------------------------
    # 28. Chat
    # ------------------------------------------------------------------
    h2(doc, "28. Chat en Tiempo Real", anchor="sec_chat")
    para(doc, "El empleado tiene acceso al sistema de chat para comunicarse con clientes y colegas del negocio.")
    h3(doc, "Acceso al chat")
    bullet(doc, "Boton flotante de chat (FloatingChatButton) en la esquina inferior derecha del dashboard.", bold_prefix="Boton Flotante")
    bullet(doc, "A traves de notificaciones de mensajes nuevos en NotificationBell.", bold_prefix="Notificaciones")
    h3(doc, "Componentes del chat")
    bullet(doc, "Lista de conversaciones con ultimo mensaje, hora, indicador de no leidos.", bold_prefix="ConversationList")
    bullet(doc, "Ventana de mensajes con burbujas de texto, indicador de escritura, confirmacion de lectura.", bold_prefix="ChatWindow")
    bullet(doc, "Campo de texto con soporte para adjuntos de archivos (bucket chat-attachments).", bold_prefix="ChatInput")
    h3(doc, "Funcionalidades")
    bullet(doc, "Envio y recepcion de mensajes en tiempo real via Supabase Realtime.")
    bullet(doc, "Adjuntos de archivos (imagenes, documentos).")
    bullet(doc, "Indicadores de escritura (typing indicators).")
    bullet(doc, "Confirmacion de lectura (read receipts).")
    bullet(doc, "Notificaciones por email de mensajes no leidos (Edge Function send-unread-chat-emails).")
    h3(doc, "Preferencia de mensajes")
    para(doc, "El empleado puede activar o desactivar la recepcion de mensajes de clientes desde Configuraciones -> Preferencias de Mensajes. Si esta desactivado, el empleado NO aparecera en la lista de empleados disponibles para chatear en el BusinessProfile del cliente.")

    # ------------------------------------------------------------------
    # 29. Notificaciones
    # ------------------------------------------------------------------
    h2(doc, "29. Notificaciones", anchor="sec_notificaciones")
    para(doc, "El empleado recibe notificaciones in-app y por email sobre eventos relevantes a su trabajo.")
    h3(doc, "Tipos de notificacion que recibe un empleado")
    simple_table(doc,
        ["Tipo", "Descripcion", "Accion al clic"],
        [
            ["appointment_new", "Nueva cita asignada", "Navega a Mis Citas"],
            ["appointment_confirmed", "Cita confirmada por cliente", "Navega a Mis Citas"],
            ["appointment_cancelled", "Cita cancelada", "Navega a Mis Citas"],
            ["appointment_rescheduled", "Cita reprogramada", "Navega a Mis Citas"],
            ["appointment_reminder", "Recordatorio de cita proxima", "Navega a Mis Citas"],
            ["absence_approved", "Solicitud de ausencia aprobada", "Navega a Mis Ausencias"],
            ["absence_rejected", "Solicitud de ausencia rechazada", "Navega a Mis Ausencias"],
            ["chat_message", "Nuevo mensaje de chat", "Abre la conversacion"],
            ["employee_request_approved", "Solicitud de vinculacion aprobada", "Navega a Mis Empleos"],
            ["employee_request_rejected", "Solicitud de vinculacion rechazada", "Navega a Mis Empleos"],
            ["vacancy_application_accepted", "Aplicacion a vacante aceptada", "Navega a Mis Aplicaciones"],
            ["vacancy_application_rejected", "Aplicacion a vacante rechazada", "Navega a Mis Aplicaciones"],
        ]
    )
    h3(doc, "Campana de notificaciones")
    para(doc, "En el header del dashboard se muestra NotificationBell con un contador de notificaciones no leidas. Al hacer clic se abre NotificationCenter con la lista completa de notificaciones, ordenadas por fecha descendente. Cada notificacion se puede marcar como leida individualmente o todas a la vez.")
    h3(doc, "Cambio automatico de rol")
    para(doc, "Si el empleado recibe una notificacion que requiere otro rol (ej: una notificacion de cita como cliente), el sistema cambia automaticamente el rol activo antes de navegar. Esto se gestiona por notificationRoleMapping.ts que mapea 30+ tipos de notificacion al rol requerido.")

    # ------------------------------------------------------------------
    # 30. Configuraciones
    # ------------------------------------------------------------------
    h2(doc, "30. Configuraciones del Empleado", anchor="sec_settings")
    para(doc, "Las configuraciones del empleado se acceden desde el menu de usuario en el header. Se muestran dentro de CompleteUnifiedSettings con la pestana 'Preferencias de Empleado' activa.")
    h3(doc, "Secciones de configuracion")
    simple_table(doc,
        ["Seccion", "Descripcion"],
        [
            ["Perfil publico compartible", "URL unica del perfil profesional con botones de copiar y abrir"],
            ["Disponibilidad y horarios", "Toggle de disponibilidad, notificaciones, recordatorios, horario visual 7 dias"],
            ["Preferencias de mensajes", "Toggle para recibir/rechazar mensajes de clientes (por negocio)"],
            ["Informacion profesional", "Resumen profesional (min 50 chars), anos de experiencia, tipo de trabajo"],
            ["Expectativas salariales", "Salario minimo y maximo esperado en COP"],
            ["Especializaciones", "Badges dinamicos con add/remove"],
            ["Idiomas", "Badges dinamicos con add/remove"],
            ["Certificaciones", "CRUD completo: nombre, emisor, fechas, credencial ID/URL"],
            ["Enlaces externos", "Portfolio, LinkedIn, GitHub"],
        ]
    )

    # ------------------------------------------------------------------
    # 31. Perfil Profesional
    # ------------------------------------------------------------------
    h2(doc, "31. Perfil Profesional Publico", anchor="sec_perfil_prof")
    para(doc, "El empleado tiene un perfil profesional publico accesible en /profesional/{userId}. Este perfil es visible para cualquier persona sin necesidad de autenticacion y puede ser compartido como enlace directo.")
    h3(doc, "Informacion del perfil")
    bullet(doc, "Nombre completo y avatar.")
    bullet(doc, "Resumen profesional.")
    bullet(doc, "Anos de experiencia.")
    bullet(doc, "Especializaciones y habilidades.")
    bullet(doc, "Certificaciones con detalle.")
    bullet(doc, "Idiomas.")
    bullet(doc, "Rating promedio y numero de resenas.")
    bullet(doc, "Enlaces a portfolio, LinkedIn y GitHub.")
    h3(doc, "Compartir")
    para(doc, "Desde Configuraciones, el empleado puede copiar la URL de su perfil publico o abrirla directamente en una nueva pestana para previsualizar como se ve.")

    # ------------------------------------------------------------------
    # 32. Certificaciones, Idiomas y Especializaciones
    # ------------------------------------------------------------------
    h2(doc, "32. Certificaciones, Idiomas y Especializaciones", anchor="sec_cert")
    h3(doc, "Certificaciones")
    para(doc, "Sistema CRUD completo para certificaciones profesionales. Cada certificacion tiene: nombre de la certificacion, nombre del emisor, fecha de emision, fecha de expiracion (opcional), ID de credencial (opcional), URL de credencial (opcional).")
    h3(doc, "Idiomas y Especializaciones")
    para(doc, "Sistema de badges dinamicos. El empleado puede agregar nuevos idiomas o especializaciones escribiendo el nombre y presionando Enter o haciendo clic en el boton de agregar. Para eliminar, hace clic en la X del badge. Los cambios se guardan al presionar el boton Guardar de la seccion correspondiente.")
    para(doc, "Estos datos se almacenan en la tabla employee_profiles como arrays JSONB.")

    # ------------------------------------------------------------------
    # 33. Preferencias de mensajes
    # ------------------------------------------------------------------
    h2(doc, "33. Preferencias de Mensajes de Clientes", anchor="sec_mensajes_pref")
    para(doc, "El empleado puede controlar si recibe mensajes de clientes a traves de un toggle en la seccion de Configuraciones. Esta preferencia es independiente por negocio — un empleado puede tener el toggle activo en un negocio y desactivo en otro.")
    h3(doc, "Comportamiento")
    bullet(doc, "Si allow_client_messages = true: el empleado aparece en la lista de empleados disponibles para chatear en el BusinessProfile.")
    bullet(doc, "Si allow_client_messages = false: el empleado NO aparece en esa lista. Los clientes no pueden iniciar conversaciones con el.")
    bullet(doc, "El valor por defecto es true (todos los empleados reciben mensajes a menos que lo desactiven).")
    para(doc, "El filtrado se hace a nivel de base de datos mediante el hook useBusinessEmployeesForChat, que consulta solo empleados con allow_client_messages = true.")

    # ------------------------------------------------------------------
    # 34. Finalizar empleo
    # ------------------------------------------------------------------
    h2(doc, "34. Finalizar Empleo", anchor="sec_fin_empleo")
    para(doc, "El empleado puede finalizar su vinculacion con un negocio desde la tarjeta de empleo. Esta accion es irreversible sin la intervencion de un administrador.")
    h3(doc, "Flujo")
    numbered(doc, "Desde la tarjeta de empleo, el empleado hace clic en el menu contextual y selecciona 'Marcar como Finalizado'.")
    numbered(doc, "Se abre ConfirmEndEmploymentDialog que muestra 4 consecuencias.")
    numbered(doc, "El empleado debe activar un checkbox de confirmacion obligatorio.")
    numbered(doc, "Al hacer clic en Confirmar: business_employees.is_active = false, termination_date = ahora.")
    numbered(doc, "Se desactivan automaticamente todos los employee_services del empleado en ese negocio.")
    numbered(doc, "La pagina se recarga.")
    h3(doc, "Consecuencias mostradas")
    numbered(doc, "No recibiras mas citas en este negocio.")
    numbered(doc, "Tus servicios seran desactivados.")
    numbered(doc, "Necesitas un administrador para reactivar tu vinculacion.")
    numbered(doc, "Tu historial de citas y datos se conservan.")
    h3(doc, "Restriccion para propietarios")
    para(doc, "La opcion 'Marcar como Finalizado' esta deshabilitada si el empleado es el propietario (owner) del negocio. Un propietario no puede autodesvincularse de su propio negocio.")
    screenshot_placeholder(doc, "ConfirmEndEmploymentDialog con checkbox de confirmacion")

    # ------------------------------------------------------------------
    # 35. Permisos del Empleado
    # ------------------------------------------------------------------
    h2(doc, "35. Permisos del Empleado", anchor="sec_permisos")
    para(doc, "Los permisos del empleado son asignados por el administrador del negocio usando templates o asignacion individual. El empleado no puede modificar sus propios permisos.")
    h3(doc, "Permisos relevantes para el empleado")
    simple_table(doc,
        ["Permiso", "Descripcion", "Donde se usa"],
        [
            ["employees.edit_own_schedule", "Editar su propio horario de trabajo", "WorkScheduleEditor"],
            ["employees.edit_own_profile", "Editar su perfil de empleado", "CompleteUnifiedSettings"],
            ["employees.request_time_off", "Solicitar ausencias y vacaciones", "TimeOffRequestModal"],
            ["appointments.view", "Ver sus citas asignadas", "EmployeeAppointmentsPage"],
        ]
    )
    para(doc, "El administrador puede asignar permisos adicionales al empleado, como ver reportes, gestionar servicios, etc. Los permisos se verifican via PermissionGate y el hook usePermissions.")

    # ------------------------------------------------------------------
    # 36. Modales y Dialogos
    # ------------------------------------------------------------------
    h2(doc, "36. Modales y Dialogos del Empleado", anchor="sec_modales")
    simple_table(doc,
        ["Modal", "Disparador", "Proposito"],
        [
            ["PhoneRequiredModal", "Automatico al entrar al dashboard sin telefono", "Verificar telefono via SMS OTP"],
            ["EmploymentDetailModal", "Click en tarjeta de empleo", "6 tabs con detalles completos del empleo"],
            ["ConfirmEndEmploymentDialog", "Menu > Marcar como Finalizado", "Confirmar desvinculacion del negocio"],
            ["TimeOffRequestModal", "Menu > Solicitar ausencia (por tipo)", "Formulario rapido de ausencia"],
            ["AbsenceRequestModal", "Boton en seccion Mis Ausencias", "Formulario avanzado con calendarios"],
            ["LocationTransferModal", "Boton 'Programar traslado' en sede", "Programar cambio de sede"],
            ["ApplicationFormModal", "Boton 'Aplicar' en vacante", "Formulario de aplicacion a vacante"],
            ["MyApplicationsModal", "Boton 'Mis Aplicaciones'", "Historial de aplicaciones enviadas"],
            ["ClientProfileModal", "Click en tarjeta de cliente", "Perfil y historial del cliente"],
            ["EmployeeAppointmentModal", "Click en cita (lista o calendario)", "Detalle de cita (solo lectura)"],
        ]
    )
    screenshot_placeholder(doc, "Ejemplo de modal — AbsenceRequestModal abierto")

    # ------------------------------------------------------------------
    # 37. Estados
    # ------------------------------------------------------------------
    h2(doc, "37. Estados de Carga, Error y Vacio", anchor="sec_estados")
    simple_table(doc,
        ["Componente", "Loading", "Error", "Vacio"],
        [
            ["Mis Empleos", "Skeleton cards", "Alert con mensaje", "Mensaje invitando a unirse a un negocio"],
            ["Mis Citas", "Spinner central", "Toast de error", "Mensaje 'No tienes citas' con icono Calendar"],
            ["Mis Clientes", "Skeleton grid", "Alert con retry", "Mensaje 'Aun no has atendido clientes'"],
            ["Mis Ausencias", "Spinner", "Toast error", "Mensaje por pestana vacia"],
            ["Marketplace Vacantes", "Skeleton cards", "Alert", "Mensaje 'No hay vacantes disponibles'"],
            ["Chat", "Skeleton conversaciones", "Reconnect button", "Mensaje 'No hay conversaciones'"],
            ["Horario", "Spinner", "Toast", "Se muestran valores por defecto"],
            ["Servicios", "Spinner", "Alert", "Mensaje 'No hay servicios disponibles'"],
        ]
    )

    # ------------------------------------------------------------------
    # 38. Limitaciones por plan
    # ------------------------------------------------------------------
    h2(doc, "38. Limitaciones por Plan", anchor="sec_plan_limits")
    para(doc, "Las funcionalidades del rol Empleado estan disponibles en todos los planes. Las restricciones de plan aplican al negocio (administrador), no al empleado individual.")
    simple_table(doc,
        ["Restriccion del Negocio", "Gratis", "Basico", "Pro"],
        [
            ["Empleados permitidos", "1", "10", "Ilimitados"],
            ["Sedes", "1", "3", "Ilimitadas"],
            ["Servicios", "5", "50", "Ilimitados"],
            ["Citas/mes", "30", "500", "Ilimitadas"],
        ]
    )
    para(doc, "Si el negocio esta en plan Gratuito con maximo 1 empleado, solo el propietario puede operar como empleado. Al actualizar al plan Basico, se pueden agregar hasta 10 empleados adicionales.")
    callout(doc, "El empleado no paga",
            "Los empleados no tienen costos asociados en Gestabiz. El plan es una responsabilidad del administrador del negocio. Todas las funcionalidades de empleado estan disponibles independientemente del plan del negocio, salvo que el cupo de empleados se haya agotado.",
            color=ACCENT)

    # ------------------------------------------------------------------
    # 39. Glosario
    # ------------------------------------------------------------------
    h2(doc, "39. Glosario de Terminos del Empleado", anchor="sec_glosario")
    simple_table(doc,
        ["Termino", "Definicion"],
        [
            ["business_employees", "Tabla que vincula un usuario con un negocio como empleado. Contiene rol, sede, horario, salario, estado."],
            ["employee_services", "Relacion muchos a muchos entre empleado y servicios del negocio. Incluye nivel de experiencia y comision."],
            ["work_schedules", "Horario semanal del empleado: 7 registros (uno por dia) con dia, hora inicio, hora fin, activo/inactivo."],
            ["employee_absences", "Solicitudes de ausencia del empleado con tipo, fechas, razon, estado (pending/approved/rejected/cancelled)."],
            ["vacation_balance", "Balance anual de vacaciones: total, usados, pendientes, restantes."],
            ["employee_profiles", "Perfil profesional extendido: resumen, experiencia, certificaciones, idiomas, especializaciones."],
            ["employee_join_requests", "Solicitudes de vinculacion a negocios. Estados: pending, approved, rejected."],
            ["job_applications", "Aplicaciones a vacantes laborales. Estados: pending, reviewing, in_selection_process, accepted, rejected, withdrawn, completed."],
            ["PermissionGate", "Componente React que envuelve botones y verifica permisos antes de renderizar o habilitar la accion."],
            ["effectiveBusinessId", "ID del negocio seleccionado actualmente en el dashboard del empleado. Scopa todas las queries."],
            ["transfer_status", "Estado de un traslado de sede: pending (en espera), completed (ejecutado), cancelled (anulado)."],
            ["allow_client_messages", "Flag booleano en business_employees que controla si el empleado recibe mensajes de clientes."],
            ["lunch_break_start/end", "Hora de inicio y fin del almuerzo del empleado. Bloquea slots de citas en ese rango."],
            ["expertise_level", "Nivel de experiencia del empleado en un servicio (1-5 estrellas)."],
            ["commission_percentage", "Porcentaje de comision que recibe el empleado por cada servicio realizado (0-100%)."],
        ]
    )

    return doc


# ===========================================================================
# PROPUESTA DE VALOR — PARTE 2: EXPERIENCIA DEL EMPLEADO
# ===========================================================================

def build_proposal_part2() -> Document:
    doc = setup_document(
        title="Propuesta de Valor",
        subtitle="Por que tu equipo elegira Gestabiz"
    )

    # =====================================================================
    # INDICE
    # =====================================================================
    h1(doc, "Indice — Parte 2: Experiencia del Empleado", anchor="toc_prop")
    para(doc, "Este documento presenta la propuesta de valor de Gestabiz enfocada en la experiencia del profesional — el empleado que ofrece servicios. Es el segundo de cinco volumenes.")

    prop_toc = [
        ("prop_resumen",     "Parte 1 — Resumen Ejecutivo"),
        ("prop_problema",    "1. El Problema: Profesionales Desorganizados"),
        ("prop_solucion",    "2. La Solucion: Un Espacio de Trabajo Profesional"),
        ("prop_ventajas",    "3. Ventajas Competitivas para el Empleado"),
        ("prop_comparativa", "4. Comparativa con la Competencia"),
        ("prop_detalle",     "Parte 2 — Detalle Comercial"),
        ("prop_productividad", "5. Aumento de Productividad"),
        ("prop_autonomia",   "6. Autonomia sin Perder Control"),
        ("prop_vacantes",    "7. Marketplace de Oportunidades Laborales"),
        ("prop_ausencias",   "8. Gestion de Ausencias Digitalizada"),
        ("prop_traslados",   "9. Traslados de Sede sin Friccion"),
        ("prop_perfil",      "10. Perfil Profesional y Reputacion"),
        ("prop_multiempleo", "11. Multi-empleo: Trabaja en Varios Negocios"),
        ("prop_comunicacion","12. Comunicacion Integrada"),
        ("prop_testimonios", "13. Escenarios de Uso Real"),
        ("prop_faq",         "14. Preguntas Frecuentes"),
        ("prop_cta",         "15. Llamada a la Accion"),
    ]
    for anchor, label in prop_toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor, label)

    page_break(doc)

    # =====================================================================
    # PARTE 1 — RESUMEN EJECUTIVO
    # =====================================================================
    h1(doc, "Parte 1 — Resumen Ejecutivo", anchor="prop_resumen")

    # 1. El Problema
    h2(doc, "1. El Problema: Profesionales Desorganizados", anchor="prop_problema")
    para(doc, "Hoy en dia, miles de profesionales de servicios en Colombia — peluqueros, esteticistas, fisioterapeutas, entrenadores, medicos y mas — enfrentan problemas criticos que afectan su productividad y calidad de vida:")
    bullet(doc, "Dependen de grupos de WhatsApp para saber su agenda del dia.")
    bullet(doc, "No tienen visibilidad sobre sus proximas citas hasta que alguien les avisa.")
    bullet(doc, "Solicitar vacaciones o una licencia implica hablar con el jefe verbalmente, sin registro.")
    bullet(doc, "No tienen control sobre su horario: el dueno decide todo sin consultar.")
    bullet(doc, "No pueden demostrar su experiencia profesional de forma unificada.")
    bullet(doc, "Si trabajan en varios sitios, tienen que gestionar todo por separado.")

    callout(doc, "El costo de la desorganizacion",
            "Un profesional de servicios pierde en promedio 45 minutos diarios coordinando su agenda por WhatsApp y llamadas. Eso son 16 horas al mes — el equivalente a 2 dias de trabajo completos perdidos.",
            color=DANGER)

    # 2. La Solucion
    h2(doc, "2. La Solucion: Un Espacio de Trabajo Profesional", anchor="prop_solucion")
    para(doc, "Gestabiz le entrega a cada profesional un espacio de trabajo digital completo, diseñado para que pueda operar con autonomia, profesionalismo y eficiencia. No es otra app de mensajes: es una herramienta de trabajo real.")
    h3(doc, "6 Pilares de la Solucion para el Empleado")
    numbered(doc, "Agenda digital en tiempo real: ve sus citas del dia, semana y mes sin preguntar a nadie.")
    numbered(doc, "Gestion de ausencias formalizada: solicita vacaciones y licencias con un flujo digital con aprobacion, balance y festivos integrados.")
    numbered(doc, "Horario configurable: define sus propios dias de trabajo, horarios y hora de almuerzo.")
    numbered(doc, "Perfil profesional publico: certificaciones, especializaciones, reviews y enlace compartible.")
    numbered(doc, "Multi-empleo centralizado: gestiona todos sus trabajos desde una sola app.")
    numbered(doc, "Marketplace de vacantes: explora oportunidades laborales y aplica sin salir de la plataforma.")
    screenshot_placeholder(doc, "Dashboard del Empleado — Todas las herramientas en un solo lugar")

    # 3. Ventajas Competitivas
    h2(doc, "3. Ventajas Competitivas para el Empleado", anchor="prop_ventajas")
    simple_table(doc,
        ["Funcion", "Sin Gestabiz", "Con Gestabiz"],
        [
            ["Ver agenda", "Preguntar al jefe por WhatsApp", "Dashboard con citas en tiempo real"],
            ["Solicitar vacaciones", "Pedir verbalmente, sin registro", "Formulario digital con aprobacion y balance"],
            ["Saber cuantos dias de vacaciones quedan", "Calculo manual o no lo sabe", "Widget automatico con barra de progreso"],
            ["Mostrar certificaciones", "CV en PDF desactualizado", "Perfil profesional publico y actualizable"],
            ["Trabajar en varios sitios", "Multiples apps y agendas", "Una app con selector de negocio"],
            ["Buscar nuevo empleo", "Portales externos (Computrabajo, etc.)", "Marketplace integrado con matching"],
            ["Chatear con clientes", "WhatsApp personal", "Chat profesional dentro de la plataforma"],
            ["Cambiar de sede", "Proceso informal, citas perdidas", "Traslado programado con manejo de citas"],
        ]
    )

    # 4. Comparativa
    h2(doc, "4. Comparativa con la Competencia", anchor="prop_comparativa")
    para(doc, "Comparamos la experiencia del empleado en Gestabiz vs las principales plataformas del mercado:")
    simple_table(doc,
        ["Funcion", "Calendly", "Booksy", "Fresha", "Gestabiz"],
        [
            ["Dashboard propio", "No", "Limitado", "Si", "Si (completo)"],
            ["Agenda en tiempo real", "No", "Si", "Si", "Si + Realtime"],
            ["Solicitud de ausencias", "No", "No", "Basico", "Avanzado (5 tipos, aprobacion, balance)"],
            ["Marketplace de vacantes", "No", "No", "No", "Si (con matching)"],
            ["Perfil profesional publico", "No", "Si (basico)", "Si (basico)", "Si (completo, compartible)"],
            ["Multi-empleo", "No", "No", "No", "Si"],
            ["Traslado de sede", "No", "No", "No", "Si (con preaviso)"],
            ["Chat integrado", "No", "Si", "Si", "Si (con preferencias)"],
            ["Gestion de comisiones", "No", "No", "Si", "Si (por servicio)"],
            ["Certificaciones digitales", "No", "No", "No", "Si (CRUD completo)"],
        ]
    )
    callout(doc, "Gestabiz: 10 de 10 funciones",
            "Mientras que competidores como Calendly, Booksy y Fresha cubren entre 2 y 6 de las 10 funciones clave para el empleado, Gestabiz las cubre TODAS. Y no con funcionalidad basica, sino con flujos completos, validaciones y experiencia premium.",
            color=ACCENT)

    page_break(doc)

    # =====================================================================
    # PARTE 2 — DETALLE COMERCIAL
    # =====================================================================
    h1(doc, "Parte 2 — Detalle Comercial", anchor="prop_detalle")

    # 5. Productividad
    h2(doc, "5. Aumento de Productividad", anchor="prop_productividad")
    para(doc, "Con Gestabiz, el empleado pasa de ser un participante pasivo a un profesional autonomo con herramientas reales:")
    h3(doc, "Metricas de mejora")
    simple_table(doc,
        ["Metrica", "Sin Gestabiz", "Con Gestabiz", "Mejora"],
        [
            ["Tiempo coordinando agenda/dia", "45 minutos", "5 minutos", "-89%"],
            ["Citas perdidas por falta de informacion/mes", "3-5", "0-1", "-80%"],
            ["Tiempo solicitando ausencia", "1-2 dias (ida y vuelta)", "5 minutos", "-99%"],
            ["Visibilidad del horario", "Solo si pregunta", "24/7 en la app", "Total"],
            ["Acceso a metricas personales", "Nunca", "Siempre", "Nuevo"],
        ]
    )
    callout(doc, "ROI para el profesional",
            "Un empleado que ahorra 40 minutos diarios de coordinacion puede atender 1-2 citas adicionales al dia. Con un precio promedio de servicio de $35,000 COP, eso son hasta $70,000 COP diarios adicionales = $1,540,000 COP/mes de ingresos recuperados.",
            color=ACCENT)

    # 6. Autonomia
    h2(doc, "6. Autonomia sin Perder Control", anchor="prop_autonomia")
    para(doc, "Gestabiz equilibra perfectamente la autonomia del empleado con el control del administrador:")
    bullet(doc, "El empleado configura su propio horario y servicios, pero el administrador los aprueba.", bold_prefix="Horario")
    bullet(doc, "El empleado solicita ausencias, pero requiere aprobacion obligatoria.", bold_prefix="Ausencias")
    bullet(doc, "El empleado puede ver sus metricas, pero no puede modificar citas.", bold_prefix="Citas")
    bullet(doc, "El empleado controla quien le escribe, pero no puede limitar funciones del negocio.", bold_prefix="Mensajes")
    para(doc, "Este balance es clave: el empleado se siente empoderado sin que el administrador pierda visibilidad ni control sobre las operaciones del negocio.")

    # 7. Marketplace
    h2(doc, "7. Marketplace de Oportunidades Laborales", anchor="prop_vacantes")
    para(doc, "Un diferenciador unico de Gestabiz es el marketplace de vacantes integrado. Los profesionales pueden explorar oportunidades laborales publicadas por otros negocios, aplicar con CV y carta de presentacion, y ser contactados directamente dentro de la plataforma.")
    h3(doc, "Ventajas sobre portales de empleo tradicionales")
    bullet(doc, "Las vacantes son especificas de la industria de servicios (no hay ofertas de oficina o fabrica mezcladas).")
    bullet(doc, "El matching automatico considera las habilidades, experiencia y horarios del profesional.")
    bullet(doc, "La deteccion de conflictos de horario avisa al profesional si la nueva vacante se solapa con sus trabajos actuales.")
    bullet(doc, "El proceso completo ocurre dentro de Gestabiz: buscar, aplicar, ser seleccionado, onboarding.")
    bullet(doc, "El profesional no necesita crear un perfil nuevo en cada portal — su perfil profesional de Gestabiz funciona como CV vivo.")
    screenshot_placeholder(doc, "Marketplace de vacantes con filtros y cards")

    # 8. Ausencias
    h2(doc, "8. Gestion de Ausencias Digitalizada", anchor="prop_ausencias")
    para(doc, "El sistema de ausencias de Gestabiz transforma un proceso informal y propenso a errores en un flujo digital completo:")
    h3(doc, "El ciclo completo")
    numbered(doc, "El profesional abre la app y selecciona el tipo de ausencia.")
    numbered(doc, "Elige las fechas en calendarios visuales que marcan festivos y dias no laborables.")
    numbered(doc, "El sistema calcula automaticamente los dias habiles, descuenta festivos, y muestra citas afectadas.")
    numbered(doc, "Envia la solicitud con un clic.")
    numbered(doc, "Todos los administradores del negocio reciben notificacion inmediata (in-app + email).")
    numbered(doc, "El administrador aprueba o rechaza con notas explicativas.")
    numbered(doc, "El profesional recibe la respuesta al instante.")
    numbered(doc, "Si es ausencia de emergencia aprobada: las citas se cancelan automaticamente y los clientes son notificados.")
    callout(doc, "Cero conflictos, cero malentendidos",
            "Al digitalizar las ausencias, se eliminan las discusiones tipo 'yo te dije que iba a salir' o 'no me acuerdo que pidio vacaciones'. Todo queda registrado con fechas, estados y notas.",
            color=PURPLE)

    # 9. Traslados
    h2(doc, "9. Traslados de Sede sin Friccion", anchor="prop_traslados")
    para(doc, "Cuando un profesional necesita cambiar de sede dentro del mismo negocio, Gestabiz maneja todo el proceso de forma automatizada:")
    bullet(doc, "El profesional elige la nueva sede y define un periodo de preaviso (1-30 dias).")
    bullet(doc, "El sistema calcula cuantas citas se veran afectadas por el cambio.")
    bullet(doc, "Las citas despues de la fecha efectiva se cancelan automaticamente con notificacion a los clientes.")
    bullet(doc, "Las citas anteriores a la fecha se mantienen intactas.")
    bullet(doc, "El traslado se puede cancelar antes de que se haga efectivo.")
    para(doc, "Sin Gestabiz, este proceso implica llamar a cada cliente, reorganizar la agenda manualmente, y arriesgarse a que alguien quede sin atender.")

    # 10. Perfil Profesional
    h2(doc, "10. Perfil Profesional y Reputacion", anchor="prop_perfil")
    para(doc, "Cada profesional en Gestabiz construye una reputacion digital que es suya — no del negocio:")
    bullet(doc, "Perfil publico con URL compartible (/profesional/{id}).")
    bullet(doc, "Certificaciones verificables con emisor, fecha y enlace de credencial.")
    bullet(doc, "Reviews de clientes reales (solo clientes con citas completadas pueden dejar reviews).")
    bullet(doc, "Especializaciones e idiomas como badges visibles.")
    bullet(doc, "Rating acumulado a lo largo de toda su carrera en Gestabiz, no solo en un negocio.")
    para(doc, "Esto crea un 'curriculum viviente' que el profesional puede compartir con potenciales empleadores, clientes directos o en redes sociales. Es un diferenciador enorme frente a competidores que no ofrecen perfiles de empleados.")

    # 11. Multi-empleo
    h2(doc, "11. Multi-empleo: Trabaja en Varios Negocios", anchor="prop_multiempleo")
    para(doc, "Muchos profesionales de servicios trabajan en mas de un lugar: un peluquero puede estar en un salon lunes a miercoles y en otro jueves a sabado. Gestabiz es la unica plataforma del mercado que soporta esto de forma nativa:")
    bullet(doc, "Un selector de negocio en el dashboard permite cambiar de contexto al instante.")
    bullet(doc, "Cada negocio tiene su propio horario, servicios, ausencias y clientes.")
    bullet(doc, "El sistema detecta conflictos de horario al aplicar a nuevas vacantes.")
    bullet(doc, "La preferencia de negocio se persiste en localStorage para acceso rapido.")
    callout(doc, "Ningun competidor ofrece esto",
            "Ni Calendly, ni Booksy, ni Fresha, ni ningun competidor local permite que un profesional gestione multiples empleos desde una sola cuenta. En Gestabiz es una funcionalidad nativa y fluida.",
            color=ACCENT)

    # 12. Comunicacion
    h2(doc, "12. Comunicacion Integrada", anchor="prop_comunicacion")
    para(doc, "El chat de Gestabiz reemplaza los mensajes de WhatsApp personal con un canal profesional dentro de la plataforma:")
    bullet(doc, "Conversaciones separadas por contexto laboral — no se mezclan con chats personales.")
    bullet(doc, "Control de privacidad: el profesional decide si acepta mensajes de clientes.")
    bullet(doc, "Adjuntos de archivos (fotos del trabajo, documentos).")
    bullet(doc, "Notificaciones por email si hay mensajes sin leer.")
    bullet(doc, "Sin exponer el numero de telefono personal del profesional.")

    # 13. Escenarios de uso
    h2(doc, "13. Escenarios de Uso Real", anchor="prop_testimonios")

    callout(doc, "Maria — Esteticista en Medellin",
            "Trabaja en 2 salones de belleza. Antes usaba 2 apps diferentes y un grupo de WhatsApp por cada salon. Ahora abre Gestabiz, selecciona el salon y ve toda su agenda. En un salon tiene turno L-M-X y en el otro J-V-S. 'Ya no se me olvida donde tengo que ir cada dia.'",
            color=PURPLE)

    callout(doc, "Carlos — Fisioterapeuta en Bogota",
            "Necesitaba pedir 3 dias de licencia por enfermedad. Antes tenia que llamar al administrador, esperar que le respondiera, y luego confirmar. Con Gestabiz, abrio el modal de ausencias, selecciono las fechas, y en 2 minutos la solicitud estaba enviada. El admin la aprobo desde su celular. 'Todo quedo registrado sin malentendidos.'",
            color=PURPLE)

    callout(doc, "Laura — Entrenadora Personal en Cali",
            "Queria cambiar de sede dentro de su gimnasio porque la sede norte le quedaba mas cerca. Abrio Gestabiz, selecciono la nueva sede, puso 15 dias de preaviso, y el sistema cancelo automaticamente las citas de la sede vieja y notifico a sus clientes. 'Antes esto me habria tomado una semana de llamadas.'",
            color=PURPLE)

    # 14. FAQ
    h2(doc, "14. Preguntas Frecuentes", anchor="prop_faq")

    h4(doc, "El empleado tiene que pagar algo?")
    para(doc, "No. Gestabiz no cobra a los empleados. El costo lo asume el negocio a traves de su plan de suscripcion.")

    h4(doc, "Puedo usar Gestabiz si trabajo en varios negocios?")
    para(doc, "Si. Gestabiz es la unica plataforma del mercado que soporta multi-empleo de forma nativa. Puedes cambiar de negocio con un clic en el selector del dashboard.")

    h4(doc, "Que pasa si me desvinculo de un negocio?")
    para(doc, "Tu historial de citas, reviews y datos se conservan. Solo se desactiva tu vinculacion (is_active = false). Un administrador puede reactivarte si es necesario.")

    h4(doc, "Mis clientes pueden escribirme por chat?")
    para(doc, "Solo si tu lo permites. En Configuraciones puedes activar o desactivar la recepcion de mensajes de clientes. Si lo desactivas, no apareceras en la lista de empleados disponibles para chatear.")

    h4(doc, "Puedo cancelar o mover una cita?")
    para(doc, "No directamente. Las citas solo pueden ser modificadas por el cliente (reprogramar/cancelar) o por el administrador del negocio. El empleado ve su agenda en modo lectura.")

    h4(doc, "Como solicito vacaciones?")
    para(doc, "Desde la seccion Mis Ausencias, haces clic en Solicitar Ausencia, seleccionas tipo 'Vacaciones', eliges las fechas en el calendario, y envias. El sistema calcula automaticamente los dias habiles descontando festivos y fines de semana. Tu administrador recibe la solicitud al instante.")

    h4(doc, "Puedo ver cuantos dias de vacaciones me quedan?")
    para(doc, "Si. El widget de balance de vacaciones muestra tus dias totales, usados, pendientes y restantes con una barra de progreso visual.")

    h4(doc, "Que pasa si hay un conflicto de horario al aplicar a una vacante?")
    para(doc, "El sistema te avisa automaticamente si hay solapamiento con tus empleos actuales. Es informativo — puedes continuar con la aplicacion si lo deseas. Sera responsabilidad tuya y del nuevo empleador resolver el conflicto.")

    # 15. CTA
    h2(doc, "15. Llamada a la Accion", anchor="prop_cta")

    callout(doc, "Tu carrera profesional merece una herramienta profesional",
            "Si eres un profesional de servicios que quiere dejar de depender de WhatsApp, tener control real sobre su agenda, construir una reputacion digital y acceder a oportunidades laborales exclusivas — Gestabiz es tu plataforma. Registrate gratis y descubre todo lo que puedes lograr.",
            color=ACCENT)

    para(doc, "Invita a tu negocio a unirse a Gestabiz, o unete como empleado a un negocio que ya este en la plataforma. Tu agenda, tu perfil, tus certificaciones, tus reviews — todo en un solo lugar.", size=12)

    screenshot_placeholder(doc, "Dashboard del Empleado — Registro en 2 minutos")

    for _ in range(3):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_TITURING), width=Cm(3.0))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing — https://github.com/TI-Turing")
    style_run(r, size=10, italic=True, color=GREY)

    return doc


# ===========================================================================
# MAIN
# ===========================================================================

def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # --- Manual de Usuario ---
    print("Generando Manual de Usuario — Parte 2: Rol Empleado...")
    manual = build_manual_part2()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz - copilot - parte2.docx"
    manual.save(str(manual_path))
    print(f"  -> {manual_path}")

    # --- Propuesta de Valor ---
    print("Generando Propuesta de Valor — Parte 2: Experiencia del Empleado...")
    propuesta = build_proposal_part2()
    propuesta_path = DOCS_DIR / "Propuesta_Valor_Gestabiz - copilot - parte2.docx"
    propuesta.save(str(propuesta_path))
    print(f"  -> {propuesta_path}")

    print("\nListo. Ambos documentos generados exitosamente.")


if __name__ == "__main__":
    main()
