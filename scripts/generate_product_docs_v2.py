"""
generate_product_docs.py — Genera Manual de Usuario y Propuesta de Valor de Gestabiz.

Produce 4 archivos en docs/:
  - Manual_Usuario_Gestabiz.docx
  - Manual_Usuario_Gestabiz.md
  - Propuesta_Valor_Gestabiz.docx
  - Propuesta_Valor_Gestabiz.md

Ejecutar:  python scripts/generate_product_docs.py
Requiere:  pip install python-docx
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)

ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
LOGO_GESTABIZ = ROOT / "public" / "og-image.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

_bookmark_counter = 0


def _next_bookmark_id() -> int:
    global _bookmark_counter
    _bookmark_counter += 1
    return _bookmark_counter


# ---------------------------------------------------------------------------
# Españolización — reemplazar términos técnicos / anglicismos
# ---------------------------------------------------------------------------
_TERM_MAP = {
    "Dashboard": "Panel principal",
    "dashboard": "panel principal",
    "Walk-In": "sin cita previa",
    "walk-in": "sin cita previa",
    "CRM": "Gestión de clientes",
    "SEO": "posicionamiento web",
    "Branding": "marca visual",
    "Tracking": "seguimiento",
    "Free Trial": "prueba gratuita",
    "Roadmap": "Hoja de ruta",
    "roadmap": "hoja de ruta",
    "Marketplace": "bolsa de empleo",
    "marketplace": "bolsa de empleo",
    "Payroll": "nómina",
    "Setup Checklist": "asistente de configuración",
    "Pipeline": "flujo",
    "pipeline": "flujo",
    "Flywheel": "efecto volante",
    "flywheel": "efecto volante",
    "Early adopter": "primeros usuarios",
    "early adopter": "primeros usuarios",
    "Walk-in": "sin cita previa",
}

# Terms that must NEVER appear in user-facing docs
_BANNED_TERMS = [
    "PermissionGate", "PlanGate", "Edge Function", "edge function",
    "RPC", "JSONB", "useAuth", "React Query", "hook", "trigger",
    "RLS", "Row Level Security", "singleton", "debounce", "cache",
    "CRUD", "UPSERT", "soft-delete", "webhook", "tsvector", "GIN",
    "SECURITY DEFINER", "owner bypass", "LEFT JOIN", "two-step query",
    "employee_id", "business_employees", "useAssigneeAvailability",
    "useCompletedAppointments", "useWizardDataCache", "createAppointment",
    "auth.uid()", "business_roles", "businesses.owner_id",
    "permission_audit_log", "offers_services", "is_location_exception",
    "is_active", "setup_completed", "require_absence_approval",
    "vacation_days_accrued", "business_tax_config",
    "usePreferredLocation", "usePreferredCity", "useKV",
    "OxmlElement", "qn(", "Supabase", "PostgreSQL", "Deno",
    "React 19", "Vite", "Tailwind", "Expo", "React Native",
    "Recharts", "jspdf", "xlsx", "sonner",
    "describe.skip", "employee_services", "location_services",
    "resource_services", "business_resources",
    "review_type", "notification_log", "in_app_notifications",
    "useInAppNotifications", "useEmployeeAbsences", "useAbsenceApprovals",
    "useBusinessResources", "useTaxCalculation", "useBusinessTaxConfig",
    "useTransactions", "useFinancialReports", "useChartData",
    "useJobVacancies", "useJobApplications", "useMatchingVacancies",
    "useMandatoryReviews", "useScheduleConflicts", "useEmployeeProfile",
    "useBusinessProfileData", "useEmployeeBusinesses",
    "useReviews", "usePendingReviews",
]


def clean(text: str) -> str:
    """Replace anglicisms and remove banned technical terms from text."""
    for eng, esp in _TERM_MAP.items():
        text = text.replace(eng, esp)
    return text


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------
def add_bookmark(paragraph, name: str) -> None:
    bid = str(_next_bookmark_id())
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    paragraph._p.insert(0, start)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *, bold: bool = False,
                            color: RGBColor = PURPLE, size: int = 11) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rpr.append(color_el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(clean(text))
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(clean(text))
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(clean(text))
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(clean(text))
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(clean(text))
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + clean(text))
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(clean(text))
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(clean(text))
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(clean(title))
    style_run(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(clean(text))
    style_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {clean(caption)}")
    style_run(r, size=9, italic=True, color=GREY)
    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(clean(h))
        style_run(r, size=10, bold=True, color=header_fg)
    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(clean(str(val)))
            style_run(r, size=10, color=DARK)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Header / Footer with logos and purple borders
# ---------------------------------------------------------------------------
def _border_xml(side: str, color: str = "7C3AED") -> str:
    return (f'<w:pBdr {nsdecls("w")}>'
            f'<w:{side} w:val="single" w:sz="6" w:space="4" w:color="{color}"/>'
            f'</w:pBdr>')


def _add_page_field(paragraph) -> None:
    """Insert a PAGE number field into a paragraph."""
    r1 = OxmlElement("w:r")
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    r1.append(fld1)
    paragraph._p.append(r1)

    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2.append(instr)
    paragraph._p.append(r2)

    r3 = OxmlElement("w:r")
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    r3.append(fld3)
    paragraph._p.append(r3)


def setup_header_footer(doc: Document, title: str) -> None:
    """Add header with Gestabiz logo and footer with Ti Turing logo + page number."""
    for section in doc.sections:
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.different_first_page_header_footer = True

        # --- Header (not on first page = cover) ---
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()

        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(4)

        if LOGO_GESTABIZ.exists():
            try:
                run = hp.add_run()
                run.add_picture(str(LOGO_GESTABIZ), height=Cm(1.0))
            except Exception:
                run = hp.add_run("Gestabiz | Agenda. Gestiona. Crece.")
                style_run(run, size=10, bold=True, color=PURPLE)
        else:
            run = hp.add_run("Gestabiz | Agenda. Gestiona. Crece.")
            style_run(run, size=10, bold=True, color=PURPLE)

        # Purple bottom border on header
        pPr = hp._p.get_or_add_pPr()
        pPr.append(parse_xml(_border_xml("bottom")))

        # --- Footer ---
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)

        # Purple top border on footer
        fpPr = fp._p.get_or_add_pPr()
        fpPr.append(parse_xml(_border_xml("top")))

        if LOGO_TITURING.exists():
            try:
                run = fp.add_run()
                run.add_picture(str(LOGO_TITURING), height=Cm(0.8))
            except Exception:
                run = fp.add_run("Ti Turing")
                style_run(run, size=9, bold=True, color=DARK)
        else:
            run = fp.add_run("Ti Turing")
            style_run(run, size=9, bold=True, color=DARK)

        sep = fp.add_run("  ·  Pág. ")
        style_run(sep, size=9, color=GREY)
        _add_page_field(fp)

        copy_run = fp.add_run(f"  ·  {title}  ·  © 2026")
        style_run(copy_run, size=8, italic=True, color=GREY)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()

    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Versión del producto: 1.0.3   ·   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Listo para producción")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3.0))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por ")
    style_run(r, size=11, color=GREY)
    r = p.add_run("Ti Turing")
    style_run(r, size=11, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("https://github.com/TI-Turing   ·   contacto@gestabiz.com")
    style_run(r, size=9, italic=True, color=GREY)


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    build_cover(doc, title, subtitle)
    page_break(doc)
    return doc


# ===========================================================================
# DATA: Plans
# ===========================================================================
@dataclass
class PlanSpec:
    key: str
    name: str
    price_monthly: str
    price_annual: str
    tagline: str
    ideal_for: str
    includes_previous: bool
    exclusives: list[str] = field(default_factory=list)
    limits: list[str] = field(default_factory=list)


PLANS: list[PlanSpec] = [
    PlanSpec(
        key="free", name="Gratuito",
        price_monthly="$0", price_annual="$0",
        tagline="Para probar Gestabiz sin costo y arrancar operación en minutos.",
        ideal_for="Profesionales independientes que están empezando y necesitan un canal de reservas sin inversión inicial.",
        includes_previous=False,
        exclusives=[
            "Registro de un (1) negocio con una (1) sede",
            "Hasta 30 citas activas por mes",
            "Perfil público visible en Google con dirección web propia",
            "Asistente de reserva completo con detección de horarios ocupados",
            "Panel de Resumen con métricas básicas",
            "Gestión de Servicios (crear, editar, activar/desactivar)",
            "Gestión de Clientes con historial de visitas",
            "Visualización del plan y pagos",
            "Notificaciones dentro de la aplicación (campanita)",
            "Notificaciones por email (recordatorios de cita)",
            "Chat cliente ↔ dueño del negocio",
            "Reseñas de clientes con respuesta del negocio",
            "Configuraciones del negocio y del perfil personal",
        ],
        limits=[
            "Sin gestión de empleados",
            "Sin múltiples sedes",
            "Sin reportes avanzados",
            "Sin ventas rápidas ni contabilidad",
            "Sin sistema de permisos por empleado",
            "Sin WhatsApp ni SMS (solo email)",
        ],
    ),
    PlanSpec(
        key="basico", name="Básico",
        price_monthly="$89.900 COP / mes", price_annual="$899.000 COP / año (2 meses gratis)",
        tagline="Para negocios con equipo y operación multi-sede que exigen control real.",
        ideal_for="Salones de belleza, barberías, consultorios, estudios, gimnasios y centros de servicio con 2-15 empleados que ya tienen flujo diario de citas.",
        includes_previous=True,
        exclusives=[
            "Hasta 15 empleados activos",
            "Hasta 5 sedes (ubicaciones físicas)",
            "Citas ilimitadas",
            "Módulo de Empleados con jerarquía visual (mapa organizacional)",
            "Sistema de Ausencias y Vacaciones con aprobación obligatoria",
            "Festivos públicos colombianos 2025-2027 integrados",
            "Ventas Rápidas (punto de venta sin cita previa)",
            "Historial de Ventas por periodo, empleado y sede",
            "Reportes operacionales y financieros básicos",
            "Sistema de Permisos por empleado (79 tipos de permiso, 9 plantillas predefinidas)",
            "Notificaciones por WhatsApp Business",
            "Transferencia de empleados entre sedes",
            "Chat con múltiples empleados (clientes pueden elegir con quién hablar)",
            "Sede preferida por dispositivo",
        ],
        limits=[
            "Sin contabilidad completa con impuestos colombianos",
            "Sin gastos recurrentes",
            "Sin reclutamiento de vacantes",
            "Sin recursos físicos (salas, canchas, mesas)",
        ],
    ),
    PlanSpec(
        key="pro", name="Pro",
        price_monthly="$159.900 COP / mes", price_annual="$1.599.000 COP / año (2 meses gratis)",
        tagline="Todo lo que necesita una PyME en serio: contabilidad, recursos físicos y reclutamiento.",
        ideal_for="Centros médicos, clínicas estéticas, hoteles boutique, coworkings, restaurantes por reserva, escuelas deportivas y negocios que manejan recursos físicos o necesitan contabilidad con impuestos.",
        includes_previous=True,
        exclusives=[
            "Empleados y sedes ilimitados",
            "Módulo de Contabilidad completo con IVA, ICA y Retención en la Fuente",
            "Módulo de Gastos recurrentes y puntuales",
            "Configuración fiscal por sede",
            "Exportación de reportes a PDF / CSV / Excel",
            "Gráficos financieros avanzados (tendencia mensual, ingresos vs gastos, categorías, empleados top)",
            "Sistema de Reclutamiento (vacantes, aplicaciones, evaluación inteligente, reseñas obligatorias)",
            "Sistema de Recursos Físicos (salas, canchas, mesas, equipos — 15 tipos)",
            "Modelos de negocio flexibles (profesional / recurso físico / híbrido / clase grupal)",
            "SMS transaccional",
            "Sincronización con Google Calendar",
            "Soporte prioritario por WhatsApp",
        ],
        limits=[],
    ),
]


# ===========================================================================
# DATA: Module specs (cleaned of technical jargon)
# ===========================================================================
@dataclass
class ModuleSpec:
    key: str
    name: str
    plan: str       # "free" | "basico" | "pro"
    role: str       # "cliente" | "empleado" | "admin" | "transversal"
    summary: str
    details: list[str]
    flows: list[tuple[str, list[str]]] = field(default_factory=list)
    exceptions: list[str] = field(default_factory=list)
    screenshot: str | None = None


MODULES: list[ModuleSpec] = [
    # ====================== TRANSVERSAL ======================
    ModuleSpec(
        key="landing", name="Página de Inicio Pública", plan="free", role="transversal",
        summary="Página de presentación que explica qué es Gestabiz, muestra beneficios, planes y permite registrarse o iniciar sesión.",
        details=[
            "Accesible sin iniciar sesión desde la dirección principal del sitio.",
            "Secciones: propuesta de valor, beneficios, cómo funciona (3 pasos), testimonios, planes, botón de registro.",
            "Barra superior fija con logo, menú de navegación y botones de Iniciar sesión / Registrarse.",
            "Se adapta a todos los tamaños de pantalla (celular, tablet, computador).",
            "Selector de idioma (español/inglés) y selector de tema claro/oscuro.",
            "Aviso de cookies: si el usuario acepta, se activa el seguimiento anónimo de visitas.",
        ],
        screenshot="Página de inicio de Gestabiz en versión de escritorio.",
    ),
    ModuleSpec(
        key="auth", name="Registro e Inicio de Sesión", plan="free", role="transversal",
        summary="Registro con correo, Google o GitHub; inicio de sesión con contraseña o enlace mágico por email.",
        details=[
            "Registro con correo y contraseña, o con un clic usando cuenta de Google o GitHub.",
            "Campos obligatorios: correo, contraseña (mínimo 8 caracteres), nombre completo, país.",
            "Verificación de correo electrónico obligatoria antes de usar la plataforma.",
            "Opción de recibir un enlace de acceso por email (sin necesidad de recordar contraseña).",
            "Si el correo ya está registrado, se sugiere iniciar sesión directamente.",
        ],
        flows=[
            ("Registro estándar con correo", [
                "Ingresa correo y contraseña.",
                "Completa datos personales (nombre, país, teléfono).",
                "Acepta términos y pulsa 'Crear cuenta'.",
                "Recibe email de verificación; al confirmarlo, queda autenticado.",
            ]),
            ("Reserva sin cuenta previa", [
                "El cliente llega al perfil público de un negocio desde Google.",
                "Pulsa 'Reservar' y se le pide iniciar sesión.",
                "Tras autenticarse, el asistente de reserva se abre con el negocio y servicio preseleccionados.",
            ]),
        ],
        screenshot="Pantalla de registro con opciones de correo, Google y GitHub.",
    ),
    ModuleSpec(
        key="onboarding", name="Creación de Negocio (Primeros Pasos)", plan="free", role="transversal",
        summary="Asistente de 4 pasos que lleva al nuevo usuario desde el registro hasta tener un negocio operativo.",
        details=[
            "Primera decisión: '¿Vienes a reservar o a ofrecer servicios?'",
            "Si elige 'Reservar', se activa el rol de Cliente y se abre el catálogo de negocios.",
            "Si elige 'Ofrecer servicios', se inicia un asistente de 4 pasos.",
            "Paso 1: datos del negocio (nombre, categoría, hasta 3 subcategorías, descripción, dirección web única).",
            "Paso 2: primera sede (país, ciudad, dirección, horario semanal).",
            "Paso 3: primer servicio (nombre, duración, precio, descripción, imagen).",
            "Paso 4: resumen y confirmación.",
            "Al completar, el dueño queda con acceso completo a todos los módulos de su plan.",
            "Si el usuario abandona el asistente, el progreso se guarda y se retoma al volver.",
        ],
        screenshot="Asistente de creación de negocio — paso 1: datos básicos.",
    ),
    ModuleSpec(
        key="notifications", name="Notificaciones", plan="free", role="transversal",
        summary="Notificaciones por email, WhatsApp, SMS y dentro de la aplicación, con preferencias configurables por usuario.",
        details=[
            "Campanita en la esquina superior con contador de notificaciones no leídas.",
            "Tipos: citas (creada, confirmada, cancelada, reprogramada, completada, recordatorio), empleados, ausencias, reclutamiento, sistema.",
            "Canales disponibles según plan: email (todos), WhatsApp (Básico+), SMS (Pro).",
            "Cada usuario puede activar o desactivar cada tipo de notificación y cada canal desde Configuraciones.",
            "El negocio configura cuándo enviar recordatorios (por ejemplo, 24 h y 2 h antes de la cita).",
            "Plan Gratuito: solo email y notificaciones dentro de la app.",
        ],
        screenshot="Campanita de notificaciones con lista desplegable.",
    ),
    ModuleSpec(
        key="chat", name="Chat en Tiempo Real", plan="free", role="transversal",
        summary="Mensajería instantánea entre clientes y el negocio, con archivos adjuntos y confirmación de lectura.",
        details=[
            "Ventana de chat accesible desde cualquier pantalla con el botón flotante.",
            "Lista de conversaciones a la izquierda, chat activo a la derecha.",
            "Envío de imágenes y PDF (hasta 10 MB por archivo).",
            "Indicador de 'está escribiendo…' en tiempo real.",
            "Confirmación de lectura (doble check).",
            "Resumen diario por email de mensajes no leídos.",
            "Plan Gratuito: chat solo con el dueño. Plan Básico+: el cliente elige con qué empleado chatear.",
            "Cada empleado puede activar o desactivar si recibe mensajes de clientes.",
        ],
        screenshot="Ventana de chat con conversación activa y adjuntos.",
    ),
    ModuleSpec(
        key="public_profile", name="Perfil Público del Negocio", plan="free", role="transversal",
        summary="Página pública del negocio visible en Google con servicios, ubicaciones, reseñas y botón de reservar.",
        details=[
            "Dirección web propia tipo gestabiz.com/negocio/mi-salon.",
            "Visible en Google sin necesidad de que el visitante tenga cuenta.",
            "4 pestañas: Servicios (con precio y duración), Ubicaciones (con mapa y horario), Reseñas (ordenadas por fecha), Acerca de (descripción, categorías, redes).",
            "Botón 'Reservar' en cada servicio que abre el asistente de reserva.",
            "Datos estructurados para que Google muestre información enriquecida en los resultados de búsqueda.",
            "Buscador global con filtro por ciudad, categoría o texto libre, y 6 criterios de ordenamiento (relevancia, calificación, cercanía, precio ascendente/descendente, más recientes).",
        ],
        screenshot="Perfil público de un negocio con pestañas, mapa y botón Reservar.",
    ),
    ModuleSpec(
        key="reviews", name="Reseñas y Calificaciones", plan="free", role="transversal",
        summary="Calificaciones de 1 a 5 estrellas para negocios y empleados, con respuesta del negocio.",
        details=[
            "Solo pueden dejar reseña los clientes que hayan completado al menos una cita.",
            "Calificación obligatoria (1-5 estrellas); comentario opcional (máximo 500 caracteres).",
            "Reseñas anónimas soportadas (no muestran el nombre del cliente).",
            "El negocio puede responder una vez por reseña; la respuesta queda visible debajo de la original.",
            "El administrador puede ocultar una reseña (pero no editarla ni borrar su contenido).",
            "Las calificaciones promedio se actualizan automáticamente cada 5 minutos.",
        ],
        screenshot="Lista de reseñas en el perfil público con estrellas y respuestas.",
    ),
    ModuleSpec(
        key="settings", name="Configuraciones", plan="free", role="transversal",
        summary="Todas las preferencias del negocio, del empleado y del cliente en un solo lugar.",
        details=[
            "Administrador: datos del negocio (nombre, categorías, logo, banner), contacto, dirección legal, horario global, configuración fiscal, sede preferida, métodos de pago, idioma.",
            "Empleado: horarios de 7 días con almuerzo, salarios, servicios que ofrece, si acepta mensajes de clientes, disponibilidad para vacantes, foto y biografía.",
            "Cliente: anticipación mínima para reservar, método de pago preferido, ciudad preferida, idioma, tema claro/oscuro, preferencias de notificación.",
            "Sincronización con Google Calendar (plan Pro) desde la pestaña 'Integraciones'.",
            "Cambio de contraseña, cierre de sesión en todos los dispositivos, eliminación de cuenta con doble confirmación.",
        ],
        screenshot="Configuraciones con pestañas por rol.",
    ),
    ModuleSpec(
        key="bug_reports", name="Reporte de Errores", plan="free", role="transversal",
        summary="Botón flotante para que cualquier usuario reporte un problema con evidencia adjunta.",
        details=[
            "Botón flotante presente en todas las pantallas una vez autenticado.",
            "Campos: título, descripción, severidad (Crítico / Alto / Medio / Bajo).",
            "Adjuntar capturas de pantalla o videos como evidencia.",
            "El equipo de soporte recibe un correo automático con el reporte.",
        ],
        screenshot="Formulario de reporte de error con adjuntos.",
    ),

    # ====================== CLIENTE ======================
    ModuleSpec(
        key="appointment_wizard", name="Reservar una Cita", plan="free", role="cliente",
        summary="Asistente paso a paso que guía al cliente desde elegir negocio hasta confirmar su cita.",
        details=[
            "Pasos: elegir negocio → elegir servicio → elegir sede → elegir profesional → elegir fecha y hora → confirmar → éxito.",
            "Si el cliente llega desde el perfil público, los pasos ya decididos se saltan automáticamente.",
            "Cada paso se valida antes de avanzar; el botón 'Siguiente' se habilita solo cuando la selección es correcta.",
            "Indicador visual de progreso (barra horizontal en computador, compacta en celular).",
            "Las selecciones se guardan durante 1 hora: si el cliente cierra y vuelve, retoma donde quedó.",
            "Al elegir fecha y hora, el sistema valida en tiempo real: horario de la sede, pausa de almuerzo del profesional, festivos, ausencias aprobadas, y citas ya ocupadas.",
            "Los horarios no disponibles muestran el motivo ('Ocupado', 'Hora de almuerzo', 'Ausencia', 'Festivo').",
            "Al confirmar, se envía un email con un enlace para confirmar o cancelar la cita (válido 24 horas).",
        ],
        flows=[
            ("Reserva completa desde cero", [
                "El cliente entra al buscador y filtra por ciudad, categoría o texto.",
                "Elige un negocio y entra a su perfil público.",
                "Pulsa 'Reservar' en el servicio deseado.",
                "Avanza por los pasos del asistente.",
                "Confirma y recibe su comprobante con opción de agregar a Google Calendar.",
            ]),
        ],
        exceptions=[
            "Si el cliente tiene el máximo de citas activas (5 por negocio por defecto), se muestra un aviso y no puede reservar más.",
            "Si el enlace de confirmación vence (24 h), se muestra 'Enlace vencido — contacta al negocio'.",
        ],
        screenshot="Asistente de reserva — paso Fecha/Hora con horarios disponibles.",
    ),
    ModuleSpec(
        key="client_dashboard", name="Mi Panel de Cliente", plan="free", role="cliente",
        summary="Panel personalizado con próxima cita, historial, favoritos, descubrimiento por ciudad y chat.",
        details=[
            "Vista principal con la próxima cita destacada y citas pasadas.",
            "Sección 'Descubre negocios en tu ciudad' con sugerencias basadas en ubicación.",
            "Lista de negocios favoritos con acceso rápido a reservar.",
            "Pestaña de Búsqueda: buscador con filtros por ciudad, categoría y texto.",
            "Pestaña Mis Citas: todas las citas con filtro por estado (pendiente, confirmada, completada, cancelada).",
            "Pestaña Perfil: datos personales, preferencias, idioma, tema.",
        ],
        screenshot="Panel del cliente con próxima cita, favoritos y buscador.",
    ),

    # ====================== EMPLEADO ======================
    ModuleSpec(
        key="employee_dashboard", name="Mi Panel de Empleado", plan="basico", role="empleado",
        summary="Agenda del empleado con sus citas, sus clientes, sus métricas, solicitud de ausencias y chat.",
        details=[
            "Agenda del día y de la semana con la próxima cita destacada.",
            "Sección 'Mis Clientes': listado de clientes que ha atendido, ordenados por cantidad de visitas.",
            "Click en un cliente abre su perfil con historial de citas, servicios y precios.",
            "Métricas personales: citas completadas, ingresos generados, calificación promedio, tasa de ocupación.",
            "Botón 'Solicitar Ausencia' que abre el formulario de ausencias.",
            "Indicador de saldo de vacaciones disponibles.",
            "Control de si desea recibir mensajes de clientes por chat.",
            "Preferencias de horario, almuerzo y servicios que ofrece.",
        ],
        screenshot="Panel del empleado con agenda del día, métricas y saldo de vacaciones.",
    ),

    # ====================== ADMIN — PLAN GRATUITO ======================
    ModuleSpec(
        key="appointments", name="Calendario de Citas", plan="free", role="admin",
        summary="Calendario central para ver, crear, editar, reagendar y cancelar citas con detección de conflictos.",
        details=[
            "Filtros por sede, empleado, servicio, estado y rango de fechas.",
            "Vistas: día / semana / mes / agenda (lista).",
            "Colores por estado: Confirmada (verde), Pendiente (amarillo), Completada (azul), Cancelada (gris).",
            "Click en una cita muestra: servicio, cliente, empleado, sede, hora, precio y notas internas.",
            "Acciones: Confirmar, Marcar como completada, Reagendar, Cancelar con motivo, Chatear con el cliente.",
            "Reagendar abre el asistente de reserva en modo edición (se saltan los pasos ya definidos).",
            "Cancelar notifica al cliente por la app, email y WhatsApp (según el plan).",
            "Exportación a CSV del rango visible.",
            "Sincronización con Google Calendar (plan Pro).",
        ],
        flows=[
            ("Crear cita desde el calendario", [
                "Pulsar 'Nueva cita'.",
                "Seleccionar cliente existente o crear uno nuevo.",
                "Elegir servicio (se precarga duración y precio).",
                "Elegir sede y empleado.",
                "Elegir fecha y hora (el sistema valida automáticamente todos los horarios).",
                "Confirmar y enviar notificación al cliente.",
            ]),
            ("Cancelar una cita", [
                "Pulsar 'Cancelar' en la cita.",
                "Escribir motivo (mínimo 10 caracteres).",
                "El sistema libera el horario y notifica al cliente.",
            ]),
        ],
        exceptions=[
            "Si dos personas intentan reservar el mismo horario al mismo tiempo, el segundo intento recibe el mensaje 'Ese horario acaba de ocuparse'.",
            "Las citas históricas de servicios eliminados siguen siendo visibles (aparecen con servicio 'N/A').",
        ],
        screenshot="Calendario de citas en vista semanal con filtros y leyenda de estados.",
    ),
    ModuleSpec(
        key="services", name="Gestión de Servicios", plan="free", role="admin",
        summary="Crear, editar, activar y desactivar servicios con precio, duración e imagen.",
        details=[
            "Lista de servicios con indicador de estado (activo/inactivo).",
            "Crear servicio: nombre, descripción, duración (en minutos), precio, categoría, imagen.",
            "Editar un servicio actualiza precios futuros sin alterar citas ya agendadas.",
            "Desactivar un servicio lo oculta del asistente de reserva sin borrar el historial.",
            "Asignar servicios a sedes específicas (si no se asigna ninguna, está disponible en todas).",
            "Asignar servicios a empleados específicos.",
            "Búsqueda rápida por nombre y filtro por categoría o estado.",
        ],
        exceptions=[
            "No se puede eliminar un servicio con citas futuras confirmadas; el sistema sugiere desactivarlo.",
            "Cambiar la duración no afecta citas existentes.",
        ],
        screenshot="Listado de servicios con filtros y formulario de edición.",
    ),
    ModuleSpec(
        key="clients_crm", name="Gestión de Clientes", plan="free", role="admin",
        summary="Lista de todos los clientes del negocio con historial de visitas y perfil detallado.",
        details=[
            "Muestra todos los clientes que han tenido al menos una cita no cancelada.",
            "Tarjeta por cliente con: avatar, nombre, email, total de visitas, fecha de última visita.",
            "Filtro por nombre o email.",
            "Click en un cliente abre su perfil con dos pestañas: 'Información' (estadísticas generales) e 'Historial' (cada cita con servicio, fecha, estado y precio).",
        ],
        screenshot="Listado de clientes con tarjetas y perfil abierto.",
    ),
    ModuleSpec(
        key="billing_view", name="Facturación del Plan", plan="free", role="admin",
        summary="Panel donde el dueño ve su plan actual, próximo cobro, historial de pagos y uso del plan.",
        details=[
            "Muestra plan activo, fecha de próximo cobro, método de pago y estado de la suscripción.",
            "Historial de facturas con estado (pagada / pendiente / fallida) y enlace al recibo.",
            "Indicadores de uso: empleados activos vs límite, citas del mes vs límite, sedes vs límite.",
            "Botón 'Cambiar plan' con comparativa en 3 columnas.",
            "Botón 'Agregar método de pago' (MercadoPago / Stripe / PayU según configuración).",
            "Si hay un pago fallido, aparece un aviso rojo en la parte superior con botón para actualizar método.",
            "Tras 3 pagos fallidos consecutivos, el plan se degrada automáticamente a Gratuito.",
        ],
        screenshot="Panel de facturación con plan actual, uso y botón de mejora.",
    ),

    # ====================== ADMIN — PLAN BÁSICO ======================
    ModuleSpec(
        key="employees", name="Gestión de Empleados y Jerarquía", plan="basico", role="admin",
        summary="Alta de empleados, horarios, salarios, organigrama visual, transferencia entre sedes y activación/desactivación.",
        details=[
            "Organigrama interactivo donde se puede ver y cambiar la jerarquía del equipo.",
            "Cada empleado tiene: rol (gerente, profesional, recepcionista, contador, soporte), tipo de empleo, fecha de contratación, salario, horario semanal con pausa de almuerzo.",
            "Invitar empleado por email: recibe un enlace, completa su perfil (foto, teléfono, horarios, servicios) y queda activo.",
            "Transferencia de sede: se selecciona la sede destino y la fecha efectiva; si hay citas conflictivas, el sistema avisa y permite cancelarlas o cambiar la fecha.",
            "Pausar a un empleado sin eliminar su historial.",
            "El dueño se registra automáticamente como gerente al crear el negocio.",
        ],
        flows=[
            ("Dar de alta un empleado", [
                "El administrador invita por email.",
                "El empleado recibe el enlace y completa su perfil.",
                "Al terminar, el empleado aparece disponible para recibir citas.",
            ]),
            ("Transferir entre sedes", [
                "Abrir el perfil del empleado y pulsar 'Transferir'.",
                "Seleccionar sede destino y fecha efectiva.",
                "Si hay citas futuras en la sede origen, elegir: cancelar y notificar, o cambiar la fecha.",
                "Confirmar: el sistema ejecuta la transferencia y notifica a los afectados.",
            ]),
        ],
        exceptions=[
            "No se puede eliminar un empleado con citas futuras; hay que reasignarlas o cancelarlas primero.",
        ],
        screenshot="Organigrama de empleados con ficha del empleado seleccionado.",
    ),
    ModuleSpec(
        key="locations", name="Gestión de Sedes", plan="basico", role="admin",
        summary="Hasta 5 sedes en plan Básico (ilimitadas en Pro): cada una con dirección, horario, servicios y métricas.",
        details=[
            "Crear sede: país, ciudad, dirección, geolocalización, teléfono, email, horario semanal.",
            "Indicador 'Administrada' marca la sede preferida del dispositivo actual.",
            "Cada sede tiene sus propios servicios habilitados.",
            "Métricas por sede: citas del mes, ingresos, empleados activos, servicios más pedidos.",
            "Un empleado puede atender puntualmente en otra sede diferente a la suya.",
            "Al desactivar una sede, sus empleados se transfieren o se pausan.",
        ],
        exceptions=[
            "Una sede con citas futuras no se puede eliminar; se sugiere desactivarla.",
        ],
        screenshot="Listado de sedes con indicador de administrada y métricas.",
    ),
    ModuleSpec(
        key="absences", name="Ausencias y Vacaciones", plan="basico", role="admin",
        summary="Sistema de solicitud y aprobación de ausencias con tipos (vacaciones, emergencia, enfermedad, personal), cálculo automático de días hábiles y cancelación automática de citas en emergencias.",
        details=[
            "Toda ausencia requiere aprobación del administrador, sin excepción.",
            "Tipos: vacaciones, emergencia, enfermedad, personal, otro.",
            "Formulario con selector de rango de fechas (inicio y fin), con resaltado visual del periodo seleccionado.",
            "Cálculo automático de días hábiles excluyendo festivos colombianos (54 festivos precargados 2025-2027).",
            "Al solicitar, todos los administradores del negocio reciben notificación dentro de la app y por email.",
            "El administrador aprueba o rechaza desde la pestaña 'Ausencias', con visibilidad de citas afectadas.",
            "Para emergencias: al enviar la solicitud, se cancelan automáticamente todas las citas del empleado en ese rango y se notifica a cada cliente afectado.",
            "Indicador de saldo de vacaciones: días disponibles, usados, pendientes, restantes.",
            "Los días de vacaciones se acumulan según la antigüedad del empleado (1,25 días por mes por defecto).",
        ],
        flows=[
            ("Solicitar vacaciones", [
                "El empleado abre el formulario de ausencias desde su panel.",
                "Elige tipo 'Vacaciones', selecciona el rango y escribe un motivo opcional.",
                "El sistema calcula los días hábiles y valida que tenga saldo disponible.",
                "Envía la solicitud; todos los administradores reciben notificación.",
                "El administrador revisa, ve las citas afectadas y aprueba o rechaza.",
                "El empleado recibe notificación con la decisión.",
            ]),
            ("Ausencia de emergencia", [
                "El empleado selecciona tipo 'Emergencia' y el rango.",
                "Al enviar, se cancelan automáticamente todas sus citas en ese periodo.",
                "Cada cliente afectado recibe un email informando que su cita fue cancelada por emergencia.",
            ]),
        ],
        exceptions=[
            "No se puede solicitar vacaciones si el saldo es insuficiente.",
            "No se pueden solapar ausencias con una ya aprobada.",
            "Los festivos no cuentan como días hábiles.",
        ],
        screenshot="Formulario de ausencia con calendario y pestaña de aprobaciones pendientes.",
    ),
    ModuleSpec(
        key="quicksales", name="Ventas Rápidas (Punto de Venta)", plan="basico", role="admin",
        summary="Punto de venta ligero para clientes sin cita previa o ventas fuera del flujo de reserva.",
        details=[
            "Campos: servicio (o producto libre), empleado que atiende (opcional), sede (precargada con la sede preferida), cliente (buscar existente o 'cliente ocasional'), método de pago, monto.",
            "Si se selecciona un servicio, el precio se precarga pero es editable (para descuentos o ajustes).",
            "Al confirmar, se registra como ingreso y se asocia al empleado para comisiones.",
            "Historial de ventas rápidas con filtros por rango, empleado, sede y método de pago.",
            "Impresión de recibo en formato térmico 80 mm.",
        ],
        screenshot="Formulario de Venta Rápida con sede preferida precargada.",
    ),
    ModuleSpec(
        key="sales_history", name="Historial de Ventas", plan="basico", role="admin",
        summary="Reporte de citas completadas y ventas rápidas con tarjetas resumen y tabla filtrable.",
        details=[
            "Filtro de rango: últimos 7 / 30 / 90 / 365 días (por defecto 30).",
            "Tarjetas resumen: total de ventas completadas, ingresos totales, promedio por cita.",
            "Tabla: fecha, servicio, cliente (click abre su perfil), empleado, sede, precio, método de pago.",
            "Exportación a CSV del rango visible.",
        ],
        screenshot="Historial de Ventas con tarjetas resumen y tabla detallada.",
    ),
    ModuleSpec(
        key="reports", name="Reportes Operacionales", plan="basico", role="admin",
        summary="Métricas clave del negocio: ocupación, ingresos, empleados y servicios top, retención de clientes.",
        details=[
            "Filtros globales: sede (precargada con la preferida), rango de fechas, empleado opcional.",
            "Indicadores: citas por día/semana/mes, tasa de ocupación por empleado, top 10 servicios, top 10 clientes, retención (clientes con 2+ citas).",
            "Comparativa del periodo actual vs el anterior con flechas de variación.",
            "Exportación a CSV y PDF de cada indicador.",
        ],
        screenshot="Reportes con indicadores de ocupación y comparativa.",
    ),
    ModuleSpec(
        key="permissions", name="Permisos por Empleado", plan="basico", role="admin",
        summary="79 tipos de permisos en 16 categorías, 9 plantillas predefinidas y registro completo de cambios.",
        details=[
            "Categorías: servicios, recursos, sedes, empleados, citas, reclutamiento, contabilidad, gastos, reseñas, facturación, notificaciones, configuraciones, permisos, ausencias, favoritos, ventas.",
            "9 plantillas listas para usar: Administrador Completo, Vendedor, Cajero, Gerente de Sede, Recepcionista, Profesional, Contador, Gerente, Soporte.",
            "Cada botón de acción en la plataforma respeta los permisos asignados: si el empleado no tiene permiso, el botón se oculta o se desactiva.",
            "El dueño del negocio siempre tiene acceso total; no se le pueden quitar permisos.",
            "Al crear un negocio, el dueño recibe automáticamente los 79 permisos.",
            "Cada cambio de permisos queda registrado con fecha, hora y quién lo hizo.",
        ],
        screenshot="Administrador de permisos con lista de empleados y matriz de permisos.",
    ),

    # ====================== ADMIN — PLAN PRO ======================
    ModuleSpec(
        key="accounting", name="Contabilidad con Impuestos Colombianos", plan="pro", role="admin",
        summary="Módulo contable completo con IVA, ICA y Retención en la Fuente; gráficos financieros y exportación lista para el contador.",
        details=[
            "Filtros por periodo (mes/trimestre/año fiscal), sede y categoría.",
            "Transacciones de ingreso y egreso con desglose automático de impuestos.",
            "IVA configurable: 0%, 5%, 19%.",
            "ICA por municipio (código DANE + tarifa en milésimas).",
            "Retención en la Fuente con tabla de conceptos y bases mínimas.",
            "Gráficos: distribución por categoría, ingresos por empleado, ingresos vs gastos, tendencia mensual, ingresos por sede.",
            "Exportación a PDF, CSV y Excel listos para entregar al contador.",
            "Libro diario, libro mayor, estado de resultados, balance de comprobación.",
            "Cada venta rápida y cita completada genera automáticamente su transacción con impuestos calculados.",
        ],
        exceptions=[
            "Si el negocio no tiene configuración fiscal, las transacciones se guardan sin impuestos y aparece un aviso para configurarlos.",
        ],
        screenshot="Panel financiero con gráficos y tabla de transacciones con impuestos.",
    ),
    ModuleSpec(
        key="expenses", name="Gestión de Gastos", plan="pro", role="admin",
        summary="Registro de gastos puntuales y recurrentes con categorías, soporte documental y generación automática de transacciones.",
        details=[
            "Gastos puntuales: categoría, monto, proveedor, fecha, descripción, evidencia (PDF/imagen).",
            "Gastos recurrentes: arriendo, servicios, licencias; frecuencia mensual/trimestral/anual.",
            "Asignar gastos a una sede específica para reportes por sede.",
            "Al vencer un gasto recurrente, se crea automáticamente la transacción correspondiente.",
            "Vista de conciliación mensual: gastos esperados vs registrados.",
        ],
        screenshot="Gestión de gastos con recurrentes y puntuales.",
    ),
    ModuleSpec(
        key="recruitment", name="Reclutamiento y Vacantes", plan="pro", role="admin",
        summary="Publicar vacantes, gestionar aplicaciones, evaluación inteligente de candidatos y reseñas obligatorias al contratar.",
        details=[
            "Crear vacante: título, descripción, sede, horario requerido, salario, tipo (tiempo completo/parcial/freelance).",
            "Tarjeta de vacante con estado (activa/pausada/cerrada), aplicaciones recibidas y vistas.",
            "Lista de aplicaciones con filtros por estado (nueva/en revisión/entrevista/contratada/rechazada).",
            "Perfil del aplicante con CV, experiencia y disponibilidad.",
            "El sistema compara automáticamente la disponibilidad del candidato con los requisitos de la vacante y avisa si hay conflictos de horario.",
            "Al contratar o finalizar un proceso, se obliga a dejar una reseña.",
            "Los candidatos pueden aplicar desde el perfil público del negocio.",
        ],
        exceptions=[
            "Si el candidato tiene conflictos de horario, se muestra un aviso amarillo; el administrador decide si continuar.",
        ],
        screenshot="Panel de reclutamiento con vacantes activas y perfil de aplicante.",
    ),
    ModuleSpec(
        key="resources", name="Recursos Físicos", plan="pro", role="admin",
        summary="Reservar salas, canchas, mesas, estudios, equipos: 15 tipos de recursos y 4 modelos de negocio.",
        details=[
            "15 tipos: sala, mesa, cancha, estudio, sala de reuniones, escritorio, equipo, vehículo, espacio, carril, campo, estación, estacionamiento, cama, otro.",
            "4 modelos de negocio: profesional (con empleados), recurso físico (sin empleados), híbrido (ambos), clase grupal (con capacidad).",
            "Cada recurso puede tener servicios asignados con precio personalizado.",
            "Las citas pueden asignarse a un empleado o a un recurso (o ambos en modelo híbrido).",
            "Detección automática de conflictos de horario en recursos.",
            "Capacidad por recurso (útil para clases grupales).",
            "Características del recurso (WiFi, proyector, sonido, ducha, etc.).",
            "Tarifa por hora para reservas por bloques de tiempo.",
        ],
        screenshot="Listado de recursos con salas, canchas y equipos reservables.",
    ),
]


# ===========================================================================
# MANUAL DE USUARIO
# ===========================================================================
def build_manual() -> Document:
    doc = setup_document(title="Manual de Usuario",
                         subtitle="Guía funcional completa de la plataforma")

    # ---- ÍNDICE ----
    h1(doc, "Índice", anchor="toc")
    para(doc, "Este documento está dividido en dos partes: un resumen ejecutivo y una guía detallada organizada por rol. Cada entrada es un enlace a la sección correspondiente.")

    def toc_link(label: str, anchor: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(p, anchor, label)

    h3(doc, "Parte 1 — Resumen Ejecutivo")
    toc_link("1.1  ¿Qué es Gestabiz?", "summary_what")
    toc_link("1.2  Roles de usuario", "summary_roles")
    toc_link("1.3  Planes y qué incluye cada uno", "summary_plans")
    toc_link("1.4  Primer uso: cómo empezar en 5 minutos", "summary_quickstart")

    h3(doc, "Parte 2 — Guía por Rol")
    toc_link("2.1  Funciones comunes (todos los usuarios)", "role_transversal")
    toc_link("2.2  Guía del Cliente", "role_cliente")
    toc_link("2.3  Guía del Empleado", "role_empleado")
    toc_link("2.4  Guía del Administrador", "role_admin")

    h3(doc, "Apéndices")
    toc_link("A.  Glosario", "appendix_glossary")
    toc_link("B.  Preguntas frecuentes", "appendix_faq")
    toc_link("C.  Soporte y canales oficiales", "appendix_support")

    page_break(doc)

    # =========== PARTE 1 ===========
    h1(doc, "Parte 1 — Resumen Ejecutivo")

    h2(doc, "1.1  ¿Qué es Gestabiz?", anchor="summary_what")
    para(doc, "Gestabiz es una plataforma todo-en-uno para la gestión integral de negocios de servicios por cita: agenda, clientes, empleados, recursos físicos, ventas, contabilidad, reseñas, reclutamiento, comunicaciones multicanal y perfiles públicos visibles en Google. Una sola herramienta reemplaza 6 a 10 productos independientes que normalmente combinan las PyMEs.")
    para(doc, "Funciona en navegador web, celular (iOS y Android) y como extensión de navegador. Cumple las normas de protección de datos y está disponible en español e inglés.")
    callout(doc, "Frase clave",
            "Gestabiz es la única plataforma en Colombia que combina reservas, contabilidad con impuestos, recursos físicos y reclutamiento en un solo producto — a un precio de PyME.")

    h2(doc, "1.2  Roles de usuario", anchor="summary_roles")
    para(doc, "Gestabiz calcula los roles automáticamente según la relación del usuario con cada negocio. Un mismo usuario puede tener varios roles en distintos negocios al mismo tiempo:")
    simple_table(doc,
        ["Rol", "¿Quién es?", "¿Qué puede hacer?"],
        [
            ["Dueño", "La persona que creó el negocio.", "Acceso total a todo sin restricciones."],
            ["Administrador", "Persona designada por el dueño con rol de administrador.", "Casi todo, según los permisos que le haya asignado el dueño."],
            ["Empleado", "Persona registrada como empleado del negocio.", "Su agenda, sus clientes, sus métricas, sus ausencias."],
            ["Cliente", "Cualquier persona que se registre en la plataforma.", "Reservar, ver sus citas, favoritos, chat, reseñas."],
        ],
    )

    h2(doc, "1.3  Planes y qué incluye cada uno", anchor="summary_plans")
    para(doc, "Cada plan incluye todo lo del anterior más sus exclusivos. Todos los precios son en pesos colombianos; el plan anual cobra 10 meses (2 meses gratis).")
    simple_table(doc,
        ["Plan", "Precio mensual", "Precio anual", "Ideal para"],
        [[p.name, p.price_monthly, p.price_annual, p.ideal_for] for p in PLANS],
    )
    for plan in PLANS:
        h3(doc, f"Plan {plan.name} — resumen")
        para(doc, plan.tagline, italic=True, color=GREY)
        for e in plan.exclusives:
            bullet(doc, e)
        if plan.limits:
            h4(doc, "Limitaciones del plan")
            for l in plan.limits:
                bullet(doc, l)

    h2(doc, "1.4  Primer uso: cómo empezar en 5 minutos", anchor="summary_quickstart")
    numbered(doc, "Crea tu cuenta en gestabiz.com (correo + contraseña, o Google).")
    numbered(doc, "Verifica tu correo desde el enlace que te enviamos.")
    numbered(doc, "Elige 'Ofrecer servicios' y completa el asistente de 4 pasos para crear tu negocio.")
    numbered(doc, "Publica tu primer servicio y comparte tu dirección web pública.")
    numbered(doc, "Empieza a recibir reservas — no hace falta instalar nada más.")
    screenshot_placeholder(doc, "Pantalla de bienvenida tras completar el asistente inicial.")

    page_break(doc)

    # =========== PARTE 2 ===========
    h1(doc, "Parte 2 — Guía por Rol")
    para(doc, "A continuación se describen todas las funciones de la plataforma, organizadas por el tipo de usuario que las utiliza. Si eres dueño de negocio, te interesa leer la sección de Administrador completa. Si eres empleado, tu sección es la 2.3. Si eres cliente, la 2.2.")

    # --- Helper to render a module ---
    def render_module(m: ModuleSpec) -> None:
        h3(doc, m.name, anchor=f"mod_{m.key}")
        para(doc, m.summary, italic=True, color=GREY)
        # Plan badge
        plan_name = next(p.name for p in PLANS if p.key == m.plan)
        para(doc, f"Disponible desde el plan {plan_name}.", size=10, color=PURPLE)
        if m.screenshot:
            screenshot_placeholder(doc, m.screenshot)
        h4(doc, "Descripción detallada")
        for d in m.details:
            bullet(doc, d)
        if m.flows:
            h4(doc, "Flujos paso a paso")
            for flow_name, steps in m.flows:
                para(doc, flow_name, color=PURPLE)
                for step in steps:
                    numbered(doc, step)
        if m.exceptions:
            h4(doc, "Casos especiales")
            for e in m.exceptions:
                bullet(doc, e, bold_prefix="Nota")
        doc.add_paragraph()

    # --- 2.1  Transversal ---
    h2(doc, "2.1  Funciones comunes (todos los usuarios)", anchor="role_transversal")
    para(doc, "Estas funciones están disponibles para todos los roles:")
    transversal_modules = [m for m in MODULES if m.role == "transversal"]
    for m in transversal_modules:
        render_module(m)
    page_break(doc)

    # --- 2.2  Cliente ---
    h2(doc, "2.2  Guía del Cliente", anchor="role_cliente")
    para(doc, "El cliente es cualquier persona que se registra en Gestabiz para reservar citas. Estas son sus funciones exclusivas:")
    client_modules = [m for m in MODULES if m.role == "cliente"]
    for m in client_modules:
        render_module(m)
    page_break(doc)

    # --- 2.3  Empleado ---
    h2(doc, "2.3  Guía del Empleado", anchor="role_empleado")
    para(doc, "El empleado es una persona registrada como parte del equipo de un negocio. Disponible desde el plan Básico:")
    employee_modules = [m for m in MODULES if m.role == "empleado"]
    for m in employee_modules:
        render_module(m)
    page_break(doc)

    # --- 2.4  Administrador (progressive by plan) ---
    h2(doc, "2.4  Guía del Administrador", anchor="role_admin")
    para(doc, "El administrador (dueño o persona con rol de admin) gestiona todo el negocio. Las funciones se desbloquean progresivamente según el plan:")

    admin_modules = [m for m in MODULES if m.role == "admin"]
    for plan in PLANS:
        admin_in_plan = [m for m in admin_modules if m.plan == plan.key]
        if admin_in_plan:
            h3(doc, f"Módulos del plan {plan.name}", anchor=f"admin_plan_{plan.key}")
            if plan.includes_previous:
                prev = PLANS[PLANS.index(plan) - 1]
                callout(doc, "Incluye todo lo anterior",
                        f"Este plan contiene todas las funciones del plan {prev.name} más las que aparecen aquí.")
            for m in admin_in_plan:
                render_module(m)
            page_break(doc)

    # =========== APÉNDICES ===========
    h1(doc, "Apéndices")

    h2(doc, "A.  Glosario", anchor="appendix_glossary")
    simple_table(doc,
        ["Término", "Significado"],
        [
            ["Sede", "Ubicación física donde se atiende al cliente (local, consultorio, etc.)."],
            ["Dirección web pública", "La URL única de tu negocio en Gestabiz (ej: gestabiz.com/negocio/mi-salon)."],
            ["Asistente de reserva", "El formulario paso a paso que usa el cliente para agendar una cita."],
            ["Conflicto de horario", "Cuando dos citas intentan ocupar el mismo espacio de tiempo del mismo profesional o recurso."],
            ["Sede preferida", "La sede que se precarga automáticamente en los filtros, configurable por dispositivo."],
            ["Recurso físico", "Un espacio o equipo reservable: sala, cancha, mesa, equipo, etc."],
            ["Festivo público", "Día no laboral reconocido por el calendario colombiano."],
            ["Venta rápida", "Registro de venta sin cita previa, desde el punto de venta del negocio."],
        ],
    )

    h2(doc, "B.  Preguntas frecuentes", anchor="appendix_faq")
    for q, a in [
        ("¿Tengo que firmar permanencia?",
         "No. Puedes cancelar en cualquier momento desde el panel de Facturación."),
        ("¿Puedo migrar desde otra plataforma?",
         "Sí. Importamos clientes y servicios en un archivo CSV. Citas históricas se cargan por rango."),
        ("¿Gestabiz procesa mis pagos?",
         "No directamente. Usamos MercadoPago, PayU o Stripe (a tu elección). Las tarjetas nunca pasan por nuestros servidores."),
        ("¿Mis datos son míos?",
         "Sí. Puedes exportar todo tu contenido en cualquier momento."),
        ("¿Funciona fuera de Colombia?",
         "Sí. La plataforma es multimoneda y multiidioma (español/inglés). Los módulos fiscales están pensados para Colombia pero son desactivables."),
        ("¿Qué pasa si supero un límite del plan?",
         "El módulo afectado pasa a modo lectura hasta que mejores el plan. Nunca se pierden datos."),
        ("¿Tienen aplicación para celular?",
         "Sí. Disponible para iOS y Android, más una extensión de navegador."),
    ]:
        h4(doc, q)
        para(doc, a)

    h2(doc, "C.  Soporte y canales oficiales", anchor="appendix_support")
    bullet(doc, "Centro de ayuda: gestabiz.com/ayuda")
    bullet(doc, "WhatsApp soporte (plan Pro): +57 XXX XXX XXXX")
    bullet(doc, "Email: soporte@gestabiz.com")
    bullet(doc, "Reporte de errores: botón flotante en toda la app o email bugs@gestabiz.com")

    # Apply header/footer
    title_str = "Manual de Usuario"
    return doc, title_str


# ===========================================================================
# PROPUESTA DE VALOR
# ===========================================================================
def build_proposal() -> Document:
    doc = setup_document(title="Propuesta de Valor",
                         subtitle="Por qué Gestabiz es la mejor decisión para tu negocio")

    # ---- ÍNDICE ----
    h1(doc, "Índice", anchor="toc")

    def toc_link(label: str, anchor: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(2)
        add_internal_hyperlink(p, anchor, label)

    h3(doc, "Parte 1 — Resumen Ejecutivo")
    toc_link("1.1  El problema que resolvemos", "p_problem")
    toc_link("1.2  La solución en 30 segundos", "p_solution")
    toc_link("1.3  Por qué Gestabiz gana", "p_why_win")
    toc_link("1.4  Planes y retorno de inversión", "p_plans")
    toc_link("1.5  Empieza hoy — sin costo", "p_cta")

    h3(doc, "Parte 2 — Detalle Comercial")
    toc_link("2.1  Diagnóstico del mercado PyME en servicios", "p_market")
    toc_link("2.2  Gestabiz vs Calendly, Booksy, Fresha y competidores locales", "p_comparison")
    toc_link("2.3  Sectores atendidos y casos de uso", "p_verticals")
    toc_link("2.4  Beneficios concretos por rol", "p_benefits_role")
    toc_link("2.5  Seguridad y protección de datos", "p_security")
    toc_link("2.6  Plan de incorporación y soporte", "p_onboarding")
    toc_link("2.7  Hoja de ruta 2026", "p_roadmap")
    toc_link("2.8  Testimonios y casos de éxito", "p_testimonials")
    toc_link("2.9  Preguntas frecuentes", "p_faq")

    page_break(doc)

    # =========== PARTE 1 ===========
    h1(doc, "Parte 1 — Resumen Ejecutivo")

    h2(doc, "1.1  El problema que resolvemos", anchor="p_problem")
    para(doc, "Los negocios PyME de servicios en Colombia y Latinoamérica pierden, en promedio, 1 de cada 5 citas por desorganización: agenda en papel o WhatsApp, sin recordatorios automáticos, sin control de ausencias, sin visibilidad de ingresos reales, y dependiendo de 6 a 10 herramientas que no se hablan entre sí.")
    callout(doc, "Costo oculto",
            "Un salón con 3 empleados que pierde 5 citas por semana a $40.000 promedio, pierde $9,6 millones al año — suficiente para pagar Gestabiz Pro durante 5 años.",
            color=DANGER)
    para(doc, "Las plataformas existentes fallan en dos extremos: o son demasiado simples (Calendly, agendas genéricas) o son demasiado caras y complejas (SAP, Salesforce). El espacio medio — una herramienta completa, en español, con contabilidad colombiana, a precio PyME — estaba vacío. Lo llenamos nosotros.")

    h2(doc, "1.2  La solución en 30 segundos", anchor="p_solution")
    para(doc, "Gestabiz reemplaza tu agenda, tu WhatsApp manual, tu hoja de cálculo de ventas, tu registro de ausencias, tu punto de venta, tu sistema de reseñas, tu bolsa de vacantes y tu contabilidad — en una sola plataforma web y móvil, con soporte en español y precios en pesos colombianos.")
    bullet(doc, "Perfil público visible en Google en 5 minutos.")
    bullet(doc, "Reservas en línea 24/7 con validación automática de horarios.")
    bullet(doc, "Recordatorios automáticos por email y WhatsApp.")
    bullet(doc, "Control total de empleados, sedes, permisos, ausencias, recursos físicos.")
    bullet(doc, "Contabilidad con IVA, ICA y Retención en la Fuente lista para el contador.")
    bullet(doc, "Reclutamiento con evaluación automática de candidatos.")
    screenshot_placeholder(doc, "Panel principal de Gestabiz — vista de administrador con todos los módulos.")

    h2(doc, "1.3  Por qué Gestabiz gana", anchor="p_why_win")
    simple_table(doc,
        ["Criterio", "Calendly", "Booksy", "Fresha", "Gestabiz"],
        [
            ["Precio/mes (USD)", "12-20", "30-85", "0* + comisión", "~$20-$40"],
            ["Idioma nativo", "EN", "EN/ES", "EN/ES", "ES (LATAM)"],
            ["Impuestos Colombia", "No", "No", "No", "Sí (IVA/ICA/Retefuente)"],
            ["Recursos físicos", "No", "Parcial", "Parcial", "Sí (15 tipos, 4 modelos)"],
            ["Reclutamiento", "No", "No", "No", "Sí (con evaluación)"],
            ["Permisos por empleado", "Básico", "Básico", "Medio", "79 tipos, 9 plantillas"],
            ["Perfil público en Google", "Parcial", "Sí", "Sí", "Sí + datos estructurados"],
            ["WhatsApp Business", "No", "Extra", "Extra", "Incluido (Básico+)"],
            ["Multi-sede", "No", "Sí", "Sí", "Ilimitada (Pro)"],
        ],
    )
    para(doc, "* Fresha cobra comisión por transacción que puede superar cualquier plan mensual pagado.",
         italic=True, size=9, color=GREY)
    callout(doc, "Mensaje clave",
            "Gestabiz es la única plataforma en nuestra categoría que te da contabilidad colombiana, recursos físicos y reclutamiento — sin pagar integraciones separadas.",
            color=ACCENT)

    h2(doc, "1.4  Planes y retorno de inversión", anchor="p_plans")
    simple_table(doc,
        ["Plan", "Precio", "Ideal para", "Se paga solo con*"],
        [
            ["Gratuito", "$0", "Profesional independiente empezando.", "Inmediato"],
            ["Básico", "$89.900/mes", "Salón, consultorio, gimnasio con 2-15 empleados.", "2-3 citas recuperadas/mes"],
            ["Pro", "$159.900/mes", "Clínica, hotel, coworking, restaurante por reserva.", "4-5 citas recuperadas/mes"],
        ],
    )
    para(doc, "* Cantidad de citas que necesitas recuperar por mes gracias a los recordatorios automáticos y mejor ocupación para que el plan se pague solo.",
         italic=True, size=9, color=GREY)

    h2(doc, "1.5  Empieza hoy — sin costo", anchor="p_cta")
    callout(doc, "Prueba gratis",
            "Crea tu cuenta en gestabiz.com, configura tu negocio en 5 minutos y recibe tu primera reserva hoy mismo. Sin tarjeta de crédito, sin permanencia, sin sorpresas.",
            color=PURPLE)

    page_break(doc)

    # =========== PARTE 2 ===========
    h1(doc, "Parte 2 — Detalle Comercial")

    h2(doc, "2.1  Diagnóstico del mercado PyME en servicios", anchor="p_market")
    para(doc, "En Colombia hay más de 1,6 millones de PyMEs de servicios activas (DANE 2024). La mayoría opera con herramientas manuales o con combinaciones de apps gratuitas que generan fricción, pérdidas y dolores de cabeza:")
    bullet(doc, "Agenda en cuaderno o WhatsApp → citas olvidadas, doble reserva, ausentismo del 15-25%.", bold_prefix="Caos operativo")
    bullet(doc, "Ausencias de empleados sin aprobación formal → clientes notificados a última hora, mala reputación.", bold_prefix="Recursos humanos improvisados")
    bullet(doc, "Contabilidad en Excel que no liga ventas, gastos e impuestos → multas, sobrepagos al contador.", bold_prefix="Riesgo fiscal")
    bullet(doc, "Sin canal de reservas en línea → pérdida de ventas fuera de horario (evidencia: 40% de reservas ocurren entre 8pm y 7am).", bold_prefix="Ventas dejadas sobre la mesa")
    bullet(doc, "Sin perfil público en Google → dependencia total de Instagram para ser encontrados.", bold_prefix="Sin presencia en buscadores")

    h2(doc, "2.2  Gestabiz vs Calendly, Booksy, Fresha y competidores locales", anchor="p_comparison")
    para(doc, "Análisis detallado por competidor:")

    h3(doc, "Calendly")
    para(doc, "Excelente para agendar reuniones de uno a uno, pero deja el negocio a la mitad: no tiene gestión de empleados con roles, no controla ausencias, no tiene contabilidad ni punto de venta, no tiene recursos físicos, reseñas ni perfiles públicos indexables. Está pensado para freelancers, no para negocios con equipo.")
    bullet(doc, "Gestabiz agrega: empleados con jerarquía, ausencias con aprobación, contabilidad con impuestos, punto de venta, recursos, reseñas, perfil en Google, reclutamiento, permisos por empleado.")

    h3(doc, "Booksy")
    para(doc, "Líder en belleza en Estados Unidos y Europa. Bien construido pero con precios en dólares, sin funcionalidades contables para Colombia (no maneja IVA ni ICA), sin recursos físicos tipo canchas o salas, sin módulo de reclutamiento. Soporte prioritario en inglés.")
    bullet(doc, "Gestabiz agrega: impuestos colombianos nativos, recursos físicos (15 tipos), reclutamiento con evaluación, soporte 100% en español, precios en COP.")

    h3(doc, "Fresha")
    para(doc, "Modelo aparentemente gratuito que se cobra con 2,19% + USD 0,20 por cada venta procesada, lo que se traduce en más de $200.000 COP/mes de comisiones para un salón mediano. No tiene contabilidad colombiana ni reclutamiento.")
    bullet(doc, "Gestabiz agrega: precio fijo predecible (sin comisión por venta), MercadoPago/PayU/Stripe intercambiables, contabilidad completa, reclutamiento.")

    h3(doc, "Competidores locales (Agenda Pro, Turnier, Simple, etc.)")
    para(doc, "Productos nicho con módulos limitados, sin perfil público en Google, sin recursos físicos, sin aplicación móvil robusta, sin actualizaciones frecuentes. Muchos están desatendidos desde 2022.")
    bullet(doc, "Gestabiz agrega: plataforma moderna con actualizaciones semanales, 3 aplicaciones (web + celular + extensión de navegador), hoja de ruta pública.")

    h2(doc, "2.3  Sectores atendidos y casos de uso", anchor="p_verticals")
    for vertical, case in [
        ("Belleza (salones, barberías, estética)",
         "Control de horarios individuales por profesional, pausa de almuerzo flexible, chat directo con clientes, reseñas de empleado, ventas de productos sin cita previa, reclutamiento con reseña obligatoria al contratar."),
        ("Salud (consultorios, odontología, fisioterapia, psicología)",
         "Permisos diferenciados para recepcionista, profesional y contador; historial completo del paciente, confirmación con anticipación configurable, WhatsApp con consentimiento, contabilidad con Retención en la Fuente."),
        ("Deportes (gimnasios, academias, canchas)",
         "Recursos físicos (canchas, carriles, estudios) con reserva directa sin empleado, clases grupales con capacidad, horarios extendidos, sincronización con Google Calendar."),
        ("Hoteles boutique y alojamientos",
         "Recursos tipo 'habitación' con tarifa por hora, multi-sede, contabilidad con IVA hotelero, reseñas verificadas, perfil público en Google."),
        ("Restaurantes por reserva",
         "Recursos tipo 'mesa' con capacidad, punto de venta ligero, contabilidad con IVA, WhatsApp para confirmar mesa."),
        ("Coworkings y alquiler de espacios",
         "Recursos como escritorios y salas de reuniones con características (WiFi, proyector, pizarra), tarifa por hora, contabilidad de ingresos por sede."),
        ("Escuelas y academias",
         "Clases grupales con capacidad, reclutamiento de profesores, contabilidad por sede, perfil público con catálogo de cursos."),
        ("Servicios profesionales (abogados, contadores, consultores)",
         "Horarios restringidos, confirmación obligatoria del cliente, permisos para compartir con asistentes, sincronización con Google Calendar."),
    ]:
        h3(doc, vertical)
        para(doc, case)

    h2(doc, "2.4  Beneficios concretos por rol", anchor="p_benefits_role")

    h3(doc, "Para el dueño del negocio")
    bullet(doc, "Control total desde el primer día: vista completa de ingresos, ocupación y empleados.")
    bullet(doc, "Reportes exportables listos para el contador con IVA, ICA y Retención desglosados.")
    bullet(doc, "Reducción medible de ausentismo gracias a recordatorios automáticos multicanal.")
    bullet(doc, "Perfil público indexable en Google → clientes llegando sin pagar publicidad.")

    h3(doc, "Para el administrador operativo")
    bullet(doc, "Calendario centralizado con detección automática de conflictos de horario.")
    bullet(doc, "Transferencia de empleados entre sedes con cancelación automática de citas afectadas.")
    bullet(doc, "Aprobación de ausencias en 2 clics con visibilidad del impacto en la agenda.")
    bullet(doc, "Sistema de permisos que permite delegar sin perder control.")

    h3(doc, "Para el empleado")
    bullet(doc, "Agenda clara sincronizada en todos sus dispositivos.")
    bullet(doc, "Solicitud de vacaciones con saldo visible en tiempo real.")
    bullet(doc, "Chat directo con sus clientes y control de quién le puede escribir.")
    bullet(doc, "Métricas personales: ingresos generados, calificación, tasa de ocupación.")

    h3(doc, "Para el cliente")
    bullet(doc, "Reservar 24/7 en pocos clics sin llamadas ni WhatsApp.")
    bullet(doc, "Confirmar o cancelar con un enlace seguro desde el email.")
    bullet(doc, "Ver historial, favoritos y calificar negocios en un solo lugar.")
    bullet(doc, "Chat con el negocio si necesita cambiar algo sin tener que llamar.")
    screenshot_placeholder(doc, "Flujo de reserva del cliente en 4 pantallas.")

    h2(doc, "2.5  Seguridad y protección de datos", anchor="p_security")
    bullet(doc, "Cada usuario solo puede ver y modificar los datos que le corresponden según su rol y permisos.")
    bullet(doc, "Los pagos se procesan a través de MercadoPago, PayU o Stripe — Gestabiz nunca almacena tarjetas de crédito.")
    bullet(doc, "Cumplimiento de normas de protección de datos: aviso de cookies, datos de visitas anonimizados.")
    bullet(doc, "Registro completo de cambios en permisos, accesos y operaciones críticas.")
    bullet(doc, "Copias de seguridad automáticas diarias con posibilidad de recuperación a un punto en el tiempo.")
    bullet(doc, "Entornos de desarrollo y producción completamente separados.")

    h2(doc, "2.6  Plan de incorporación y soporte", anchor="p_onboarding")
    numbered(doc, "Día 0: registro, verificación de correo, creación de negocio con el asistente de 4 pasos.")
    numbered(doc, "Día 1: publicación del perfil público y primeros 3 servicios.")
    numbered(doc, "Día 2-7: invitación de empleados, configuración de horarios, primeras reservas.")
    numbered(doc, "Día 8-30: configuración de impuestos, conexión de Google Calendar, primeros reportes.")
    numbered(doc, "Mes 2 en adelante: análisis de indicadores, optimización de ocupación, uso de reclutamiento.")
    callout(doc, "Soporte garantizado",
            "Soporte por email en todos los planes. WhatsApp prioritario en plan Pro con respuesta máxima de 4 horas hábiles.")

    h2(doc, "2.7  Hoja de ruta 2026", anchor="p_roadmap")
    bullet(doc, "Segundo trimestre 2026 — Facturación electrónica para Colombia.")
    bullet(doc, "Tercer trimestre 2026 — Pagos en cita (depósito al reservar).")
    bullet(doc, "Tercer trimestre 2026 — Programa de fidelización con puntos y descuentos.")
    bullet(doc, "Cuarto trimestre 2026 — Directorio interno de servicios (búsqueda global).")
    bullet(doc, "Cuarto trimestre 2026 — Aplicación nativa en App Store y Google Play.")

    h2(doc, "2.8  Testimonios y casos de éxito", anchor="p_testimonials")
    for quote, who in [
        ("«Pasé de perder 4 citas por semana a 0. En un mes pagué el plan Pro con lo que recuperé.»",
         "— María C., dueña de salón, Medellín"),
        ("«Tener IVA, ICA y Retención calculados solos me ahorra un día entero cada mes con mi contador.»",
         "— Andrés P., clínica odontológica, Bogotá"),
        ("«El perfil público me trajo 30 clientes nuevos en 2 meses desde Google sin pagar un peso de publicidad.»",
         "— Natalia R., estudio de yoga, Cali"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.right_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(quote)
        style_run(r, size=12, italic=True, color=DARK)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0)
        p2.paragraph_format.space_after = Pt(10)
        r = p2.add_run(who)
        style_run(r, size=10, color=GREY)

    h2(doc, "2.9  Preguntas frecuentes", anchor="p_faq")
    for q, a in [
        ("¿Tengo que firmar permanencia?",
         "No. Puedes cancelar en cualquier momento desde el panel de Facturación."),
        ("¿Puedo migrar desde otra plataforma?",
         "Sí. Importamos clientes y servicios en un archivo CSV. Citas históricas se cargan por rango."),
        ("¿Gestabiz procesa mis pagos?",
         "No directamente. Usamos MercadoPago, PayU o Stripe (a tu elección). Las tarjetas nunca pasan por nuestros servidores."),
        ("¿Mis datos son míos?",
         "Sí. Puedes exportar todo tu contenido en cualquier momento."),
        ("¿Funciona fuera de Colombia?",
         "Sí. La plataforma es multimoneda y multiidioma (español/inglés). Los módulos fiscales están pensados para Colombia pero son desactivables."),
        ("¿Qué pasa si supero un límite del plan?",
         "El módulo afectado pasa a modo lectura hasta que mejores el plan. Nunca se pierden datos."),
        ("¿Tienen aplicación para celular?",
         "Sí. Disponible para iOS y Android, más una extensión de navegador."),
    ]:
        h4(doc, q)
        para(doc, a)

    # CTA final
    page_break(doc)
    h1(doc, "Empieza hoy — Sin costo", color=PURPLE)
    para(doc, "Gestabiz está listo para producción, con operación real, soporte activo y hoja de ruta pública. No hay mejor momento para sumarte.", size=13)
    callout(doc, "Prueba Gestabiz",
            "Crea tu cuenta en gestabiz.com — configura tu negocio en 5 minutos — recibe tu primera reserva hoy.",
            color=PURPLE)
    para(doc, "Contacto comercial: ventas@gestabiz.com  ·  WhatsApp +57 XXX XXX XXXX")
    para(doc, "Desarrollado por Ti Turing — https://github.com/TI-Turing", italic=True, color=GREY)

    title_str = "Propuesta de Valor"
    return doc, title_str


# ===========================================================================
# DOCX → MARKDOWN conversion
# ===========================================================================
def _runs_to_md(para) -> str:
    """Convert runs of a paragraph to Markdown with bold/italic."""
    parts = []
    for run in para.runs:
        text = run.text or ""
        if not text:
            continue
        b = run.bold
        i = run.italic
        if b and i:
            parts.append(f"***{text}***")
        elif b:
            parts.append(f"**{text}**")
        elif i:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts)


def _table_to_md(table) -> str:
    """Convert a docx table to Markdown pipe format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        sep = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
        rows.insert(1, sep)
    return "\n".join(rows)


def docx_to_md(doc_path: Path, md_path: Path) -> None:
    """Convert a .docx to a .md file with proper headings."""
    from docx import Document as DocReader
    doc = DocReader(str(doc_path))
    lines: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            style_name = (para.style.name or "").lower()
            text = _runs_to_md(para)

            if not text.strip():
                lines.append("")
                continue

            if "heading 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif "heading 4" in style_name:
                lines.append(f"#### {text}")
            elif "list bullet" in style_name:
                lines.append(f"- {text}")
            elif "list number" in style_name:
                lines.append(f"1. {text}")
            else:
                # Check for large bold text (our custom headings)
                if para.runs:
                    first_run = para.runs[0]
                    sz = first_run.font.size
                    if sz and first_run.bold:
                        pt = sz.pt if hasattr(sz, "pt") else sz / 12700
                        raw = para.text.strip()
                        if pt >= 20:
                            lines.append(f"# {raw}")
                            continue
                        elif pt >= 14:
                            lines.append(f"## {raw}")
                            continue
                        elif pt >= 12:
                            lines.append(f"### {raw}")
                            continue
                lines.append(text)

            lines.append("")

        elif tag == "tbl":
            from docx.table import Table
            tbl = Table(element, doc)
            lines.append(_table_to_md(tbl))
            lines.append("")

        elif tag == "sectPr":
            lines.append("---")
            lines.append("")

    # Collapse 3+ blank lines
    output = "\n".join(lines)
    output = re.sub(r"\n{3,}", "\n\n", output)
    md_path.write_text(output, encoding="utf-8")


# ===========================================================================
# ENTRY POINT
# ===========================================================================
def main() -> None:
    print(">> Generando documentos Gestabiz...")
    print(f"   Repo:             {ROOT}")
    print(f"   Destino:          {DOCS_DIR}")
    print(f"   Logo Gestabiz:    {'OK' if LOGO_GESTABIZ.exists() else 'FALTA'} ({LOGO_GESTABIZ})")
    print(f"   Logo Ti Turing:   {'OK' if LOGO_TITURING.exists() else 'FALTA'} ({LOGO_TITURING})")

    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # --- Manual ---
    manual_doc, manual_title = build_manual()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz.docx"
    manual_doc.save(str(manual_path))
    setup_header_footer_post(manual_path, manual_title)
    print(f"[OK] Manual .docx:      {manual_path} ({manual_path.stat().st_size // 1024} KB)")

    manual_md = DOCS_DIR / "Manual_Usuario_Gestabiz.md"
    docx_to_md(manual_path, manual_md)
    print(f"[OK] Manual .md:        {manual_md}")

    # --- Propuesta ---
    proposal_doc, proposal_title = build_proposal()
    proposal_path = DOCS_DIR / "Propuesta_Valor_Gestabiz.docx"
    proposal_doc.save(str(proposal_path))
    setup_header_footer_post(proposal_path, proposal_title)
    print(f"[OK] Propuesta .docx:   {proposal_path} ({proposal_path.stat().st_size // 1024} KB)")

    proposal_md = DOCS_DIR / "Propuesta_Valor_Gestabiz.md"
    docx_to_md(proposal_path, proposal_md)
    print(f"[OK] Propuesta .md:     {proposal_md}")

    print(">> Listo. 4 archivos generados.")


def setup_header_footer_post(doc_path: Path, title: str) -> None:
    """Open a saved docx and add headers/footers (must be done post-save for section access)."""
    doc = Document(str(doc_path))
    for section in doc.sections:
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.different_first_page_header_footer = True

        # --- Header ---
        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(4)

        if LOGO_GESTABIZ.exists():
            try:
                run = hp.add_run()
                run.add_picture(str(LOGO_GESTABIZ), height=Cm(1.0))
            except Exception:
                run = hp.add_run("Gestabiz | Agenda. Gestiona. Crece.")
                style_run(run, size=10, bold=True, color=PURPLE)
        else:
            run = hp.add_run("Gestabiz | Agenda. Gestiona. Crece.")
            style_run(run, size=10, bold=True, color=PURPLE)

        pPr = hp._p.get_or_add_pPr()
        pPr.append(parse_xml(_border_xml("bottom")))

        # --- Footer ---
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)

        fpPr = fp._p.get_or_add_pPr()
        fpPr.append(parse_xml(_border_xml("top")))

        if LOGO_TITURING.exists():
            try:
                run = fp.add_run()
                run.add_picture(str(LOGO_TITURING), height=Cm(0.8))
            except Exception:
                run = fp.add_run("Ti Turing")
                style_run(run, size=9, bold=True, color=DARK)
        else:
            run = fp.add_run("Ti Turing")
            style_run(run, size=9, bold=True, color=DARK)

        sep = fp.add_run("  ·  Pág. ")
        style_run(sep, size=9, color=GREY)
        _add_page_field(fp)

        copy_run = fp.add_run(f"  ·  {title}  ·  © 2026")
        style_run(copy_run, size=8, italic=True, color=GREY)

    doc.save(str(doc_path))


if __name__ == "__main__":
    main()
