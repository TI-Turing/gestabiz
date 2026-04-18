#!/usr/bin/env python3
"""
Generador de documentos Gestabiz — Version Copilot — Parte 3: Rol Admin (Plan Gratuito)
========================================================================================
Genera:
  - docs/Manual_Usuario_Gestabiz - copilot - parte3.docx
  - docs/Propuesta_Valor_Gestabiz - copilot - parte3.docx

Requisitos: python-docx >= 1.0.0
Ejecucion:  python scripts/generate_copilot_docs_part3.py
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths & Brand Colors
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
LOGO_GESTABIZ = ROOT / "src" / "assets" / "images" / "logo_gestabiz.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def add_bookmark(paragraph, name: str) -> None:
    tag = OxmlElement("w:bookmarkStart")
    tag.set(qn("w:id"), str(id(name) % 2**31))
    tag.set(qn("w:name"), name)
    paragraph._p.append(tag)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(id(name) % 2**31))
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *,
                           bold: bool = False, color: str = "7C3AED") -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# High-level helpers
# ---------------------------------------------------------------------------

def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + text)
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(title)
    style_run(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {caption}")
    style_run(r, size=9, italic=True, color=GREY)
    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=header_fg)
    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, size=10, color=DARK)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup + cover
# ---------------------------------------------------------------------------

def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run(f"{title}  |  Desarrollado por Ti Turing  |  (c) 2026")
        style_run(r, size=9, italic=True, color=GREY)
    build_cover(doc, title, subtitle)
    return doc


def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Parte 3 de 5: Administrador — Plan Gratuito")
    style_run(r, size=14, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version del producto: 1.0.3   |   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fase Beta completada   |   Listo para produccion")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing")
    style_run(r, size=12, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("https://github.com/TI-Turing")
    style_run(r, size=10, italic=True, color=GREY)

    page_break(doc)


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  MANUAL DE USUARIO — PARTE 3: ADMINISTRADOR PLAN GRATUITO              ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

def build_manual_part3() -> Document:
    doc = setup_document(
        "Manual de Usuario — Gestabiz",
        "Guia funcional completa del sistema de gestion de citas y negocios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE DE CONTENIDOS", anchor="toc_p3")
    para(doc, "A continuacion se listan todas las secciones cubiertas en esta parte. "
         "Haga clic en cualquier titulo para navegar directamente al detalle.", italic=True, color=GREY)

    toc_items = [
        ("sec_resumen_admin", "1. Resumen Ejecutivo — Administrador Plan Gratuito"),
        ("sec_acceso_admin", "2. Acceso al Panel de Administracion"),
        ("sec_registro_negocio", "3. Registro de un Nuevo Negocio"),
        ("sec_onboarding", "4. Asistente de Configuracion Inicial (Onboarding)"),
        ("sec_dashboard_overview", "5. Dashboard — Pestana Resumen"),
        ("sec_stats_cards", "6. Tarjetas de Estadisticas"),
        ("sec_info_negocio", "7. Tarjeta de Informacion del Negocio"),
        ("sec_qr_modal", "8. Generacion de Codigo QR"),
        ("sec_perfil_publico", "9. Vista de Perfil Publico del Negocio"),
        ("sec_health_panel", "10. Panel de Salud Operacional"),
        ("sec_sidebar", "11. Barra Lateral (Sidebar) y Navegacion"),
        ("sec_servicios", "12. Gestion de Servicios"),
        ("sec_srv_crear", "13. Crear un Servicio"),
        ("sec_srv_editar", "14. Editar un Servicio"),
        ("sec_srv_eliminar", "15. Eliminar un Servicio"),
        ("sec_srv_reactivar", "16. Reactivar un Servicio"),
        ("sec_srv_asignaciones", "17. Asignaciones de Servicios (Sedes y Empleados)"),
        ("sec_sedes", "18. Gestion de Sedes"),
        ("sec_sede_crear", "19. Crear una Sede"),
        ("sec_sede_editar", "20. Editar una Sede"),
        ("sec_sede_eliminar", "21. Eliminar una Sede"),
        ("sec_sede_multimedia", "22. Multimedia de Sedes (Imagenes y Videos)"),
        ("sec_sede_principal", "23. Sede Principal — Reglas"),
        ("sec_citas_calendario", "24. Calendario de Citas"),
        ("sec_cal_navegacion", "25. Navegacion del Calendario"),
        ("sec_cal_filtros", "26. Filtros del Calendario"),
        ("sec_cal_colores", "27. Codificacion de Colores por Estado de Cita"),
        ("sec_cal_modal_detalle", "28. Modal de Detalle de Cita"),
        ("sec_cal_completar", "29. Flujo Completar Cita (Calculo Fiscal)"),
        ("sec_cal_acciones", "30. Acciones sobre una Cita"),
        ("sec_cal_pendientes", "31. Citas Pendientes y En Proceso"),
        ("sec_clientes", "32. Gestion de Clientes (CRM Basico)"),
        ("sec_cliente_perfil", "33. Modal de Perfil de Cliente"),
        ("sec_facturacion", "34. Facturacion y Planes"),
        ("sec_billing_free", "35. Vista de Plan Gratuito"),
        ("sec_billing_trial", "36. Prueba Gratuita (Free Trial)"),
        ("sec_billing_pricing", "37. Pagina de Precios y Comparativa"),
        ("sec_config", "38. Configuraciones del Negocio"),
        ("sec_config_info", "39. Informacion del Negocio"),
        ("sec_config_branding", "40. Logo y Banner (Branding)"),
        ("sec_config_notificaciones", "41. Configuracion de Notificaciones"),
        ("sec_config_tracking", "42. Tracking de Notificaciones"),
        ("sec_config_calendario", "43. Dias Cerrados"),
        ("sec_config_chat", "44. Configuracion del Chat"),
        ("sec_config_general", "45. Configuracion General (Tema e Idioma)"),
        ("sec_config_perfil", "46. Edicion de Perfil Personal"),
        ("sec_config_danger", "47. Zona de Peligro — Desactivar Cuenta"),
        ("sec_limites_plan", "48. Limites del Plan Gratuito — Resumen"),
        ("sec_modulos_bloqueados", "49. Modulos Bloqueados y Pantalla de Upgrade"),
        ("sec_plan_limit_banner", "50. Banner de Limite Alcanzado"),
        ("sec_permisos_admin", "51. Permisos del Administrador (PermissionGate)"),
        ("sec_multi_negocio", "52. Gestion de Multiples Negocios"),
        ("sec_cambio_rol", "53. Cambio de Rol (Admin / Empleado / Cliente)"),
        ("sec_estados_carga", "54. Estados de Carga, Error y Vacios"),
        ("sec_glosario_admin", "55. Glosario de Terminos"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 1 — RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 1 — RESUMEN EJECUTIVO", anchor="parte1_resumen_p3")

    # --- 1. Resumen ejecutivo ─────────────────────────────────────────────
    h2(doc, "1. Resumen Ejecutivo — Administrador Plan Gratuito", anchor="sec_resumen_admin")

    para(doc, "El rol de Administrador en Gestabiz es el centro de operaciones de un negocio. "
         "Desde el panel de administracion, el dueno o administrador de un negocio puede gestionar "
         "cada aspecto de su operacion: servicios, sedes, citas, clientes y configuraciones. "
         "El Plan Gratuito ofrece una experiencia completa para negocios que inician, con limites "
         "razonables que permiten validar la plataforma sin ningun costo.")

    callout(doc, "Ideal para emprendedores",
            "El Plan Gratuito esta disenado para negocios que recien comienzan o quieren probar Gestabiz "
            "sin compromiso. Incluye todas las herramientas fundamentales: gestion de citas, servicios, "
            "una sede, un catalogo de clientes y un perfil publico visible en Google.",
            color=ACCENT)

    h3(doc, "Lo que incluye el Plan Gratuito")
    simple_table(doc, ["Recurso", "Limite", "Descripcion"], [
        ["Sedes", "1", "Una ubicacion fisica del negocio"],
        ["Empleados", "1", "El propio dueno del negocio"],
        ["Citas por mes", "50", "Hasta 50 citas mensuales agendadas"],
        ["Clientes visibles", "50", "Hasta 50 clientes en el catalogo CRM"],
        ["Servicios", "15", "Hasta 15 servicios ofrecidos"],
        ["Perfil publico", "Incluido", "Pagina indexable en Google con SEO completo"],
        ["QR del negocio", "Incluido", "Codigo QR descargable para compartir"],
        ["Dashboard de resumen", "Incluido", "Estadisticas basicas del negocio"],
        ["Facturacion y planes", "Incluido", "Gestion de suscripcion y upgrades"],
        ["Configuraciones", "Incluido", "Logo, banner, notificaciones, chat, dias cerrados"],
        ["Recordatorios por email", "Incluido", "Recordatorios automaticos a clientes"],
    ])

    h3(doc, "Modulos no disponibles en el Plan Gratuito")
    para(doc, "Los siguientes modulos requieren un plan superior. Al intentar acceder, el sistema "
         "mostrara una pantalla informativa con opcion de actualizar el plan.")
    simple_table(doc, ["Modulo", "Plan requerido"], [
        ["Gestion de Empleados", "Basico"],
        ["Ausencias y Vacaciones", "Basico"],
        ["Historial de Ventas", "Basico"],
        ["Ventas Rapidas", "Basico"],
        ["Reportes", "Basico"],
        ["Permisos Granulares", "Basico"],
        ["Gastos y Egresos", "Pro"],
        ["Portal de Reclutamiento", "Pro"],
        ["Recursos Fisicos", "Pro"],
    ])

    screenshot_placeholder(doc, "Panel de administracion — Plan Gratuito con modulos bloqueados")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 2 — DETALLE EXHAUSTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 2 — DETALLE EXHAUSTIVO", anchor="parte2_detalle_p3")

    # --- 2. Acceso al Panel ───────────────────────────────────────────────
    h2(doc, "2. Acceso al Panel de Administracion", anchor="sec_acceso_admin")

    para(doc, "Para acceder al panel de administracion, el usuario debe haber iniciado sesion y tener "
         "al menos un negocio registrado como propietario (owner) o haber sido asignado como administrador "
         "de un negocio existente.")

    h3(doc, "Requisitos de acceso")
    bullet(doc, "Cuenta activa con email verificado.")
    bullet(doc, "Ser propietario del negocio (owner_id en tabla businesses) o tener rol admin en business_roles.")
    bullet(doc, "Los roles se calculan automaticamente — no se almacenan en la base de datos.")

    h3(doc, "Flujo de acceso")
    numbered(doc, "El usuario inicia sesion con email/contrasena o Google.")
    numbered(doc, "El sistema calcula los roles disponibles: si el usuario es owner de algun negocio, el rol Admin estara disponible.")
    numbered(doc, "Si el usuario tiene multiples roles disponibles (Admin, Empleado, Cliente), puede cambiar entre ellos desde el selector de roles en el encabezado.")
    numbered(doc, "Al seleccionar el rol Admin, el sistema carga el AdminDashboard con el negocio correspondiente.")
    numbered(doc, "Si el usuario administra varios negocios, puede cambiar entre ellos desde el dropdown del encabezado.")

    callout(doc, "Roles dinamicos",
            "Los roles en Gestabiz NO se guardan en la base de datos. Se calculan en tiempo real "
            "a partir de la relacion del usuario con los negocios. Un mismo usuario puede ser Admin "
            "en un negocio, Empleado en otro y Cliente en cualquiera.",
            color=PURPLE)

    screenshot_placeholder(doc, "Selector de roles en el encabezado del dashboard")

    page_break(doc)

    # --- 3. Registro de Negocio ───────────────────────────────────────────
    h2(doc, "3. Registro de un Nuevo Negocio", anchor="sec_registro_negocio")

    para(doc, "Cuando un usuario accede al rol Admin sin tener ningun negocio registrado, el sistema "
         "muestra automaticamente el formulario de registro de negocio. Tambien es posible crear un "
         "nuevo negocio desde el dropdown del encabezado con la opcion 'Crear nuevo negocio'.")

    h3(doc, "Campos del formulario")
    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"], [
        ["Nombre del negocio", "Texto", "Si", "Nombre comercial del negocio"],
        ["Categoria", "Selector con busqueda", "Si", "15 categorias principales (Salud, Belleza, etc.)"],
        ["Subcategorias", "3 campos de texto", "No", "Hasta 3 subcategorias libres"],
        ["Tipo de entidad", "Radio: Individual / Empresa", "Si", "Determina campos adicionales"],
        ["Razon social", "Texto", "Solo si es empresa", "Nombre legal de la empresa"],
        ["NIT / RUT", "Texto", "No", "Identificacion fiscal"],
        ["Numero de registro", "Texto", "Solo si es empresa", "Registro mercantil"],
        ["Modelo de negocio", "Radio grid (4 opciones)", "Si", "Define tipo de operacion"],
        ["Descripcion", "Texto area", "No", "Descripcion del negocio"],
        ["Telefono", "Telefono con prefijo", "No", "Telefono de contacto"],
        ["Email", "Email", "No", "Correo electronico del negocio"],
        ["Sitio web", "URL", "No", "Pagina web del negocio"],
        ["Pais", "Selector (default: Colombia)", "Si", "Pais de operacion"],
        ["Region / Departamento", "Selector cascada", "Si", "Region dentro del pais"],
        ["Ciudad", "Selector cascada", "Si", "Ciudad dentro de la region"],
    ])

    h3(doc, "Modelo de negocio — 4 opciones")
    simple_table(doc, ["Opcion", "Descripcion", "Ejemplo"], [
        ["Profesionales", "El negocio asigna profesionales a citas", "Salon de belleza, clinica"],
        ["Recursos Fisicos", "El negocio reserva recursos como salas o mesas", "Hotel, restaurante"],
        ["Hibrido", "Combinacion de profesionales y recursos", "Spa con masajistas y salas"],
        ["Clases Grupales", "El negocio ofrece clases o actividades en grupo", "Gimnasio, academia"],
    ])

    callout(doc, "Nota importante",
            "Despues de crear el negocio, se abrira automaticamente el asistente de configuracion "
            "inicial para guiarle en la configuracion de servicios, sedes y empleados.",
            color=ACCENT)

    h3(doc, "Botones disponibles")
    bullet(doc, "Crear Negocio — envia el formulario y crea el negocio en la base de datos. El boton cambia a 'Creando...' mientras se procesa.", bold_prefix="Crear Negocio")
    bullet(doc, "Cancelar — cierra el formulario sin guardar. Solo visible si ya hay otro negocio registrado.", bold_prefix="Cancelar")

    h3(doc, "Flujo post-registro")
    numbered(doc, "El negocio se crea con resource_model seleccionado y is_configured = false.")
    numbered(doc, "El owner es registrado automaticamente como empleado tipo 'manager' mediante un trigger.")
    numbered(doc, "Se asignan automaticamente los 79 permisos del sistema al owner.")
    numbered(doc, "Se muestra el asistente de configuracion inicial (Setup Checklist).")

    screenshot_placeholder(doc, "Formulario de registro de nuevo negocio")

    page_break(doc)

    # --- 4. Onboarding ────────────────────────────────────────────────────
    h2(doc, "4. Asistente de Configuracion Inicial (Onboarding)", anchor="sec_onboarding")

    para(doc, "El asistente de configuracion (Setup Checklist) aparece en la pestana Resumen del dashboard "
         "cuando el negocio aun no esta completamente configurado. Muestra una barra de progreso y una "
         "lista de pasos requeridos y opcionales.")

    h3(doc, "Pasos del asistente")
    simple_table(doc, ["#", "Paso", "Requerido", "Accion"], [
        ["1", "Crear al menos una sede", "Si", "Navega a Sedes"],
        ["2", "Configurar servicios", "Si", "Navega a Servicios"],
        ["3", "Asignar profesionales o recursos", "Si", "Navega a Empleados"],
        ["4", "Subir logo del negocio", "No", "Navega a Configuraciones"],
        ["5", "Escribir descripcion del negocio", "No", "Navega a Configuraciones"],
        ["6", "Agregar telefono de contacto", "No", "Navega a Configuraciones"],
    ])

    para(doc, "Cada paso muestra un icono de verificacion cuando esta completado. La barra de progreso "
         "se actualiza en tiempo real. Cuando todos los pasos requeridos estan completos, el negocio "
         "cambia a estado 'Publico' (badge verde) y se vuelve visible en busquedas y en el perfil publico.")

    h3(doc, "Condicion de negocio completamente configurado")
    bullet(doc, "El campo is_configured del negocio es verdadero.")
    bullet(doc, "Tiene al menos 1 sede activa.")
    bullet(doc, "Tiene al menos 1 servicio activo.")
    bullet(doc, "Las sedes tienen servicios asignados.")
    bullet(doc, "Los empleados tienen servicios asignados.")
    bullet(doc, "No hay problemas operacionales detectados (sin alertas en el panel de salud).")

    callout(doc, "Badge de visibilidad",
            "Cuando el negocio esta completamente configurado, aparece un badge 'Publico' verde en la parte "
            "superior del dashboard. Si falta algun paso, el badge dira 'No publico' en rojo y el negocio "
            "no sera visible para los clientes en las busquedas.",
            color=ACCENT)

    screenshot_placeholder(doc, "Asistente de configuracion inicial con barra de progreso")

    page_break(doc)

    # --- 5. Dashboard Overview ────────────────────────────────────────────
    h2(doc, "5. Dashboard — Pestana Resumen", anchor="sec_dashboard_overview")

    para(doc, "La pestana Resumen es la primera vista que ve el administrador al entrar al panel. Proporciona "
         "una vision general del negocio con estadisticas en tiempo real, accesos rapidos y el estado "
         "de configuracion del negocio.")

    h3(doc, "Secciones de la pestana Resumen")
    numbered(doc, "Setup Checklist (si el negocio no esta completamente configurado).")
    numbered(doc, "Panel de Salud Operacional (si hay problemas detectados).")
    numbered(doc, "Grid de 9 tarjetas de estadisticas en 3 filas.")
    numbered(doc, "Tarjeta de Informacion del Negocio.")

    screenshot_placeholder(doc, "Vista general del dashboard — Pestana Resumen")

    # --- 6. Stats Cards ───────────────────────────────────────────────────
    h2(doc, "6. Tarjetas de Estadisticas", anchor="sec_stats_cards")

    para(doc, "El dashboard muestra 9 tarjetas de estadisticas organizadas en 3 filas. Cada tarjeta "
         "muestra un valor numerico con icono y descripcion.")

    simple_table(doc, ["Fila", "Tarjeta", "Descripcion", "Clickeable"], [
        ["1", "Citas Hoy", "Numero de citas programadas para el dia actual", "Si — abre calendario"],
        ["1", "Proximas Citas", "Conteo de citas futuras confirmadas", "No"],
        ["1", "Citas Completadas", "Citas completadas en el mes actual", "No"],
        ["2", "Citas Canceladas", "Citas canceladas en el mes actual", "No"],
        ["2", "Empleados", "Miembros activos del equipo (excluye owner)", "No"],
        ["2", "Sedes", "Total de ubicaciones del negocio", "No"],
        ["2", "Servicios", "Servicios activos configurados", "No"],
        ["3", "Ingresos del Mes", "Ingresos COP del mes (citas completadas)", "No"],
        ["3", "Valor Promedio por Cita", "Promedio COP por cita completada", "No"],
    ])

    callout(doc, "Tarjeta interactiva",
            "La tarjeta 'Citas Hoy' es la unica clickeable. Al hacer clic, navega directamente al calendario "
            "de citas con los filtros guardados en el almacenamiento local del navegador.",
            color=PURPLE)

    h3(doc, "Estado de carga")
    para(doc, "Mientras se cargan las estadisticas, se muestra un grid de 6 tarjetas con efecto skeleton "
         "(placeholders animados) para mantener la estructura visual.")

    screenshot_placeholder(doc, "Grid de estadisticas del dashboard")

    # --- 7. Info del Negocio ──────────────────────────────────────────────
    h2(doc, "7. Tarjeta de Informacion del Negocio", anchor="sec_info_negocio")

    para(doc, "Esta tarjeta muestra la informacion principal del negocio y ofrece acciones rapidas.")

    h3(doc, "Informacion mostrada")
    bullet(doc, "Logo del negocio (o iniciales como fallback).")
    bullet(doc, "Nombre del negocio.")
    bullet(doc, "Categoria y subcategorias.")
    bullet(doc, "Descripcion del negocio.")
    bullet(doc, "Telefono y email de contacto.")

    h3(doc, "Acciones disponibles")
    simple_table(doc, ["Boton", "Accion", "Resultado"], [
        ["Copiar enlace", "Copia la URL del perfil publico al portapapeles", "Toast de confirmacion"],
        ["Abrir perfil publico", "Abre el perfil en una nueva pestana", "Navega a /negocio/{slug}"],
        ["Generar QR", "Abre el modal de generacion de QR", "Ver seccion 8"],
        ["Ver perfil / Ocultar perfil", "Embebe el perfil publico dentro del dashboard", "Toggle inline"],
    ])

    screenshot_placeholder(doc, "Tarjeta de informacion del negocio con acciones rapidas")

    # --- 8. QR Modal ──────────────────────────────────────────────────────
    h2(doc, "8. Generacion de Codigo QR", anchor="sec_qr_modal")

    para(doc, "El modal de generacion de QR permite crear un codigo QR descargable que los clientes "
         "pueden escanear para acceder al perfil del negocio o iniciar una reserva directamente.")

    h3(doc, "Opciones del QR")
    simple_table(doc, ["Modo", "URL generada", "Uso recomendado"], [
        ["Perfil publico (default)", "https://gestabiz.com/negocio/{slug}", "Tarjetas de visita, flyers"],
        ["Reservas directas", "https://gestabiz.com/negocio/{slug}?book=true", "Punto de venta, mostrador"],
    ])

    h3(doc, "Flujo de uso")
    numbered(doc, "Abrir el modal desde la tarjeta de informacion del negocio (boton 'Generar QR').")
    numbered(doc, "Seleccionar el modo: perfil publico o reservas directas (checkbox toggle).")
    numbered(doc, "El QR se genera automaticamente con el logo de Gestabiz centrado.")
    numbered(doc, "Hacer clic en 'Descargar PNG' para guardar la imagen.")
    numbered(doc, "Cerrar el modal con el boton X o haciendo clic fuera.")

    callout(doc, "QR con reserva directa",
            "Cuando el cliente escanea el QR en modo 'Reservas directas', se abre automaticamente el "
            "perfil del negocio con el wizard de reserva listo para comenzar. Esto reduce la friccion "
            "en un 57% comparado con el flujo normal.",
            color=ACCENT)

    screenshot_placeholder(doc, "Modal de generacion de codigo QR")

    page_break(doc)

    # --- 9. Perfil Publico ────────────────────────────────────────────────
    h2(doc, "9. Vista de Perfil Publico del Negocio", anchor="sec_perfil_publico")

    para(doc, "El administrador puede visualizar como se ve su negocio para los clientes sin salir "
         "del dashboard. El boton 'Ver perfil' embebe el perfil publico completo de forma inline.")

    bullet(doc, "El perfil incluye 4 pestanas: Servicios, Ubicaciones, Resenas, Acerca de.")
    bullet(doc, "Muestra informacion real del negocio: logo, banner, servicios con precios, sedes con mapa.")
    bullet(doc, "Las resenas son anonimas y muestran calificacion de 1 a 5 estrellas.")
    bullet(doc, "El boton 'Abrir perfil publico' abre la misma pagina en una pestana nueva.")

    callout(doc, "SEO integrado",
            "Cada perfil publico tiene una URL amigable (/negocio/{slug}), meta tags dinamicos, Open Graph, "
            "Twitter Card y datos estructurados JSON-LD. El negocio aparece automaticamente en Google "
            "una vez que esta completamente configurado.",
            color=PURPLE)

    screenshot_placeholder(doc, "Perfil publico del negocio embebido en el dashboard")

    # --- 10. Health Panel ─────────────────────────────────────────────────
    h2(doc, "10. Panel de Salud Operacional", anchor="sec_health_panel")

    para(doc, "El Panel de Salud Operacional (AssignmentHealthPanel) aparece automaticamente en la "
         "pestana Resumen cuando el sistema detecta problemas que afectan la operacion del negocio.")

    h3(doc, "Problemas detectados")
    bullet(doc, "Servicios sin empleados asignados — ningun profesional puede atender ese servicio.", bold_prefix="Servicios sin asignar")
    bullet(doc, "Sedes sin servicios asignados — la sede no ofrece ningun servicio a los clientes.", bold_prefix="Sedes sin servicios")
    bullet(doc, "Empleados sin supervisor — el empleado no tiene un jefe directo asignado.", bold_prefix="Sin supervisor")
    bullet(doc, "Empleados sin horario — el empleado no tiene dias/horas de trabajo configurados.", bold_prefix="Sin horario")
    bullet(doc, "Empleados sin servicios — el profesional no tiene servicios que pueda ofrecer.", bold_prefix="Sin servicios")

    para(doc, "Cada problema se muestra como una alerta con icono y descripcion, junto con un enlace "
         "al modulo correspondiente para resolverlo.")

    callout(doc, "Visibilidad condicional",
            "El Panel de Salud solo aparece cuando hay problemas reales. Si todo esta correctamente "
            "configurado, esta seccion se oculta automaticamente.",
            color=ACCENT)

    screenshot_placeholder(doc, "Panel de salud operacional con alertas activas")

    page_break(doc)

    # --- 11. Sidebar ──────────────────────────────────────────────────────
    h2(doc, "11. Barra Lateral (Sidebar) y Navegacion", anchor="sec_sidebar")

    para(doc, "La barra lateral izquierda del panel de administracion muestra todos los modulos "
         "disponibles. Los modulos desbloqueados aparecen primero; los bloqueados se muestran al final "
         "con un icono de candado y un badge indicando el plan requerido.")

    h3(doc, "Modulos accesibles en Plan Gratuito")
    simple_table(doc, ["Icono", "Modulo", "Descripcion"], [
        ["LayoutDashboard", "Resumen", "Vista general con estadisticas y configuracion"],
        ["Calendar", "Citas", "Calendario visual con gestion completa de citas"],
        ["Briefcase", "Servicios", "CRUD de servicios con asignaciones"],
        ["UserCheck", "Clientes", "Catalogo CRM con historial de visitas"],
        ["CreditCard", "Facturacion", "Gestion de plan, pagos y upgrades"],
        ["MapPin", "Sedes", "Gestion de ubicaciones fisicas del negocio"],
    ])

    h3(doc, "Modulos bloqueados (con icono de candado)")
    simple_table(doc, ["Modulo", "Plan requerido", "Badge mostrado"], [
        ["Empleados", "Basico", "Basico"],
        ["Ausencias", "Basico", "Basico"],
        ["Ventas", "Basico", "Basico"],
        ["Ventas Rapidas", "Basico", "Basico"],
        ["Reportes", "Basico", "Basico"],
        ["Permisos", "Basico", "Basico"],
        ["Gastos", "Pro", "Pro"],
        ["Reclutamiento", "Pro", "Pro"],
        ["Recursos", "Pro", "Pro"],
    ])

    para(doc, "Al hacer clic en un modulo bloqueado, se muestra una pantalla con icono de candado, "
         "el nombre del modulo, el plan requerido y un boton 'Ver planes y precios' que navega "
         "directamente a la seccion de facturacion.")

    screenshot_placeholder(doc, "Sidebar con modulos desbloqueados y bloqueados")

    page_break(doc)

    # --- 12-17. Servicios ─────────────────────────────────────────────────
    h2(doc, "12. Gestion de Servicios", anchor="sec_servicios")

    para(doc, "El modulo de Servicios permite al administrador crear, editar, eliminar y reactivar los "
         "servicios que ofrece su negocio. Cada servicio tiene precio, duracion, descripcion y puede "
         "tener una imagen asociada.")

    h3(doc, "Limite del Plan Gratuito")
    callout(doc, "Maximo 15 servicios visibles",
            "En el Plan Gratuito, se muestran hasta 15 servicios. Si se crean mas, los excedentes se almacenan "
            "en la base de datos pero no son visibles. Un banner informativo indica cuantos servicios "
            "adicionales estan ocultos y ofrece actualizar el plan.",
            color=DANGER)

    h3(doc, "Vista de la lista de servicios")
    bullet(doc, "Grid responsive: 1 columna (movil), 2 columnas (tablet), 3 columnas (escritorio).")
    bullet(doc, "Cada tarjeta muestra: imagen de fondo con gradiente, nombre, descripcion (2 lineas), precio en COP, duracion en minutos.")
    bullet(doc, "Badges de alerta: 'Sin sedes' (si el servicio no esta asignado a ninguna sede), 'Sin empleados' (si no tiene profesionales asignados).")
    bullet(doc, "Badge 'Inactivo' para servicios desactivados (solo visibles con el toggle activo).")
    bullet(doc, "Toggle 'Mostrar inactivos' para incluir servicios desactivados en la lista.")
    bullet(doc, "Click en la tarjeta abre el modal de perfil del servicio.")

    h3(doc, "Boton de accion principal")
    para(doc, "El boton 'Nuevo Servicio' (icono +) esta protegido por el sistema de permisos (PermissionGate). "
         "Si el administrador no tiene el permiso 'services.create', el boton no se mostrara.")

    screenshot_placeholder(doc, "Lista de servicios con tarjetas")

    # --- 13. Crear Servicio ───────────────────────────────────────────────
    h2(doc, "13. Crear un Servicio", anchor="sec_srv_crear")

    para(doc, "Al hacer clic en 'Nuevo Servicio', se abre un formulario en un dialogo modal.")

    h3(doc, "Campos del formulario")
    simple_table(doc, ["Campo", "Tipo", "Requerido", "Validacion"], [
        ["Nombre", "Texto", "Si", "No vacio, se limpia espacios (trim)"],
        ["Descripcion", "Texto area", "No", "Maximo 500 caracteres con contador"],
        ["Duracion (minutos)", "Numero", "Si", "Mayor a 0"],
        ["Precio", "Numero (COP)", "Si", "Mayor o igual a 0, separadores de miles"],
        ["Categoria", "Selector", "No", "De la lista de categorias"],
        ["Comision %", "Numero", "No", "Oculto si el negocio tiene 1 o menos empleados"],
        ["Imagen", "Carga de archivo", "No", "Upload diferido a bucket de Supabase"],
        ["Activo", "Toggle", "No", "Default: activado"],
    ])

    h3(doc, "Flujo de creacion")
    numbered(doc, "Llenar los campos requeridos (nombre, duracion, precio).")
    numbered(doc, "Opcionalmente subir una imagen y agregar descripcion.")
    numbered(doc, "Hacer clic en 'Crear Servicio'.")
    numbered(doc, "Si la creacion es exitosa, se abre automaticamente el dialogo de asignaciones.")
    numbered(doc, "Seleccionar las sedes y empleados que ofreceran este servicio (checkboxes).")
    numbered(doc, "Guardar las asignaciones.")
    numbered(doc, "Toast de confirmacion: 'Servicio creado exitosamente'.")

    callout(doc, "Proteccion contra doble clic",
            "El boton de guardar se deshabilita automaticamente mientras se procesa la solicitud "
            "para evitar la creacion duplicada de servicios.",
            color=PURPLE)

    screenshot_placeholder(doc, "Formulario de creacion de servicio")

    # --- 14. Editar ───────────────────────────────────────────────────────
    h2(doc, "14. Editar un Servicio", anchor="sec_srv_editar")

    para(doc, "Para editar un servicio existente, el administrador hace clic en el boton de edicion "
         "(icono de lapiz) en la tarjeta del servicio. Se abre el mismo formulario pre-llenado con los "
         "datos actuales del servicio.")

    bullet(doc, "Todos los campos son editables, incluyendo la imagen.")
    bullet(doc, "Si se cambia la imagen, la anterior se elimina del almacenamiento si es una URL de Supabase.")
    bullet(doc, "Las asignaciones de sedes y empleados se cargan automaticamente.")
    bullet(doc, "El boton de guardar esta protegido por el permiso 'services.edit'.")
    bullet(doc, "Si se guarda con imagen nueva, se aplica cache-bust para que el navegador muestre la imagen actualizada.")

    # --- 15. Eliminar ─────────────────────────────────────────────────────
    h2(doc, "15. Eliminar un Servicio", anchor="sec_srv_eliminar")

    para(doc, "La eliminacion de un servicio es una operacion critica que tiene efectos sobre las citas "
         "programadas.")

    h3(doc, "Flujo de eliminacion")
    numbered(doc, "Hacer clic en el boton de eliminacion (icono de papelera) en la tarjeta del servicio.")
    numbered(doc, "El sistema muestra una ventana de confirmacion indicando cuantas citas pendientes hay.")
    numbered(doc, "Si el administrador confirma:")
    numbered(doc, "  a) Se cancelan todas las citas pendientes asociadas al servicio.")
    numbered(doc, "  b) Se envia notificacion de cancelacion por email a cada cliente afectado.")
    numbered(doc, "  c) Se eliminan las asignaciones del servicio (location_services y employee_services).")
    numbered(doc, "  d) El servicio se desactiva (soft-delete: is_active = false).")
    numbered(doc, "Toast de confirmacion con el numero de citas canceladas.")

    callout(doc, "Soft-delete",
            "Los servicios eliminados no se borran permanentemente. Se marcan como inactivos (is_active = false) "
            "y pueden reactivarse posteriormente. Esto preserva el historial de citas y transacciones.",
            color=ACCENT)

    # --- 16. Reactivar ────────────────────────────────────────────────────
    h2(doc, "16. Reactivar un Servicio", anchor="sec_srv_reactivar")

    para(doc, "Los servicios desactivados pueden reactivarse desde la lista de servicios.")

    numbered(doc, "Activar el toggle 'Mostrar inactivos' en la parte superior de la lista.")
    numbered(doc, "Localizar el servicio con badge 'Inactivo'.")
    numbered(doc, "Hacer clic en el boton de reactivacion (icono de flecha circular).")
    numbered(doc, "El servicio se reactiva (is_active = true) y se abre el editor para reasignar sedes y empleados.")
    numbered(doc, "Toast: 'Selecciona las sedes para este servicio'.")

    # --- 17. Asignaciones ─────────────────────────────────────────────────
    h2(doc, "17. Asignaciones de Servicios (Sedes y Empleados)", anchor="sec_srv_asignaciones")

    para(doc, "Cada servicio debe estar asignado a al menos una sede y un empleado para ser reservable. "
         "Las asignaciones se gestionan desde un dialogo con checkboxes.")

    h3(doc, "Dialogo de asignaciones")
    bullet(doc, "Lista de sedes del negocio con checkboxes.")
    bullet(doc, "Lista de empleados con checkboxes.")
    bullet(doc, "Opcion de comision personalizada por combinacion sede-empleado.")
    bullet(doc, "Al guardar: se eliminan todas las asignaciones anteriores y se crean las nuevas.")

    callout(doc, "Sin asignaciones = sin reservas",
            "Si un servicio no tiene sedes o empleados asignados, los clientes NO podran reservar citas "
            "para ese servicio. El sistema mostrara badges de alerta amarillos ('Sin sedes', 'Sin empleados') "
            "en la tarjeta del servicio para advertir al administrador.",
            color=DANGER)

    screenshot_placeholder(doc, "Dialogo de asignaciones de servicio con checkboxes")

    page_break(doc)

    # --- 18-23. Sedes ─────────────────────────────────────────────────────
    h2(doc, "18. Gestion de Sedes", anchor="sec_sedes")

    para(doc, "El modulo de Sedes permite al administrador gestionar las ubicaciones fisicas de su negocio. "
         "En el Plan Gratuito, el limite es de 1 sede.")

    callout(doc, "Limite del Plan Gratuito: 1 sede",
            "Solo se puede crear y visualizar una sede. El boton 'Agregar Sede' se deshabilita cuando el "
            "limite se alcanza, mostrando un tooltip explicativo. Si por alguna razon existen mas sedes en la "
            "base de datos, solo la primera sera visible y un banner informativo indicara cuantas estan ocultas.",
            color=DANGER)

    h3(doc, "Vista de la lista de sedes")
    bullet(doc, "Grid responsive: 1/2/3 columnas segun el tamano de pantalla.")
    bullet(doc, "Cada tarjeta muestra: banner como fondo completo con gradiente, nombre, direccion, telefono, email, horarios.")
    bullet(doc, "Badges: 'Principal' (sede principal del negocio), 'Administrada' (sede preferida del admin), 'Sin servicios asignados' (alerta amarilla), 'Inactiva'.")
    bullet(doc, "Conteo de imagenes asociadas a la sede.")
    bullet(doc, "Click en la tarjeta abre el modal de perfil de la sede.")

    screenshot_placeholder(doc, "Tarjeta de sede con banner e informacion")

    # --- 19. Crear Sede ───────────────────────────────────────────────────
    h2(doc, "19. Crear una Sede", anchor="sec_sede_crear")

    para(doc, "Si el negocio no tiene sedes (primera configuracion) o si no ha alcanzado el limite del plan, "
         "el boton 'Agregar Sede' estara disponible.")

    h3(doc, "Campos del formulario")
    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"], [
        ["Nombre", "Texto", "Si", "Nombre de la sede (ej: 'Sede Norte')"],
        ["Direccion", "Texto", "No", "Direccion completa"],
        ["Ciudad", "Selector cascada", "No", "Ciudad de la sede"],
        ["Estado / Region", "Selector cascada", "No", "Departamento o estado"],
        ["Pais", "Selector", "No", "Default: Colombia"],
        ["Codigo postal", "Texto", "No", "Codigo postal de la zona"],
        ["Telefono", "Telefono con prefijo", "No", "Telefono de la sede"],
        ["Email", "Email", "No", "Correo de la sede"],
        ["Descripcion", "Texto area", "No", "Descripcion de la sede"],
        ["Horarios de negocio", "Selector de horas", "No", "Horario de apertura y cierre por dia"],
        ["Imagenes", "Carga multiple", "No", "Fotos de la sede"],
        ["Activa", "Toggle", "No", "Default: activada"],
        ["Sede principal", "Toggle", "No", "La primera sede se marca automaticamente"],
        ["Atender en festivos", "Toggle", "No", "Si la sede opera en dias festivos"],
    ])

    h3(doc, "Regla de primera sede")
    para(doc, "La primera sede creada en el negocio se marca automaticamente como sede principal (is_primary = true). "
         "Esta regla no se puede modificar hasta que exista una segunda sede.")

    screenshot_placeholder(doc, "Formulario de creacion de sede")

    # --- 20. Editar Sede ──────────────────────────────────────────────────
    h2(doc, "20. Editar una Sede", anchor="sec_sede_editar")

    para(doc, "Al editar una sede, se abre un dialogo con dos pestanas: 'Informacion' y 'Gastos' (LocationExpenseConfig).")

    bullet(doc, "Todos los campos del formulario de creacion estan disponibles para edicion.")
    bullet(doc, "Si se cambia la direccion de la sede, el sistema detecta una reubicacion y notifica automaticamente a los clientes con citas pendientes en esa sede.")
    bullet(doc, "La notificacion de reubicacion se envia por email mediante la Edge Function 'send-notification' con tipo 'appointment_location_update'.")

    callout(doc, "Deteccion de reubicacion",
            "Gestabiz detecta automaticamente si la direccion de la sede cambio. Si hay clientes con citas "
            "pendientes en esa sede, cada uno recibe un email informativo sobre el cambio de ubicacion.",
            color=PURPLE)

    # --- 21. Eliminar Sede ────────────────────────────────────────────────
    h2(doc, "21. Eliminar una Sede", anchor="sec_sede_eliminar")

    h3(doc, "Flujo de eliminacion")
    numbered(doc, "Hacer clic en el boton de eliminacion (icono de papelera) en la tarjeta de la sede.")
    numbered(doc, "El sistema muestra una ventana de confirmacion con el conteo de citas activas.")
    numbered(doc, "Si el administrador confirma:")
    numbered(doc, "  a) Se cancelan todas las citas activas en la sede.")
    numbered(doc, "  b) Se envia notificacion de cancelacion por email a cada cliente afectado.")
    numbered(doc, "  c) La sede se elimina de la base de datos.")
    numbered(doc, "Toast con el conteo de citas canceladas.")

    callout(doc, "Sugerencia del sistema",
            "Si la eliminacion es por reubicacion, el sistema sugiere editar la direccion en lugar de "
            "eliminar la sede: 'Si es una reubicacion, considera editar la direccion.'",
            color=ACCENT)

    # --- 22. Multimedia ───────────────────────────────────────────────────
    h2(doc, "22. Multimedia de Sedes (Imagenes y Videos)", anchor="sec_sede_multimedia")

    para(doc, "Cada sede puede tener multiples imagenes y videos asociados, gestionados desde el editor de la sede.")

    bullet(doc, "Subida de imagenes y videos mediante el componente MediaUploader.")
    bullet(doc, "Seleccion de imagen como banner de la sede con herramienta de recorte (BannerCropper).")
    bullet(doc, "Toggle de video principal para destacar un video sobre los demas.")
    bullet(doc, "Edicion de descripciones de media inline con guardado individual.")
    bullet(doc, "Upload diferido: los archivos se cargan al hacer clic en guardar, no al seleccionarlos.")

    # --- 23. Sede Principal ───────────────────────────────────────────────
    h2(doc, "23. Sede Principal — Reglas", anchor="sec_sede_principal")

    para(doc, "Solo una sede puede ser la sede principal del negocio. Las reglas son:")

    bullet(doc, "Solo una sede puede tener is_primary = true a la vez.")
    bullet(doc, "Al desmarcar la sede principal, el sistema obliga a seleccionar una nueva sede principal desde un selector.")
    bullet(doc, "Al marcar una nueva sede como principal, se desmarca automaticamente la anterior.")
    bullet(doc, "La primera sede creada se marca como principal automaticamente.")

    page_break(doc)

    # --- 24-31. Calendario ────────────────────────────────────────────────
    h2(doc, "24. Calendario de Citas", anchor="sec_citas_calendario")

    para(doc, "El calendario de citas es el modulo central del administrador. Muestra una cuadricula "
         "diaria de 24 horas con columnas para cada empleado, donde se posicionan las citas como "
         "bloques de color segun su estado.")

    h3(doc, "Estructura del calendario")
    bullet(doc, "Columna izquierda fija: horas del dia (00:00 a 23:00).")
    bullet(doc, "N columnas de empleados: una por cada profesional activo.")
    bullet(doc, "Cada columna tiene un color pastel alternado (8 colores disponibles: azul, purpura, rosa, verde, amarillo, indigo, rojo, teal).")
    bullet(doc, "Header de empleados: avatar, nombre y opcionalmente badges de servicios que ofrece.")
    bullet(doc, "Horas fuera del horario operativo: fondo semitransparente oscuro.")
    bullet(doc, "Hora de almuerzo: fondo gris, cursor 'no permitido', texto 'Almuerzo' centrado en italica.")
    bullet(doc, "Linea de hora actual: linea azul horizontal visible solo si es el dia actual.")

    callout(doc, "Zona horaria",
            "El calendario opera en zona horaria de Colombia (America/Bogota, GMT-5). Todas las horas "
            "se muestran y calculan en esta zona, independientemente de la ubicacion del usuario.",
            color=PURPLE)

    h3(doc, "Modo maximizado")
    para(doc, "El calendario puede expandirse a pantalla completa presionando el boton de maximizar. "
         "En este modo, ocupa toda la ventana del navegador. Se puede volver al tamano normal con el "
         "boton de minimizar o presionando la tecla Escape.")

    screenshot_placeholder(doc, "Calendario de citas con columnas de empleados")

    # --- 25. Navegacion ───────────────────────────────────────────────────
    h2(doc, "25. Navegacion del Calendario", anchor="sec_cal_navegacion")

    simple_table(doc, ["Boton", "Accion", "Descripcion"], [
        ["Flecha izquierda", "Dia anterior", "Retrocede un dia en el calendario"],
        ["Flecha derecha", "Dia siguiente", "Avanza un dia en el calendario"],
        ["Hoy", "Ir al dia actual", "Boton primario que regresa al dia de hoy"],
        ["Icono de calendario", "Selector de fecha", "Abre un mini-calendario flotante para elegir fecha"],
    ])

    h3(doc, "Mini-calendario flotante")
    bullet(doc, "Navegacion mensual con flechas izquierda/derecha.")
    bullet(doc, "Grid 7x6 con inicio de semana en lunes.")
    bullet(doc, "Dias del mes actual resaltados; dias de otros meses con opacidad reducida.")
    bullet(doc, "Dia actual: fondo con tinte del color primario y texto en negrita.")
    bullet(doc, "Dia seleccionado: fondo color primario con texto blanco.")
    bullet(doc, "Boton 'Ir a hoy' en la parte inferior.")
    bullet(doc, "Se cierra automaticamente al seleccionar un dia o al hacer clic fuera.")

    h3(doc, "Auto-scroll inteligente")
    para(doc, "Al cargar el calendario, se desplaza automaticamente a la posicion mas relevante:")
    numbered(doc, "Si es hoy: se centra en la hora actual.")
    numbered(doc, "Si hay horario operativo: se posiciona al inicio del horario laboral.")
    numbered(doc, "Si no hay horario definido: se mantiene en la parte superior.")

    screenshot_placeholder(doc, "Mini-calendario flotante para seleccion de fecha")

    # --- 26. Filtros ──────────────────────────────────────────────────────
    h2(doc, "26. Filtros del Calendario", anchor="sec_cal_filtros")

    para(doc, "El calendario tiene un panel colapsable de filtros con 3 selectores multi-opcion.")

    simple_table(doc, ["Filtro", "Opciones", "Default", "Comportamiento vacio"], [
        ["Estado", "Pendiente, Confirmada, Completada, En progreso, Cancelada, No asistio", "Confirmada + Pendiente", "No muestra nada"],
        ["Servicio", "Todos los servicios activos del negocio", "Todos", "Muestra todos"],
        ["Profesional", "Todos los empleados activos", "Todos", "Muestra todos"],
    ])

    h3(doc, "Caracteristicas de los filtros")
    bullet(doc, "Cada selector muestra 'N seleccionado(s)' o 'Todos'/'Ninguno'.")
    bullet(doc, "Opciones con checkboxes individuales y un boton 'Seleccionar Todos'.")
    bullet(doc, "Los filtros se guardan automaticamente en el almacenamiento local del navegador.")
    bullet(doc, "Boton 'Limpiar' para resetear todos los filtros a sus valores por defecto.")
    bullet(doc, "Panel colapsable con animacion del icono chevron.")

    h3(doc, "Filtro de sede preferida")
    para(doc, "Si el administrador tiene una sede preferida configurada, el calendario filtra automaticamente "
         "las citas de esa sede. Este filtro se aplica de forma transparente sin ser visible en el panel.")

    screenshot_placeholder(doc, "Panel de filtros del calendario")

    # --- 27. Colores ──────────────────────────────────────────────────────
    h2(doc, "27. Codificacion de Colores por Estado de Cita", anchor="sec_cal_colores")

    simple_table(doc, ["Estado", "Color tema claro", "Color tema oscuro", "Efecto adicional"], [
        ["Pendiente", "Amarillo suave con borde amarillo", "Amarillo oscuro/30 con borde oscuro", "—"],
        ["Confirmada", "Azul suave con borde azul", "Azul oscuro/30 con borde azul", "—"],
        ["Completada", "Verde suave con borde verde", "Verde oscuro/30 con borde verde", "—"],
        ["En progreso", "Purpura suave con borde purpura", "Purpura oscuro/30 con borde purpura", "—"],
        ["No asistio", "Gris atenuado", "Gris con opacidad 50%", "—"],
        ["Cancelada", "Rojo suave con borde rojo", "Rojo oscuro/20 con borde rojo", "Texto tachado, opacidad 60%"],
    ])

    para(doc, "Adicionalmente, si una cita tiene estado 'Pendiente' pero el cliente ya confirmo por email, "
         "se muestra un badge verde adicional 'Cliente confirmo'.")

    screenshot_placeholder(doc, "Calendario con citas de diferentes colores por estado")

    # --- 28. Modal Detalle ────────────────────────────────────────────────
    h2(doc, "28. Modal de Detalle de Cita", anchor="sec_cal_modal_detalle")

    para(doc, "Al hacer clic en cualquier bloque de cita en el calendario, se abre un modal con "
         "la informacion completa de la cita y botones de accion.")

    h3(doc, "Informacion mostrada")
    bullet(doc, "Nombre del cliente.")
    bullet(doc, "Nombre del servicio.")
    bullet(doc, "Horario de inicio y fin (formato Colombia).")
    bullet(doc, "Precio del servicio en COP.")
    bullet(doc, "Empleado asignado.")
    bullet(doc, "Notas de la cita (si existen).")
    bullet(doc, "Badge de estado con codificacion de colores.")
    bullet(doc, "Si la cita esta completada y pagada: desglose financiero (ingreso bruto, comision, otras deducciones, neto).")
    bullet(doc, "Campo de propina opcional (input numerico, paso 1000, minimo 0).")

    screenshot_placeholder(doc, "Modal de detalle de cita con informacion completa")

    # --- 29. Completar Cita ───────────────────────────────────────────────
    h2(doc, "29. Flujo Completar Cita (Calculo Fiscal)", anchor="sec_cal_completar")

    para(doc, "Completar una cita es una operacion importante que genera transacciones fiscales "
         "automaticamente. El sistema calcula impuestos, comisiones y registra todo en el sistema contable.")

    h3(doc, "Flujo paso a paso")
    numbered(doc, "El administrador hace clic en 'Completar' en el modal de detalle de la cita.")
    numbered(doc, "El sistema verifica que la cita tenga un servicio asociado (validacion obligatoria).")
    numbered(doc, "Se obtiene el porcentaje de comision del servicio desde la base de datos.")
    numbered(doc, "Se calculan los montos: ingreso bruto = precio del servicio, comision = bruto x comision%, neto = bruto - comision - deducciones.")
    numbered(doc, "Se actualiza la cita con estado 'completada', estado de pago 'pagado' y los montos calculados.")
    numbered(doc, "Se calculan los impuestos incluidos segun la configuracion fiscal del negocio:")
    bullet(doc, "IVA 5%: subtotal = monto / 1.05")
    bullet(doc, "IVA 19%: subtotal = monto / 1.19")
    bullet(doc, "Sin impuesto: subtotal = monto")
    numbered(doc, "Se crean hasta 3 transacciones en el sistema contable:")
    bullet(doc, "Ingreso: ingreso bruto del servicio (con desglose fiscal).")
    bullet(doc, "Egreso (si hay comision): comision del empleado.")
    bullet(doc, "Ingreso (si hay propina): propina registrada.")
    numbered(doc, "Toast con detalle: 'Ingreso bruto: X COP, Comision: Y COP, Neto: Z COP'.")

    callout(doc, "Automatizacion contable",
            "Al completar una cita, Gestabiz genera automaticamente todas las transacciones fiscales "
            "necesarias. El administrador no necesita registrar ingresos manualmente — el sistema lo hace.",
            color=ACCENT)

    screenshot_placeholder(doc, "Toast con detalle fiscal al completar cita")

    # --- 30. Acciones sobre Cita ──────────────────────────────────────────
    h2(doc, "30. Acciones sobre una Cita", anchor="sec_cal_acciones")

    para(doc, "El modal de detalle de cita ofrece 5 acciones disponibles cuando la cita NO esta "
         "completada ni cancelada:")

    simple_table(doc, ["Accion", "Color", "Efecto", "Notificacion"], [
        ["Confirmar Cita", "Primario", "Cambia estado a 'confirmada', registra fecha de confirmacion", "Toast exito"],
        ["Reenviar Email", "Contorno", "Reenvia email de confirmacion al cliente via Edge Function", "Toast exito"],
        ["Completar", "Verde", "Ver seccion 29 — genera transacciones fiscales", "Toast con detalle COP"],
        ["No Asistio", "Ambar", "Cambia estado a 'no_show', agrega nota 'Cliente no se presento'", "Toast advertencia"],
        ["Cancelar", "Rojo", "Cambia estado a 'cancelada'", "Toast exito"],
    ])

    para(doc, "Todos los botones se deshabilitan automaticamente mientras se procesa una accion "
         "(indicador de procesamiento) para evitar acciones duplicadas.")

    callout(doc, "Solo citas activas",
            "Los botones de accion solo aparecen si la cita NO tiene estado 'completada' o 'cancelada'. "
            "Para citas ya finalizadas, el modal es de solo lectura.",
            color=PURPLE)

    # --- 31. Pendientes ───────────────────────────────────────────────────
    h2(doc, "31. Citas Pendientes y En Proceso", anchor="sec_cal_pendientes")

    para(doc, "Debajo del grid del calendario, pueden aparecer dos secciones adicionales:")

    h3(doc, "En Proceso")
    bullet(doc, "Muestra citas confirmadas cuyo rango de horario incluye el momento actual.")
    bullet(doc, "Badge azul con el conteo de citas en proceso.")
    bullet(doc, "Boton 'Gestionar' que abre el modal de detalle.")

    h3(doc, "Pendientes de Confirmar")
    bullet(doc, "Muestra citas confirmadas cuya hora de fin ya paso (overdueAppointments).")
    bullet(doc, "Badge naranja con el conteo de citas vencidas.")
    bullet(doc, "Botones rapidos: 'Completada' (verde) y 'Sin Asistencia' (ambar) sin abrir el modal.")

    screenshot_placeholder(doc, "Secciones de citas pendientes y en proceso debajo del calendario")

    page_break(doc)

    # --- 32-33. Clientes ──────────────────────────────────────────────────
    h2(doc, "32. Gestion de Clientes (CRM Basico)", anchor="sec_clientes")

    para(doc, "El modulo de Clientes ofrece una vista centralizada de todos los clientes que han "
         "tenido al menos una cita no cancelada en el negocio.")

    callout(doc, "Limite del Plan Gratuito: 50 clientes visibles",
            "En el Plan Gratuito se muestran hasta 50 clientes. Los datos de todos los clientes se almacenan "
            "en la base de datos, pero solo los primeros 50 son visibles. Un banner indica cuantos estan ocultos.",
            color=DANGER)

    h3(doc, "Datos mostrados por cliente")
    bullet(doc, "Avatar (imagen de perfil o iniciales como fallback).")
    bullet(doc, "Nombre completo.")
    bullet(doc, "Email de contacto.")
    bullet(doc, "Total de citas realizadas (conteo).")
    bullet(doc, "Fecha de la ultima visita.")

    h3(doc, "Funcionalidades")
    bullet(doc, "Busqueda local por nombre o email (campo de texto con filtrado en tiempo real).")
    bullet(doc, "Grid responsive: 1/2/3 columnas segun pantalla.")
    bullet(doc, "Ordenamiento: por total de visitas descendente (los clientes mas frecuentes primero).")
    bullet(doc, "Click en un cliente abre el Modal de Perfil de Cliente (seccion 33).")

    h3(doc, "Consulta de datos (patron two-step)")
    para(doc, "Los datos de clientes se obtienen en dos pasos por limitaciones de la estructura de la base de datos:")
    numbered(doc, "Se consultan las citas no canceladas del negocio (hasta 1000 registros).")
    numbered(doc, "Se agrupan por client_id, calculando: total de visitas, visitas completadas, ultima visita.")
    numbered(doc, "Se consultan los perfiles de los clientes unicos en lote (batch fetch).")
    numbered(doc, "Se combinan los datos de citas con los perfiles para generar el listado final.")

    screenshot_placeholder(doc, "Lista de clientes con tarjetas de resumen")

    # --- 33. Perfil Cliente ───────────────────────────────────────────────
    h2(doc, "33. Modal de Perfil de Cliente", anchor="sec_cliente_perfil")

    para(doc, "El modal de perfil de cliente muestra informacion detallada de un cliente especifico, "
         "incluyendo estadisticas y un historial completo de citas.")

    h3(doc, "Tarjetas de estadisticas (3)")
    simple_table(doc, ["Estadistica", "Descripcion"], [
        ["Visitas", "Numero total de citas completadas"],
        ["Total citas", "Numero total de citas (todos los estados)"],
        ["Ingresos", "Total COP generado por citas completadas"],
    ])

    h3(doc, "Informacion de contacto")
    bullet(doc, "Email del cliente.")
    bullet(doc, "Telefono (si esta registrado).")
    bullet(doc, "Mensaje 'Sin informacion de contacto' si no hay datos.")

    h3(doc, "Pestanas")
    bullet(doc, "Informacion: fecha de primera y ultima visita.", bold_prefix="Informacion")
    bullet(doc, "Historial (N): lista cronologica de citas con servicio, fecha, hora, estado y precio.", bold_prefix="Historial")

    h3(doc, "Estados de cita con badges de color")
    simple_table(doc, ["Estado", "Etiqueta", "Estilo"], [
        ["scheduled", "Programada", "Contorno"],
        ["confirmed", "Confirmada", "Secundario"],
        ["completed", "Completada", "Default (primario)"],
        ["cancelled", "Cancelada", "Destructivo (rojo)"],
        ["in_progress", "En progreso", "Contorno"],
        ["no_show", "No asistio", "Destructivo (rojo)"],
        ["rescheduled", "Reprogramada", "Contorno"],
    ])

    screenshot_placeholder(doc, "Modal de perfil de cliente con historial de citas")

    page_break(doc)

    # --- 34-37. Facturacion ───────────────────────────────────────────────
    h2(doc, "34. Facturacion y Planes", anchor="sec_facturacion")

    para(doc, "El modulo de Facturacion es el centro de gestion de suscripciones del negocio. "
         "Desde aqui el administrador puede ver su plan actual, activar pruebas gratuitas, comparar "
         "planes y gestionar pagos.")

    # --- 35. Vista Free ───────────────────────────────────────────────────
    h2(doc, "35. Vista de Plan Gratuito", anchor="sec_billing_free")

    para(doc, "Cuando el negocio esta en el Plan Gratuito, la pantalla de facturacion muestra:")

    numbered(doc, "Banner de prueba gratuita (si es elegible para el mes gratis del Plan Basico).")
    numbered(doc, "Tarjeta con las caracteristicas incluidas en el plan gratuito (con iconos de verificacion).")
    numbered(doc, "Comparativa de 2 columnas: Plan Basico ($89,900/mes, badge 'Mas Popular') y Plan Pro ($159,900/mes).")
    numbered(doc, "Cada plan muestra hasta 6 caracteristicas principales y un boton de activacion.")

    h3(doc, "Comparativa de planes")
    simple_table(doc, ["Caracteristica", "Gratuito", "Basico ($89,900)", "Pro ($159,900)"], [
        ["Sedes", "1", "3", "10"],
        ["Empleados", "1", "6", "15"],
        ["Citas/mes", "50", "Ilimitadas", "Ilimitadas"],
        ["Clientes", "50", "Ilimitados", "Ilimitados"],
        ["Servicios", "15", "Ilimitados", "Ilimitados"],
        ["Chat", "No", "Si", "Si"],
        ["Ventas rapidas", "No", "Si", "Si"],
        ["Sistema contable", "No", "No", "Si"],
        ["Reclutamiento", "No", "No", "Si"],
        ["Recursos fisicos", "No", "No", "Si"],
    ])

    screenshot_placeholder(doc, "Pantalla de facturacion — vista Plan Gratuito")

    # --- 36. Free Trial ───────────────────────────────────────────────────
    h2(doc, "36. Prueba Gratuita (Free Trial)", anchor="sec_billing_trial")

    para(doc, "Los negocios nuevos pueden acceder a un mes gratuito del Plan Basico.")

    h3(doc, "Reglas de la prueba gratuita")
    bullet(doc, "Solo se puede activar UNA VEZ por negocio.")
    bullet(doc, "Al hacer clic en 'Activar Mes Gratis', se muestra una ventana de confirmacion con advertencia en rojo.")
    bullet(doc, "Texto de advertencia: 'Solo puedes utilizar el mes gratis una sola vez'.")
    bullet(doc, "Boton de confirmacion explicito: 'Si, activar prueba'.")
    bullet(doc, "Una vez activada, el negocio tiene acceso completo al Plan Basico durante 30 dias.")
    bullet(doc, "Al finalizar la prueba, el negocio regresa al Plan Gratuito automaticamente si no se activa un plan de pago.")

    callout(doc, "Sin compromiso",
            "La prueba gratuita no requiere metodo de pago y no genera cargos automaticos. Al finalizar, "
            "el negocio simplemente regresa al Plan Gratuito sin perder datos.",
            color=ACCENT)

    # --- 37. Pricing ──────────────────────────────────────────────────────
    h2(doc, "37. Pagina de Precios y Comparativa", anchor="sec_billing_pricing")

    para(doc, "La pagina de precios se muestra al hacer clic en 'Activar plan' desde las comparativas "
         "o al navegar a la seccion de facturacion.")

    h3(doc, "Funcionalidades")
    bullet(doc, "Toggle mensual/anual con badge 'Ahorra 17%' para el ciclo anual.")
    bullet(doc, "Precios recalculados segun el ciclo de facturacion seleccionado.")
    bullet(doc, "Campo de codigo de descuento con boton 'Aplicar'. Si es valido, muestra precio tachado + precio final.")
    bullet(doc, "3 tarjetas de planes en grid responsive.")
    bullet(doc, "El Plan Gratuito muestra 'Ya estas en el plan gratuito' (deshabilitado).")
    bullet(doc, "El Plan Pro muestra 'Proximamente' (deshabilitado).")

    h3(doc, "Flujo de activacion de plan")
    numbered(doc, "Seleccionar plan (click en boton CTA de la tarjeta).")
    numbered(doc, "El sistema crea una sesion de checkout (Stripe, PayU o MercadoPago segun configuracion).")
    numbered(doc, "Redireccion a la pasarela de pago externa.")
    numbered(doc, "Procesamiento del pago.")
    numbered(doc, "Webhook procesa la confirmacion y activa el plan.")
    numbered(doc, "El usuario regresa a Gestabiz con el plan activo.")

    h3(doc, "FAQ incluidas (4 preguntas)")
    simple_table(doc, ["Pregunta", "Respuesta resumida"], [
        ["Puedo cambiar de plan?", "Si, upgrades prorrateados, downgrades al final del periodo"],
        ["Hay prueba gratuita?", "14 dias gratis en Plan Basico, cancelar sin cargo"],
        ["Metodos de pago?", "Visa, Mastercard, Amex via Stripe"],
        ["Puedo cancelar?", "Si, sin penalidades ni costos ocultos"],
    ])

    para(doc, "Enlace de soporte al final: 'Necesitas ayuda?' con boton de email a soporte@gestabiz.com.")

    screenshot_placeholder(doc, "Pagina de precios con toggle mensual/anual")

    page_break(doc)

    # --- 38-47. Configuraciones ───────────────────────────────────────────
    h2(doc, "38. Configuraciones del Negocio", anchor="sec_config")

    para(doc, "Las configuraciones del administrador se dividen en multiples sub-secciones accesibles "
         "desde la pagina de ajustes del negocio. Incluyen informacion general, branding, notificaciones, "
         "tracking, dias cerrados y chat.")

    # --- 39. Info ─────────────────────────────────────────────────────────
    h2(doc, "39. Informacion del Negocio", anchor="sec_config_info")

    para(doc, "La seccion de informacion permite editar los datos principales del negocio.")

    h3(doc, "Campos editables")
    simple_table(doc, ["Campo", "Tipo", "Validacion"], [
        ["Nombre del negocio", "Texto", "Requerido, no vacio"],
        ["Descripcion", "Texto area", "Opcional"],
        ["Telefono", "Telefono con prefijo", "Opcional"],
        ["Email", "Email", "Opcional"],
        ["Sitio web", "URL", "Opcional"],
        ["Direccion", "Texto", "Opcional"],
        ["Ciudad", "Selector", "Opcional"],
        ["NIT / RUT", "Texto", "Opcional"],
        ["Razon social", "Texto", "Opcional"],
        ["Atender en festivos", "Toggle", "Default segun negocio"],
    ])

    para(doc, "El boton 'Guardar' esta protegido por el permiso 'settings.edit_business'. Si el usuario no "
         "tiene este permiso, el boton estara deshabilitado (modo disable del PermissionGate).")

    # --- 40. Branding ─────────────────────────────────────────────────────
    h2(doc, "40. Logo y Banner (Branding)", anchor="sec_config_branding")

    h3(doc, "Logo del negocio")
    simple_table(doc, ["Aspecto", "Detalle"], [
        ["Tamano maximo", "2 MB"],
        ["Formatos aceptados", "JPEG, PNG, WebP"],
        ["Proporcion", "Cuadrado (1:1) con herramienta de recorte"],
        ["Fallback", "Iniciales del negocio como avatar"],
    ])

    h3(doc, "Banner del negocio")
    simple_table(doc, ["Aspecto", "Detalle"], [
        ["Tamano maximo", "5 MB"],
        ["Formatos aceptados", "JPEG, PNG, WebP"],
        ["Proporcion", "Horizontal (16:9) con herramienta de recorte"],
        ["Fallback", "Texto 'Sin banner aun' centrado"],
    ])

    h3(doc, "Flujo de carga")
    numbered(doc, "Hacer clic en el area del logo o banner.")
    numbered(doc, "Seleccionar archivo desde el explorador de archivos.")
    numbered(doc, "Se valida el tipo y tamano del archivo.")
    numbered(doc, "Se abre la herramienta de recorte con la proporcion correspondiente.")
    numbered(doc, "Confirmar el recorte.")
    numbered(doc, "La imagen se sube al almacenamiento y se actualiza la URL en la base de datos.")
    numbered(doc, "Se aplica cache-bust para que la nueva imagen se muestre inmediatamente.")

    screenshot_placeholder(doc, "Seccion de branding con logo y banner")

    # --- 41. Notificaciones ───────────────────────────────────────────────
    h2(doc, "41. Configuracion de Notificaciones", anchor="sec_config_notificaciones")

    para(doc, "El administrador puede configurar los canales de notificacion, prioridades y tipos "
         "de notificaciones para su negocio.")

    h3(doc, "Canales disponibles")
    simple_table(doc, ["Canal", "Estado", "Descripcion"], [
        ["Email", "Disponible", "Notificaciones por correo electronico via Brevo"],
        ["SMS", "No disponible aun", "Switch deshabilitado con mensaje informativo"],
        ["WhatsApp", "Disponible", "Mensajes via WhatsApp Business API"],
    ])

    h3(doc, "Prioridad de canales")
    para(doc, "Los canales se ordenan por prioridad con botones de flecha arriba/abajo. "
         "El canal con prioridad #1 se usa primero; si falla, se intenta con el #2, luego el #3.")

    h3(doc, "Sistema de respaldo (Fallback)")
    para(doc, "Toggle para activar el intento automatico del siguiente canal si el primero falla. "
         "Recomendado mantener activado para maximizar la entrega de notificaciones.")

    h3(doc, "Tiempos de recordatorio")
    para(doc, "Lista configurada de tiempos antes de la cita para enviar recordatorios. "
         "Ejemplo: 24 horas y 2 horas antes de la cita. Formato legible con badge de minutos.")

    h3(doc, "Tipos de notificacion (6)")
    simple_table(doc, ["Tipo", "Descripcion"], [
        ["Recordatorio de cita", "Aviso antes de una cita programada"],
        ["Confirmacion de cita", "Confirmacion al crear una cita"],
        ["Cancelacion de cita", "Aviso al cancelar una cita"],
        ["Cita reprogramada", "Aviso al cambiar horario de una cita"],
        ["Solicitud de empleado", "Cuando un empleado solicita unirse al negocio"],
        ["Resumen diario", "Resumen de la actividad del dia anterior"],
    ])

    h3(doc, "Avanzado")
    bullet(doc, "Maximo de reintentos: configurable de 1 a 5 intentos por notificacion.")

    h3(doc, "Defaults automaticos")
    para(doc, "Si es la primera vez que se accede a la configuracion, el sistema crea automaticamente "
         "los defaults: Email y WhatsApp habilitados, SMS deshabilitado, prioridad WhatsApp > Email > SMS, "
         "recordatorios a las 24h y 2h antes, zona horaria America/Bogota, ventana de envio 08:00-22:00.")

    screenshot_placeholder(doc, "Configuracion de canales de notificacion")

    # --- 42. Tracking ─────────────────────────────────────────────────────
    h2(doc, "42. Tracking de Notificaciones", anchor="sec_config_tracking")

    para(doc, "La seccion de tracking permite monitorear el estado de todas las notificaciones enviadas "
         "por el negocio, con estadisticas, graficos y filtros avanzados.")

    h3(doc, "Estadisticas principales (4 tarjetas)")
    simple_table(doc, ["Tarjeta", "Descripcion"], [
        ["Total", "Numero total de notificaciones enviadas"],
        ["Exitosas", "Notificaciones entregadas correctamente"],
        ["Fallidas", "Notificaciones que no pudieron entregarse"],
        ["Tasa de exito", "Porcentaje de entrega exitosa"],
    ])

    h3(doc, "Graficos (3)")
    bullet(doc, "Grafico circular por canal: distribucion email/sms/whatsapp.", bold_prefix="Por Canal")
    bullet(doc, "Grafico circular por estado: enviadas/fallidas/pendientes.", bold_prefix="Por Estado")
    bullet(doc, "Grafico de barras: top 5 tipos de notificacion mas frecuentes.", bold_prefix="Top 5 Tipos")

    h3(doc, "Filtros (6)")
    simple_table(doc, ["Filtro", "Opciones"], [
        ["Canal", "Todos, Email, SMS, WhatsApp"],
        ["Estado", "Todos, Enviada, Fallida, Pendiente"],
        ["Tipo", "Todos + 17 tipos de notificacion"],
        ["Fecha desde", "Selector de fecha"],
        ["Fecha hasta", "Selector de fecha"],
        ["Buscar", "Por email o telefono del destinatario"],
    ])

    h3(doc, "Exportacion CSV")
    para(doc, "Boton 'Exportar' que descarga un archivo CSV con todas las notificaciones filtradas. "
         "Columnas: Fecha, Tipo, Canal, Destinatario, Estado, Error, Reintentos, ID Externo.")

    h3(doc, "Tabla de detalle")
    para(doc, "Tabla con las ultimas 500 notificaciones, ordenadas por fecha descendente. "
         "Columnas: Fecha (formato dd/mm/yyyy HH:mm), Tipo, Canal, Destinatario, Estado (con badge), Error.")

    screenshot_placeholder(doc, "Dashboard de tracking de notificaciones con graficos")

    # --- 43. Dias Cerrados ────────────────────────────────────────────────
    h2(doc, "43. Dias Cerrados", anchor="sec_config_calendario")

    para(doc, "El administrador puede definir dias en los que el negocio estara cerrado. "
         "Los clientes no podran reservar citas en dias cerrados.")

    h3(doc, "Agregar dia cerrado")
    numbered(doc, "Seleccionar fecha en el calendario (dias pasados y ya cerrados estan deshabilitados).")
    numbered(doc, "Seleccionar alcance: 'Todas las sedes' o una sede especifica.")
    numbered(doc, "Opcionalmente escribir una razon (maximo 120 caracteres).")
    numbered(doc, "Hacer clic en 'Agregar dia cerrado'.")

    h3(doc, "Tabla de dias cerrados")
    simple_table(doc, ["Columna", "Descripcion"], [
        ["Fecha", "Fecha en formato espanol"],
        ["Sede", "'Todas' o nombre de la sede especifica"],
        ["Razon", "Texto libre o vacio"],
        ["Acciones", "Boton de papelera para eliminar el dia cerrado"],
    ])

    para(doc, "Solo se muestran dias cerrados futuros, ordenados por fecha ascendente.")

    screenshot_placeholder(doc, "Configuracion de dias cerrados con calendario")

    # --- 44. Chat ─────────────────────────────────────────────────────────
    h2(doc, "44. Configuracion del Chat", anchor="sec_config_chat")

    para(doc, "El administrador puede configurar el sistema de chat del negocio.")

    bullet(doc, "Toggle principal: 'Permitir chat con profesionales' — activa o desactiva el chat entre clientes y empleados del negocio.")
    bullet(doc, "Administrador de chat por sede: permite asignar un empleado especifico como administrador del chat para cada sede activa.")
    bullet(doc, "Candidatos: solo empleados con nivel de jerarquia 0 o 1 (managers).")
    bullet(doc, "Default: 'Usar propietario (por defecto)'.")

    callout(doc, "Nota sobre el Plan Gratuito",
            "Aunque el chat se puede configurar desde las opciones, el modulo de chat como tal no "
            "esta incluido en el Plan Gratuito. Para habilitar la mensajeria en tiempo real con clientes, "
            "se requiere al menos el Plan Basico.",
            color=PURPLE)

    # --- 45. General ──────────────────────────────────────────────────────
    h2(doc, "45. Configuracion General (Tema e Idioma)", anchor="sec_config_general")

    bullet(doc, "Tema: selector entre Claro, Oscuro y Sistema (sigue la preferencia del navegador).", bold_prefix="Tema")
    bullet(doc, "Idioma: actualmente solo espanol esta disponible. El selector muestra las opciones pero esta deshabilitado.", bold_prefix="Idioma")

    # --- 46. Perfil Personal ──────────────────────────────────────────────
    h2(doc, "46. Edicion de Perfil Personal", anchor="sec_config_perfil")

    para(doc, "La pestana 'Perfil' en las configuraciones permite al administrador editar su "
         "informacion personal (no la del negocio).")

    bullet(doc, "Nombre completo.")
    bullet(doc, "Foto de perfil (upload a bucket avatars).")
    bullet(doc, "Telefono personal.")
    bullet(doc, "Tipo de documento y numero.")
    bullet(doc, "Genero.")
    bullet(doc, "Fecha de nacimiento.")

    # --- 47. Danger Zone ──────────────────────────────────────────────────
    h2(doc, "47. Zona de Peligro — Desactivar Cuenta", anchor="sec_config_danger")

    para(doc, "La Zona de Peligro permite al usuario desactivar permanentemente su cuenta. "
         "Es un proceso de dos pasos con multiples confirmaciones para evitar errores.")

    h3(doc, "Paso 1 — Consentimiento")
    numbered(doc, "Se muestra advertencia en rojo: 'Esta accion desactivara tu cuenta'.")
    numbered(doc, "El usuario debe ingresar su email de confirmacion (debe coincidir exactamente).")
    numbered(doc, "El usuario debe marcar un checkbox: 'Entiendo que esta accion es irreversible'.")
    numbered(doc, "Boton 'Continuar con la desactivacion' — deshabilitado hasta completar ambos requisitos.")

    h3(doc, "Paso 2 — Confirmacion final")
    numbered(doc, "El usuario debe escribir exactamente: DESACTIVAR CUENTA (sensible a mayusculas/minusculas).")
    numbered(doc, "Boton 'Desactivar mi cuenta permanentemente' (rojo) — deshabilitado hasta que el texto coincida.")

    h3(doc, "Ejecucion")
    numbered(doc, "Se llama a la funcion RPC 'deactivate_user_account'.")
    numbered(doc, "Se cierra la sesion del usuario.")
    numbered(doc, "Se limpia todo el almacenamiento local.")
    numbered(doc, "Se redirige a la pagina de login despues de 2 segundos.")

    callout(doc, "Irreversible",
            "Una vez desactivada la cuenta, no es posible recuperarla. Todos los datos del usuario "
            "seran inaccesibles. Si el usuario administra negocios, estos quedaran sin administrador.",
            color=DANGER)

    screenshot_placeholder(doc, "Zona de peligro — proceso de desactivacion de cuenta")

    page_break(doc)

    # --- 48-51. Limites y Permisos ────────────────────────────────────────
    h2(doc, "48. Limites del Plan Gratuito — Resumen", anchor="sec_limites_plan")

    simple_table(doc, ["Recurso", "Limite", "Comportamiento al exceder"], [
        ["Sedes", "1", "Boton 'Agregar Sede' deshabilitado con tooltip. Banner informativo."],
        ["Empleados", "1", "Excedentes no mostrados. Banner informativo."],
        ["Citas/mes", "50", "Controlado por quotaInfo. Banner informativo."],
        ["Clientes visibles", "50", "Solo 50 visibles. Todos almacenados en BD. Banner informativo."],
        ["Servicios", "15", "Solo 15 visibles. Excedentes ocultos. Banner informativo."],
    ])

    callout(doc, "Los datos siempre se conservan",
            "Aunque el Plan Gratuito limita la visualizacion, todos los datos se almacenan en la base de datos "
            "sin restriccion. Al actualizar a un plan superior, toda la informacion estara inmediatamente disponible.",
            color=ACCENT)

    # --- 49. Modulos Bloqueados ───────────────────────────────────────────
    h2(doc, "49. Modulos Bloqueados y Pantalla de Upgrade", anchor="sec_modulos_bloqueados")

    para(doc, "Cuando el administrador intenta acceder a un modulo que no esta incluido en su plan, "
         "el sistema muestra una pantalla de upgrade en lugar del contenido del modulo.")

    h3(doc, "Pantalla de upgrade (PlanGate)")
    bullet(doc, "Icono de candado grande con efecto sparkle.")
    bullet(doc, "Texto: 'Modulo disponible en Plan {Basico/Pro}'.")
    bullet(doc, "Boton CTA: 'Ver planes y precios' que navega a /app/admin/billing.")

    h3(doc, "Modos de bloqueo")
    simple_table(doc, ["Modo", "Comportamiento"], [
        ["upgrade (default)", "Pantalla completa con candado, nombre del modulo y boton CTA"],
        ["hide", "No renderiza nada — el modulo desaparece completamente"],
        ["disable", "Contenido visible pero opaco (40%), sin interactividad, overlay con texto"],
    ])

    para(doc, "En el sidebar, los modulos bloqueados muestran un icono de candado y un badge con "
         "el nombre del plan requerido (ej: 'Basico', 'Pro').")

    # --- 50. Banner Limite ────────────────────────────────────────────────
    h2(doc, "50. Banner de Limite Alcanzado", anchor="sec_plan_limit_banner")

    para(doc, "Cuando la cantidad de registros excede el limite del plan, aparece un banner informativo "
         "al final de la lista del modulo correspondiente.")

    bullet(doc, "Fondo ambar suave con borde ambar.")
    bullet(doc, "Icono de candado + texto: '{N} {recurso} adicionales disponibles al actualizar'.")
    bullet(doc, "Sub-texto: 'Actualiza al Plan {Basico/Pro} para ver todos los registros'.")
    bullet(doc, "Boton 'Actualizar plan' que navega a la seccion de facturacion.")

    screenshot_placeholder(doc, "Banner de limite alcanzado en la lista de servicios")

    # --- 51. Permisos ─────────────────────────────────────────────────────
    h2(doc, "51. Permisos del Administrador (PermissionGate)", anchor="sec_permisos_admin")

    para(doc, "El sistema de permisos granulares protege cada boton de accion en la interfaz. "
         "Como propietario (owner) del negocio, el administrador tiene bypass total — todos los "
         "permisos son concedidos automaticamente sin consulta a la base de datos.")

    h3(doc, "Categorias de permisos relevantes para el admin")
    simple_table(doc, ["Categoria", "Permisos incluidos"], [
        ["services.*", "create, edit, delete, view"],
        ["locations.*", "create, edit, delete, view"],
        ["appointments.*", "create, edit, delete, cancel"],
        ["settings.*", "edit, edit_business"],
        ["notifications.*", "manage"],
        ["billing.*", "manage, view"],
    ])

    para(doc, "Si el negocio tiene administradores adicionales (no owners), estos dependen de los "
         "permisos asignados individualmente. Los botones protegidos por PermissionGate se ocultan "
         "(mode=hide) o se deshabilitan (mode=disable) segun el permiso y el modo configurado.")

    page_break(doc)

    # --- 52-55. Miscelaneos ───────────────────────────────────────────────
    h2(doc, "52. Gestion de Multiples Negocios", anchor="sec_multi_negocio")

    para(doc, "Un usuario puede ser propietario de varios negocios. El sistema permite cambiar entre "
         "ellos de forma fluida.")

    h3(doc, "Dropdown de negocios")
    bullet(doc, "Ubicado en el encabezado del AdminDashboard.")
    bullet(doc, "Muestra el nombre del negocio activo.")
    bullet(doc, "Lista desplegable con todos los negocios del usuario.")
    bullet(doc, "Opcion 'Crear nuevo negocio' al final de la lista.")

    h3(doc, "Flujo de cambio")
    numbered(doc, "Hacer clic en el nombre del negocio en el encabezado.")
    numbered(doc, "Seleccionar otro negocio de la lista.")
    numbered(doc, "El sistema guarda la seleccion en localStorage y recarga todos los datos del dashboard.")
    numbered(doc, "Todos los modulos se actualizan con los datos del nuevo negocio.")

    # --- 53. Cambio de Rol ────────────────────────────────────────────────
    h2(doc, "53. Cambio de Rol (Admin / Empleado / Cliente)", anchor="sec_cambio_rol")

    para(doc, "El usuario puede cambiar entre sus roles disponibles en cualquier momento.")

    bullet(doc, "Selector de roles en el encabezado (junto al avatar).")
    bullet(doc, "Roles disponibles se calculan automaticamente segun la relacion con negocios.")
    bullet(doc, "Al cambiar de rol, se recarga el dashboard correspondiente (AdminDashboard, EmployeeDashboard o ClientDashboard).")
    bullet(doc, "El rol activo se guarda en localStorage y se restaura al iniciar sesion.")

    # --- 54. Estados ──────────────────────────────────────────────────────
    h2(doc, "54. Estados de Carga, Error y Vacios", anchor="sec_estados_carga")

    para(doc, "Cada modulo del administrador maneja los siguientes estados:")

    simple_table(doc, ["Estado", "Indicador visual", "Descripcion"], [
        ["Carga", "Skeleton / Spinner", "Placeholders animados o indicador circular mientras se cargan datos"],
        ["Error", "Toast rojo / Sentry", "Mensaje de error con registro en Sentry para monitoreo"],
        ["Vacio sin datos", "Icono + mensaje + CTA", "Icono grande, texto descriptivo y boton para crear el primer registro"],
        ["Vacio con filtros", "Texto sugerencia", "Sugiere cambiar filtros o activar opciones adicionales"],
        ["Error de permisos", "Mensaje 403", "Deteccion especifica de errores RLS/permisos con mensaje en espanol"],
    ])

    # --- 55. Glosario ─────────────────────────────────────────────────────
    h2(doc, "55. Glosario de Terminos", anchor="sec_glosario_admin")

    simple_table(doc, ["Termino", "Definicion"], [
        ["Owner / Propietario", "Usuario que creo el negocio. Tiene bypass total de permisos."],
        ["Admin", "Usuario con rol administrativo asignado. Sus permisos dependen de las plantillas."],
        ["Sede", "Ubicacion fisica del negocio donde se prestan los servicios."],
        ["Sede Principal", "La sede por defecto del negocio. Solo una puede serlo."],
        ["Sede Preferida", "La sede que el admin selecciona como filtro predeterminado."],
        ["Setup Checklist", "Asistente de configuracion que guia al admin para completar su negocio."],
        ["PlanGate", "Componente que bloquea modulos no incluidos en el plan actual."],
        ["PlanLimitBanner", "Banner que informa cuantos registros estan ocultos por limite del plan."],
        ["PermissionGate", "Componente que protege botones de accion segun permisos del usuario."],
        ["Soft-delete", "Eliminacion logica (is_active = false) que preserva el historial."],
        ["Two-step query", "Patron de consulta en dos pasos para obtener datos relacionados."],
        ["Cache-bust", "Tecnica para forzar al navegador a cargar una imagen actualizada."],
        ["Slug", "URL amigable del negocio (ej: salon-belleza-medellin)."],
        ["COP", "Pesos colombianos — moneda predeterminada del sistema."],
        ["Edge Function", "Funcion serverless en Supabase para logica del lado del servidor."],
        ["Webhook", "Notificacion automatica de la pasarela de pago al confirmar transaccion."],
        ["RPC", "Remote Procedure Call — funcion SQL invocada desde el cliente."],
        ["Free Trial", "Periodo de prueba gratuito de 30 dias del Plan Basico."],
    ])

    page_break(doc)

    # ── Cierre ────────────────────────────────────────────────────────────
    para(doc, "Fin de la Parte 3 — Administrador Plan Gratuito.", size=12, italic=True,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Continua en la Parte 4: Administrador Plan Basico.", size=12, italic=True,
         color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ╔═══════════════════════════════════════════════════════════════════════════╗
# ║  PROPUESTA DE VALOR — PARTE 3: ADMINISTRADOR PLAN GRATUITO             ║
# ╚═══════════════════════════════════════════════════════════════════════════╝

def build_proposal_part3() -> Document:
    doc = setup_document(
        "Propuesta de Valor — Gestabiz",
        "La plataforma todo-en-uno para negocios de servicios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE", anchor="toc_prop_p3")
    toc_items = [
        ("prop_problema", "1. El Problema: Empezar un negocio sin herramientas"),
        ("prop_solucion", "2. La Solucion: Gestabiz Plan Gratuito"),
        ("prop_funcionalidades", "3. Funcionalidades Incluidas Sin Costo"),
        ("prop_limites", "4. Limites Transparentes — Sin Sorpresas"),
        ("prop_onboarding", "5. De 0 a Operando en 15 Minutos"),
        ("prop_citas", "6. Calendario de Citas Profesional"),
        ("prop_crm", "7. CRM Basico — Conoce a tus Clientes"),
        ("prop_seo", "8. Presencia Digital: SEO, QR y Perfil Publico"),
        ("prop_notificaciones", "9. Notificaciones Automaticas que Retienen Clientes"),
        ("prop_contable", "10. Automatizacion Contable desde la Primera Cita"),
        ("prop_comparativa", "11. Comparativa: Gestabiz vs Competencia"),
        ("prop_crecimiento", "12. Ruta de Crecimiento Transparente"),
        ("prop_trial", "13. Prueba Sin Riesgo: 30 Dias del Plan Basico"),
        ("prop_escenarios", "14. Escenarios Reales"),
        ("prop_faq", "15. Preguntas Frecuentes"),
        ("prop_cta", "16. Siguiente Paso"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # CONTENIDO DE LA PROPUESTA
    # ══════════════════════════════════════════════════════════════════════

    # --- 1. Problema ──────────────────────────────────────────────────────
    h1(doc, "1. El Problema: Empezar un Negocio Sin Herramientas", anchor="prop_problema")

    para(doc, "Cada dia, miles de emprendedores en Colombia abren negocios de servicios: salones de belleza, "
         "barberias, consultorios medicos, centros de estetica, estudios de yoga. Y todos enfrentan el mismo "
         "problema: gestionar citas, clientes y servicios con herramientas improvisadas.")

    bullet(doc, "Agendas de papel que se pierden o se mojan.")
    bullet(doc, "Mensajes de WhatsApp que se mezclan con conversaciones personales.")
    bullet(doc, "Hojas de Excel que nadie actualiza despues de la primera semana.")
    bullet(doc, "Aplicaciones gratuitas que cobran por la cuarta funcionalidad.")
    bullet(doc, "Ninguna visibilidad online — los clientes no encuentran el negocio en Google.")

    callout(doc, "El costo real de no tener sistema",
            "Un negocio sin sistema de gestion pierde en promedio 3-5 citas por semana por errores de agenda, "
            "olvidos o doble agendamiento. Con una cita promedio de $50,000 COP, eso equivale a "
            "$600,000 - $1,000,000 COP mensuales en ingresos perdidos.",
            color=DANGER)

    screenshot_placeholder(doc, "Emprendedor gestionando citas en papel vs Gestabiz")

    # --- 2. Solucion ──────────────────────────────────────────────────────
    h1(doc, "2. La Solucion: Gestabiz Plan Gratuito", anchor="prop_solucion")

    para(doc, "Gestabiz elimina todos estos problemas desde el primer dia, sin costo. El Plan Gratuito "
         "no es una demo limitada ni un periodo de prueba — es un plan real con herramientas reales "
         "que un emprendedor puede usar indefinidamente para gestionar su negocio.")

    h3(doc, "6 Pilares del Plan Gratuito")
    numbered(doc, "Calendario de citas profesional — visual, con codificacion de colores y gestion completa.")
    numbered(doc, "Gestion de servicios — hasta 15 servicios con precios, duraciones e imagenes.")
    numbered(doc, "Sede del negocio — una ubicacion completa con horarios, multimedia y direccion.")
    numbered(doc, "Catalogo de clientes — CRM basico con historial de visitas y estadisticas.")
    numbered(doc, "Perfil publico SEO — pagina indexada en Google con URL amigable y codigo QR.")
    numbered(doc, "Automatizacion contable — calculo automatico de impuestos al completar citas.")

    callout(doc, "Para siempre, no por 14 dias",
            "A diferencia de Calendly, Booksy y Fresha que ofrecen funciones basicas limitadas en tiempo, "
            "el Plan Gratuito de Gestabiz es permanente. Sin fecha de vencimiento, sin tarjeta de credito, "
            "sin sorpresas.",
            color=ACCENT)

    # --- 3. Funcionalidades ───────────────────────────────────────────────
    h1(doc, "3. Funcionalidades Incluidas Sin Costo", anchor="prop_funcionalidades")

    simple_table(doc, ["Funcionalidad", "Que resuelve", "Valor estimado"], [
        ["Calendario de citas", "Doble agendamiento, olvidos, desorden", "$30,000/mes en apps"],
        ["15 servicios con precios", "Catalogo profesional, sin hojas de calculo", "$15,000/mes"],
        ["1 sede completa", "Horarios, direccion, multimedia", "Incluido"],
        ["50 clientes con historial", "Conocer frecuencia y preferencias", "$20,000/mes en CRM"],
        ["Perfil publico SEO", "Visibilidad en Google sin pagar ads", "$50,000/mes en SEO"],
        ["Codigo QR descargable", "Material de marketing instantaneo", "$10,000 por diseno"],
        ["50 citas mensuales", "Capacidad para negocio en crecimiento", "Incluido"],
        ["Recordatorios por email", "Reducir no-shows en 40%", "$25,000/mes"],
        ["Calculo fiscal automatico", "IVA, comisiones, propinas sin calculadora", "Invaluable"],
        ["Dashboard de estadisticas", "Saber como va el negocio en tiempo real", "Incluido"],
        ["Tema claro/oscuro", "Comodidad visual", "Incluido"],
    ])

    callout(doc, "Valor total estimado: $150,000+ COP/mes",
            "Si tuvieras que contratar estas funcionalidades por separado con diferentes herramientas, "
            "pagarias mas de $150,000 COP mensuales. Con Gestabiz: $0 COP.",
            color=PURPLE)

    # --- 4. Limites ───────────────────────────────────────────────────────
    h1(doc, "4. Limites Transparentes — Sin Sorpresas", anchor="prop_limites")

    para(doc, "Creemos en la transparencia total. Estos son los limites del Plan Gratuito y por que "
         "son mas que suficientes para empezar:")

    simple_table(doc, ["Recurso", "Limite", "Por que es suficiente"], [
        ["Sedes", "1", "La mayoria de negocios nuevos operan desde una sola ubicacion"],
        ["Empleados", "1", "El emprendedor es su propio profesional al inicio"],
        ["Citas/mes", "50", "50 citas = ~2.5 citas/dia, ideal para un profesional independiente"],
        ["Clientes", "50", "Los primeros 50 clientes son los mas valiosos — los fieles"],
        ["Servicios", "15", "15 servicios cubren la oferta completa de cualquier negocio de servicios"],
    ])

    callout(doc, "Tus datos nunca se pierden",
            "Aunque el plan limita la visualizacion, todos los datos se almacenan sin restriccion en la base de datos. "
            "Al actualizar el plan, toda tu informacion historica estara disponible inmediatamente.",
            color=ACCENT)

    # --- 5. Onboarding ────────────────────────────────────────────────────
    h1(doc, "5. De 0 a Operando en 15 Minutos", anchor="prop_onboarding")

    para(doc, "Gestabiz fue disenada para que cualquier emprendedor, sin conocimientos tecnicos, "
         "pueda configurar su negocio y empezar a recibir citas en menos de 15 minutos.")

    h3(doc, "El proceso paso a paso")
    numbered(doc, "Registrarse con email o Google (2 minutos).")
    numbered(doc, "Crear el negocio: nombre, categoria y datos basicos (3 minutos).")
    numbered(doc, "Seguir el asistente de configuracion: crear sede, agregar servicios (5 minutos).")
    numbered(doc, "Subir logo y completar el perfil (3 minutos).")
    numbered(doc, "Compartir el enlace publico o el QR con los primeros clientes (2 minutos).")

    callout(doc, "Asistente inteligente",
            "El Setup Checklist guia al administrador paso a paso con una barra de progreso visual. "
            "Cada paso completado acerca al negocio al estado 'Publico' — visible para todos los clientes.",
            color=PURPLE)

    screenshot_placeholder(doc, "Asistente de configuracion mostrando progreso 4 de 6 pasos")

    # --- 6. Calendario ────────────────────────────────────────────────────
    h1(doc, "6. Calendario de Citas Profesional", anchor="prop_citas")

    para(doc, "El calendario de citas de Gestabiz es una herramienta visual poderosa que reemplaza "
         "agendas fisicas, hojas de calculo y mensajes de WhatsApp.")

    h3(doc, "Ventajas competitivas del calendario")
    bullet(doc, "Vista diaria con columnas por empleado y codificacion de colores por estado.")
    bullet(doc, "Linea de hora actual visible cuando es el dia de hoy.")
    bullet(doc, "Mini-calendario flotante para navegar rapidamente a cualquier fecha.")
    bullet(doc, "Filtros avanzados persistentes: por estado, servicio y profesional.")
    bullet(doc, "Modo pantalla completa para ver todas las citas sin distracciones.")
    bullet(doc, "Deteccion automatica de citas vencidas y en proceso.")
    bullet(doc, "Acciones rapidas sin salir del calendario: confirmar, completar, cancelar, marcar no-show.")
    bullet(doc, "Calculo fiscal automatico al completar una cita — ingreso, comision, IVA y propina.")

    callout(doc, "Mas que un calendario: es tu centro de operaciones",
            "Con Gestabiz no solo ves tus citas — las gestionas. Confirma, completa, cancela y registra "
            "ingresos todo desde una sola vista. Ni Calendly ni Google Calendar ofrecen esto.",
            color=ACCENT)

    screenshot_placeholder(doc, "Calendario de citas con citas de multiples colores")

    # --- 7. CRM ───────────────────────────────────────────────────────────
    h1(doc, "7. CRM Basico — Conoce a tus Clientes", anchor="prop_crm")

    para(doc, "Con el catalogo de clientes integrado, el administrador conoce a cada cliente sin "
         "necesidad de herramientas externas.")

    bullet(doc, "Automatico: cada cliente que reserva una cita se agrega al catalogo sin accion manual.")
    bullet(doc, "Historial completo: total de visitas, visitas completadas, fecha de primera y ultima visita.")
    bullet(doc, "Ingresos por cliente: cuanto ha generado cada cliente en COP.")
    bullet(doc, "Busqueda instantanea por nombre o email.")
    bullet(doc, "Perfil detallado con historial cronologico de todas las citas.")

    callout(doc, "La base para fidelizar",
            "Conocer la frecuencia de visita de tus clientes te permite identificar quienes son los mas fieles, "
            "quienes dejaron de venir y quienes son nuevos. Esta informacion es oro para cualquier negocio.",
            color=PURPLE)

    screenshot_placeholder(doc, "Grid de clientes con estadisticas de visitas")

    # --- 8. SEO y QR ──────────────────────────────────────────────────────
    h1(doc, "8. Presencia Digital: SEO, QR y Perfil Publico", anchor="prop_seo")

    para(doc, "Uno de los mayores diferenciadores de Gestabiz es que cada negocio obtiene una pagina "
         "web profesional optimizada para Google, sin necesidad de pagar por diseno web ni hosting.")

    h3(doc, "Lo que obtiene cada negocio")
    bullet(doc, "URL amigable: gestabiz.com/negocio/{nombre-del-negocio}")
    bullet(doc, "Meta tags dinamicos: titulo, descripcion, imagen para Google.")
    bullet(doc, "Open Graph y Twitter Card: previsualizacion atractiva al compartir en redes sociales.")
    bullet(doc, "JSON-LD structured data: datos que Google entiende para mostrar en resultados enriquecidos.")
    bullet(doc, "Sitemap.xml automatico: indexacion rapida en buscadores.")
    bullet(doc, "Codigo QR descargable: 2 modos — perfil publico o reserva directa.")

    callout(doc, "El equivalente a una pagina web profesional",
            "Tener un perfil publico indexado en Google con meta tags, Open Graph y datos estructurados "
            "normalmente cuesta $200,000 - $500,000 COP en diseno web + hosting mensual. "
            "Con Gestabiz: incluido gratis.",
            color=ACCENT)

    screenshot_placeholder(doc, "Perfil publico del negocio con servicios y resenas")

    # --- 9. Notificaciones ────────────────────────────────────────────────
    h1(doc, "9. Notificaciones Automaticas que Retienen Clientes", anchor="prop_notificaciones")

    para(doc, "Las notificaciones automaticas son una de las herramientas mas poderosas para reducir "
         "ausencias (no-shows) y mejorar la experiencia del cliente.")

    h3(doc, "Incluido en el Plan Gratuito")
    bullet(doc, "Recordatorios por email antes de cada cita (configurable: 24h, 2h antes).")
    bullet(doc, "Confirmacion automatica al crear una cita.")
    bullet(doc, "Notificacion al cancelar o reprogramar una cita.")
    bullet(doc, "Dashboard de tracking para monitorear entregas, fallos y tasas de exito.")

    callout(doc, "Reduce no-shows en un 40%",
            "Segun estudios del sector, los recordatorios automaticos reducen las ausencias en un 40%. "
            "Con una cita promedio de $50,000 COP, eso puede significar $800,000+ COP al mes de ingresos "
            "que no se pierden.",
            color=ACCENT)

    # --- 10. Contable ─────────────────────────────────────────────────────
    h1(doc, "10. Automatizacion Contable desde la Primera Cita", anchor="prop_contable")

    para(doc, "Cada vez que el administrador completa una cita en el calendario, Gestabiz genera "
         "automaticamente las transacciones contables correspondientes.")

    bullet(doc, "Ingreso bruto del servicio con desglose fiscal (IVA 5%, 19% o sin impuesto).")
    bullet(doc, "Comision del empleado (si aplica) registrada como egreso.")
    bullet(doc, "Propinas registradas como ingreso adicional.")
    bullet(doc, "Todo en pesos colombianos con formato de miles.")

    callout(doc, "Sin calculadora, sin Excel, sin contador (para lo basico)",
            "Para un emprendedor individual, el calculo automatico de impuestos y comisiones desde el "
            "calendario elimina la necesidad de llevar cuentas en Excel. Gestabiz lo hace en cada cita.",
            color=PURPLE)

    page_break(doc)

    # --- 11. Comparativa ──────────────────────────────────────────────────
    h1(doc, "11. Comparativa: Gestabiz vs Competencia", anchor="prop_comparativa")

    para(doc, "Comparamos el Plan Gratuito de Gestabiz con los planes gratuitos de los principales "
         "competidores del mercado.")

    simple_table(doc, ["Funcionalidad", "Gestabiz Free", "Calendly Free", "Booksy Free", "Fresha Free"], [
        ["Citas mensuales", "50", "Sin limite (1 tipo)", "Limitadas", "Ilimitadas"],
        ["Servicios", "15", "1 tipo de evento", "Limitados", "Ilimitados"],
        ["Sede/ubicacion", "1 completa", "No aplica", "1", "1"],
        ["CRM de clientes", "50 clientes + historial", "No incluido", "Basico", "Basico"],
        ["Perfil publico SEO", "Si, con JSON-LD", "Pagina de booking", "Si", "Si"],
        ["Codigo QR", "Si, 2 modos", "No", "No", "No"],
        ["Recordatorios email", "Si, configurables", "Si", "Si (limitados)", "Si"],
        ["Calculo fiscal", "Si, automatico", "No", "No", "No"],
        ["Dashboard estadisticas", "9 metricas", "Basico", "Basico", "Basico"],
        ["Notif. tracking", "Si, con graficos", "No", "No", "No"],
        ["Tema claro/oscuro", "Si", "No", "No", "No"],
        ["Colombia (COP)", "Nativo", "Solo USD", "Parcial", "Si"],
    ])

    callout(doc, "Ventaja Gestabiz",
            "Mientras Calendly se limita a 1 tipo de evento y Booksy restringe funciones clave, "
            "Gestabiz ofrece 15 servicios, CRM con historial, perfil SEO, QR, calculo fiscal automatico "
            "y dashboard completo. Todo gratis. Todo en espanol. Todo en pesos colombianos.",
            color=ACCENT)

    screenshot_placeholder(doc, "Tabla comparativa Gestabiz vs competidores")

    # --- 12. Crecimiento ──────────────────────────────────────────────────
    h1(doc, "12. Ruta de Crecimiento Transparente", anchor="prop_crecimiento")

    para(doc, "Gestabiz esta disenado para crecer con tu negocio. Los limites del Plan Gratuito estan "
         "pensados para que, cuando necesites mas, el upgrade sea natural — no forzado.")

    h3(doc, "Cuando es momento de actualizar")
    simple_table(doc, ["Senal", "Solucion", "Plan requerido"], [
        ["Mas de 50 citas/mes", "Citas ilimitadas", "Basico"],
        ["Contratar primer empleado", "Gestion de hasta 6 empleados", "Basico"],
        ["Abrir segunda sede", "Hasta 3 sedes", "Basico"],
        ["Necesitar reportes", "Dashboard de reportes completo", "Basico"],
        ["Ventas walk-in frecuentes", "Modulo de ventas rapidas", "Basico"],
        ["Controlar gastos", "Sistema contable completo", "Pro"],
        ["Publicar vacantes", "Portal de reclutamiento", "Pro"],
        ["Reservar recursos fisicos", "Salas, mesas, canchas", "Pro"],
    ])

    callout(doc, "Sin presion, sin urgencia",
            "Gestabiz nunca bloqueara tus datos ni degradara tu experiencia para forzar un upgrade. "
            "Cuando los limites del Plan Gratuito se sientan cortos, es porque tu negocio esta creciendo "
            "— y eso es exactamente lo que queremos celebrar.",
            color=PURPLE)

    # --- 13. Trial ────────────────────────────────────────────────────────
    h1(doc, "13. Prueba Sin Riesgo: 30 Dias del Plan Basico", anchor="prop_trial")

    para(doc, "Ademas del Plan Gratuito permanente, Gestabiz ofrece una prueba gratuita de 30 dias "
         "del Plan Basico para que pruebes todo sin compromiso.")

    bullet(doc, "Sin tarjeta de credito requerida.")
    bullet(doc, "Acceso completo a todas las funcionalidades del Plan Basico.")
    bullet(doc, "Al finalizar, regresa automaticamente al Plan Gratuito sin perder datos.")
    bullet(doc, "Solo se puede activar una vez por negocio — aprovechala.")

    # --- 14. Escenarios ───────────────────────────────────────────────────
    h1(doc, "14. Escenarios Reales", anchor="prop_escenarios")

    h3(doc, "Escenario 1: Diana — Estilista independiente en Bogota")
    para(doc, "Diana acaba de abrir su salon de belleza en el barrio Chapinero. Trabaja sola, ofrece 8 servicios "
         "(corte, color, manicure, pedicure, etc.) y atiende entre 6 y 10 clientes por semana.")
    bullet(doc, "Plan Gratuito le alcanza perfectamente: 1 sede, 8 servicios, ~40 citas/mes.")
    bullet(doc, "Con el perfil publico, sus clientes la encuentran en Google buscando 'salon de belleza Chapinero'.")
    bullet(doc, "El QR lo imprime y lo pone en el espejo del salon — los clientes reservan desde su celular.")
    bullet(doc, "Los recordatorios por email redujeron sus no-shows de 5 a 1 por semana.")
    bullet(doc, "Sin pagar un peso. Sin contratos. Sin complicaciones.")

    h3(doc, "Escenario 2: Roberto — Barbero en Medellin")
    para(doc, "Roberto tiene una barberia en el centro de Medellin. Ofrece 5 servicios y atiende "
         "unos 8 clientes por dia.")
    bullet(doc, "Con el Plan Gratuito gestiona 40-50 citas mensuales facilmente.")
    bullet(doc, "El CRM le muestra que el 30% de sus clientes no vuelven despues de la primera visita.")
    bullet(doc, "Decidio enviar un WhatsApp personal a esos clientes — ahora el 60% regresa.")
    bullet(doc, "Cuando contrate a un segundo barbero, activara el Plan Basico.")

    h3(doc, "Escenario 3: Laura — Instructora de yoga en Cali")
    para(doc, "Laura da clases de yoga en un estudio rentado. Ofrece 4 tipos de clases y atiende "
         "grupos de hasta 15 personas.")
    bullet(doc, "El Plan Gratuito le permite gestionar sus reservaciones sin papel.")
    bullet(doc, "El perfil publico atrae nuevos estudiantes que buscan 'yoga Cali' en Google.")
    bullet(doc, "Probo el mes gratis del Plan Basico y ahora esta evaluando contratar una asistente.")

    screenshot_placeholder(doc, "Collage de negocios usando Gestabiz en el Plan Gratuito")

    page_break(doc)

    # --- 15. FAQ ──────────────────────────────────────────────────────────
    h1(doc, "15. Preguntas Frecuentes", anchor="prop_faq")

    faq = [
        ("El Plan Gratuito es temporal?",
         "No. El Plan Gratuito es permanente e indefinido. No tiene fecha de vencimiento ni requiere tarjeta de credito. Puedes usarlo para siempre."),
        ("Que pasa si supero los 50 clientes?",
         "Tus datos nunca se pierden. Los clientes adicionales se almacenan en la base de datos pero no se muestran hasta que actualices el plan. Un banner te informa cuantos estan ocultos."),
        ("Puedo tener mas de un negocio en el Plan Gratuito?",
         "Si. Puedes crear multiples negocios y cada uno tendra sus propios limites del Plan Gratuito de forma independiente."),
        ("Como funciona la prueba gratuita del Plan Basico?",
         "Al activar la prueba, obtienes acceso completo al Plan Basico durante 30 dias sin pagar nada. Al finalizar, si no activas un plan de pago, regresas al Plan Gratuito sin perder datos."),
        ("Mis datos estan seguros?",
         "Si. Gestabiz usa Supabase (PostgreSQL) con seguridad a nivel de fila (RLS). Cada tabla tiene politicas de seguridad que garantizan que solo puedas acceder a tus propios datos."),
        ("Que diferencia hay entre owner y admin?",
         "El owner es quien creo el negocio — tiene permisos totales sin restriccion. Un admin es un usuario al que se le asigno el rol administrativo con permisos que pueden ser personalizados."),
        ("Puedo exportar mis datos?",
         "En el Plan Gratuito puedes exportar el tracking de notificaciones en CSV. Los planes superiores agregan exportacion de reportes financieros en PDF, CSV y Excel."),
        ("Gestabiz funciona en mi celular?",
         "Si. La aplicacion web es totalmente responsive y se adapta a cualquier tamano de pantalla. Adicionalmente, hay una aplicacion movil nativa en desarrollo con Expo/React Native."),
    ]
    for q, a in faq:
        h3(doc, f"P: {q}")
        para(doc, f"R: {a}")

    page_break(doc)

    # --- 16. CTA ──────────────────────────────────────────────────────────
    h1(doc, "16. Siguiente Paso", anchor="prop_cta")

    para(doc, "Empieza hoy. Sin costo. Sin compromiso. Sin tarjeta de credito.", size=14)

    for _ in range(2):
        doc.add_paragraph()

    callout(doc, "Crea tu cuenta gratuita en gestabiz.com",
            "1. Registrate con tu email o Google\n"
            "2. Crea tu negocio en 3 minutos\n"
            "3. Sigue el asistente de configuracion\n"
            "4. Comparte tu enlace o QR con tus clientes\n"
            "5. Empieza a recibir citas — hoy mismo",
            color=ACCENT)

    for _ in range(2):
        doc.add_paragraph()

    para(doc, "Gestabiz — La plataforma todo-en-uno para negocios de servicios.",
         size=14, italic=True, color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Desarrollado por Ti Turing | soporte@gestabiz.com | gestabiz.com",
         size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass

    page_break(doc)

    para(doc, "Fin de la Parte 3 — Propuesta de Valor: Administrador Plan Gratuito.", size=12, italic=True,
         color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Continua en la Parte 4: Administrador Plan Basico.", size=12, italic=True,
         color=PURPLE, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # --- Manual de Usuario ---
    print("Generando Manual de Usuario — Parte 3: Admin Plan Gratuito...")
    manual = build_manual_part3()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz - copilot - parte3.docx"
    manual.save(str(manual_path))
    print(f"  -> {manual_path}")

    # --- Propuesta de Valor ---
    print("Generando Propuesta de Valor — Parte 3: Admin Plan Gratuito...")
    propuesta = build_proposal_part3()
    propuesta_path = DOCS_DIR / "Propuesta_Valor_Gestabiz - copilot - parte3.docx"
    propuesta.save(str(propuesta_path))
    print(f"  -> {propuesta_path}")

    print("\nListo. Ambos documentos generados exitosamente.")


if __name__ == "__main__":
    main()
