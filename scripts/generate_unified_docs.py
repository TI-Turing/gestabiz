#!/usr/bin/env python3
"""
generate_unified_docs.py
========================
Unifica las 5 partes de documentación Gestabiz en dos documentos finales
con header (logo Gestabiz) y footer (logo Ti Turing), más versiones .md.

Salida:
  docs/Manual_Usuario_Gestabiz.docx
  docs/Propuesta_Valor_Gestabiz.docx
  docs/Manual_Usuario_Gestabiz.md
  docs/Propuesta_Valor_Gestabiz.md
"""

import os, sys, subprocess, re
from pathlib import Path
from typing import List

# ── Ensure dependencies ────────────────────────────────────────────
def _ensure(pkg, import_name=None):
    try:
        __import__(import_name or pkg)
    except ImportError:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "-q", pkg],
            stdout=subprocess.DEVNULL,
        )

_ensure("python-docx", "docx")
_ensure("docxcompose")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docxcompose.composer import Composer

# ── Paths & Constants ──────────────────────────────────────────────
WORKSPACE       = Path(__file__).resolve().parent.parent
DOCS_DIR        = WORKSPACE / "docs"
LOGO_GESTABIZ   = WORKSPACE / "public" / "og-image.png"       # 1200×630 PNG
LOGO_TI_TURING  = WORKSPACE / "src" / "assets" / "images" / "tt" / "1.png"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK   = RGBColor(0x1F, 0x1F, 0x2E)
GREY   = RGBColor(0x6B, 0x72, 0x80)

# Document configurations
DOCS_CONFIG = [
    {
        "prefix":    "Manual_Usuario_Gestabiz",
        "label":     "MANUAL DE USUARIO",
        "old_cover": "Rol Cliente",
        "new_cover": "Todos los Roles y Planes",
    },
    {
        "prefix":    "Propuesta_Valor_Gestabiz",
        "label":     "PROPUESTA DE VALOR",
        "old_cover": "Rol Cliente",
        "new_cover": "Plataforma Integral para PyMEs de Servicios",
    },
]


# ═══════════════════════════════════════════════════════════════════
# 1. LOCATE PART FILES
# ═══════════════════════════════════════════════════════════════════

def get_part_files(prefix: str) -> List[Path]:
    """
    Locate all 5 parts for a document type.
    Part 1 is named '{prefix} - copilot.docx' (no 'parteN').
    Parts 2-5 are named '{prefix} - copilot - parteN.docx'.
    """
    parts: List[Path] = []

    # Part 1: special naming
    p1 = DOCS_DIR / f"{prefix} - copilot.docx"
    if p1.exists():
        parts.append(p1)
    else:
        # Fallback: maybe it does have "parte1"
        p1_alt = DOCS_DIR / f"{prefix} - copilot - parte1.docx"
        if p1_alt.exists():
            parts.append(p1_alt)
        else:
            print(f"  !! Parte 1 no encontrada")

    # Parts 2-5
    for i in range(2, 6):
        p = DOCS_DIR / f"{prefix} - copilot - parte{i}.docx"
        if p.exists():
            parts.append(p)
        else:
            print(f"  !! Parte {i} no encontrada: {p.name}")

    return parts


# ═══════════════════════════════════════════════════════════════════
# 2. MERGE DOCUMENTS
# ═══════════════════════════════════════════════════════════════════

def merge_parts(part_files: List[Path], output: Path):
    """Merge multiple .docx files into one using docxcompose."""
    master = Document(str(part_files[0]))
    composer = Composer(master)
    for pf in part_files[1:]:
        composer.append(Document(str(pf)))
    composer.save(str(output))


# ═══════════════════════════════════════════════════════════════════
# 3. UPDATE COVER PAGE
# ═══════════════════════════════════════════════════════════════════

def update_cover(doc_path: Path, old_subtitle: str, new_subtitle: str):
    """Replace 'Parte 1 de 5' with 'Versión Completa' on the first cover."""
    doc = Document(str(doc_path))
    changed = False
    for para in doc.paragraphs[:80]:          # cover is in first ~80 paragraphs
        for run in para.runs:
            if "Parte 1 de 5" in run.text:
                run.text = run.text.replace("Parte 1 de 5", u"Versión Completa")
                changed = True
            if old_subtitle in run.text:
                run.text = run.text.replace(old_subtitle, new_subtitle)
                changed = True
    if changed:
        doc.save(str(doc_path))


# ═══════════════════════════════════════════════════════════════════
# 4. HEADER & FOOTER WITH LOGOS
# ═══════════════════════════════════════════════════════════════════

def _border_xml(side: str, color: str = "7C3AED") -> str:
    """Create XML for a single paragraph border (top or bottom)."""
    return (
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:{side} w:val="single" w:sz="6" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )


def _add_page_field(paragraph):
    """Insert a PAGE field code (auto page number) into the paragraph."""
    run = paragraph.add_run()
    run._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    )
    run2 = paragraph.add_run()
    run2._element.append(
        parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
        )
    )
    run3 = paragraph.add_run()
    run3._element.append(
        parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    )


def setup_header_footer(doc_path: Path):
    """Add Gestabiz logo in header and Ti Turing logo + page number in footer."""
    doc = Document(str(doc_path))

    for section in doc.sections:
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)

        # ── HEADER: Gestabiz logo ─────────────────────────────
        hdr = section.header
        hdr.is_linked_to_previous = False

        hp = hdr.paragraphs[0]
        hp.clear()
        # Remove extra paragraphs in header
        for extra in list(hdr.paragraphs[1:]):
            extra._element.getparent().remove(extra._element)

        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if LOGO_GESTABIZ.exists():
            run = hp.add_run()
            run.add_picture(str(LOGO_GESTABIZ), height=Cm(1.0))
        else:
            run = hp.add_run("Gestabiz")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = PURPLE
            run2 = hp.add_run("  |  Agenda. Gestiona. Crece.")
            run2.font.size = Pt(8)
            run2.font.color.rgb = GREY

        # Purple line under header
        hp._element.get_or_add_pPr().append(parse_xml(_border_xml("bottom")))

        # ── FOOTER: Ti Turing logo + page number ──────────────
        ftr = section.footer
        ftr.is_linked_to_previous = False

        fp = ftr.paragraphs[0]
        fp.clear()
        for extra in list(ftr.paragraphs[1:]):
            extra._element.getparent().remove(extra._element)

        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Purple line above footer
        fp._element.get_or_add_pPr().append(parse_xml(_border_xml("top")))

        if LOGO_TI_TURING.exists():
            run = fp.add_run()
            run.add_picture(str(LOGO_TI_TURING), height=Cm(0.8))
        else:
            run = fp.add_run("Ti Turing")
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = PURPLE

        # Separator + page number
        sep = fp.add_run("   \u00b7   P\u00e1g. ")
        sep.font.size = Pt(8)
        sep.font.color.rgb = GREY
        _add_page_field(fp)

    doc.save(str(doc_path))


# ═══════════════════════════════════════════════════════════════════
# 5. DOCX → MARKDOWN CONVERSION
# ═══════════════════════════════════════════════════════════════════

def _runs_to_md(para) -> str:
    """Convert a paragraph's runs to Markdown inline formatting."""
    parts = []
    for r in para.runs:
        t = r.text
        if not t:
            continue
        # Skip if it looks like an image placeholder (empty run with drawing)
        if r.bold and r.italic:
            parts.append(f"***{t}***")
        elif r.bold:
            parts.append(f"**{t}**")
        elif r.italic:
            parts.append(f"*{t}*")
        else:
            parts.append(t)
    return "".join(parts) or para.text


def _table_to_md(table) -> List[str]:
    """Convert a docx Table to Markdown table lines."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return []

    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append("")

    md = []
    md.append("| " + " | ".join(rows[0]) + " |")
    md.append("| " + " | ".join(["---"] * ncols) + " |")
    for row in rows[1:]:
        md.append("| " + " | ".join(row) + " |")
    return md


def docx_to_md(doc_path: Path, md_path: Path):
    """Convert a merged .docx to a well-formatted Markdown file."""
    doc = Document(str(doc_path))

    # Build element → python-docx object maps for ordered traversal
    para_map  = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    lines: List[str] = []
    prev_was_blank = False

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        eid = id(child)

        # ── Paragraph ──────────────────────────────────────────
        if tag == "p" and eid in para_map:
            p = para_map[eid]
            text = p.text.strip()
            sn = (p.style.name or "") if p.style else ""

            if not text:
                if not prev_was_blank:
                    lines.append("")
                    prev_was_blank = True
                continue
            prev_was_blank = False

            # Heading styles
            if "Heading 1" in sn or sn == "Title":
                lines.append(f"\n# {text}\n")
            elif "Heading 2" in sn or sn == "Subtitle":
                lines.append(f"\n## {text}\n")
            elif "Heading 3" in sn:
                lines.append(f"\n### {text}\n")
            elif "Heading 4" in sn:
                lines.append(f"\n#### {text}\n")
            # List items
            elif "List" in sn or "Bullet" in sn:
                lines.append(f"- {text}")
            # Normal paragraphs with inline formatting
            else:
                lines.append(_runs_to_md(p))

        # ── Table ──────────────────────────────────────────────
        elif tag == "tbl" and eid in table_map:
            prev_was_blank = False
            lines.append("")
            lines.extend(_table_to_md(table_map[eid]))
            lines.append("")

        # ── Section break ──────────────────────────────────────
        elif tag == "sectPr":
            if not prev_was_blank:
                lines.append("")
            lines.append("---")
            lines.append("")
            prev_was_blank = False

    # Final cleanup: collapse 3+ consecutive blank lines
    text = "\n".join(lines)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    # Remove trailing whitespace on lines
    text = "\n".join(l.rstrip() for l in text.split("\n"))
    md_path.write_text(text.strip() + "\n", encoding="utf-8")


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    print("=" * 60)
    print("  GESTABIZ — Documentos Finales Unificados")
    print("=" * 60)
    print(f"  Logo Gestabiz: {'OK' if LOGO_GESTABIZ.exists() else 'NO ENCONTRADO (usara texto)'}")
    print(f"  Logo Ti Turing: {'OK' if LOGO_TI_TURING.exists() else 'NO ENCONTRADO (usara texto)'}")

    for cfg in DOCS_CONFIG:
        prefix = cfg["prefix"]
        print(f"\n{'─' * 55}")
        print(f"  {cfg['label']}")
        print(f"{'─' * 55}")

        # Locate parts
        parts = get_part_files(prefix)
        if not parts:
            print("  ERROR: No se encontraron partes. Saltando.")
            continue
        print(f"  Partes encontradas: {len(parts)}/5")
        for p in parts:
            print(f"    - {p.name} ({p.stat().st_size // 1024} KB)")

        out_docx = DOCS_DIR / f"{prefix}.docx"
        out_md   = DOCS_DIR / f"{prefix}.md"

        # Step 1: Merge
        print(f"\n  [1/4] Unificando {len(parts)} partes...")
        merge_parts(parts, out_docx)
        print(f"         -> {out_docx.name} ({out_docx.stat().st_size // 1024} KB)")

        # Step 2: Update cover
        print("  [2/4] Actualizando portada a 'Version Completa'...")
        update_cover(out_docx, cfg["old_cover"], cfg["new_cover"])

        # Step 3: Header & Footer
        print("  [3/4] Insertando header (Gestabiz) y footer (Ti Turing)...")
        setup_header_footer(out_docx)
        print(f"         -> {out_docx.name} ({out_docx.stat().st_size // 1024} KB)")

        # Step 4: Markdown
        print("  [4/4] Generando version Markdown...")
        docx_to_md(out_docx, out_md)
        print(f"         -> {out_md.name} ({out_md.stat().st_size // 1024} KB)")

    # ── Summary ────────────────────────────────────────────────
    print(f"\n{'=' * 60}")
    print("  RESUMEN FINAL")
    print(f"{'=' * 60}")
    for cfg in DOCS_CONFIG:
        prefix = cfg["prefix"]
        for ext in [".docx", ".md"]:
            f = DOCS_DIR / f"{prefix}{ext}"
            if f.exists():
                sz = f.stat().st_size / 1024
                print(f"  {f.name:50s} {sz:>8,.0f} KB")
    print(f"{'=' * 60}")
    print("  Completado.")


if __name__ == "__main__":
    main()
