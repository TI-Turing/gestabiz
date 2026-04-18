#!/usr/bin/env python3
"""
Generador de documentos Gestabiz — Version Copilot — Parte 4: Rol Admin (Plan Basico)
========================================================================================
Genera:
  - docs/Manual_Usuario_Gestabiz - copilot - parte4.docx
  - docs/Propuesta_Valor_Gestabiz - copilot - parte4.docx

Requisitos: python-docx >= 1.0.0
Ejecucion:  python scripts/generate_copilot_docs_part4.py
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths & Brand Colors
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
LOGO_GESTABIZ = ROOT / "src" / "assets" / "images" / "logo_gestabiz.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def add_bookmark(paragraph, name: str) -> None:
    tag = OxmlElement("w:bookmarkStart")
    tag.set(qn("w:id"), str(id(name) % 2**31))
    tag.set(qn("w:name"), name)
    paragraph._p.append(tag)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(id(name) % 2**31))
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *,
                           bold: bool = False, color: str = "7C3AED") -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# High-level helpers
# ---------------------------------------------------------------------------

def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + text)
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(title)
    style_run(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {caption}")
    style_run(r, size=9, italic=True, color=GREY)
    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=header_fg)
    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, size=10, color=DARK)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup + cover
# ---------------------------------------------------------------------------

def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run(f"{title}  |  Desarrollado por Ti Turing  |  (c) 2026")
        style_run(r, size=9, italic=True, color=GREY)
    build_cover(doc, title, subtitle)
    return doc


def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Parte 4 de 5: Administrador — Plan Basico")
    style_run(r, size=14, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version del producto: 1.0.3   |   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fase Beta completada   |   Listo para produccion")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing")
    style_run(r, size=12, bold=True, color=DARK)


# ============================================================================
# MANUAL DE USUARIO — PARTE 4 (ADMIN PLAN BASICO)
# ============================================================================

def build_manual_part4() -> Document:
    doc = setup_document(
        "Manual de Usuario — Gestabiz",
        "Guia funcional completa del sistema de gestion de citas y negocios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE DE CONTENIDOS", anchor="toc_p4")
    para(doc, "A continuacion se listan todas las secciones cubiertas en esta parte. "
         "Haga clic en cualquier titulo para navegar directamente al detalle.", italic=True, color=GREY)

    toc_items = [
        ("sec_resumen_basico", "1. Resumen Ejecutivo — Plan Basico"),
        ("sec_que_incluye", "2. Que incluye el Plan Basico"),
        ("sec_limites_comparativa", "3. Comparativa de Limites: Gratuito vs. Basico"),
        ("sec_activar_basico", "4. Como Activar el Plan Basico"),
        ("sec_free_trial", "5. Prueba Gratuita de 30 Dias"),
        ("sec_empleados", "6. Gestion de Empleados"),
        ("sec_emp_dashboard", "7. Dashboard de Empleados — Vista General"),
        ("sec_emp_stats", "8. Tarjetas de Estadisticas de Empleados"),
        ("sec_emp_filtros", "9. Filtros y Busqueda de Empleados"),
        ("sec_emp_vistas", "10. Modos de Vista (Lista y Mapa Jerarquico)"),
        ("sec_emp_invitar", "11. Invitar Empleados (Codigo QR de 6 Digitos)"),
        ("sec_emp_solicitudes", "12. Solicitudes de Ingreso Pendientes"),
        ("sec_emp_setup", "13. Empleados Pendientes de Configuracion"),
        ("sec_emp_perfil", "14. Modal de Perfil de Empleado"),
        ("sec_emp_horarios", "15. Horarios del Empleado (7 Dias)"),
        ("sec_emp_almuerzo", "16. Hora de Almuerzo del Empleado"),
        ("sec_emp_salario", "17. Configuracion de Salario"),
        ("sec_emp_servicios", "18. Asignacion de Servicios al Empleado"),
        ("sec_emp_activar", "19. Activar y Desactivar Empleados"),
        ("sec_emp_limite", "20. Limite de Empleados en Plan Basico (6)"),
        ("sec_ausencias", "21. Gestion de Ausencias y Vacaciones (Admin)"),
        ("sec_aus_politica", "22. Politica de Aprobacion Obligatoria"),
        ("sec_aus_pendientes", "23. Pestana Pendientes — Aprobar o Rechazar"),
        ("sec_aus_card", "24. Tarjeta de Aprobacion de Ausencia — Detalle"),
        ("sec_aus_tipos", "25. Tipos de Ausencia y Colores"),
        ("sec_aus_citas", "26. Citas Afectadas por una Ausencia"),
        ("sec_aus_historial", "27. Pestana Historial de Ausencias"),
        ("sec_aus_vacaciones", "28. Widget de Balance de Vacaciones"),
        ("sec_aus_festivos", "29. Festivos Publicos y su Impacto"),
        ("sec_aus_notificaciones", "30. Notificaciones de Ausencias"),
        ("sec_ventas", "31. Historial de Ventas"),
        ("sec_ventas_resumen", "32. Tarjetas de Resumen de Ventas"),
        ("sec_ventas_filtros", "33. Filtros de Periodo y Busqueda"),
        ("sec_ventas_lista", "34. Lista de Ventas Completadas"),
        ("sec_ventas_cliente", "35. Ver Perfil de Cliente desde Ventas"),
        ("sec_rapidas", "36. Ventas Rapidas (Walk-In)"),
        ("sec_rapidas_form", "37. Formulario de Venta Rapida — Campos"),
        ("sec_rapidas_metodos", "38. Metodos de Pago Disponibles"),
        ("sec_rapidas_auto", "39. Autocompletado de Monto por Servicio"),
        ("sec_rapidas_stats", "40. Estadisticas de Ventas Rapidas"),
        ("sec_rapidas_historial", "41. Historial de Ultimas 10 Ventas"),
        ("sec_rapidas_contable", "42. Integracion Contable de Ventas Rapidas"),
        ("sec_reportes", "43. Reportes Financieros"),
        ("sec_rep_dashboard", "44. Dashboard Financiero Interactivo"),
        ("sec_rep_graficos", "45. Graficos Disponibles"),
        ("sec_rep_filtros", "46. Filtros del Dashboard Financiero"),
        ("sec_rep_export", "47. Exportacion de Reportes (Nota de Plan)"),
        ("sec_permisos", "48. Gestion de Permisos"),
        ("sec_perm_usuarios", "49. Tabla de Usuarios y Roles"),
        ("sec_perm_acciones", "50. Acciones sobre Permisos"),
        ("sec_perm_templates", "51. Plantillas de Permisos (3 Predefinidas)"),
        ("sec_perm_stats", "52. Estadisticas de Usuarios y Roles"),
        ("sec_limites_ampliados", "53. Limites Ampliados — Sedes, Servicios, Citas, Clientes"),
        ("sec_sedes_3", "54. Gestion de hasta 3 Sedes"),
        ("sec_servicios_ilim", "55. Servicios Ilimitados"),
        ("sec_citas_ilim", "56. Citas Ilimitadas por Mes"),
        ("sec_clientes_ilim", "57. Clientes Ilimitados"),
        ("sec_features_pago", "58. Funcionalidades Adicionales de Planes de Pago"),
        ("sec_whatsapp", "59. Recordatorios por WhatsApp"),
        ("sec_gcal", "60. Sincronizacion con Google Calendar"),
        ("sec_modulos_pro", "61. Modulos Bloqueados — Solo en Plan Pro"),
        ("sec_upgrade_pro", "62. Como Actualizar al Plan Pro"),
        ("sec_soporte", "63. Soporte y Ayuda"),
        ("sec_glosario_basico", "64. Glosario de Terminos del Plan Basico"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 1 — RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 1 — RESUMEN EJECUTIVO", anchor="parte1_basico")
    para(doc, "Esta seccion presenta a grandes rasgos las capacidades exclusivas del Plan Basico de Gestabiz. "
         "Para el detalle exhaustivo de cada funcionalidad, consulte la Parte 2.")

    # --- 1. Resumen Ejecutivo ---
    h2(doc, "1. Resumen Ejecutivo — Plan Basico", anchor="sec_resumen_basico")
    para(doc, "El Plan Basico de Gestabiz ($89,900 COP/mes o $899,000 COP/ano con 2 meses gratis) "
         "esta disenado para negocios que necesitan crecer mas alla de las limitaciones del Plan Gratuito. "
         "Incluye todo lo del Plan Gratuito mas seis modulos exclusivos y limites ampliados que "
         "permiten gestionar un equipo de hasta 6 empleados, multiples sedes y un volumen ilimitado "
         "de citas, clientes y servicios.")

    para(doc, "Los seis modulos exclusivos del Plan Basico son:")
    bullet(doc, "Gestionar hasta 6 empleados con invitaciones por codigo QR, horarios por dia, "
           "salarios, asignacion de servicios y jerarquia organizacional.", bold_prefix="Empleados")
    bullet(doc, "Aprobar o rechazar solicitudes de ausencias y vacaciones de los empleados, "
           "con calculo automatico de balance de vacaciones y bloqueo de citas.", bold_prefix="Ausencias")
    bullet(doc, "Consultar el historial completo de citas completadas con ingresos, "
           "filtrado por periodo y busqueda por cliente o servicio.", bold_prefix="Historial de Ventas")
    bullet(doc, "Registrar ventas presenciales (walk-in) de clientes que llegan sin cita previa, "
           "con integracion contable automatica.", bold_prefix="Ventas Rapidas")
    bullet(doc, "Acceder a un dashboard financiero interactivo con graficos de ingresos, "
           "gastos, distribucion por categoria y tendencias mensuales.", bold_prefix="Reportes Financieros")
    bullet(doc, "Controlar que pueden hacer los empleados en la plataforma mediante "
           "plantillas de permisos predefinidas.", bold_prefix="Permisos")

    para(doc, "Ademas, los limites se amplian significativamente:")
    simple_table(doc, ["Recurso", "Plan Gratuito", "Plan Basico"],
                 [["Sedes", "1", "3"],
                  ["Empleados", "1 (solo owner)", "6"],
                  ["Citas por mes", "50", "Ilimitadas"],
                  ["Clientes visibles", "50", "Ilimitados"],
                  ["Servicios", "15", "Ilimitados"]])

    callout(doc, "Incluye todo lo del Plan Gratuito",
            "El Plan Basico hereda todas las funcionalidades descritas en la Parte 3 de este manual "
            "(Dashboard, Servicios, Sedes, Calendario, Clientes CRM, Facturacion, Configuraciones, "
            "Perfil Publico, QR, etc.) sin restricciones adicionales. Esta parte documenta "
            "unicamente las funcionalidades NUEVAS y los limites ampliados.",
            color=ACCENT)

    screenshot_placeholder(doc, "Sidebar del admin con modulos del Plan Basico desbloqueados")

    # --- 2. Que incluye ---
    h2(doc, "2. Que incluye el Plan Basico", anchor="sec_que_incluye")
    para(doc, "El Plan Basico de Gestabiz incluye tres grandes categorias de mejoras respecto al Plan Gratuito:")

    h3(doc, "A. Modulos exclusivos desbloqueados")
    para(doc, "Seis modulos que estaban bloqueados en el Plan Gratuito se desbloquean "
         "al activar el Plan Basico. Cada uno aparece en la barra lateral sin el icono de candado "
         "y es completamente funcional.")
    simple_table(doc, ["Modulo", "Icono Sidebar", "Descripcion Corta"],
                 [["Empleados", "Users", "Gestion de hasta 6 empleados con jerarquia"],
                  ["Ausencias", "CalendarOff", "Aprobacion de ausencias y vacaciones"],
                  ["Ventas", "BarChart3", "Historial de citas completadas con ingresos"],
                  ["Ventas Rapidas", "ShoppingCart", "Registro de ventas walk-in"],
                  ["Reportes", "FileText", "Dashboard financiero con graficos"],
                  ["Permisos", "Shield", "Gestion de permisos por plantillas"]])

    h3(doc, "B. Limites ampliados en modulos existentes")
    para(doc, "Los modulos que ya estaban disponibles en el Plan Gratuito ahora operan con limites "
         "mucho mas generosos o directamente sin limites.")

    h3(doc, "C. Funcionalidades adicionales de planes de pago")
    para(doc, "Al ser un plan de pago, el Plan Basico tambien habilita:")
    bullet(doc, "Recordatorios automaticos de citas a clientes por WhatsApp Business API.")
    bullet(doc, "Sincronizacion bidireccional con Google Calendar para el admin y empleados.")
    bullet(doc, "Soporte prioritario por email.")

    page_break(doc)

    # --- 3. Comparativa ---
    h2(doc, "3. Comparativa de Limites: Gratuito vs. Basico", anchor="sec_limites_comparativa")
    para(doc, "La siguiente tabla muestra la comparativa completa de limites y funcionalidades "
         "entre el Plan Gratuito y el Plan Basico.")
    simple_table(doc, ["Caracteristica", "Plan Gratuito", "Plan Basico"],
                 [["Precio mensual", "$0 COP", "$89,900 COP"],
                  ["Precio anual", "$0 COP", "$899,000 COP (2 meses gratis)"],
                  ["Sedes", "1", "3"],
                  ["Empleados", "1 (solo owner)", "6"],
                  ["Citas por mes", "50", "Ilimitadas"],
                  ["Clientes visibles", "50", "Ilimitados"],
                  ["Servicios", "15", "Ilimitados"],
                  ["Dashboard Resumen", "Si", "Si"],
                  ["Calendario de Citas", "Si", "Si"],
                  ["Gestion de Servicios", "Si (hasta 15)", "Si (ilimitados)"],
                  ["Gestion de Sedes", "Si (1 sede)", "Si (hasta 3)"],
                  ["CRM Clientes", "Si (50 visibles)", "Si (ilimitados)"],
                  ["Perfil Publico y SEO", "Si", "Si"],
                  ["Codigo QR", "Si", "Si"],
                  ["Facturacion y Planes", "Si (vista)", "Si (completa)"],
                  ["Configuraciones", "Si", "Si"],
                  ["Gestion de Empleados", "No", "Si (hasta 6)"],
                  ["Ausencias y Vacaciones", "No", "Si"],
                  ["Historial de Ventas", "No", "Si"],
                  ["Ventas Rapidas", "No", "Si"],
                  ["Reportes Financieros", "No", "Si (vista, sin exportacion)"],
                  ["Permisos", "No", "Si (3 plantillas)"],
                  ["Gestion de Gastos", "No", "No (Plan Pro)"],
                  ["Reclutamiento", "No", "No (Plan Pro)"],
                  ["Recursos Fisicos", "No", "No (Plan Pro)"],
                  ["Exportacion de Reportes", "No", "No (Plan Pro)"],
                  ["Permisos Granulares (79)", "No", "No (Plan Pro)"],
                  ["WhatsApp Recordatorios", "No", "Si"],
                  ["Google Calendar Sync", "No", "Si"],
                  ["Soporte", "Comunidad", "Email prioritario"]])

    screenshot_placeholder(doc, "Pagina de precios con los tres planes comparados")

    page_break(doc)

    # --- 4. Activar Plan ---
    h2(doc, "4. Como Activar el Plan Basico", anchor="sec_activar_basico")
    para(doc, "Existen dos rutas para activar el Plan Basico desde el panel de administracion:")

    h3(doc, "Ruta 1: Desde la Barra Lateral")
    numbered(doc, "Haga clic en cualquier modulo bloqueado en la barra lateral (por ejemplo, 'Empleados').")
    numbered(doc, "El sistema muestra una pantalla de upgrade con un icono de candado grande, "
             "un badge ambar con estrellas y el texto 'Modulo disponible en Plan Basico'.")
    numbered(doc, "Haga clic en el boton 'Ver planes y precios'.")
    numbered(doc, "El sistema navega a la seccion de Facturacion (/app/admin/billing) "
             "donde se muestra la pagina de precios.")
    numbered(doc, "Seleccione el Plan Basico (mensual o anual) y complete el pago.")

    h3(doc, "Ruta 2: Desde Facturacion")
    numbered(doc, "Navegue a Facturacion en la barra lateral.")
    numbered(doc, "En la vista del plan actual, haga clic en 'Cambiar plan' o 'Actualizar'.")
    numbered(doc, "Seleccione el Plan Basico y complete el proceso de pago.")

    callout(doc, "Metodos de pago aceptados",
            "Gestabiz soporta tres pasarelas de pago: Stripe (tarjetas internacionales), "
            "PayU Latam (tarjetas colombianas, PSE, efectivo) y MercadoPago (Argentina, Brasil, Mexico, Chile). "
            "La pasarela activa depende de la configuracion del negocio.",
            color=PURPLE)

    screenshot_placeholder(doc, "Pantalla de upgrade (PlanGate) al intentar acceder a un modulo bloqueado")

    # --- 5. Free Trial ---
    h2(doc, "5. Prueba Gratuita de 30 Dias", anchor="sec_free_trial")
    para(doc, "Todo nuevo negocio tiene la opcion de activar una prueba gratuita del Plan Basico "
         "por 30 dias antes de decidir si desea pagar. Esta prueba se puede activar una sola vez "
         "por cada owner (propietario del negocio).")

    h3(doc, "Reglas de la Prueba Gratuita")
    bullet(doc, "Duracion: exactamente 30 dias calendario desde la activacion.")
    bullet(doc, "Activacion: una sola vez por owner, independientemente del numero de negocios.")
    bullet(doc, "Funcionalidades: acceso completo a todos los modulos y limites del Plan Basico.")
    bullet(doc, "Al finalizar: el negocio vuelve automaticamente al Plan Gratuito si no se paga.")
    bullet(doc, "Datos conservados: los datos creados durante la prueba se conservan, pero los "
           "modulos exclusivos se bloquean nuevamente.")
    bullet(doc, "Notificacion: el sistema avisa con 7 dias de anticipacion antes de que expire.")

    h3(doc, "Como Activar la Prueba")
    numbered(doc, "En la pagina de Facturacion, localice el banner 'Prueba gratuita de 30 dias'.")
    numbered(doc, "Haga clic en 'Activar prueba gratuita'.")
    numbered(doc, "El sistema activa inmediatamente el Plan Basico sin requerir datos de pago.")
    numbered(doc, "Un contador muestra los dias restantes de la prueba en la seccion de Facturacion.")

    callout(doc, "Sin compromiso",
            "No se requiere tarjeta de credito ni datos de pago para activar la prueba gratuita. "
            "Al terminar los 30 dias, simplemente se revierte al Plan Gratuito automaticamente.",
            color=ACCENT)

    screenshot_placeholder(doc, "Banner de prueba gratuita en la seccion de Facturacion")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 2 — DETALLE EXHAUSTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 2 — DETALLE EXHAUSTIVO DE FUNCIONALIDADES", anchor="parte2_basico")
    para(doc, "A continuacion se documenta en detalle cada modulo exclusivo del Plan Basico, "
         "incluyendo todos los elementos de interfaz, flujos normales y alternos, reglas de negocio, "
         "validaciones, estados y excepciones.")

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: EMPLEADOS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "GESTION DE EMPLEADOS", anchor="sec_empleados", color=PURPLE)
    para(doc, "El modulo de Empleados permite al administrador gestionar un equipo de hasta 6 personas, "
         "organizado en una jerarquia visual con niveles definidos. Desde aqui se invitan nuevos miembros, "
         "se configuran horarios, salarios, servicios ofrecidos y se gestionan las solicitudes de ingreso.")

    # --- 7. Dashboard ---
    h2(doc, "7. Dashboard de Empleados — Vista General", anchor="sec_emp_dashboard")
    para(doc, "Al acceder al modulo de Empleados desde la barra lateral, se presenta un dashboard completo "
         "con las siguientes secciones visibles:")
    bullet(doc, "Encabezado con titulo 'Gestion de Empleados' y subtitulo descriptivo.")
    bullet(doc, "Botones de cambio de vista: Lista (icono List) y Mapa jerarquico (icono Network).")
    bullet(doc, "Boton 'Invitar' (icono QrCode) para generar codigos de invitacion.")
    bullet(doc, "Boton 'Solicitudes' (icono Users) para ver solicitudes de ingreso pendientes.")
    bullet(doc, "Tarjetas de estadisticas (4 cards en fila).")
    bullet(doc, "Barra de filtros con busqueda, filtro por nivel y tipo.")
    bullet(doc, "Contenido principal segun la vista seleccionada (lista o mapa).")
    bullet(doc, "Banner de limite de plan (si se excede el limite de 6 empleados).")

    screenshot_placeholder(doc, "Dashboard de gestion de empleados con vista de lista")

    # --- 8. Stats Cards ---
    h2(doc, "8. Tarjetas de Estadisticas de Empleados", anchor="sec_emp_stats")
    para(doc, "En la parte superior del modulo se muestran 4 tarjetas con estadisticas en tiempo real:")

    simple_table(doc, ["Tarjeta", "Icono", "Dato Mostrado", "Descripcion"],
                 [["Total Empleados", "Users", "Numero total", "Cuenta todos los empleados registrados"],
                  ["Por Nivel", "(grid 5 col)", "Prop./Admin/Ger./Lider/Pers.", "Distribucion por nivel jerarquico"],
                  ["Ocupacion Promedio", "—", "Porcentaje (ej. 75.3%)", "Promedio de ocupacion del equipo"],
                  ["Rating Promedio", "Star (amarilla)", "Puntuacion (ej. 4.5)", "Calificacion promedio de reviews"]])

    h4(doc, "Detalle de la tarjeta Por Nivel")
    para(doc, "La tarjeta de distribucion por nivel muestra un grid de 5 columnas con las abreviaturas "
         "y la cantidad de empleados en cada nivel:")
    simple_table(doc, ["Nivel", "Abreviatura", "Significado"],
                 [["0", "Prop.", "Propietario (owner)"],
                  ["1", "Admin", "Administrador"],
                  ["2", "Ger.", "Gerente de sede"],
                  ["3", "Lider", "Lider de equipo"],
                  ["4", "Pers.", "Personal operativo"]])

    # --- 9. Filtros ---
    h2(doc, "9. Filtros y Busqueda de Empleados", anchor="sec_emp_filtros")
    para(doc, "La barra de filtros permite refinar la lista de empleados con los siguientes controles:")

    bullet(doc, "Boton expandible con badge que cuenta filtros activos. Al hacer clic se despliega "
           "el panel completo de filtros (FiltersPanel).", bold_prefix="Boton Filtros")
    bullet(doc, "Solo visible si hay filtros aplicados. Resetea todos los filtros a su estado inicial.", bold_prefix="Boton Limpiar filtros")
    bullet(doc, "Toggle que incluye o excluye empleados que no tienen sede asignada.", bold_prefix="Incluir personal sin sede")
    bullet(doc, "Muestra la cantidad de empleados visibles despues de aplicar filtros.", bold_prefix="Contador")

    h3(doc, "Campos del Panel de Filtros (FiltersPanel)")
    simple_table(doc, ["Campo", "Tipo", "Opciones", "Pre-seleccion"],
                 [["Busqueda", "Input texto", "Nombre o email", "Vacio"],
                  ["Nivel jerarquico", "Select", "Todos / Prop. / Admin / Ger. / Lider / Personal", "Todos"],
                  ["Tipo de empleado", "Select", "Todos / Profesional / Recepcionista / Contador / etc.", "Todos"],
                  ["Departamento", "Select", "Segun configuracion", "Todos"],
                  ["Sede", "Select", "Todas / cada sede del negocio", "Sede preferida (si configurada)"]])

    callout(doc, "Pre-seleccion inteligente de sede",
            "Si el administrador tiene una sede preferida configurada (ver Parte 3 — Sede Preferida), "
            "el filtro de sede se pre-selecciona automaticamente con esa sede al abrir el modulo.",
            color=PURPLE)

    # --- 10. Vistas ---
    h2(doc, "10. Modos de Vista (Lista y Mapa Jerarquico)", anchor="sec_emp_vistas")

    h3(doc, "Vista de Lista")
    para(doc, "Muestra a los empleados en formato de tarjetas (cards) con la siguiente informacion por cada uno:")
    bullet(doc, "Avatar del empleado (foto o iniciales).")
    bullet(doc, "Nombre completo y email.")
    bullet(doc, "Badge de rol (Admin, Profesional, Recepcionista, etc.).")
    bullet(doc, "Badge de nivel jerarquico con color.")
    bullet(doc, "Nombre de la sede asignada.")
    bullet(doc, "Estado (Activo en verde o Inactivo en rojo).")
    bullet(doc, "Botones de accion: Ver perfil, Editar, Asignar supervisor.")

    h3(doc, "Vista de Mapa Jerarquico")
    para(doc, "Muestra la estructura organizacional del negocio en un diagrama de arbol visual "
         "(HierarchyMapView) donde:")
    bullet(doc, "Cada nodo es un empleado con su avatar, nombre y nivel.")
    bullet(doc, "Las lineas conectan supervisores con sus reportes directos.")
    bullet(doc, "Se puede navegar por niveles para ver la estructura completa.")
    bullet(doc, "Hacer clic en un nodo abre el modal de perfil del empleado.")

    screenshot_placeholder(doc, "Vista de mapa jerarquico de empleados con conexiones")

    # --- 11. Invitar ---
    h2(doc, "11. Invitar Empleados (Codigo QR de 6 Digitos)", anchor="sec_emp_invitar")
    para(doc, "El sistema de invitacion permite al administrador generar codigos unicos que los nuevos "
         "empleados pueden usar para unirse al negocio.")

    h3(doc, "Flujo Normal de Invitacion")
    numbered(doc, "El administrador hace clic en el boton 'Invitar' (icono QrCode) en el header del modulo.")
    numbered(doc, "Se abre un dialog (modal) con el titulo 'Invitar empleado' y la descripcion "
             "'Se generara un codigo de 6 digitos que caduca en 24 horas'.")
    numbered(doc, "El administrador hace clic en 'Generar codigo'.")
    numbered(doc, "El sistema genera un codigo aleatorio de 6 digitos y lo muestra en grande en el dialog.")
    numbered(doc, "Debajo se muestran los codigos activos con su fecha de expiracion.")
    numbered(doc, "El administrador comparte el codigo con el nuevo empleado (por WhatsApp, SMS, verbal, etc.).")
    numbered(doc, "El empleado ingresa el codigo en su app para unirse al negocio.")
    numbered(doc, "El sistema crea la solicitud de ingreso que el admin puede aprobar.")

    h3(doc, "Reglas del Codigo de Invitacion")
    bullet(doc, "El codigo es de 6 digitos numericos aleatorios.")
    bullet(doc, "Expira automaticamente a las 24 horas de su generacion.")
    bullet(doc, "Se pueden tener multiples codigos activos simultaneamente.")
    bullet(doc, "Cada codigo es de un solo uso — una vez utilizado, se invalida.")
    bullet(doc, "El boton 'Generar codigo' se deshabilita mientras hay una generacion pendiente.")

    callout(doc, "Seguridad del codigo",
            "El codigo de invitacion solo permite al empleado solicitar unirse al negocio, no le da "
            "acceso automatico. El administrador debe aprobar la solicitud antes de que el empleado "
            "pueda operar en el negocio.",
            color=PURPLE)

    screenshot_placeholder(doc, "Dialog de invitacion con codigo generado de 6 digitos")

    # --- 12. Solicitudes ---
    h2(doc, "12. Solicitudes de Ingreso Pendientes", anchor="sec_emp_solicitudes")
    para(doc, "Cuando un empleado utiliza un codigo de invitacion, se genera una solicitud de ingreso "
         "que el administrador debe gestionar.")

    h3(doc, "Acceso")
    para(doc, "El boton 'Solicitudes' en el header del modulo de Empleados tiene un estado visual "
         "que indica si hay solicitudes pendientes (variant='default' cuando hay, 'outline' cuando no).")

    h3(doc, "Flujo de Aprobacion")
    numbered(doc, "Al hacer clic en 'Solicitudes', se muestra la lista de solicitudes pendientes.")
    numbered(doc, "Cada solicitud muestra: nombre del solicitante, email, fecha de solicitud y el codigo usado.")
    numbered(doc, "El administrador puede Aprobar o Rechazar cada solicitud.")
    numbered(doc, "Al aprobar: el empleado se registra en business_employees con status='approved' "
             "y recibe una notificacion.")
    numbered(doc, "Al rechazar: el solicitante recibe notificacion de rechazo.")

    # --- 13. Setup pendiente ---
    h2(doc, "13. Empleados Pendientes de Configuracion", anchor="sec_emp_setup")
    para(doc, "Cuando un empleado nuevo es aprobado pero aun no tiene un supervisor (jefe directo) "
         "asignado, el sistema muestra una alerta prominente.")

    h3(doc, "Alerta Visual")
    para(doc, "En la parte superior del modulo aparece una tarjeta con borde ambar y un icono de "
         "triangulo de alerta con el siguiente contenido:")
    bullet(doc, "Titulo: '{N} empleado(s) pendiente(s) de configuracion'.")
    bullet(doc, "Badge: 'Sin jefe directo' (ambar).")
    bullet(doc, "Texto explicativo: 'No apareceran disponibles para recibir citas hasta que se les "
           "asigne un jefe directo'.")

    h3(doc, "Accion de Asignar Jefe Directo")
    para(doc, "Por cada empleado pendiente se muestra:")
    bullet(doc, "Avatar e informacion del empleado (nombre y email).")
    bullet(doc, "Un select desplegable con la lista de supervisores validos (empleados de nivel "
           "jerarquico superior al del empleado pendiente).")
    bullet(doc, "Botones de confirmacion (Check verde) y cancelacion (X rojo).")

    callout(doc, "Regla de negocio critica",
            "Un empleado que no tiene jefe directo asignado NO aparecera disponible para recibir "
            "citas en el wizard de reserva. Es indispensable completar esta configuracion para que "
            "el empleado pueda operar normalmente.",
            color=DANGER)

    # --- 14. Perfil de empleado ---
    h2(doc, "14. Modal de Perfil de Empleado", anchor="sec_emp_perfil")
    para(doc, "Al hacer clic en 'Ver perfil' de cualquier empleado, se abre un modal completo "
         "con toda la informacion detallada del empleado.")

    h3(doc, "Contenido del Modal")
    bullet(doc, "Foto de perfil (avatar) o iniciales si no tiene foto.")
    bullet(doc, "Nombre completo, email y telefono.")
    bullet(doc, "Rol dentro del negocio (Admin, Profesional, Recepcionista, etc.).")
    bullet(doc, "Nivel jerarquico con badge de color.")
    bullet(doc, "Sede asignada.")
    bullet(doc, "Fecha de contratacion (hire_date).")
    bullet(doc, "Estado: Activo o Inactivo.")
    bullet(doc, "Lista de servicios que ofrece.")
    bullet(doc, "Horarios configurados por dia de la semana.")
    bullet(doc, "Informacion salarial (si tiene permisos para verla).")
    bullet(doc, "Balance de vacaciones (dias disponibles, usados, pendientes).")
    bullet(doc, "Rating promedio de reviews de clientes.")

    screenshot_placeholder(doc, "Modal de perfil de empleado con informacion completa")

    # --- 15. Horarios ---
    h2(doc, "15. Horarios del Empleado (7 Dias)", anchor="sec_emp_horarios")
    para(doc, "Cada empleado tiene un horario configurable para los 7 dias de la semana. "
         "Esta configuracion se realiza desde el panel de Configuraciones del empleado "
         "(CompleteUnifiedSettings, tab Preferencias de Empleado).")

    h3(doc, "Campos por cada dia")
    simple_table(doc, ["Campo", "Tipo", "Ejemplo", "Requerido"],
                 [["Dia activo", "Toggle (on/off)", "Lunes: activo", "Si"],
                  ["Hora inicio", "Time picker", "08:00", "Si (si activo)"],
                  ["Hora fin", "Time picker", "18:00", "Si (si activo)"]])

    h3(doc, "Reglas")
    bullet(doc, "Si un dia esta desactivado, el empleado no esta disponible para citas ese dia.")
    bullet(doc, "La hora de fin debe ser mayor que la hora de inicio.")
    bullet(doc, "Los horarios se validan contra los horarios de la sede asignada.")
    bullet(doc, "Los cambios se guardan inmediatamente al confirmar.")

    # --- 16. Almuerzo ---
    h2(doc, "16. Hora de Almuerzo del Empleado", anchor="sec_emp_almuerzo")
    para(doc, "Cada empleado puede tener configurada una hora de almuerzo que bloquea automaticamente "
         "los slots de cita durante ese periodo.")

    h3(doc, "Campos")
    simple_table(doc, ["Campo", "Tipo", "Ejemplo"],
                 [["Inicio almuerzo", "Time picker", "12:00"],
                  ["Fin almuerzo", "Time picker", "13:00"]])

    h3(doc, "Impacto en el Sistema")
    bullet(doc, "Los slots de cita durante la hora de almuerzo se muestran deshabilitados en el wizard de reserva.")
    bullet(doc, "Se muestra un tooltip 'Hora de almuerzo' en los slots bloqueados.")
    bullet(doc, "La hora de almuerzo NO aplica retroactivamente a citas historicas "
           "(para dias pasados, isLunchBreak retorna false).")
    bullet(doc, "Solo se almacena en business_employees.lunch_break_start y lunch_break_end.")

    # --- 17. Salario ---
    h2(doc, "17. Configuracion de Salario", anchor="sec_emp_salario")
    para(doc, "La configuracion salarial del empleado se gestiona desde el panel de Configuraciones "
         "y esta protegida por permisos (solo usuarios con employees.edit_salary pueden verla o editarla).")

    h3(doc, "Campos de Salario")
    simple_table(doc, ["Campo", "Tipo", "Descripcion"],
                 [["Salario base", "Input numerico", "Monto mensual en COP"],
                  ["Tipo de salario", "Select", "Fijo / Variable / Mixto"],
                  ["Comision", "Input porcentaje", "Porcentaje sobre ventas (si aplica)"],
                  ["Frecuencia de pago", "Select", "Quincenal / Mensual"]])

    callout(doc, "Permiso requerido",
            "Solo los usuarios con el permiso 'employees.edit_salary' pueden ver y editar la informacion "
            "salarial. Esto se controla a traves del sistema de permisos (ver seccion 48).",
            color=PURPLE)

    # --- 18. Servicios ---
    h2(doc, "18. Asignacion de Servicios al Empleado", anchor="sec_emp_servicios")
    para(doc, "Cada empleado puede ser asignado a uno o mas servicios del negocio. Solo los empleados "
         "asignados a un servicio aparecen como opciones en el wizard de reserva para ese servicio.")

    h3(doc, "Flujo de Asignacion")
    numbered(doc, "Acceda al perfil del empleado o a la seccion de asignaciones.")
    numbered(doc, "Se muestra la lista de todos los servicios del negocio con checkboxes.")
    numbered(doc, "Marque los servicios que el empleado puede realizar.")
    numbered(doc, "Confirme la asignacion.")
    numbered(doc, "El sistema guarda la relacion en la tabla employee_services.")
    numbered(doc, "Ahora, cuando un cliente seleccione ese servicio en el wizard de reserva, "
             "el empleado aparecera como opcion.")

    h3(doc, "Regla de negocio")
    bullet(doc, "Un empleado puede estar asignado a multiples servicios.")
    bullet(doc, "Un servicio puede tener multiples empleados asignados.")
    bullet(doc, "Si un empleado no esta asignado a ningun servicio, no aparecera en el wizard "
           "de reserva para ningun servicio.")
    bullet(doc, "La asignacion se almacena en la tabla employee_services con employee_id y service_id.")

    # --- 19. Activar/Desactivar ---
    h2(doc, "19. Activar y Desactivar Empleados", anchor="sec_emp_activar")
    para(doc, "El administrador puede activar o desactivar empleados en cualquier momento.")

    h3(doc, "Impacto de Desactivar")
    bullet(doc, "El empleado desactivado NO aparece en el wizard de reserva para nuevas citas.")
    bullet(doc, "Las citas existentes del empleado NO se cancelan automaticamente (deben gestionarse manualmente).")
    bullet(doc, "El empleado desactivado no puede acceder a su panel de empleado del negocio.")
    bullet(doc, "El empleado sigue existiendo en la base de datos (soft delete mediante is_active = false).")

    h3(doc, "Reactivacion")
    bullet(doc, "El administrador puede reactivar un empleado en cualquier momento.")
    bullet(doc, "Al reactivar, el empleado vuelve a ser visible en el wizard de reserva y puede acceder a su panel.")

    # --- 20. Limite ---
    h2(doc, "20. Limite de Empleados en Plan Basico (6)", anchor="sec_emp_limite")
    para(doc, "El Plan Basico permite gestionar hasta 6 empleados activos simultaneamente.")

    h3(doc, "Que ocurre al alcanzar el limite")
    bullet(doc, "Se muestra el banner PlanLimitBanner en la parte inferior del modulo con el texto: "
           "'{N} empleados no se muestran por limite del plan actual'.")
    bullet(doc, "El banner incluye un boton 'Actualizar plan' que navega a la pagina de facturacion.")
    bullet(doc, "Los empleados que exceden el limite NO se eliminan, simplemente no se muestran en la interfaz.")
    bullet(doc, "Al actualizar al Plan Pro (15 empleados), todos los empleados vuelven a ser visibles.")

    callout(doc, "Los datos estan seguros",
            "Aunque un empleado no sea visible por el limite del plan, sus datos permanecen intactos "
            "en la base de datos. Al actualizar el plan, toda la informacion estara disponible.",
            color=ACCENT)

    screenshot_placeholder(doc, "Banner de limite de empleados con boton de actualizacion")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: AUSENCIAS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "GESTION DE AUSENCIAS Y VACACIONES", anchor="sec_ausencias", color=PURPLE)
    para(doc, "El modulo de Ausencias permite al administrador gestionar las solicitudes de ausencia "
         "y vacaciones de todos los empleados del negocio. Este modulo es exclusivo del Plan Basico "
         "y superiores.")

    # --- 22. Politica ---
    h2(doc, "22. Politica de Aprobacion Obligatoria", anchor="sec_aus_politica")
    para(doc, "En Gestabiz, la aprobacion de ausencias es SIEMPRE obligatoria. Esta es una regla "
         "de negocio no negociable que se aplica a todos los negocios sin excepcion.")

    callout(doc, "REGLA CRITICA: Aprobacion siempre requerida",
            "Ningun empleado puede tomar una ausencia o vacacion sin la aprobacion previa de un "
            "administrador. Esta politica se implementa a nivel de base de datos con el campo "
            "require_absence_approval = true (no parametrizable). Aplica para todos los tipos "
            "de ausencia: vacaciones, emergencia, incapacidad, personal y otros.",
            color=DANGER)

    h3(doc, "Flujo Completo de una Solicitud")
    numbered(doc, "El empleado solicita una ausencia desde su panel (AbsenceRequestModal).")
    numbered(doc, "El sistema envia notificacion in-app y email a TODOS los administradores y gerentes del negocio.")
    numbered(doc, "La solicitud aparece en la pestana 'Pendientes' del modulo de Ausencias del admin.")
    numbered(doc, "El administrador revisa la solicitud, opcionalmente agrega notas, y aprueba o rechaza.")
    numbered(doc, "Si se aprueba: el balance de vacaciones se actualiza, las citas en el rango se cancelan "
             "automaticamente y el empleado recibe notificacion.")
    numbered(doc, "Si se rechaza: el empleado recibe notificacion con la razon (si se proporciono).")

    # --- 23. Pendientes ---
    h2(doc, "23. Pestana Pendientes — Aprobar o Rechazar", anchor="sec_aus_pendientes")
    para(doc, "La pestana 'Pendientes' muestra todas las solicitudes de ausencia que requieren "
         "la accion del administrador. El numero entre parentesis indica la cantidad de solicitudes pendientes.")

    h3(doc, "Estructura de la Vista")
    bullet(doc, "Tabs con 2 pestanas: 'Pendientes ({N})' y 'Historial ({N})'.")
    bullet(doc, "Grid de tarjetas (AbsenceApprovalCard) para cada solicitud pendiente.")
    bullet(doc, "Estado de carga: spinner animado centrado.")
    bullet(doc, "Estado vacio: contenedor con borde punteado y texto 'No hay solicitudes pendientes'.")

    screenshot_placeholder(doc, "Pestana de solicitudes de ausencia pendientes")

    # --- 24. Tarjeta de aprobacion ---
    h2(doc, "24. Tarjeta de Aprobacion de Ausencia — Detalle", anchor="sec_aus_card")
    para(doc, "Cada solicitud de ausencia se presenta en una tarjeta (card) con toda la informacion "
         "necesaria para que el administrador tome una decision informada.")

    h3(doc, "Informacion Mostrada en la Tarjeta")
    simple_table(doc, ["Elemento", "Icono", "Descripcion"],
                 [["Nombre empleado", "User", "Nombre completo del solicitante"],
                  ["Email empleado", "—", "Correo electronico del empleado"],
                  ["Tipo de ausencia", "(badge color)", "Vacaciones, Emergencia, Incapacidad, Personal u Otro"],
                  ["Rango de fechas", "Calendar", "Fecha inicio — Fecha fin"],
                  ["Duracion", "(badge)", "'{N} dias' calculado automaticamente"],
                  ["Fecha solicitud", "Clock", "Fecha y hora cuando se creo la solicitud"],
                  ["Razon", "(fondo gris)", "Texto de la razon proporcionada por el empleado"],
                  ["Notas del empleado", "(italica)", "Notas adicionales opcionales"],
                  ["Citas afectadas", "AlertCircle", "Cantidad de citas que seran canceladas (si aplica)"]])

    h3(doc, "Botones de Accion")
    para(doc, "Cada tarjeta tiene hasta 4 botones dependiendo del estado:")

    simple_table(doc, ["Boton", "Icono", "Color", "Accion"],
                 [["Agregar Nota", "—", "Outline (gris)", "Muestra textarea para notas del admin"],
                  ["Aprobar", "CheckCircle", "Verde (green-600)", "Aprueba la ausencia"],
                  ["Rechazar", "XCircle", "Rojo (destructive)", "Rechaza la ausencia"],
                  ["Cancelar", "—", "Outline (gris)", "Oculta el textarea de notas"]])

    h3(doc, "Textarea de Notas del Administrador")
    bullet(doc, "Label: 'Notas para el Empleado (opcional)'.")
    bullet(doc, "Placeholder: 'Comentarios adicionales...'.")
    bullet(doc, "2 filas de altura.")
    bullet(doc, "Las notas se envian junto con la decision al empleado.")

    h3(doc, "Estados de Carga de los Botones")
    bullet(doc, "Mientras se procesa la accion, el boton muestra 'Aprobando...' o 'Rechazando...' con spinner.")
    bullet(doc, "Todos los botones se deshabilitan durante el procesamiento para evitar doble envio.")

    # --- 25. Tipos ---
    h2(doc, "25. Tipos de Ausencia y Colores", anchor="sec_aus_tipos")
    para(doc, "El sistema soporta 5 tipos de ausencia, cada uno con un color distintivo en la interfaz:")

    simple_table(doc, ["Tipo", "Label en Espanol", "Color del Badge", "Uso Tipico"],
                 [["vacation", "Vacaciones", "Azul", "Descanso planificado"],
                  ["emergency", "Emergencia", "Rojo", "Situacion imprevista urgente"],
                  ["sick_leave", "Incapacidad", "Amarillo", "Enfermedad o certificado medico"],
                  ["personal", "Personal", "Morado", "Asuntos personales"],
                  ["other", "Otro", "Gris", "Cualquier otra razon"]])

    # --- 26. Citas afectadas ---
    h2(doc, "26. Citas Afectadas por una Ausencia", anchor="sec_aus_citas")
    para(doc, "Cuando un empleado solicita una ausencia, el sistema calcula automaticamente "
         "cuantas citas activas (no canceladas) caen dentro del rango de fechas solicitado.")

    h3(doc, "Visualizacion")
    bullet(doc, "Si hay citas afectadas, se muestra una tarjeta con fondo amarillo y un icono AlertCircle.")
    bullet(doc, "El texto indica: '{N} cita(s) sera(n) cancelada(s)'.")
    bullet(doc, "Subtexto: 'Los clientes recibiran notificacion por email y en la app'.")

    h3(doc, "Que ocurre al aprobar con citas afectadas")
    numbered(doc, "Al aprobar la ausencia, el sistema invoca automaticamente la Edge Function "
             "'cancel-appointments-on-emergency-absence'.")
    numbered(doc, "Todas las citas no canceladas del empleado en el rango de fechas se cancelan.")
    numbered(doc, "Cada cliente afectado recibe notificacion por email y notificacion in-app.")
    numbered(doc, "Los slots de cita del empleado se bloquean automaticamente en el wizard de reserva "
             "para el rango de la ausencia aprobada.")

    callout(doc, "Cancelacion automatica",
            "Las citas se cancelan automaticamente al aprobar la ausencia, no al solicitarla. "
            "Esto permite al administrador evaluar el impacto antes de tomar la decision.",
            color=PURPLE)

    # --- 27. Historial ---
    h2(doc, "27. Pestana Historial de Ausencias", anchor="sec_aus_historial")
    para(doc, "La pestana 'Historial' muestra todas las solicitudes de ausencia ya procesadas "
         "(aprobadas y rechazadas) en formato de grid de tarjetas.")

    bullet(doc, "Las tarjetas son identicas a las de la pestana Pendientes pero sin botones de accion.")
    bullet(doc, "Se muestra el estado final (Aprobada o Rechazada) con el color correspondiente.")
    bullet(doc, "Se incluyen las notas del administrador si las hubo.")
    bullet(doc, "Estado vacio: 'No hay historial de ausencias'.")

    # --- 28. Widget vacaciones ---
    h2(doc, "28. Widget de Balance de Vacaciones", anchor="sec_aus_vacaciones")
    para(doc, "El widget de balance de vacaciones (VacationDaysWidget) muestra el estado actual "
         "del derecho a vacaciones de cada empleado.")

    h3(doc, "Elementos del Widget")
    bullet(doc, "Header: icono Calendar + 'Vacaciones {ano}' + badge '15 dias totales' (default Colombia).")
    bullet(doc, "Numero prominente en el centro: dias disponibles restantes (texto 4xl azul).")
    bullet(doc, "Barra de progreso tricolor: verde (usados), amarillo (pendientes), azul (disponibles).")
    bullet(doc, "Grid de 3 columnas con iconos:")

    simple_table(doc, ["Metrica", "Icono", "Color", "Descripcion"],
                 [["Usados", "CheckCircle", "Verde", "Dias de vacaciones ya tomados"],
                  ["Pendientes", "Clock", "Amarillo", "Dias solicitados pendientes de aprobacion"],
                  ["Libres", "Calendar", "Azul", "Dias disponibles para solicitar"]])

    h3(doc, "Reglas de Calculo")
    bullet(doc, "Los 15 dias son el default de vacaciones por ano en Colombia.")
    bullet(doc, "El balance se calcula automaticamente: Total - Usados - Pendientes = Libres.")
    bullet(doc, "Los festivos publicos se excluyen del calculo de dias de ausencia.")
    bullet(doc, "El widget se actualiza en tiempo real al aprobar o rechazar ausencias.")

    # --- 29. Festivos ---
    h2(doc, "29. Festivos Publicos y su Impacto", anchor="sec_aus_festivos")
    para(doc, "El sistema tiene precargados 54 festivos colombianos para los anos 2025-2027, "
         "que impactan tanto el calculo de vacaciones como la disponibilidad de citas.")

    h3(doc, "Impacto de los Festivos")
    bullet(doc, "En el calculo de ausencias: los festivos se excluyen del conteo de dias de ausencia. "
           "Si una ausencia de 5 dias incluye 1 festivo, solo se cuentan 4 dias de vacaciones.")
    bullet(doc, "En el wizard de reserva: los slots de cita en dias festivos se bloquean automaticamente.")
    bullet(doc, "En el calendario del admin: los dias festivos se marcan con indicador visual.")

    h3(doc, "Tipos de Festivos")
    bullet(doc, "Fijos (13 por ano): Ano Nuevo, Dia del Trabajo, Independencia, Navidad, etc.")
    bullet(doc, "Moviles (5 por ano): basados en la fecha de Pascua — Carnaval, Semana Santa, Corpus Christi, etc.")

    # --- 30. Notificaciones ---
    h2(doc, "30. Notificaciones de Ausencias", anchor="sec_aus_notificaciones")
    para(doc, "El sistema genera notificaciones automaticas en los siguientes eventos de ausencias:")

    simple_table(doc, ["Evento", "Destinatarios", "Canales", "Contenido"],
                 [["Solicitud creada", "TODOS los admins/gerentes", "In-app + Email", "Nombre empleado, tipo, rango, razon"],
                  ["Ausencia aprobada", "Empleado solicitante", "In-app + Email", "Confirmacion + notas del admin"],
                  ["Ausencia rechazada", "Empleado solicitante", "In-app + Email", "Razon del rechazo + notas"],
                  ["Citas canceladas", "Clientes afectados", "In-app + Email", "Empleado no disponible, reagendar"]])

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: VENTAS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "HISTORIAL DE VENTAS", anchor="sec_ventas", color=PURPLE)
    para(doc, "El modulo de Historial de Ventas permite al administrador consultar todas las citas "
         "completadas con sus ingresos asociados. Es una herramienta clave para hacer seguimiento "
         "de la facturacion y el rendimiento del negocio.")

    # --- 32. Resumen ---
    h2(doc, "32. Tarjetas de Resumen de Ventas", anchor="sec_ventas_resumen")
    para(doc, "En la parte superior del modulo se muestran 3 tarjetas con metricas clave del periodo seleccionado:")

    simple_table(doc, ["Tarjeta", "Icono", "Formato", "Descripcion"],
                 [["Citas completadas", "CheckCircle2", "Numero entero", "Total de citas con status=completed"],
                  ["Ingresos en periodo", "DollarSign", "COP (ej. $1.250.000)", "Suma de precios de todas las citas completadas"],
                  ["Promedio por cita", "TrendingUp", "COP (ej. $45.000)", "Division de ingresos / cantidad"]])

    # --- 33. Filtros ---
    h2(doc, "33. Filtros de Periodo y Busqueda", anchor="sec_ventas_filtros")
    para(doc, "El modulo ofrece dos controles de filtrado:")

    h3(doc, "Filtro de Periodo")
    simple_table(doc, ["Opcion", "Rango", "Default"],
                 [["Ultimos 7 dias", "7 dias atras — hoy", "No"],
                  ["Ultimos 30 dias", "30 dias atras — hoy", "Si (default)"],
                  ["Ultimos 90 dias", "90 dias atras — hoy", "No"],
                  ["Ultimo ano", "365 dias atras — hoy", "No"]])

    h3(doc, "Campo de Busqueda")
    bullet(doc, "Icono de lupa (Search) a la izquierda.")
    bullet(doc, "Placeholder: 'Buscar cliente o servicio...'.")
    bullet(doc, "Filtra en tiempo real por nombre de cliente o nombre de servicio.")
    bullet(doc, "La busqueda es local (client-side) sobre los datos ya cargados.")

    # --- 34. Lista ---
    h2(doc, "34. Lista de Ventas Completadas", anchor="sec_ventas_lista")
    para(doc, "La lista principal muestra cada venta completada como una tarjeta individual con "
         "la siguiente informacion:")

    bullet(doc, "Fecha de la cita (dia y mes abreviado) en la columna izquierda.")
    bullet(doc, "Separador vertical visual.")
    bullet(doc, "Nombre del servicio realizado + hora formateada.")
    bullet(doc, "Boton del cliente: avatar con iniciales + nombre (clickeable para abrir perfil).")
    bullet(doc, "Precio del servicio en formato COP.")

    h3(doc, "Limite de Carga")
    para(doc, "El sistema carga un maximo de 500 citas completadas por consulta. Para negocios "
         "con mayor volumen, se recomienda usar filtros de periodo mas cortos.")

    h3(doc, "Estados")
    bullet(doc, "Cargando: spinner (LoadingSpinner) centrado.")
    bullet(doc, "Sin resultados con busqueda: icono CheckCircle2 + 'No se encontraron resultados para la busqueda'.")
    bullet(doc, "Sin ventas en periodo: icono CheckCircle2 + 'No hay ventas completadas en este periodo'.")

    # --- 35. Perfil cliente ---
    h2(doc, "35. Ver Perfil de Cliente desde Ventas", anchor="sec_ventas_cliente")
    para(doc, "Al hacer clic en el nombre de un cliente en la lista de ventas, se abre el "
         "ClientProfileModal con toda la informacion del cliente y su historial de citas.")

    para(doc, "El modal muestra dos pestanas: 'Informacion' (datos de contacto, estadisticas de visitas, "
         "fecha de primera y ultima visita) e 'Historial ({N})' (lista de todas las citas del cliente "
         "con servicio, fecha, estado y precio). Para mas detalle sobre este modal, consulte la "
         "Parte 3 del manual — seccion 33.")

    screenshot_placeholder(doc, "Historial de ventas con tarjetas de resumen y lista de citas")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: VENTAS RAPIDAS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "VENTAS RAPIDAS (WALK-IN)", anchor="sec_rapidas", color=PURPLE)
    para(doc, "El modulo de Ventas Rapidas permite registrar ventas de clientes que llegan al negocio "
         "sin cita previa (walk-in). Es una herramienta esencial para capturar ingresos que "
         "de otro modo no quedarian registrados en el sistema.")

    # --- 37. Formulario ---
    h2(doc, "37. Formulario de Venta Rapida — Campos", anchor="sec_rapidas_form")
    para(doc, "El formulario de registro de venta rapida contiene los siguientes campos:")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Icono", "Validacion"],
                 [["Nombre del Cliente", "Input texto", "Si", "User", "HTML required"],
                  ["Telefono", "Input tel + prefijo", "No", "—", "Selector de prefijo pais"],
                  ["Documento", "Input texto", "No", "—", "Libre"],
                  ["Correo Electronico", "Input email", "No", "—", "Validacion type=email"],
                  ["Servicio", "Select desplegable", "Si", "Package", "HTML required"],
                  ["Sede", "Select desplegable", "Si", "MapPin", "HTML required"],
                  ["Empleado que atendio", "Select desplegable", "No", "—", "Opcional"],
                  ["Monto Pagado (COP)", "Input numerico", "Si", "—", "min=0, step=100"],
                  ["Metodo de Pago", "Select desplegable", "Si", "CreditCard", "3 opciones"],
                  ["Notas", "Textarea (2 filas)", "No", "—", "Libre"]])

    h3(doc, "Comportamiento Especial del Campo Sede")
    para(doc, "El campo de sede utiliza un sistema de doble cache para pre-seleccionar automaticamente:")
    numbered(doc, "Primero busca en localStorage la ultima sede usada para ventas rapidas "
             "(clave: quick-sale-location-{businessId}).")
    numbered(doc, "Si no encuentra, usa la sede preferida del negocio (usePreferredLocation).")
    numbered(doc, "Al cambiar de sede, el valor se guarda en ambos caches para la proxima venta.")

    # --- 38. Metodos ---
    h2(doc, "38. Metodos de Pago Disponibles", anchor="sec_rapidas_metodos")
    simple_table(doc, ["Valor", "Icono", "Label", "Descripcion"],
                 [["cash", "Banknote", "Efectivo", "Pago en billetes o monedas"],
                  ["card", "CreditCard", "Tarjeta", "Pago con tarjeta debito o credito"],
                  ["transfer", "Landmark", "Transferencia", "Transferencia bancaria o Nequi/Daviplata"]])

    # --- 39. Autocompletado ---
    h2(doc, "39. Autocompletado de Monto por Servicio", anchor="sec_rapidas_auto")
    para(doc, "Cuando el administrador selecciona un servicio en el formulario, el sistema "
         "autocompleta automaticamente el campo 'Monto Pagado' con el precio configurado "
         "de ese servicio. El administrador puede modificar el monto si el precio real "
         "fue diferente (descuento, propina, etc.).")

    # --- 40. Stats ---
    h2(doc, "40. Estadisticas de Ventas Rapidas", anchor="sec_rapidas_stats")
    para(doc, "En la parte superior de la pagina se muestran 3 tarjetas con estadisticas "
         "actualizadas en tiempo real:")

    simple_table(doc, ["Tarjeta", "Periodo", "Formato"],
                 [["Hoy", "Solo transacciones de hoy", "COP (ej. $350.000)"],
                  ["7 dias", "Ultimos 7 dias", "COP (ej. $1.500.000)"],
                  ["30 dias", "Ultimos 30 dias", "COP (ej. $5.200.000)"]])

    para(doc, "Estas estadisticas se recalculan automaticamente cada vez que se registra una nueva venta "
         "o se recarga la pagina.")

    # --- 41. Historial ---
    h2(doc, "41. Historial de Ultimas 10 Ventas", anchor="sec_rapidas_historial")
    para(doc, "Debajo del formulario se muestra una lista con las ultimas 10 ventas rapidas registradas. "
         "Cada venta muestra:")

    bullet(doc, "Nombre del cliente.")
    bullet(doc, "Telefono (con icono Phone) si fue proporcionado.")
    bullet(doc, "Documento (con icono IdCard) si fue proporcionado.")
    bullet(doc, "Email (con icono Mail) si fue proporcionado.")
    bullet(doc, "Notas (con icono PencilLine) si fueron proporcionadas.")
    bullet(doc, "Metodo de pago con icono correspondiente (Banknote, CreditCard o Landmark).")
    bullet(doc, "Monto y fecha de la transaccion.")

    h3(doc, "Estados")
    bullet(doc, "Cargando: spinner animado.")
    bullet(doc, "Sin ventas: icono Package + 'No hay ventas registradas aun' + subtexto motivacional.")

    screenshot_placeholder(doc, "Formulario de venta rapida con estadisticas y historial")

    # --- 42. Integracion contable ---
    h2(doc, "42. Integracion Contable de Ventas Rapidas", anchor="sec_rapidas_contable")
    para(doc, "Cada venta rapida registrada se almacena automaticamente como una transaccion contable "
         "en la tabla transactions con los siguientes valores:")

    simple_table(doc, ["Campo", "Valor"],
                 [["type", "income"],
                  ["category", "service_sale"],
                  ["is_verified", "true"],
                  ["verified_by", "ID del admin que registro la venta"],
                  ["metadata (JSONB)", "client_name, client_phone, client_document, client_email, service_id, notes, source='quick_sale'"]])

    para(doc, "Esta integracion permite que las ventas rapidas aparezcan automaticamente en el "
         "modulo de Reportes Financieros y en cualquier exportacion contable futura (Plan Pro).")

    callout(doc, "Proteccion por permisos",
            "El boton de envio del formulario esta protegido por PermissionGate con el permiso "
            "'sales.create'. Solo usuarios con este permiso pueden registrar ventas rapidas.",
            color=PURPLE)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: REPORTES
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "REPORTES FINANCIEROS", anchor="sec_reportes", color=PURPLE)
    para(doc, "El modulo de Reportes ofrece un dashboard financiero interactivo con graficos "
         "que permiten visualizar los ingresos, gastos y tendencias del negocio. "
         "En el Plan Basico, los reportes estan disponibles en modo de visualizacion "
         "(sin exportacion a PDF/CSV/Excel, que es exclusiva del Plan Pro).")

    # --- 44. Dashboard ---
    h2(doc, "44. Dashboard Financiero Interactivo", anchor="sec_rep_dashboard")
    para(doc, "El dashboard financiero (EnhancedFinancialDashboard) se carga de forma diferida "
         "(lazy loading) para optimizar el rendimiento. Incluye tarjetas de resumen, graficos "
         "interactivos y filtros avanzados.")

    h3(doc, "Acceso")
    para(doc, "El acceso al dashboard esta protegido por PermissionGate con el permiso "
         "'reports.view_financial'. Usuarios sin este permiso veran un mensaje de acceso denegado.")

    h3(doc, "Tarjetas de Resumen Financiero")
    para(doc, "En la parte superior del dashboard se muestran tarjetas con metricas clave:")
    bullet(doc, "Ingresos totales del periodo seleccionado.")
    bullet(doc, "Gastos totales del periodo (si aplica).")
    bullet(doc, "Balance neto (ingresos - gastos).")
    bullet(doc, "Cantidad de transacciones.")

    # --- 45. Graficos ---
    h2(doc, "45. Graficos Disponibles", anchor="sec_rep_graficos")
    para(doc, "El dashboard incluye 5 tipos de graficos interactivos (implementados con la libreria Recharts):")

    simple_table(doc, ["Grafico", "Componente", "Tipo", "Que Muestra"],
                 [["Ingresos vs Gastos", "IncomeVsExpenseChart", "Barras agrupadas", "Comparacion mensual de ingresos y gastos"],
                  ["Tendencia Mensual", "MonthlyTrendChart", "Linea", "Evolucion de ingresos a lo largo del tiempo"],
                  ["Distribucion por Categoria", "CategoryPieChart", "Torta/Pie", "Proporcion de ingresos por tipo de servicio"],
                  ["Ingresos por Empleado", "EmployeeRevenueChart", "Barras horizontales", "Ranking de ingresos generados por cada empleado"],
                  ["Ingresos por Sede", "LocationBarChart", "Barras verticales", "Comparacion de ingresos entre sedes"]])

    para(doc, "Cada grafico es interactivo: al pasar el cursor sobre un elemento se muestra un tooltip "
         "con el detalle numerico. Algunos graficos permiten hacer clic para filtrar los datos.")

    screenshot_placeholder(doc, "Dashboard financiero con graficos de ingresos y distribucion")

    # --- 46. Filtros ---
    h2(doc, "46. Filtros del Dashboard Financiero", anchor="sec_rep_filtros")
    para(doc, "El dashboard ofrece multiples opciones de filtrado:")

    bullet(doc, "Periodo de tiempo: 7 dias, 30 dias, 90 dias, 12 meses, personalizado.")
    bullet(doc, "Sede: filtrar por sede especifica o todas las sedes.")
    bullet(doc, "Categoria de servicio: filtrar por tipo de servicio.")
    bullet(doc, "Empleado: filtrar por empleado especifico.")
    bullet(doc, "Tipo de transaccion: ingresos, gastos, o ambos.")

    para(doc, "Los filtros se aplican en tiempo real y actualizan todos los graficos y tarjetas "
         "simultaneamente sin recargar la pagina.")

    # --- 47. Exportacion ---
    h2(doc, "47. Exportacion de Reportes (Nota de Plan)", anchor="sec_rep_export")

    callout(doc, "Funcion exclusiva del Plan Pro",
            "La exportacion de reportes a PDF, CSV y Excel NO esta disponible en el Plan Basico. "
            "El dashboard muestra la informacion completa en pantalla, pero para descargar los "
            "datos en formato de archivo, es necesario actualizar al Plan Pro ($159,900 COP/mes). "
            "El subtitulo del modulo menciona 'exportacion a PDF/CSV/Excel' como referencia a la "
            "funcionalidad completa disponible en el plan superior.",
            color=PURPLE)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULO: PERMISOS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "GESTION DE PERMISOS", anchor="sec_permisos", color=PURPLE)
    para(doc, "El modulo de Permisos permite al administrador controlar que acciones pueden realizar "
         "los empleados dentro de la plataforma. En el Plan Basico, la gestion se realiza mediante "
         "3 plantillas predefinidas de permisos. La asignacion granular de permisos individuales "
         "(79 tipos disponibles) es exclusiva del Plan Pro.")

    # --- 49. Tabla de usuarios ---
    h2(doc, "49. Tabla de Usuarios y Roles", anchor="sec_perm_usuarios")
    para(doc, "La pestana principal del modulo muestra una tabla con todos los usuarios que tienen "
         "acceso al negocio (administradores y empleados).")

    h3(doc, "Columnas de la Tabla")
    simple_table(doc, ["Columna", "Contenido"],
                 [["Usuario", "Avatar + nombre + badge 'Owner' (si aplica) + email"],
                  ["Rol", "Badge: Admin (default) o Empleado (secondary)"],
                  ["Tipo", "Badge outline: 'Presta servicios' o 'Staff soporte'"],
                  ["Permisos", "Icono Shield + cantidad de permisos activos (o 'Todos' si es owner)"],
                  ["Estado", "Badge: 'Activo' (verde) o 'Inactivo' (rojo)"],
                  ["Acciones", "Botones: Editar permisos + Eliminar permisos"]])

    h3(doc, "Filtros de la Tabla")
    bullet(doc, "Input de busqueda con icono Search: 'Buscar por nombre o email...'.")
    bullet(doc, "Select de rol: 'Todos los roles' / 'Administradores' / 'Empleados'.")

    h3(doc, "Tarjetas de Estadisticas")
    para(doc, "Debajo de la tabla se muestran 3 tarjetas:")
    simple_table(doc, ["Tarjeta", "Icono", "Dato"],
                 [["Total Usuarios", "Users", "Cantidad total"],
                  ["Administradores", "Crown", "Cantidad de admins"],
                  ["Empleados", "UserCheck", "Cantidad de empleados"]])

    h3(doc, "Estados")
    bullet(doc, "Cargando: texto centrado 'Cargando usuarios...'.")
    bullet(doc, "Sin resultados: texto centrado 'No se encontraron usuarios'.")

    screenshot_placeholder(doc, "Tabla de usuarios con roles y acciones en el modulo de permisos")

    # --- 50. Acciones ---
    h2(doc, "50. Acciones sobre Permisos", anchor="sec_perm_acciones")
    para(doc, "Desde la tabla de usuarios se pueden realizar las siguientes acciones:")

    h3(doc, "Editar Permisos de un Usuario")
    bullet(doc, "Icono: Edit (lapiz).")
    bullet(doc, "Protegido por PermissionGate: 'permissions.edit' (mode=hide).")
    bullet(doc, "Abre el editor de permisos (PermissionEditor) para el usuario seleccionado.")
    bullet(doc, "En el Plan Basico, el editor permite aplicar una de las 3 plantillas predefinidas.")

    h3(doc, "Eliminar Permisos de un Usuario")
    bullet(doc, "Icono: Trash2 (papelera roja).")
    bullet(doc, "Protegido por PermissionGate: 'permissions.delete' (mode=hide).")
    bullet(doc, "Solo disponible para usuarios que NO son el owner del negocio.")
    bullet(doc, "Muestra un AlertDialog de confirmacion antes de proceder.")
    bullet(doc, "Al confirmar: marca todos los permisos del usuario como inactivos (is_active = false).")
    bullet(doc, "No elimina al usuario del negocio, solo revoca sus permisos.")
    bullet(doc, "Toast de confirmacion o error segun el resultado.")

    h3(doc, "Asignar Rol a Nuevo Usuario")
    bullet(doc, "Boton en el header: 'Asignar Rol' con icono UserPlus.")
    bullet(doc, "Protegido por PermissionGate: 'permissions.assign_role' (mode=hide).")
    bullet(doc, "Abre el componente RoleAssignment para buscar y asignar rol a un usuario existente.")

    callout(doc, "Proteccion del Owner",
            "El propietario (owner) del negocio tiene TODOS los permisos por defecto (bypass total) "
            "y no puede ser editado ni eliminado. Su columna de permisos muestra 'Todos' en lugar "
            "de un numero. No aparece el boton de eliminar para el owner.",
            color=ACCENT)

    # --- 51. Templates ---
    h2(doc, "51. Plantillas de Permisos (3 Predefinidas)", anchor="sec_perm_templates")
    para(doc, "En el Plan Basico, los permisos se gestionan mediante plantillas predefinidas que agrupan "
         "conjuntos de permisos comunes para roles tipicos de un negocio de servicios.")

    simple_table(doc, ["Plantilla", "Permisos Incluidos", "Ideal Para"],
                 [["Recepcionista", "appointments.create, appointments.edit, appointments.cancel, "
                   "clients.view, services.view", "Personal de recepcion que gestiona citas"],
                  ["Profesional", "appointments.view, appointments.edit, appointments.cancel_own, "
                   "reviews.view, services.view, absences.request", "Empleados que prestan servicios directamente"],
                  ["Contador", "accounting.view_reports, accounting.create, accounting.edit, "
                   "expenses.view, billing.view", "Personal encargado de la contabilidad del negocio"]])

    h3(doc, "Aplicar una Plantilla")
    numbered(doc, "En el editor de permisos del usuario, seleccione la plantilla deseada.")
    numbered(doc, "El sistema muestra un preview de los permisos que se asignaran.")
    numbered(doc, "Confirme la aplicacion.")
    numbered(doc, "Los permisos de la plantilla se asignan al usuario via el servicio PermissionRPCService.applyTemplate.")
    numbered(doc, "Toast de confirmacion con el nombre de la plantilla aplicada.")

    callout(doc, "Permisos granulares — Plan Pro",
            "Para asignar permisos individuales de los 79 tipos disponibles (en lugar de usar "
            "plantillas predefinidas), es necesario actualizar al Plan Pro. El Plan Pro tambien "
            "incluye 9 plantillas (en lugar de 3) y la pestaña de auditoria de cambios.",
            color=PURPLE)

    # --- 52. Stats ---
    h2(doc, "52. Estadisticas de Usuarios y Roles", anchor="sec_perm_stats")
    para(doc, "Debajo de la tabla de usuarios se muestran 3 tarjetas estadisticas que brindan "
         "una vision rapida de la composicion del equipo:")
    bullet(doc, "Total Usuarios (icono Users): cuenta total de personas con acceso al negocio.")
    bullet(doc, "Administradores (icono Crown): cantidad de usuarios con rol de administrador.")
    bullet(doc, "Empleados (icono UserCheck): cantidad de usuarios con rol de empleado.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # LIMITES AMPLIADOS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "LIMITES AMPLIADOS EN MODULOS EXISTENTES", anchor="sec_limites_ampliados", color=PURPLE)
    para(doc, "Ademas de los 6 modulos exclusivos, el Plan Basico amplia significativamente los "
         "limites de los modulos que ya estaban disponibles en el Plan Gratuito.")

    # --- 54. Sedes ---
    h2(doc, "54. Gestion de hasta 3 Sedes", anchor="sec_sedes_3")
    para(doc, "El Plan Gratuito permite unicamente 1 sede. Con el Plan Basico, el negocio puede "
         "gestionar hasta 3 sedes simultaneamente.")

    h3(doc, "Que cambia con multiples sedes")
    bullet(doc, "El boton 'Agregar sede' en el modulo de Sedes se habilita hasta que se alcancen las 3.")
    bullet(doc, "Cada sede puede tener su propia direccion, horarios de apertura/cierre, telefono, "
           "email y coordenadas GPS.")
    bullet(doc, "Los servicios se pueden asignar de forma independiente a cada sede.")
    bullet(doc, "Los empleados se asignan a una sede principal pero pueden trabajar en otras "
           "(usando is_location_exception en citas).")
    bullet(doc, "El filtro de sede aparece en multiples modulos para facilitar la gestion multi-sede: "
           "Calendario, Empleados, Ventas Rapidas, Reportes.")
    bullet(doc, "La sede preferida se puede configurar en Configuraciones para pre-seleccionar "
           "automaticamente en los filtros.")

    h3(doc, "Banner de limite")
    para(doc, "Al alcanzar las 3 sedes, el sistema muestra el PlanLimitBanner indicando que se alcanzo "
         "el limite y ofreciendo la opcion de actualizar al Plan Pro (hasta 10 sedes).")

    # --- 55. Servicios ---
    h2(doc, "55. Servicios Ilimitados", anchor="sec_servicios_ilim")
    para(doc, "El Plan Gratuito limita a 15 servicios. Con el Plan Basico, no hay limite en la cantidad "
         "de servicios que un negocio puede crear y gestionar.")

    bullet(doc, "Se elimina el banner PlanLimitBanner del modulo de Servicios.")
    bullet(doc, "Todos los servicios creados son visibles y funcionales sin restriccion.")
    bullet(doc, "La gestion de servicios (crear, editar, eliminar, reactivar, asignar a sedes y empleados) "
           "funciona identicamente a como se describe en la Parte 3 del manual.")

    # --- 56. Citas ---
    h2(doc, "56. Citas Ilimitadas por Mes", anchor="sec_citas_ilim")
    para(doc, "El Plan Gratuito limita a 50 citas por mes. Con el Plan Basico, los clientes pueden "
         "agendar citas ilimitadas sin restriccion mensual.")

    bullet(doc, "El wizard de reserva no muestra ningun mensaje de limite alcanzado.")
    bullet(doc, "El calendario del admin muestra todas las citas sin restriccion.")
    bullet(doc, "Las estadisticas del dashboard reflejan el volumen real de citas sin truncar.")

    # --- 57. Clientes ---
    h2(doc, "57. Clientes Ilimitados", anchor="sec_clientes_ilim")
    para(doc, "El Plan Gratuito limita a 50 clientes visibles en el CRM. Con el Plan Basico, "
         "todos los clientes son visibles sin restriccion.")

    bullet(doc, "El modulo ClientsManager muestra todos los clientes con al menos una cita no cancelada.")
    bullet(doc, "Se elimina el banner PlanLimitBanner del modulo de Clientes.")
    bullet(doc, "La busqueda y el modal de perfil de cliente funcionan sobre la totalidad de los datos.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # FUNCIONALIDADES ADICIONALES DE PAGO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "FUNCIONALIDADES ADICIONALES DE PLANES DE PAGO", anchor="sec_features_pago", color=PURPLE)
    para(doc, "Al ser un plan de pago, el Plan Basico habilita funcionalidades que no estan "
         "disponibles en ningun plan gratuito.")

    h2(doc, "59. Recordatorios por WhatsApp", anchor="sec_whatsapp")
    para(doc, "Los clientes pueden recibir recordatorios automaticos de sus citas a traves de "
         "WhatsApp Business API. Esto reduce significativamente las inasistencias (no-shows).")

    bullet(doc, "Los recordatorios se envian en los intervalos configurados por el negocio "
           "(por ejemplo, 24h y 1h antes de la cita).")
    bullet(doc, "El contenido del mensaje incluye: nombre del negocio, servicio, fecha, hora, "
           "direccion de la sede y enlace para confirmar o cancelar.")
    bullet(doc, "La configuracion se realiza desde Configuraciones > Notificaciones > Canal WhatsApp.")
    bullet(doc, "Requisito: el negocio debe tener configuradas las credenciales de WhatsApp Business API.")

    h2(doc, "60. Sincronizacion con Google Calendar", anchor="sec_gcal")
    para(doc, "El administrador y los empleados pueden sincronizar sus citas de Gestabiz con "
         "Google Calendar para tener una vista unificada de su agenda.")

    bullet(doc, "Sincronizacion bidireccional: las citas creadas en Gestabiz aparecen en Google Calendar "
           "y viceversa.")
    bullet(doc, "Autenticacion via OAuth 2.0 con cuenta de Google.")
    bullet(doc, "Seleccion de calendario especifico para la sincronizacion.")
    bullet(doc, "Colores y descripciones automaticas segun el tipo de cita.")
    bullet(doc, "Configuracion desde Configuraciones > Integraciones > Google Calendar.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # MODULOS BLOQUEADOS
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "MODULOS BLOQUEADOS — SOLO EN PLAN PRO", anchor="sec_modulos_pro", color=PURPLE)
    para(doc, "Los siguientes modulos NO estan disponibles en el Plan Basico y requieren "
         "una actualizacion al Plan Pro ($159,900 COP/mes):")

    simple_table(doc, ["Modulo", "Icono", "Descripcion", "Precio Pro"],
                 [["Gastos (Expenses)", "Wallet", "Gestion de gastos fijos y recurrentes del negocio", "$159,900/mes"],
                  ["Reclutamiento", "BriefcaseBusiness", "Publicar vacantes, recibir aplicaciones, matching inteligente", "$159,900/mes"],
                  ["Recursos Fisicos", "Box", "Gestionar salas, mesas, canchas y otros recursos reservables", "$159,900/mes"],
                  ["Exportacion de Reportes", "Download", "Descargar reportes financieros en PDF, CSV y Excel", "$159,900/mes"],
                  ["Permisos Granulares", "Shield", "Asignar 79 permisos individuales + 9 plantillas + auditoria", "$159,900/mes"]])

    para(doc, "Al intentar acceder a cualquiera de estos modulos, el sistema muestra la pantalla "
         "de PlanGate con el icono de candado y un boton para ver los planes disponibles.")

    h2(doc, "62. Como Actualizar al Plan Pro", anchor="sec_upgrade_pro")
    numbered(doc, "Navegue a Facturacion en la barra lateral.")
    numbered(doc, "Haga clic en 'Cambiar plan' o 'Actualizar'.")
    numbered(doc, "Seleccione el Plan Pro ($159,900 COP/mes o $1,599,000 COP/ano).")
    numbered(doc, "Complete el proceso de pago con la pasarela configurada.")
    numbered(doc, "Los modulos se desbloquean inmediatamente tras la confirmacion del pago.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # SOPORTE Y GLOSARIO
    # ══════════════════════════════════════════════════════════════════════

    h2(doc, "63. Soporte y Ayuda", anchor="sec_soporte")
    para(doc, "Los usuarios del Plan Basico cuentan con soporte prioritario por email. "
         "Para reportar errores, la app incluye un boton flotante de reporte de bugs "
         "(FloatingBugReportButton) que permite enviar reportes con capturas de pantalla "
         "y descripcion detallada del problema.")

    bullet(doc, "Email de soporte: soporte@gestabiz.com.")
    bullet(doc, "Tiempo de respuesta estimado: 24-48 horas habiles.")
    bullet(doc, "Reporte de bugs: boton flotante en la esquina inferior derecha de la aplicacion.")
    bullet(doc, "Severidades de bug: Critico, Alto, Medio, Bajo.")

    h2(doc, "64. Glosario de Terminos del Plan Basico", anchor="sec_glosario_basico")

    simple_table(doc, ["Termino", "Definicion"],
                 [["Walk-in", "Cliente que llega al negocio sin cita previa"],
                  ["PlanGate", "Pantalla de bloqueo que aparece al intentar acceder a un modulo no incluido en el plan"],
                  ["PlanLimitBanner", "Banner informativo que indica cuando se ha alcanzado el limite del plan"],
                  ["Jerarquia", "Estructura organizacional del equipo con niveles (Owner > Admin > Gerente > Lider > Personal)"],
                  ["Plantilla de permisos", "Conjunto predefinido de permisos agrupados para un rol comun"],
                  ["PermissionGate", "Componente invisible que oculta o deshabilita botones segun los permisos del usuario"],
                  ["Balance de vacaciones", "Calculo automatico de dias de vacaciones disponibles, usados y pendientes"],
                  ["Ausencia aprobada", "Solicitud de ausencia que fue aceptada por un administrador"],
                  ["Free Trial", "Prueba gratuita de 30 dias del Plan Basico sin requerir datos de pago"],
                  ["CRUD", "Create, Read, Update, Delete — las 4 operaciones basicas sobre cualquier entidad"],
                  ["Edge Function", "Funcion ejecutada en el servidor (Supabase) para operaciones que requieren privilegios"],
                  ["Soft delete", "Desactivacion logica (is_active = false) en lugar de eliminacion fisica de datos"],
                  ["COP", "Peso colombiano — moneda utilizada en toda la plataforma"],
                  ["Owner bypass", "El propietario del negocio tiene todos los permisos sin verificacion adicional"],
                  ["Two-step query", "Patron de consulta: primero citas, luego perfiles y servicios por separado"]])

    # ── Footer ────────────────────────────────────────────────────────────
    page_break(doc)
    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing — www.tituring.com")
    style_run(r, size=11, bold=True, color=PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gestabiz v1.0.3 — Manual de Usuario — Parte 4 de 5")
    style_run(r, size=10, italic=True, color=GREY)

    return doc


# ============================================================================
# PROPUESTA DE VALOR — PARTE 4 (ADMIN PLAN BASICO)
# ============================================================================

def build_proposal_part4() -> Document:
    doc = setup_document(
        "Propuesta de Valor — Gestabiz",
        "Por que Gestabiz es la mejor opcion para tu negocio de servicios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE DE CONTENIDOS", anchor="toc_pv4")
    para(doc, "Haga clic en cualquier titulo para navegar directamente a la seccion.",
         italic=True, color=GREY)

    toc_items = [
        ("pv_resumen", "1. Tu Negocio Crece — Tu Herramienta Tambien"),
        ("pv_que_desbloqueas", "2. Que Desbloqueas con el Plan Basico"),
        ("pv_equipo", "3. Gestiona tu Equipo como un Profesional"),
        ("pv_ausencias", "4. Ausencias y Vacaciones Bajo Control"),
        ("pv_ventas_hist", "5. Historial de Ventas — Conoce tus Numeros"),
        ("pv_ventas_rapidas", "6. Ventas Rapidas — Ningun Ingreso se Pierde"),
        ("pv_reportes", "7. Reportes Financieros — Decisiones con Datos"),
        ("pv_permisos", "8. Permisos — Tu Equipo, Tus Reglas"),
        ("pv_sin_limites", "9. Sin Limites que Frenen tu Crecimiento"),
        ("pv_whatsapp", "10. WhatsApp + Google Calendar — Productividad al Maximo"),
        ("pv_roi", "11. Retorno de Inversion — Los Numeros Hablan"),
        ("pv_comparativa", "12. Comparativa: Plan Gratuito vs. Basico"),
        ("pv_vs_competencia", "13. Gestabiz Basico vs. La Competencia"),
        ("pv_escenarios", "14. Escenarios Reales de Negocios"),
        ("pv_trial", "15. Prueba Gratis 30 Dias — Sin Riesgo"),
        ("pv_faq", "16. Preguntas Frecuentes"),
        ("pv_cta", "17. Activa tu Plan Basico Hoy"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 1 — RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 1 — POR QUE ACTUALIZAR AL PLAN BASICO", anchor="parte1_pv4")

    h2(doc, "1. Tu Negocio Crece — Tu Herramienta Tambien", anchor="pv_resumen")
    para(doc, "Empezaste con Gestabiz gratis y ya comprobaste que funciona. Tus clientes reservan "
         "en linea, tu calendario se organizo y tu perfil publico atrae nuevos clientes. "
         "Pero ahora necesitas mas. Mas empleados, mas sedes, mas control. "
         "El Plan Basico esta disenado exactamente para este momento.")

    para(doc, "Por $89,900 COP al mes (menos de $3,000 COP al dia), desbloqueas las herramientas "
         "que separan a un negocio que sobrevive de uno que crece con proposito.")

    callout(doc, "$89,900 COP/mes — Menos que un almuerzo diario",
            "El costo del Plan Basico es menor que un almuerzo de $3,000 COP al dia. "
            "A cambio, obtienes herramientas que generan decenas de veces ese valor en eficiencia, "
            "control y nuevos ingresos capturados.",
            color=ACCENT)

    h2(doc, "2. Que Desbloqueas con el Plan Basico", anchor="pv_que_desbloqueas")
    para(doc, "Seis modulos profesionales que transforman tu operacion:")

    simple_table(doc, ["Modulo", "Problema que Resuelve", "Valor Estimado"],
                 [["Empleados (hasta 6)", "Gestion manual en cuadernos/WhatsApp", "$50,000/mes en ahorro de tiempo"],
                  ["Ausencias", "No saber quien falta ni cuando", "$30,000/mes en citas no perdidas"],
                  ["Historial de Ventas", "No saber cuanto vendes realmente", "Visibilidad total de ingresos"],
                  ["Ventas Rapidas", "Clientes walk-in no registrados", "$80,000+/mes en ingresos capturados"],
                  ["Reportes Financieros", "Decisiones basadas en intuicion", "Decisiones basadas en datos"],
                  ["Permisos", "Empleados con acceso a todo", "Seguridad y control total"]])

    para(doc, "Ademas, se eliminan los limites que te frenaban:")
    bullet(doc, "Sedes: de 1 a 3 — abre tu segunda y tercera sucursal.", bold_prefix="x3")
    bullet(doc, "Empleados: de 1 a 6 — contrata y gestiona tu equipo.", bold_prefix="x6")
    bullet(doc, "Citas: de 50/mes a ilimitadas — sin techo para crecer.", bold_prefix="Ilimitadas")
    bullet(doc, "Clientes: de 50 a ilimitados — ve a todos tus clientes.", bold_prefix="Ilimitados")
    bullet(doc, "Servicios: de 15 a ilimitados — ofrece todo tu catalogo.", bold_prefix="Ilimitados")

    screenshot_placeholder(doc, "Dashboard del admin con todos los modulos del Plan Basico activos")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 2 — DETALLE DE CADA MODULO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 2 — DETALLE DE CADA FUNCIONALIDAD", anchor="parte2_pv4")

    # --- Equipo ---
    h2(doc, "3. Gestiona tu Equipo como un Profesional", anchor="pv_equipo")
    para(doc, "Cuando tienes un solo empleado (tu), la gestion es facil. Pero apenas sumas un segundo, "
         "un tercero, un sexto miembro al equipo, necesitas estructura. "
         "Gestabiz te da esa estructura sin que tengas que inventarla.")

    h3(doc, "Lo que puedes hacer")
    bullet(doc, "Invitar empleados con un codigo de 6 digitos que puedes compartir por WhatsApp — "
           "se registran solos y tu apruebas.", bold_prefix="Invitacion inteligente")
    bullet(doc, "Organiza a tu equipo en una estructura visual: Owner > Admin > Gerente > Lider > Personal. "
           "Cada persona sabe de quien depende.", bold_prefix="Jerarquia visual")
    bullet(doc, "Configura el horario de cada empleado dia por dia, incluyendo hora de almuerzo. "
           "Las citas respetan automaticamente estos horarios.", bold_prefix="Horarios personalizados")
    bullet(doc, "Define que servicios ofrece cada profesional. Los clientes solo ven "
           "empleados que pueden atender lo que necesitan.", bold_prefix="Servicios asignados")
    bullet(doc, "Registra salarios, tipo de pago (fijo, variable, mixto) y comisiones. "
           "Informacion sensible protegida por permisos.", bold_prefix="Configuracion salarial")

    callout(doc, "Piensa en esto",
            "Sin Gestabiz, gestionar horarios de 6 empleados con 15 servicios distintos en 3 sedes "
            "requiere una hoja de calculo imposible. Con Gestabiz, esta todo automatizado, visual "
            "y en tiempo real.",
            color=PURPLE)

    screenshot_placeholder(doc, "Vista de mapa jerarquico con la estructura del equipo")

    # --- Ausencias ---
    h2(doc, "4. Ausencias y Vacaciones Bajo Control", anchor="pv_ausencias")
    para(doc, "Cuando un empleado falta sin aviso, las consecuencias son costosas: clientes "
         "molestos, citas canceladas de ultimo minuto y perdida de ingresos. "
         "El modulo de Ausencias de Gestabiz elimina este caos.")

    h3(doc, "Como funciona")
    numbered(doc, "El empleado solicita la ausencia desde su celular (tipo, fechas, razon).")
    numbered(doc, "Tu recibes notificacion instantanea en la app y en tu correo.")
    numbered(doc, "Ves cuantas citas se veran afectadas ANTES de decidir.")
    numbered(doc, "Apruebas o rechazas con un clic. Puedes agregar notas.")
    numbered(doc, "Si apruebas: las citas se cancelan automaticamente y los clientes son notificados.")
    numbered(doc, "El balance de vacaciones se actualiza automaticamente (15 dias/ano, excluye festivos).")

    h3(doc, "Beneficios concretos")
    bullet(doc, "Cero sorpresas: siempre sabes quien esta ausente y cuando.")
    bullet(doc, "Cero citas perdidas: el sistema bloquea automaticamente los horarios de empleados ausentes.")
    bullet(doc, "Cero calculo manual: el balance de vacaciones se calcula solo, excluyendo festivos colombianos.")
    bullet(doc, "Registro historico: puedes consultar todas las ausencias pasadas en cualquier momento.")

    screenshot_placeholder(doc, "Tarjeta de aprobacion de ausencia con citas afectadas")

    # --- Ventas Historial ---
    h2(doc, "5. Historial de Ventas — Conoce tus Numeros", anchor="pv_ventas_hist")
    para(doc, "No puedes mejorar lo que no mides. El Historial de Ventas te muestra exactamente "
         "cuanto estas generando, que servicios son los mas rentables y cuales son tus mejores clientes.")

    bullet(doc, "Citas completadas, ingresos totales y promedio por cita — todo de un vistazo.", bold_prefix="3 metricas clave")
    bullet(doc, "Filtra por 7, 30, 90 o 365 dias para ver tendencias.", bold_prefix="Filtro de periodo")
    bullet(doc, "Encuentra citas especificas por nombre de cliente o servicio.", bold_prefix="Busqueda rapida")
    bullet(doc, "Haz clic en cualquier cliente para ver su perfil completo con historial.", bold_prefix="Perfil del cliente")

    callout(doc, "Ejemplo practico",
            "Descubres que el servicio 'Keratina Premium' genera el 40% de tus ingresos pero solo "
            "lo ofrece 1 empleado. Solucion: capacitar a otro profesional para duplicar la capacidad. "
            "Sin datos, nunca lo habrias sabido.",
            color=ACCENT)

    # --- Ventas Rapidas ---
    h2(doc, "6. Ventas Rapidas — Ningun Ingreso se Pierde", anchor="pv_ventas_rapidas")
    para(doc, "No todos los clientes reservan por anticipado. Muchos simplemente llegan. "
         "Sin Ventas Rapidas, esos ingresos se pierden en el limbo de los pagos en efectivo "
         "sin registro. Con Gestabiz, cada peso queda contabilizado.")

    h3(doc, "Valor real para tu negocio")
    bullet(doc, "Un salon tipico recibe 3-5 clientes walk-in por dia. Si el servicio promedio "
           "cuesta $40,000 COP, son $120,000 - $200,000 COP/dia que antes no se registraban.")
    bullet(doc, "Cada venta rapida se integra automaticamente con la contabilidad — no hay doble trabajo.")
    bullet(doc, "3 metodos de pago (efectivo, tarjeta, transferencia) para capturar todos los ingresos.")
    bullet(doc, "Estadisticas en tiempo real: cuanto vendiste hoy, esta semana y este mes.")
    bullet(doc, "Ultimas 10 ventas visibles para seguimiento inmediato.")

    para(doc, "El formulario es rapido: nombre del cliente, servicio (autocompleta el precio), "
         "metodo de pago y listo. Menos de 30 segundos por registro.")

    screenshot_placeholder(doc, "Formulario de ventas rapidas con estadisticas del dia")

    # --- Reportes ---
    h2(doc, "7. Reportes Financieros — Decisiones con Datos", anchor="pv_reportes")
    para(doc, "El dashboard financiero interactivo te da una radiografia de tu negocio "
         "con graficos claros que cualquiera puede entender.")

    h3(doc, "5 graficos que cuentan la historia de tu negocio")
    bullet(doc, "Ingresos vs Gastos (barras): comparativa mensual para saber si estás creciendo.", bold_prefix="Barras")
    bullet(doc, "Tendencia Mensual (linea): evolucion temporal de tus ingresos.", bold_prefix="Tendencia")
    bullet(doc, "Distribucion por Categoria (torta): que servicios generan mas ingresos.", bold_prefix="Torta")
    bullet(doc, "Ingresos por Empleado (barras horizontales): quien genera mas valor.", bold_prefix="Ranking")
    bullet(doc, "Ingresos por Sede (barras): rendimiento comparativo entre sedes.", bold_prefix="Por sede")

    para(doc, "Todos los graficos son interactivos: al pasar el cursor ves el detalle numerico. "
         "Filtra por periodo, sede, empleado o tipo para analizar exactamente lo que necesitas.")

    callout(doc, "Nota sobre exportacion",
            "El Plan Basico incluye visualizacion completa de todos los reportes en pantalla. "
            "La exportacion a PDF, CSV y Excel esta disponible en el Plan Pro.",
            color=PURPLE)

    # --- Permisos ---
    h2(doc, "8. Permisos — Tu Equipo, Tus Reglas", anchor="pv_permisos")
    para(doc, "No todos los empleados necesitan ver todo. La recepcionista necesita gestionar citas "
         "pero no deberia ver los salarios. El contador necesita acceso a reportes pero no a "
         "la gestion de empleados. Los permisos te dan ese control.")

    h3(doc, "3 plantillas predefinidas para empezar de inmediato")
    simple_table(doc, ["Plantilla", "Para quien", "Que puede hacer"],
                 [["Recepcionista", "Personal de recepcion", "Crear/editar/cancelar citas, ver clientes y servicios"],
                  ["Profesional", "Empleados que atienden", "Ver sus citas, cancelar propias, solicitar ausencias, ver reviews"],
                  ["Contador", "Encargado de finanzas", "Ver reportes financieros, crear transacciones, ver gastos y facturacion"]])

    para(doc, "Aplicar una plantilla toma literalmente 3 clics: abrir permisos del empleado, "
         "seleccionar plantilla, confirmar. Listo.")

    page_break(doc)

    # --- Sin limites ---
    h2(doc, "9. Sin Limites que Frenen tu Crecimiento", anchor="pv_sin_limites")
    para(doc, "Los limites del Plan Gratuito estan disenados para probar la plataforma. "
         "Pero un negocio real no puede operar con 50 citas al mes o 1 sola sede. "
         "El Plan Basico elimina estas barreras:")

    simple_table(doc, ["Lo que cambia", "Antes (Gratuito)", "Ahora (Basico)", "Impacto Real"],
                 [["Sedes", "1", "3", "Abre tu segunda sucursal"],
                  ["Empleados", "Solo tu", "Hasta 6", "Contrata tu equipo"],
                  ["Citas/mes", "50", "Ilimitadas", "Sin techo para crecer"],
                  ["Clientes", "50 visibles", "Ilimitados", "Ve a todos tus clientes"],
                  ["Servicios", "15", "Ilimitados", "Ofrece todo tu catalogo"]])

    # --- WhatsApp + GCal ---
    h2(doc, "10. WhatsApp + Google Calendar — Productividad al Maximo", anchor="pv_whatsapp")
    para(doc, "Dos integraciones que hacen una gran diferencia en tu dia a dia:")

    h3(doc, "Recordatorios por WhatsApp")
    para(doc, "Los clientes reciben recordatorios automaticos de sus citas directamente en WhatsApp. "
         "Esto reduce las inasistencias hasta un 40%, lo que significa mas ingresos reales "
         "sin esfuerzo adicional de tu parte.")

    h3(doc, "Google Calendar")
    para(doc, "Sincroniza tus citas de Gestabiz con Google Calendar. Tu y tus empleados ven todo "
         "en un solo lugar, sin duplicar informacion ni olvidar compromisos.")

    # --- ROI ---
    h2(doc, "11. Retorno de Inversion — Los Numeros Hablan", anchor="pv_roi")
    para(doc, "Veamos un calculo conservador para un salon de belleza tipico en Colombia:")

    simple_table(doc, ["Concepto", "Sin Gestabiz Basico", "Con Gestabiz Basico", "Diferencia"],
                 [["Clientes walk-in registrados", "0 (se pierden)", "3/dia x $40K = $120K/dia", "+$2,400,000/mes"],
                  ["Inasistencias (no-shows)", "~20% de citas", "~12% (WhatsApp reminder)", "+$300,000/mes"],
                  ["Horas admin gestionando equipo", "~5h/semana", "~1h/semana", "16h libres/mes"],
                  ["Ausencias no controladas", "~2/mes sin aviso", "0 sin aprobacion", "+$200,000/mes"],
                  ["Costo del plan", "—", "$89,900/mes", "-$89,900/mes"],
                  ["", "", "", ""],
                  ["RETORNO NETO ESTIMADO", "", "", "+$2,810,100/mes"]])

    callout(doc, "31x de retorno",
            "Por cada peso que inviertes en el Plan Basico, recuperas aproximadamente 31 pesos "
            "en ingresos capturados, tiempo ahorrado y perdidas evitadas. Este calculo es conservador — "
            "negocios con mas volumen veran retornos aun mayores.",
            color=ACCENT)

    page_break(doc)

    # --- Comparativa planes ---
    h2(doc, "12. Comparativa: Plan Gratuito vs. Basico", anchor="pv_comparativa")
    para(doc, "Si todavia no estas seguro, esta comparativa lado a lado te ayudara a decidir:")

    simple_table(doc, ["Aspecto", "Plan Gratuito", "Plan Basico"],
                 [["Para quien", "Emprendedor solo", "Negocio en crecimiento"],
                  ["Empleados", "Solo tu", "Tu + 5 empleados"],
                  ["Sedes", "1 ubicacion", "Hasta 3 ubicaciones"],
                  ["Limite de citas", "50/mes", "Ilimitadas"],
                  ["Gestion de equipo", "No", "Si (jerarquia, horarios, salarios)"],
                  ["Control de ausencias", "No", "Si (aprobacion, vacaciones)"],
                  ["Registro de walk-ins", "No", "Si (ventas rapidas)"],
                  ["Reportes graficos", "No", "Si (5 tipos de graficos)"],
                  ["Control de accesos", "No", "Si (3 plantillas de permisos)"],
                  ["WhatsApp reminders", "No", "Si"],
                  ["Google Calendar", "No", "Si"],
                  ["Soporte", "Comunidad", "Email prioritario"],
                  ["Precio", "Gratis", "$89,900 COP/mes"]])

    # --- Vs Competencia ---
    h2(doc, "13. Gestabiz Basico vs. La Competencia", anchor="pv_vs_competencia")
    para(doc, "El Plan Basico de Gestabiz ofrece mas funcionalidades a un precio menor que "
         "las alternativas mas conocidas del mercado:")

    simple_table(doc, ["Funcion", "Gestabiz Basico", "Calendly Teams", "Booksy Premium", "Fresha Plus"],
                 [["Precio/mes", "$89,900 COP (~$22 USD)", "$96 USD (4 users)", "$30 USD (1 user)", "Comision variable"],
                  ["Empleados incluidos", "6", "4", "1 (extra: $30/ea)", "Ilimitados (comision)"],
                  ["Gestion de ausencias", "Si", "No", "No", "Basico"],
                  ["Ventas walk-in", "Si", "No", "Si", "Si"],
                  ["Reportes graficos", "Si (5 tipos)", "Basico", "Basico", "Si"],
                  ["CRM integrado", "Si", "No", "Si", "Si"],
                  ["Permisos por rol", "Si (3 templates)", "No", "No", "No"],
                  ["WhatsApp reminders", "Si", "No", "No", "No"],
                  ["Google Calendar", "Si", "Si", "Si", "Si"],
                  ["Multi-sede", "Si (3)", "No", "No", "Si (comision)"],
                  ["Festivos automaticos", "Si (Colombia)", "No", "No", "No"],
                  ["Moneda local (COP)", "Si", "USD", "USD", "Varia"],
                  ["Idioma espanol nativo", "Si", "Parcial", "No", "No"]])

    callout(doc, "4x mas barato que Calendly con 50% mas funcionalidades",
            "Gestabiz Basico cuesta $22 USD/mes para 6 usuarios. Calendly Teams cuesta $96 USD/mes "
            "para solo 4 usuarios y no incluye gestion de ausencias, ventas walk-in, permisos ni "
            "CRM de clientes. Booksy cobra $30 USD POR CADA usuario adicional.",
            color=ACCENT)

    page_break(doc)

    # --- Escenarios ---
    h2(doc, "14. Escenarios Reales de Negocios", anchor="pv_escenarios")
    para(doc, "Estos son escenarios basados en negocios reales de servicios en Colombia que "
         "se beneficiarian del Plan Basico:")

    h3(doc, "Escenario 1: Carlos — Barberia en Chapinero, Bogota")
    para(doc, "Carlos tiene una barberia con 4 barberos. Usa un grupo de WhatsApp para coordinar "
         "los horarios y una libreta para anotar quienes llegan sin cita. Cada mes pierde "
         "la cuenta de cuanto gana cada empleado.")
    para(doc, "Con Gestabiz Basico: Carlos invita a sus 4 barberos con codigo QR. Cada uno "
         "tiene su horario configurado y los clientes reservan directamente con su barbero "
         "preferido. Las ventas walk-in se registran en 30 segundos. Al final del mes, Carlos "
         "abre Reportes y ve exactamente cuanto genero cada barbero. Precio: $89,900/mes. "
         "Valor capturado: +$3,000,000/mes en walk-ins registrados.")

    h3(doc, "Escenario 2: Maria — Centro de Estetica en El Poblado, Medellin")
    para(doc, "Maria tiene 2 sedes con 3 esteticistas en total. Las ausencias por enfermedad "
         "o citas medicas le causan cancelaciones de ultimo momento que enojan a sus clientes.")
    para(doc, "Con Gestabiz Basico: Maria configura sus 2 sedes y 3 empleadas. Cuando Luisa "
         "necesita faltar, solicita la ausencia desde su celular. Maria ve que hay 5 citas "
         "afectadas y decide aprobar la ausencia — las 5 clientes reciben notificacion "
         "automatica para reagendar. Cero llamadas telefonicas, cero clientes enojados.")

    h3(doc, "Escenario 3: Andres — Gym Funcional en Cali")
    para(doc, "Andres tiene un gimnasio con 5 entrenadores y ofrece 20 tipos de clases diferentes. "
         "Con el Plan Gratuito solo podia registrar 15 servicios y atendia maximo 50 citas al mes.")
    para(doc, "Con Gestabiz Basico: Andres registra todos sus 20 tipos de clases sin limite. "
         "Sus 5 entrenadores tienen horarios propios y los 200+ clientes mensuales reservan "
         "sin restriccion. Los reportes le muestran que la clase de CrossFit de los martes "
         "a las 6pm tiene 95% de ocupacion — senalando que deberia abrir un segundo horario.")

    screenshot_placeholder(doc, "Ejemplo de reportes financieros con datos reales de un negocio")

    page_break(doc)

    # --- Trial ---
    h2(doc, "15. Prueba Gratis 30 Dias — Sin Riesgo", anchor="pv_trial")
    para(doc, "No tienes que tomar la decision hoy. Gestabiz te ofrece 30 dias gratis del Plan Basico "
         "para que lo pruebes con tu equipo real, tus clientes reales y tu operacion real.")

    h3(doc, "Sin letra pequena")
    bullet(doc, "No se requiere tarjeta de credito.")
    bullet(doc, "No hay cobro automatico al terminar.")
    bullet(doc, "Acceso completo a todas las funcionalidades del Plan Basico.")
    bullet(doc, "Tus datos se conservan al terminar la prueba (los modulos se bloquean pero los datos estan intactos).")
    bullet(doc, "Una sola activacion por owner — usala cuando estes listo para probar en serio.")

    callout(doc, "30 dias son suficientes para ver resultados",
            "En 30 dias puedes: invitar a todo tu equipo, registrar al menos 100 ventas, "
            "gestionar 2-3 ausencias, generar tu primer reporte financiero completo y configurar "
            "los permisos de tu equipo. Al final sabras exactamente cuanto vale Gestabiz para tu negocio.",
            color=ACCENT)

    # --- FAQ ---
    h2(doc, "16. Preguntas Frecuentes", anchor="pv_faq")

    h3(doc, "Puedo volver al Plan Gratuito despues de pagar?")
    para(doc, "Si. Puedes cancelar tu suscripcion en cualquier momento desde Facturacion. "
         "Al cancelar, tu negocio vuelve al Plan Gratuito con los limites originales. "
         "Tus datos no se eliminan — solo los modulos exclusivos se bloquean.")

    h3(doc, "Que pasa con mis empleados si bajo al Plan Gratuito?")
    para(doc, "Los empleados siguen registrados en la base de datos pero el modulo de Empleados "
         "se bloquea. Solo veras 1 empleado (tu como owner). Al reactivar el Plan Basico, "
         "todos los empleados vuelven a estar disponibles.")

    h3(doc, "El Plan Basico incluye soporte tecnico?")
    para(doc, "Si. El Plan Basico incluye soporte prioritario por email con tiempo de respuesta "
         "de 24-48 horas habiles. Ademas, la app incluye un boton de reporte de bugs integrado.")

    h3(doc, "Puedo cambiar de plan en cualquier momento?")
    para(doc, "Si. Puedes actualizar al Plan Pro o volver al Gratuito en cualquier momento "
         "desde la seccion de Facturacion. Los cambios se aplican inmediatamente.")

    h3(doc, "Cuantas sedes puedo tener con el Plan Basico?")
    para(doc, "Hasta 3 sedes. Si necesitas mas, el Plan Pro permite hasta 10 sedes.")

    h3(doc, "Puedo asignar permisos individuales a cada empleado?")
    para(doc, "En el Plan Basico puedes aplicar 3 plantillas de permisos predefinidas "
         "(Recepcionista, Profesional, Contador). Para asignar permisos individuales "
         "de los 79 tipos disponibles, necesitas el Plan Pro.")

    h3(doc, "La prueba gratuita requiere tarjeta de credito?")
    para(doc, "No. La prueba de 30 dias se activa sin datos de pago. Al terminar, "
         "simplemente vuelves al Plan Gratuito automaticamente.")

    h3(doc, "Hay descuento por pago anual?")
    para(doc, "Si. El Plan Basico anual cuesta $899,000 COP — equivalente a 10 meses, "
         "obteniendo 2 meses gratis (ahorro de $179,800 COP).")

    page_break(doc)

    # --- CTA ---
    h2(doc, "17. Activa tu Plan Basico Hoy", anchor="pv_cta")

    para(doc, "Tu negocio ya demostro que necesita mas. Tus clientes estan ahi, tu equipo "
         "esta listo, y la herramienta esta esperandote. Da el siguiente paso.")

    callout(doc, "ACTIVA TU PLAN BASICO",
            "Ingresa a tu panel de administracion → Facturacion → Selecciona Plan Basico → "
            "Completa el pago. En menos de 2 minutos tendras acceso a todas las funcionalidades. "
            "O prueba gratis por 30 dias sin compromiso.\n\n"
            "www.gestabiz.com | soporte@gestabiz.com | WhatsApp: +57 XXX XXX XXXX",
            color=ACCENT)

    para(doc, "El Plan Basico de Gestabiz no es un gasto — es una inversion que se paga sola "
         "en la primera semana de uso.", italic=True, color=PURPLE)

    # ── Footer ────────────────────────────────────────────────────────────
    page_break(doc)
    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing — www.tituring.com")
    style_run(r, size=11, bold=True, color=PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gestabiz v1.0.3 — Propuesta de Valor — Parte 4 de 5")
    style_run(r, size=10, italic=True, color=GREY)

    return doc


# ============================================================================
# MAIN
# ============================================================================

def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # --- Manual de Usuario ---
    print("Generando Manual de Usuario — Parte 4: Admin Plan Basico...")
    manual = build_manual_part4()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz - copilot - parte4.docx"
    manual.save(str(manual_path))
    print(f"  -> {manual_path}")

    # --- Propuesta de Valor ---
    print("Generando Propuesta de Valor — Parte 4: Admin Plan Basico...")
    propuesta = build_proposal_part4()
    propuesta_path = DOCS_DIR / "Propuesta_Valor_Gestabiz - copilot - parte4.docx"
    propuesta.save(str(propuesta_path))
    print(f"  -> {propuesta_path}")

    print("\nListo. Ambos documentos generados exitosamente.")


if __name__ == "__main__":
    main()
