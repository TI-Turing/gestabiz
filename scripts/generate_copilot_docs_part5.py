#!/usr/bin/env python3
"""
Generador de documentos Gestabiz — Version Copilot — Parte 5: Rol Admin (Plan Pro)
========================================================================================
Genera:
  - docs/Manual_Usuario_Gestabiz - copilot - parte5.docx
  - docs/Propuesta_Valor_Gestabiz - copilot - parte5.docx

Requisitos: python-docx >= 1.0.0
Ejecucion:  python scripts/generate_copilot_docs_part5.py
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Paths & Brand Colors
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
LOGO_GESTABIZ = ROOT / "src" / "assets" / "images" / "logo_gestabiz.png"
LOGO_TITURING = ROOT / "src" / "assets" / "images" / "tt" / "1.png"

PURPLE = RGBColor(0x7C, 0x3A, 0xED)
DARK = RGBColor(0x1F, 0x1F, 0x2E)
GREY = RGBColor(0x6B, 0x72, 0x80)
ACCENT = RGBColor(0x10, 0xB9, 0x81)
DANGER = RGBColor(0xE1, 0x1D, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def add_bookmark(paragraph, name: str) -> None:
    tag = OxmlElement("w:bookmarkStart")
    tag.set(qn("w:id"), str(id(name) % 2**31))
    tag.set(qn("w:name"), name)
    paragraph._p.append(tag)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(id(name) % 2**31))
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str, *,
                           bold: bool = False, color: str = "7C3AED") -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rpr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    if bold:
        rpr.append(OxmlElement("w:b"))
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, color: str = "D1D5DB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


# ---------------------------------------------------------------------------
# High-level helpers
# ---------------------------------------------------------------------------

def style_run(run, *, size: int = 11, bold: bool = False, italic: bool = False,
              color: RGBColor = DARK, font: str = "Calibri") -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def h1(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=22, bold=True, color=color)


def h2(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = DARK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=16, bold=True, color=color)


def h3(doc: Document, text: str, *, anchor: str | None = None, color: RGBColor = PURPLE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if anchor:
        add_bookmark(p, anchor)
    run = p.add_run(text)
    style_run(run, size=13, bold=True, color=color)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, size=11, bold=True, color=DARK)


def para(doc: Document, text: str, *, size: int = 11, italic: bool = False,
         color: RGBColor = DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = align
    run = p.add_run(text)
    style_run(run, size=size, italic=italic, color=color)


def bullet(doc: Document, text: str, *, size: int = 11, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        style_run(r, size=size, bold=True, color=PURPLE)
        r2 = p.add_run(" — " + text)
        style_run(r2, size=size, color=DARK)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=DARK)


def numbered(doc: Document, text: str, *, size: int = 11) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, color=DARK)


def callout(doc: Document, title: str, text: str, *, color: RGBColor = PURPLE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    bg = "F5F3FF" if color == PURPLE else ("ECFDF5" if color == ACCENT else "FEF2F2")
    set_cell_shading(cell, bg)
    set_cell_border(cell, color=hex_color, size="8")
    cell.width = Cm(16)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(title)
    style_run(r, size=11, bold=True, color=color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=DARK)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    img_cell = table.cell(0, 0)
    img_cell.width = Cm(15)
    set_cell_shading(img_cell, "F3F4F6")
    set_cell_border(img_cell, color="9CA3AF", size="8")
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    label_p = img_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_p.paragraph_format.space_after = Pt(0)
    r = label_p.add_run("[ Espacio reservado para captura de pantalla ]")
    style_run(r, size=10, italic=True, color=GREY)
    for _ in range(3):
        ep = img_cell.add_paragraph()
        ep.paragraph_format.space_after = Pt(0)
    cap_cell = table.cell(1, 0)
    set_cell_shading(cap_cell, "FFFFFF")
    set_cell_border(cap_cell, color="FFFFFF", size="2")
    cp = cap_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run(f"Figura: {caption}")
    style_run(r, size=9, italic=True, color=GREY)
    doc.add_paragraph()


def simple_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
                 *, header_bg: str = "7C3AED",
                 header_fg: RGBColor = RGBColor(0xFF, 0xFF, 0xFF)) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, header_bg)
        set_cell_border(c, color="FFFFFF", size="4")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color=header_fg)
    for ridx, row in enumerate(rows):
        zebra = "F9FAFB" if ridx % 2 == 0 else "FFFFFF"
        for cidx, val in enumerate(row):
            c = table.rows[1 + ridx].cells[cidx]
            set_cell_shading(c, zebra)
            set_cell_border(c, color="E5E7EB", size="4")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            style_run(r, size=10, color=DARK)
    doc.add_paragraph()


def page_break(doc: Document) -> None:
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document setup + cover
# ---------------------------------------------------------------------------

def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run(f"{title}  |  Desarrollado por Ti Turing  |  (c) 2026")
        style_run(r, size=9, italic=True, color=GREY)
    build_cover(doc, title, subtitle)
    return doc


def build_cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    if LOGO_GESTABIZ.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_GESTABIZ), width=Cm(5.5))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("GESTABIZ")
    style_run(r, size=44, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(title)
    style_run(r, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    style_run(r, size=13, italic=True, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Parte 5 de 5: Administrador — Plan Pro")
    style_run(r, size=14, bold=True, color=PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version del producto: 1.0.3   |   Abril 2026")
    style_run(r, size=11, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fase Beta completada   |   Listo para produccion")
    style_run(r, size=11, color=GREY, italic=True)

    for _ in range(5):
        doc.add_paragraph()

    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing")
    style_run(r, size=12, bold=True, color=DARK)


# ============================================================================
# MANUAL DE USUARIO — PARTE 5 (ADMIN PLAN PRO)
# ============================================================================

def build_manual_part5() -> Document:
    doc = setup_document(
        "Manual de Usuario — Gestabiz",
        "Guia funcional completa del sistema de gestion de citas y negocios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE DE CONTENIDOS", anchor="toc_p5")
    para(doc, "A continuacion se listan todas las secciones cubiertas en esta parte. "
         "Haga clic en cualquier titulo para navegar directamente al detalle.", italic=True, color=GREY)

    toc_items = [
        ("sec_resumen_pro", "1. Resumen Ejecutivo — Plan Pro"),
        ("sec_que_incluye_pro", "2. Que incluye el Plan Pro"),
        ("sec_comparativa_3planes", "3. Comparativa de Limites: Gratuito vs. Basico vs. Pro"),
        ("sec_activar_pro", "4. Como Activar el Plan Pro"),
        ("sec_gastos", "5. Modulo de Gastos y Egresos"),
        ("sec_gastos_stats", "6. Tarjetas de Estadisticas de Gastos"),
        ("sec_gastos_unicos", "7. Pestana Egresos Unicos"),
        ("sec_gastos_recurrentes", "8. Pestana Egresos Recurrentes"),
        ("sec_gastos_resumen", "9. Pestana Resumen por Categoria"),
        ("sec_gastos_categorias", "10. Las 48 Categorias de Gastos"),
        ("sec_gastos_form", "11. Formulario de Nuevo Egreso — Campos"),
        ("sec_gastos_recurrente_cfg", "12. Configuracion de Egresos Recurrentes"),
        ("sec_gastos_metodos", "13. Metodos de Pago para Egresos"),
        ("sec_gastos_permisos", "14. Permisos del Modulo de Gastos"),
        ("sec_reclutamiento", "15. Modulo de Reclutamiento"),
        ("sec_recl_dashboard", "16. Dashboard de Reclutamiento — Vista General"),
        ("sec_recl_crear", "17. Crear Nueva Vacante — 4 Tarjetas"),
        ("sec_recl_info", "18. Tarjeta 1: Informacion Basica de la Vacante"),
        ("sec_recl_detalles", "19. Tarjeta 2: Detalles de la Vacante"),
        ("sec_recl_compensacion", "20. Tarjeta 3: Compensacion y Ubicacion"),
        ("sec_recl_estado", "21. Tarjeta 4: Estado de la Vacante"),
        ("sec_recl_lista", "22. Lista de Vacantes — Filtros y Busqueda"),
        ("sec_recl_card", "23. Tarjeta de Vacante — Detalle Visual"),
        ("sec_recl_apps", "24. Gestion de Aplicaciones"),
        ("sec_recl_tabs", "25. Las 6 Pestanas de Estado de Aplicaciones"),
        ("sec_recl_stats", "26. 7 Tarjetas de Estadisticas de Aplicaciones"),
        ("sec_recl_app_card", "27. Tarjeta de Aplicacion — Detalle"),
        ("sec_recl_acciones", "28. Acciones por Estado de Aplicacion"),
        ("sec_recl_seleccionar", "29. Modal de Seleccionar Empleado"),
        ("sec_recl_rechazar", "30. Dialogo de Rechazo"),
        ("sec_recl_perfil", "31. Modal de Perfil del Aplicante — 3 Pestanas"),
        ("sec_recl_reviews", "32. Reviews Obligatorias al Contratar y Finalizar"),
        ("sec_recl_notificaciones", "33. Notificaciones Automaticas de Reclutamiento"),
        ("sec_recl_permisos", "34. Permisos del Modulo de Reclutamiento"),
        ("sec_recursos", "35. Modulo de Recursos Fisicos"),
        ("sec_rec_dashboard", "36. Vista General de Recursos"),
        ("sec_rec_tipos", "37. Los 15 Tipos de Recursos"),
        ("sec_rec_filtros", "38. Filtros por Tipo de Recurso"),
        ("sec_rec_tabla", "39. Tabla de Recursos — Columnas"),
        ("sec_rec_crear", "40. Crear o Editar Recurso — Modal"),
        ("sec_rec_disponibilidad", "41. Validacion de Disponibilidad de Recursos"),
        ("sec_rec_modelo", "42. Condicion de Modelo de Negocio"),
        ("sec_rec_permisos", "43. Permisos del Modulo de Recursos"),
        ("sec_nomina", "44. Sistema de Nomina y Salarios"),
        ("sec_nom_config", "45. Configuracion de Salario del Empleado"),
        ("sec_nom_frecuencia", "46. Frecuencias de Pago Disponibles"),
        ("sec_nom_dia", "47. Dia de Pago Configurable"),
        ("sec_nom_recurrente", "48. Generacion Automatica de Egreso Recurrente"),
        ("sec_nom_permisos", "49. Permisos del Sistema de Nomina"),
        ("sec_jerarquia", "50. Mapa Jerarquico del Equipo"),
        ("sec_jer_niveles", "51. Los 5 Niveles Jerarquicos"),
        ("sec_jer_zoom", "52. Controles del Mapa: Zoom, Expandir, Pantalla Completa"),
        ("sec_jer_editar", "53. Editar Nivel Jerarquico de un Empleado"),
        ("sec_facturacion_pro", "54. Dashboard de Facturacion — Vista Pro"),
        ("sec_fact_estados", "55. Estados del Dashboard de Facturacion"),
        ("sec_fact_uso", "56. Estadisticas de Uso del Plan"),
        ("sec_fact_historial", "57. Historial de Pagos"),
        ("sec_fact_cancelar", "58. Flujo de Cancelacion"),
        ("sec_pricing", "59. Pagina de Precios — Vista Completa"),
        ("sec_pricing_planes", "60. Grid de 3 Planes"),
        ("sec_pricing_descuento", "61. Codigos de Descuento"),
        ("sec_pricing_faq", "62. Preguntas Frecuentes en Pricing"),
        ("sec_limites_pro", "63. Limites Ampliados del Plan Pro"),
        ("sec_sedes_10", "64. Gestion de hasta 10 Sedes"),
        ("sec_empleados_15", "65. Hasta 15 Empleados"),
        ("sec_proximamente", "66. Funcionalidades Marcadas como Proximamente"),
        ("sec_permisos_individuales", "67. Editor de Permisos Individuales (Proximamente)"),
        ("sec_audit_trail", "68. Historial de Auditoria (Proximamente)"),
        ("sec_upgrade_basico_pro", "69. Como Actualizar de Basico a Pro"),
        ("sec_soporte_pro", "70. Soporte y Ayuda del Plan Pro"),
        ("sec_glosario_pro", "71. Glosario de Terminos del Plan Pro"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 1 — RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 1 — RESUMEN EJECUTIVO", anchor="parte1_pro")
    para(doc, "Esta seccion presenta a grandes rasgos las capacidades exclusivas del Plan Pro de Gestabiz. "
         "Para el detalle exhaustivo de cada funcionalidad, consulte la Parte 2.")

    # --- 1. Resumen Ejecutivo ---
    h2(doc, "1. Resumen Ejecutivo — Plan Pro", anchor="sec_resumen_pro")
    para(doc, "El Plan Pro de Gestabiz ($159,900 COP/mes o $1,599,000 COP/ano con 2 meses gratis — "
         "17% de ahorro) esta disenado para negocios en crecimiento que necesitan control total sobre "
         "sus operaciones: gastos, reclutamiento de personal, gestion de recursos fisicos (salas, canchas, "
         "mesas, equipos) y nomina automatizada. Incluye todo lo del Plan Basico ($89,900/mes) mas "
         "cuatro modulos exclusivos y limites ampliados que permiten gestionar hasta 10 sedes y "
         "15 empleados con citas, clientes y servicios ilimitados.")

    para(doc, "Los cuatro modulos exclusivos del Plan Pro son:")
    bullet(doc, "Registrar, categorizar y controlar todos los gastos del negocio con 48 categorias, "
           "egresos unicos y recurrentes, y resumen por categoria. Control total de tus finanzas.", bold_prefix="Gastos y Egresos")
    bullet(doc, "Publicar vacantes, recibir aplicaciones, gestionar candidatos con flujo de 6 estados, "
           "seleccion con review obligatoria y notificaciones automaticas.", bold_prefix="Reclutamiento")
    bullet(doc, "Gestionar salas, canchas, mesas, equipos, vehiculos y 15 tipos mas de recursos "
           "fisicos con disponibilidad, precios por hora y amenidades.", bold_prefix="Recursos Fisicos")
    bullet(doc, "Configurar salario base, frecuencia de pago, dia de pago y generar automaticamente "
           "egresos recurrentes de nomina para cada empleado.", bold_prefix="Nomina / Payroll")

    para(doc, "Ademas, los limites se amplian significativamente respecto al Plan Basico:")
    simple_table(doc, ["Recurso", "Plan Gratuito", "Plan Basico", "Plan Pro"],
                 [["Sedes", "1", "3", "10"],
                  ["Empleados", "1 (solo owner)", "6", "15"],
                  ["Citas por mes", "50", "Ilimitadas", "Ilimitadas"],
                  ["Clientes visibles", "50", "Ilimitados", "Ilimitados"],
                  ["Servicios", "15", "Ilimitados", "Ilimitados"],
                  ["Gastos/Egresos", "No disponible", "No disponible", "Ilimitados"],
                  ["Vacantes activas", "No disponible", "No disponible", "Ilimitadas"],
                  ["Recursos fisicos", "No disponible", "No disponible", "Ilimitados"],
                  ["Nomina/Payroll", "No disponible", "No disponible", "Incluido"]])

    callout(doc, "Incluye TODO lo del Plan Gratuito y del Plan Basico",
            "El Plan Pro hereda todas las funcionalidades descritas en la Parte 3 (Plan Gratuito: "
            "Dashboard, Servicios, Sedes, Calendario, CRM Clientes, Facturacion, Configuraciones, "
            "Perfil Publico, QR) y en la Parte 4 (Plan Basico: Empleados, Ausencias, Historial de Ventas, "
            "Ventas Rapidas, Reportes Financieros, Permisos) sin restricciones adicionales. "
            "Esta parte documenta unicamente las funcionalidades NUEVAS y los limites ampliados del Plan Pro.",
            color=ACCENT)

    screenshot_placeholder(doc, "Sidebar del admin con todos los modulos del Plan Pro desbloqueados")

    # --- 2. Que incluye ---
    h2(doc, "2. Que incluye el Plan Pro", anchor="sec_que_incluye_pro")
    para(doc, "El Plan Pro de Gestabiz incluye tres grandes categorias de mejoras respecto al Plan Basico:")

    h3(doc, "A. Cuatro modulos exclusivos desbloqueados")
    para(doc, "Cuatro modulos que estaban bloqueados tanto en el Plan Gratuito como en el Plan Basico "
         "se desbloquean al activar el Plan Pro. Cada uno aparece en la barra lateral sin candado "
         "y es completamente funcional.")
    simple_table(doc, ["Modulo", "Icono Sidebar", "Descripcion Corta"],
                 [["Gastos", "Wallet", "Control completo de egresos con 48 categorias"],
                  ["Reclutamiento", "BriefcaseBusiness", "Publicacion de vacantes y gestion de candidatos"],
                  ["Recursos", "Box", "Gestion de recursos fisicos (salas, canchas, mesas, etc.)"],
                  ["Nomina", "(Sub-modulo Empleados)", "Configuracion de salarios y egresos automaticos"]])

    h3(doc, "B. Limites ampliados")
    para(doc, "Los limites operativos se expanden para negocios mas grandes:")
    bullet(doc, "Hasta 10 sedes activas simultaneamente (vs. 3 en Plan Basico).")
    bullet(doc, "Hasta 15 empleados activos (vs. 6 en Plan Basico).")
    bullet(doc, "Citas, clientes, servicios, gastos, vacantes y recursos: todos ilimitados.")

    h3(doc, "C. Funcionalidades adicionales")
    para(doc, "El Plan Pro tambien habilita:")
    bullet(doc, "Mapa jerarquico visual del equipo con 5 niveles (Owner, Admin, Gerente, Lider, Personal).")
    bullet(doc, "Dashboard de facturacion completo con historial de pagos y flujo de cancelacion.")
    bullet(doc, "Recordatorios por WhatsApp y sincronizacion con Google Calendar (heredados del Plan Basico).")
    bullet(doc, "Proximamente: Editor de permisos individuales (79 permisos), plantillas de permisos "
           "personalizadas e historial de auditoria de permisos.")

    page_break(doc)

    # --- 3. Comparativa de 3 planes ---
    h2(doc, "3. Comparativa de Limites: Gratuito vs. Basico vs. Pro", anchor="sec_comparativa_3planes")
    para(doc, "La siguiente tabla muestra la comparativa completa entre los tres planes disponibles:")

    simple_table(doc, ["Funcionalidad", "Gratuito ($0)", "Basico ($89,900/mes)", "Pro ($159,900/mes)"],
                 [["Sedes", "1", "3", "10"],
                  ["Empleados", "1 (owner)", "6", "15"],
                  ["Citas/mes", "50", "Ilimitadas", "Ilimitadas"],
                  ["Clientes", "50", "Ilimitados", "Ilimitados"],
                  ["Servicios", "15", "Ilimitados", "Ilimitados"],
                  ["Dashboard Admin", "Incluido", "Incluido", "Incluido"],
                  ["Calendario de Citas", "Incluido", "Incluido", "Incluido"],
                  ["CRM de Clientes", "Incluido", "Incluido", "Incluido"],
                  ["Perfil Publico + QR", "Incluido", "Incluido", "Incluido"],
                  ["Notificaciones In-App", "Incluido", "Incluido", "Incluido"],
                  ["Chat en Tiempo Real", "Incluido", "Incluido", "Incluido"],
                  ["Facturacion", "Incluido", "Incluido", "Incluido"],
                  ["Configuraciones", "Incluido", "Incluido", "Incluido"],
                  ["Gestion de Empleados", "Bloqueado", "Incluido", "Incluido"],
                  ["Ausencias/Vacaciones", "Bloqueado", "Incluido", "Incluido"],
                  ["Historial de Ventas", "Bloqueado", "Incluido", "Incluido"],
                  ["Ventas Rapidas", "Bloqueado", "Incluido", "Incluido"],
                  ["Reportes Financieros", "Bloqueado", "Incluido", "Incluido"],
                  ["Permisos (plantillas)", "Bloqueado", "Incluido", "Incluido"],
                  ["Gastos y Egresos", "Bloqueado", "Bloqueado", "Incluido"],
                  ["Reclutamiento", "Bloqueado", "Bloqueado", "Incluido"],
                  ["Recursos Fisicos", "Bloqueado", "Bloqueado", "Incluido"],
                  ["Nomina / Payroll", "Bloqueado", "Bloqueado", "Incluido"],
                  ["Mapa Jerarquico", "Bloqueado", "Bloqueado", "Incluido"],
                  ["Permisos individuales", "Bloqueado", "Bloqueado", "Proximamente"],
                  ["Auditoria de permisos", "Bloqueado", "Bloqueado", "Proximamente"],
                  ["WhatsApp Reminders", "Bloqueado", "Incluido", "Incluido"],
                  ["Google Calendar Sync", "Bloqueado", "Incluido", "Incluido"]])

    callout(doc, "Nota sobre funcionalidades \"Proximamente\"",
            "Algunas funcionalidades del Plan Pro estan marcadas como \"Proximamente\" porque la interfaz de "
            "usuario aun esta en desarrollo. El backend esta listo y funcional. Se habilitaran en futuras "
            "actualizaciones sin costo adicional para los suscriptores del Plan Pro.",
            color=PURPLE)

    screenshot_placeholder(doc, "Pagina de Precios con los 3 planes")

    page_break(doc)

    # --- 4. Como activar ---
    h2(doc, "4. Como Activar el Plan Pro", anchor="sec_activar_pro")
    para(doc, "El Plan Pro se activa desde el panel de administracion del negocio. "
         "Actualmente el boton de compra muestra \"Proximamente\" con opacidad reducida y esta deshabilitado. "
         "Cuando se habilite, el flujo sera:")

    numbered(doc, "Ingresar al panel de administracion como propietario del negocio.")
    numbered(doc, "Navegar a Facturacion en la barra lateral.")
    numbered(doc, "Hacer clic en \"Mejorar Plan\" o navegar a la pagina de Precios.")
    numbered(doc, "Seleccionar el Plan Pro ($159,900/mes o $1,599,000/ano).")
    numbered(doc, "Opcionalmente ingresar un codigo de descuento.")
    numbered(doc, "Completar el pago con tarjeta de credito/debito via Stripe, PayU o MercadoPago.")
    numbered(doc, "El plan se activa inmediatamente — todos los modulos Pro se desbloquean.")

    callout(doc, "Estado Actual: Proximamente",
            "Al momento de escribir este manual, el Plan Pro aun no esta disponible para compra. "
            "El boton muestra un badge \"Proximamente\" y el plan aparece con opacidad reducida (60%) "
            "en la pagina de precios. Se habilitara en una actualizacion futura.",
            color=DANGER)

    para(doc, "Para negocios que actualmente estan en el Plan Basico, la actualizacion al Pro sera un "
         "\"upgrade\" prorrateado: se cobra la diferencia proporcional por los dias restantes del "
         "ciclo de facturacion actual.")

    screenshot_placeholder(doc, "Pagina de precios con el Plan Pro deshabilitado (Proximamente)")

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 2 — DETALLE EXHAUSTIVO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 2 — DETALLE EXHAUSTIVO DE FUNCIONALIDADES", anchor="parte2_pro")
    para(doc, "A continuacion se documenta cada funcionalidad exclusiva del Plan Pro con "
         "nivel de detalle exhaustivo: cada campo, boton, validacion, flujo normal y alterno.")

    # ═══════════════════════════════════════════════════════════════════
    # MODULO DE GASTOS Y EGRESOS
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "MODULO DE GASTOS Y EGRESOS", anchor="sec_gastos", color=DARK)

    # --- 5. Gastos ---
    h2(doc, "5. Modulo de Gastos y Egresos", anchor="sec_gastos")
    para(doc, "El modulo de Gastos y Egresos (ExpensesManagementPage) permite a los administradores "
         "registrar, categorizar y controlar todos los gastos operativos del negocio. Es exclusivo del "
         "Plan Pro y aparece en la barra lateral con el icono Wallet.")

    para(doc, "El modulo se compone de tres secciones principales:")
    bullet(doc, "Cabecera con 3 tarjetas de estadisticas de gastos (hoy, 7 dias, mes).")
    bullet(doc, "3 pestanas de contenido: Egresos Unicos, Egresos Recurrentes, Resumen por Categoria.")
    bullet(doc, "Boton \"Nuevo Egreso\" protegido por permisos (accounting.create).")

    callout(doc, "Proteccion por Plan",
            "Este modulo esta protegido por PlanGate con feature=\"expenses\". Si el negocio no tiene "
            "el Plan Pro activo, se muestra una pantalla de bloqueo invitando a actualizar el plan.",
            color=PURPLE)

    screenshot_placeholder(doc, "Vista general del modulo de Gastos y Egresos")

    # --- 6. Stats ---
    h2(doc, "6. Tarjetas de Estadisticas de Gastos", anchor="sec_gastos_stats")
    para(doc, "En la parte superior del modulo se muestran 3 tarjetas de estadisticas que resumen "
         "los gastos del negocio en tres periodos temporales:")

    simple_table(doc, ["Tarjeta", "Periodo", "Color", "Descripcion"],
                 [["Gastos de Hoy", "Dia actual", "Rojo (destructive)", "Suma de egresos registrados hoy"],
                  ["Gastos Ultimos 7 Dias", "7 dias atras", "Rojo (destructive)", "Suma de egresos en la ultima semana"],
                  ["Gastos del Mes", "Mes actual", "Rojo (destructive)", "Suma de egresos del mes en curso"]])

    para(doc, "Cada tarjeta muestra el monto total en formato COP con separadores de miles "
         "(ejemplo: $1,250,000 COP). El color rojo (variant=\"destructive\") refuerza visualmente "
         "que se trata de salidas de dinero.")

    # --- 7. Egresos Unicos ---
    h2(doc, "7. Pestana Egresos Unicos", anchor="sec_gastos_unicos")
    para(doc, "La primera pestana muestra una tabla con todos los egresos de tipo unico (no recurrente) "
         "registrados en el negocio. Estos son gastos puntuales que se registran una sola vez.")

    simple_table(doc, ["Columna", "Descripcion"],
                 [["Fecha", "Fecha en que se registro el egreso (formato dd/mm/yyyy)"],
                  ["Categoria", "Categoria del gasto (ej: Mantenimiento, Marketing, Impuestos)"],
                  ["Monto", "Monto en COP con formato de miles"],
                  ["Sede", "Sede asociada al gasto (opcional, puede estar vacia)"],
                  ["Descripcion", "Descripcion libre del gasto"],
                  ["Notas", "Notas adicionales opcionales"],
                  ["Metodo de Pago", "Metodo utilizado (efectivo, tarjeta, transferencia, etc.)"]])

    para(doc, "La tabla se ordena por fecha descendente (mas recientes primero). "
         "Los egresos unicos se almacenan en la tabla transactions con type='expense'.")

    screenshot_placeholder(doc, "Tabla de Egresos Unicos")

    # --- 8. Egresos Recurrentes ---
    h2(doc, "8. Pestana Egresos Recurrentes", anchor="sec_gastos_recurrentes")
    para(doc, "La segunda pestana muestra una tabla con todos los egresos recurrentes configurados "
         "en el negocio. Estos son gastos que se repiten periodicamente (mensual, quincenal, semanal).")

    simple_table(doc, ["Columna", "Descripcion"],
                 [["Nombre", "Nombre descriptivo del egreso recurrente (ej: \"Arriendo Local Principal\")"],
                  ["Categoria", "Categoria del gasto"],
                  ["Monto", "Monto en COP que se cobra en cada periodo"],
                  ["Frecuencia", "Periodicidad: mensual, quincenal, semanal"],
                  ["Dia", "Dia del mes/semana en que se ejecuta"],
                  ["Automatizado", "Switch indicando si el sistema lo genera automaticamente"],
                  ["Estado", "Activo o Inactivo"],
                  ["Acciones", "Editar, Activar/Desactivar, Eliminar"]])

    para(doc, "Cada fila permite 3 acciones:")
    bullet(doc, "Editar: Abre el formulario prellenado con los datos actuales del egreso.")
    bullet(doc, "Activar/Desactivar (toggle): Cambia el estado del egreso recurrente sin eliminarlo.")
    bullet(doc, "Eliminar: Elimina el egreso recurrente de forma permanente (con confirmacion).")

    para(doc, "Los egresos recurrentes se almacenan en la tabla recurring_expenses. "
         "Cuando estan configurados como automatizados, el sistema genera una transaccion "
         "de tipo expense en la fecha indicada.")

    screenshot_placeholder(doc, "Tabla de Egresos Recurrentes con acciones")

    # --- 9. Resumen por Categoria ---
    h2(doc, "9. Pestana Resumen por Categoria", anchor="sec_gastos_resumen")
    para(doc, "La tercera pestana muestra un resumen agrupado de los gastos por categoria. "
         "Cada categoria se presenta con el monto total acumulado, permitiendo al administrador "
         "identificar rapidamente en que areas se concentran los egresos del negocio.")

    para(doc, "Las categorias se presentan como tarjetas o filas agrupadas con:")
    bullet(doc, "Nombre de la categoria (ej: \"Nomina y Salarios\").")
    bullet(doc, "Monto total acumulado en COP.")
    bullet(doc, "Porcentaje del total de gastos.")
    bullet(doc, "Indicador visual de proporcion.")

    screenshot_placeholder(doc, "Resumen de Gastos por Categoria")

    # --- 10. Las 48 categorias ---
    h2(doc, "10. Las 48 Categorias de Gastos", anchor="sec_gastos_categorias")
    para(doc, "El modulo ofrece 48 categorias de gastos organizadas en 10 grupos tematicos. "
         "Al crear un nuevo egreso, el administrador selecciona una de estas categorias:")

    simple_table(doc, ["Grupo", "Categorias Incluidas"],
                 [["Nomina y Salarios", "Salarios, Bonificaciones, Prestaciones sociales, Horas extra, Liquidaciones"],
                  ["Arriendo e Inmuebles", "Arriendo, Administracion, Servicios publicos (agua, luz, gas, internet, telefono)"],
                  ["Mantenimiento", "Mantenimiento preventivo, Mantenimiento correctivo, Reparaciones, Limpieza"],
                  ["Marketing y Publicidad", "Publicidad digital, Publicidad impresa, Redes sociales, Eventos, Material POP"],
                  ["Impuestos y Obligaciones", "IVA, ICA, Retencion en la fuente, Predial, Camara de comercio, Industria y comercio"],
                  ["Seguros", "Seguro de responsabilidad civil, Seguro de empleados, Seguro de equipos, Poliza todo riesgo"],
                  ["Capacitacion", "Cursos, Certificaciones, Conferencias, Material didactico"],
                  ["Equipos y Tecnologia", "Computadores, Software, Licencias, Mobiliario, Herramientas"],
                  ["Transporte y Logistica", "Combustible, Peajes, Parqueadero, Mensajeria, Envios"],
                  ["Honorarios y Servicios", "Contabilidad, Asesoria legal, Consultoria, Servicios profesionales, Otros"]])

    para(doc, "Esta clasificacion permite generar reportes financieros detallados por categoria "
         "y facilita el analisis de costos operativos del negocio.")

    page_break(doc)

    # --- 11. Formulario de Nuevo Egreso ---
    h2(doc, "11. Formulario de Nuevo Egreso — Campos", anchor="sec_gastos_form")
    para(doc, "Al hacer clic en \"Nuevo Egreso\", se abre un formulario completo con los siguientes campos:")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Nombre", "Texto", "Solo si es recurrente", "Nombre descriptivo del egreso recurrente"],
                  ["Categoria", "Select (48 opciones)", "Si", "Categoria del gasto (ver tabla anterior)"],
                  ["Monto", "Numero (COP)", "Si (> $0)", "Monto del egreso en pesos colombianos"],
                  ["Sede", "Select (sedes del negocio)", "No", "Sede asociada al gasto (opcional)"],
                  ["Descripcion", "Textarea", "No", "Descripcion libre del gasto"],
                  ["Es Recurrente", "Checkbox", "No", "Marca si el egreso se repite periodicamente"],
                  ["Frecuencia", "Select", "Si (si recurrente)", "mensual, quincenal, semanal"],
                  ["Dia de Ejecucion", "Select", "Si (si recurrente)", "Dia del mes o semana para el egreso"],
                  ["Automatizado", "Switch", "No", "Si el sistema debe generar el egreso automaticamente"],
                  ["Metodo de Pago", "Select", "No", "Forma de pago utilizada"],
                  ["Notas", "Textarea", "No", "Notas adicionales"]])

    h4(doc, "Validaciones del formulario:")
    bullet(doc, "El monto debe ser mayor a $0 COP.")
    bullet(doc, "Si el checkbox \"Es Recurrente\" esta marcado, se habilitan los campos de Frecuencia, Dia y Automatizado.")
    bullet(doc, "Si no es recurrente, esos campos se ocultan.")
    bullet(doc, "La categoria es obligatoria — el formulario no se puede enviar sin seleccionarla.")

    h4(doc, "Flujo de guardado:")
    bullet(doc, "Si es recurrente: se ejecuta un UPSERT en la tabla recurring_expenses con los datos del formulario.")
    bullet(doc, "Si es unico (no recurrente): se ejecuta un INSERT en la tabla transactions con type='expense'.")
    bullet(doc, "En ambos casos se muestra un toast de confirmacion al guardar exitosamente.")

    screenshot_placeholder(doc, "Formulario de Nuevo Egreso con campos visibles")

    # --- 12. Egresos recurrentes config ---
    h2(doc, "12. Configuracion de Egresos Recurrentes", anchor="sec_gastos_recurrente_cfg")
    para(doc, "Los egresos recurrentes tienen configuracion adicional que permite controlar "
         "cuando y como se generan las transacciones automaticamente:")

    simple_table(doc, ["Parametro", "Opciones", "Descripcion"],
                 [["Frecuencia", "mensual, quincenal, semanal", "Con que periodicidad se repite el egreso"],
                  ["Dia de Ejecucion", "1-28 o \"ultimo dia\"", "En que dia del periodo se genera el egreso"],
                  ["Automatizado", "Si / No", "Si el sistema genera la transaccion sin intervencion manual"],
                  ["Estado", "Activo / Inactivo", "Los egresos inactivos no generan transacciones"]])

    para(doc, "Cuando un egreso recurrente esta configurado como automatizado y activo, el sistema "
         "genera automaticamente una transaccion de tipo expense en la tabla transactions en la "
         "fecha configurada. Si no esta automatizado, el administrador debe registrar el egreso "
         "manualmente cada periodo.")

    # --- 13. Metodos de Pago ---
    h2(doc, "13. Metodos de Pago para Egresos", anchor="sec_gastos_metodos")
    para(doc, "Al registrar un egreso, el administrador puede seleccionar el metodo de pago utilizado:")
    bullet(doc, "Efectivo")
    bullet(doc, "Tarjeta de Credito")
    bullet(doc, "Tarjeta de Debito")
    bullet(doc, "Transferencia Bancaria")
    bullet(doc, "PSE (Pagos Seguros en Linea)")
    bullet(doc, "Nequi")
    bullet(doc, "Daviplata")
    bullet(doc, "Cheque")
    bullet(doc, "Otro")

    para(doc, "El metodo de pago es informativo — se almacena en la transaccion pero no afecta "
         "el procesamiento del egreso.")

    # --- 14. Permisos del modulo ---
    h2(doc, "14. Permisos del Modulo de Gastos", anchor="sec_gastos_permisos")
    para(doc, "El modulo de Gastos esta protegido por los siguientes permisos:")

    simple_table(doc, ["Permiso", "Que protege", "Modo PermissionGate"],
                 [["accounting.create", "Boton \"Nuevo Egreso\"", "hide"],
                  ["expenses.create", "Creacion de egresos individuales", "hide"],
                  ["expenses.delete", "Eliminacion de egresos", "hide"]])

    para(doc, "Si el usuario no tiene el permiso accounting.create, el boton \"Nuevo Egreso\" no es "
         "visible. El propietario del negocio (owner) siempre tiene todos los permisos sin restriccion.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # MODULO DE RECLUTAMIENTO
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "MODULO DE RECLUTAMIENTO", anchor="sec_reclutamiento", color=DARK)

    # --- 15. Dashboard ---
    h2(doc, "15. Modulo de Reclutamiento", anchor="sec_reclutamiento")
    para(doc, "El modulo de Reclutamiento (RecruitmentDashboard) es un sistema completo de publicacion "
         "de vacantes laborales, recepcion de aplicaciones y gestion de candidatos con flujo de estados, "
         "reviews obligatorias y notificaciones automaticas. Es exclusivo del Plan Pro.")

    para(doc, "El modulo ofrece tres vistas principales:")
    bullet(doc, "Lista de Vacantes (VacancyList): vista principal con todas las vacantes del negocio.")
    bullet(doc, "Crear Vacante (CreateVacancy): formulario de 4 tarjetas para publicar una nueva vacante.")
    bullet(doc, "Gestion de Aplicaciones (ApplicationsManagement): vista de aplicaciones por vacante.")

    callout(doc, "Proteccion por Plan",
            "Este modulo esta protegido por PlanGate con feature=\"recruitment\". Si el negocio no tiene "
            "el Plan Pro activo, se muestra una pantalla de bloqueo.",
            color=PURPLE)

    screenshot_placeholder(doc, "Vista general del modulo de Reclutamiento")

    # --- 16. Dashboard ---
    h2(doc, "16. Dashboard de Reclutamiento — Vista General", anchor="sec_recl_dashboard")
    para(doc, "El dashboard de reclutamiento se organiza con una navegacion de pestanas "
         "y una vista activa que cambia segun la seleccion:")

    simple_table(doc, ["Pestana", "Vista Mostrada", "Descripcion"],
                 [["Vacantes Activas", "VacancyList (filtro open)", "Vacantes con estado 'open' (abiertas a aplicaciones)"],
                  ["Historial", "VacancyList (todas)", "Todas las vacantes incluyendo cerradas y pausadas"]])

    para(doc, "En la vista de vacantes activas, un boton \"Nueva Vacante\" protegido por el permiso "
         "recruitment.create_vacancy permite crear vacantes nuevas.")

    # --- 17. Crear Vacante ---
    h2(doc, "17. Crear Nueva Vacante — 4 Tarjetas", anchor="sec_recl_crear")
    para(doc, "El formulario de creacion de vacante (CreateVacancy) se organiza en 4 tarjetas (cards) "
         "que agrupan los campos por seccion logica. El administrador completa los campos y hace clic "
         "en \"Publicar Vacante\" para guardarla.")

    para(doc, "Las 4 tarjetas son:")
    numbered(doc, "Informacion Basica: titulo, descripcion, tipo de posicion y experiencia requerida.")
    numbered(doc, "Detalles: requisitos, responsabilidades y beneficios del puesto.")
    numbered(doc, "Compensacion: rango salarial, moneda, comision, sede y modalidad remota.")
    numbered(doc, "Estado: estado inicial de la vacante (abierta, pausada o cerrada).")

    screenshot_placeholder(doc, "Formulario de Crear Vacante con las 4 tarjetas")

    # --- 18. Tarjeta 1 ---
    h2(doc, "18. Tarjeta 1: Informacion Basica de la Vacante", anchor="sec_recl_info")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Titulo", "Texto", "Si", "Titulo de la vacante (ej: \"Estilista Senior\")"],
                  ["Descripcion", "Textarea", "Si", "Descripcion detallada del puesto y responsabilidades generales"],
                  ["Tipo de Posicion", "Select (4 opciones)", "Si", "full_time, part_time, contractor, temporary"],
                  ["Experiencia", "Select (3 niveles)", "Si", "junior, mid, senior"]])

    h4(doc, "Opciones de Tipo de Posicion:")
    bullet(doc, "Tiempo Completo (full_time): jornada completa de 8 horas.")
    bullet(doc, "Medio Tiempo (part_time): jornada parcial de 4 horas.")
    bullet(doc, "Contratista (contractor): contratacion por prestacion de servicios.")
    bullet(doc, "Temporal (temporary): contrato por duracion determinada.")

    h4(doc, "Opciones de Experiencia:")
    bullet(doc, "Junior: 0-2 anos de experiencia.")
    bullet(doc, "Mid: 2-5 anos de experiencia.")
    bullet(doc, "Senior: 5+ anos de experiencia.")

    # --- 19. Tarjeta 2 ---
    h2(doc, "19. Tarjeta 2: Detalles de la Vacante", anchor="sec_recl_detalles")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Requisitos", "Textarea", "No", "Lista de requisitos del puesto (habilidades, certificaciones)"],
                  ["Responsabilidades", "Textarea", "No", "Lista de responsabilidades del puesto"],
                  ["Beneficios", "Textarea", "No", "Beneficios que ofrece el negocio (descuentos, horario flexible, etc.)"]])

    para(doc, "Los tres campos son textareas de formato libre. Se recomienda usar listas con viñetas "
         "para facilitar la lectura por parte de los candidatos.")

    # --- 20. Tarjeta 3 ---
    h2(doc, "20. Tarjeta 3: Compensacion y Ubicacion", anchor="sec_recl_compensacion")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Salario Minimo", "Numero (COP)", "No", "Limite inferior del rango salarial"],
                  ["Salario Maximo", "Numero (COP)", "No", "Limite superior del rango salarial"],
                  ["Moneda", "Select", "Si", "Moneda del salario (default: COP)"],
                  ["Basado en Comision", "Switch", "No", "Indica si la compensacion incluye comisiones"],
                  ["Sede", "Select (sedes del negocio)", "No", "Sede donde se desempenara el puesto"],
                  ["Modalidad Remota", "Switch", "No", "Indica si el puesto permite trabajo remoto"]])

    para(doc, "Si el negocio tiene una sede preferida configurada, el campo de Sede se preselecciona "
         "automaticamente con esa sede. El administrador puede cambiarla si lo desea.")

    para(doc, "El campo \"Basado en Comision\" activa un indicador visual en la tarjeta de la vacante "
         "y en el formulario de aplicacion, informando al candidato que parte de la compensacion "
         "depende de comisiones por ventas o servicios.")

    # --- 21. Tarjeta 4 ---
    h2(doc, "21. Tarjeta 4: Estado de la Vacante", anchor="sec_recl_estado")

    simple_table(doc, ["Estado", "Descripcion", "Efecto"],
                 [["open (Abierta)", "Acepta nuevas aplicaciones", "Visible en vacantes activas; candidatos pueden aplicar"],
                  ["paused (Pausada)", "Temporalmente cerrada", "No acepta nuevas aplicaciones pero mantiene las existentes"],
                  ["closed (Cerrada)", "Finalizada", "No acepta aplicaciones; solo visible en historial"],
                  ["filled (Ocupada)", "Puesto cubierto", "Indica que ya se selecciono un candidato"],
                  ["expired (Expirada)", "Vencida", "La vacante supero su fecha limite"]])

    para(doc, "El estado inicial puede ser cualquiera de los tres primeros. Los estados \"filled\" "
         "y \"expired\" se asignan automaticamente cuando se selecciona un candidato o la vacante "
         "supera su fecha limite.")

    page_break(doc)

    # --- 22. Lista de Vacantes ---
    h2(doc, "22. Lista de Vacantes — Filtros y Busqueda", anchor="sec_recl_lista")
    para(doc, "La lista de vacantes (VacancyList) muestra todas las vacantes del negocio "
         "con opciones de filtrado y busqueda:")

    h4(doc, "Filtros disponibles:")
    simple_table(doc, ["Filtro", "Opciones", "Descripcion"],
                 [["Estado", "5 opciones: Todas, Abiertas, Pausadas, Cerradas, Ocupadas", "Filtra vacantes por estado actual"],
                  ["Tipo de Posicion", "4 opciones: Todos, Tiempo Completo, Medio Tiempo, Contratista, Temporal", "Filtra por tipo de contratacion"],
                  ["Busqueda", "Texto libre", "Busca en titulo y descripcion de la vacante"]])

    para(doc, "Los filtros se aplican en tiempo real — la lista se actualiza inmediatamente "
         "al cambiar cualquier filtro.")

    # --- 23. Tarjeta de Vacante ---
    h2(doc, "23. Tarjeta de Vacante — Detalle Visual", anchor="sec_recl_card")
    para(doc, "Cada vacante se muestra como una tarjeta (VacancyCard) con los siguientes elementos:")

    bullet(doc, "Cabecera con titulo de la vacante, badge de estado (color segun estado) y menu dropdown con acciones.")
    bullet(doc, "Grid de metadatos con 4 columnas: Tipo de posicion, Experiencia requerida, Rango salarial, Sede asignada.")
    bullet(doc, "Pie con estadisticas: numero de aplicaciones recibidas, fecha de publicacion.")
    bullet(doc, "Animacion highlight al crear una vacante nueva (la tarjeta destella brevemente).")

    h4(doc, "Acciones del menu dropdown:")
    bullet(doc, "Ver Aplicaciones: navega a la gestion de aplicaciones de esta vacante.")
    bullet(doc, "Editar: abre el formulario de edicion prellenado (permiso: recruitment.edit_vacancy).")
    bullet(doc, "Pausar/Reanudar: cambia entre estado open y paused.")
    bullet(doc, "Cerrar: cambia a estado closed (permiso: recruitment.edit_vacancy).")
    bullet(doc, "Eliminar: elimina la vacante permanentemente (permiso: recruitment.delete_vacancy, requiere confirmacion).")

    screenshot_placeholder(doc, "Tarjeta de Vacante con metadata y acciones")

    # --- 24. Gestion de Aplicaciones ---
    h2(doc, "24. Gestion de Aplicaciones", anchor="sec_recl_apps")
    para(doc, "La vista de Gestion de Aplicaciones (ApplicationsManagement) se accede al hacer clic "
         "en \"Ver Aplicaciones\" de una vacante. Muestra todas las aplicaciones recibidas "
         "organizadas por estado con estadisticas y filtros avanzados.")

    para(doc, "La vista se compone de:")
    bullet(doc, "7 tarjetas de estadisticas en la parte superior.")
    bullet(doc, "6 pestanas de estado para filtrar aplicaciones.")
    bullet(doc, "3 columnas de filtros adicionales.")
    bullet(doc, "Lista de ApplicationCards con acciones contextuales.")

    # --- 25. Las 6 pestanas ---
    h2(doc, "25. Las 6 Pestanas de Estado de Aplicaciones", anchor="sec_recl_tabs")

    simple_table(doc, ["Pestana", "Estado", "Color Badge", "Descripcion"],
                 [["Pendientes", "pending", "Amarillo", "Aplicaciones nuevas sin revisar"],
                  ["En Revision", "reviewing", "Azul", "Aplicaciones que estan siendo evaluadas"],
                  ["En Proceso", "in_process", "Indigo", "Candidatos en proceso de seleccion avanzado"],
                  ["Aceptadas", "accepted", "Verde", "Candidatos seleccionados y aceptados"],
                  ["Rechazadas", "rejected", "Rojo", "Candidatos descartados"],
                  ["Retiradas", "withdrawn", "Gris", "Candidatos que retiraron su aplicacion"]])

    para(doc, "Cada pestana muestra un contador con el numero de aplicaciones en ese estado. "
         "Las pestanas se actualizan en tiempo real al cambiar el estado de una aplicacion.")

    # --- 26. 7 Stats ---
    h2(doc, "26. 7 Tarjetas de Estadisticas de Aplicaciones", anchor="sec_recl_stats")

    simple_table(doc, ["Tarjeta", "Dato Mostrado"],
                 [["Total Aplicaciones", "Numero total de aplicaciones recibidas para esta vacante"],
                  ["Pendientes", "Aplicaciones sin revisar"],
                  ["En Revision", "Aplicaciones en evaluacion"],
                  ["En Proceso", "Candidatos en seleccion avanzada"],
                  ["Aceptados", "Candidatos seleccionados"],
                  ["Rechazados", "Candidatos descartados"],
                  ["Tasa de Aceptacion", "Porcentaje de aceptados sobre el total"]])

    # --- 27. ApplicationCard ---
    h2(doc, "27. Tarjeta de Aplicacion — Detalle", anchor="sec_recl_app_card")
    para(doc, "Cada aplicacion se muestra como una tarjeta (ApplicationCard) con los siguientes datos:")

    bullet(doc, "Avatar del candidato con iniciales (si no tiene foto).")
    bullet(doc, "Nombre completo, email y telefono del candidato.")
    bullet(doc, "Badge de estado con color correspondiente.")
    bullet(doc, "Informacion de la vacante a la que aplico.")
    bullet(doc, "Rango salarial esperado por el candidato.")
    bullet(doc, "Disponibilidad horaria del candidato.")
    bullet(doc, "Carta de presentacion (cover letter) si la envio.")
    bullet(doc, "Botones de accion contextuales segun el estado actual de la aplicacion.")

    screenshot_placeholder(doc, "Tarjeta de Aplicacion con datos del candidato")

    # --- 28. Acciones por Estado ---
    h2(doc, "28. Acciones por Estado de Aplicacion", anchor="sec_recl_acciones")
    para(doc, "Las acciones disponibles cambian segun el estado actual de la aplicacion:")

    simple_table(doc, ["Estado Actual", "Acciones Disponibles"],
                 [["pending (Pendiente)", "Mover a Revision, Rechazar, Ver Perfil"],
                  ["reviewing (En Revision)", "Mover a En Proceso, Rechazar, Ver Perfil"],
                  ["in_process (En Proceso)", "Seleccionar (Aceptar), Rechazar, Ver Perfil"],
                  ["accepted (Aceptada)", "Ver Perfil, (sin mas acciones)"],
                  ["rejected (Rechazada)", "Ver Perfil, (sin mas acciones)"],
                  ["withdrawn (Retirada)", "Ver Perfil, (sin mas acciones)"]])

    para(doc, "Las acciones de Mover a Revision, Mover a En Proceso, Seleccionar y Rechazar "
         "requieren el permiso recruitment.manage_applications. Sin este permiso, los botones "
         "no son visibles (mode=\"hide\").")

    page_break(doc)

    # --- 29. Modal Seleccionar Empleado ---
    h2(doc, "29. Modal de Seleccionar Empleado", anchor="sec_recl_seleccionar")
    para(doc, "Al hacer clic en \"Seleccionar\" en una aplicacion en estado \"in_process\", "
         "se abre un AlertDialog que confirma la seleccion del candidato. Este modal es "
         "critico porque tiene efectos permanentes en el sistema.")

    h4(doc, "Efectos de seleccionar un candidato:")
    numbered(doc, "El estado de la aplicacion cambia a \"accepted\".")
    numbered(doc, "El candidato se registra automaticamente como empleado en business_employees.")
    numbered(doc, "Se le asigna el rol de \"employee\" en el negocio.")
    numbered(doc, "Se genera una review obligatoria para el administrador (calificar al candidato).")
    numbered(doc, "Todas las demas aplicaciones a esta vacante se rechazan automaticamente.")

    callout(doc, "Alerta de Notificaciones Automaticas",
            "Al seleccionar un candidato, el sistema envia automaticamente notificaciones al candidato "
            "aceptado (felicitacion) y a los demas candidatos (rechazo con mensaje personalizado). "
            "El administrador debe tener esto en cuenta antes de confirmar la seleccion.",
            color=DANGER)

    screenshot_placeholder(doc, "Modal de Seleccionar Empleado con lista de efectos")

    # --- 30. Dialogo Rechazo ---
    h2(doc, "30. Dialogo de Rechazo", anchor="sec_recl_rechazar")
    para(doc, "Al rechazar una aplicacion, se abre un dialogo con un campo de texto (textarea) "
         "donde el administrador puede escribir opcionalmente el motivo del rechazo.")

    para(doc, "El motivo de rechazo:")
    bullet(doc, "Es opcional — el administrador puede rechazar sin dar motivo.")
    bullet(doc, "Si se proporciona, se almacena en la aplicacion y puede ser consultado.")
    bullet(doc, "No se envia al candidato por email (solo se notifica el cambio de estado).")
    bullet(doc, "Es util para fines internos y de auditoria del proceso de seleccion.")

    # --- 31. Perfil del Aplicante ---
    h2(doc, "31. Modal de Perfil del Aplicante — 3 Pestanas", anchor="sec_recl_perfil")
    para(doc, "Al hacer clic en \"Ver Perfil\" de cualquier aplicacion, se abre un modal "
         "(ApplicantProfileModal) con 3 pestanas de informacion detallada:")

    h3(doc, "Pestana 1: Informacion Personal")
    bullet(doc, "Nombre completo, email, telefono.")
    bullet(doc, "Documento de identidad (tipo y numero).")
    bullet(doc, "Avatar o iniciales.")
    bullet(doc, "Fecha de registro en la plataforma.")

    h3(doc, "Pestana 2: Experiencia Profesional")
    bullet(doc, "Habilidades registradas en su perfil profesional.")
    bullet(doc, "Anos de experiencia.")
    bullet(doc, "Certificaciones.")
    bullet(doc, "Historial de trabajos anteriores (si los registro).")
    bullet(doc, "CV adjunto (si lo subio — enlace de descarga).")

    h3(doc, "Pestana 3: Aplicacion Actual")
    bullet(doc, "Estado actual de la aplicacion con badge de color.")
    bullet(doc, "Fecha de aplicacion.")
    bullet(doc, "Carta de presentacion completa.")
    bullet(doc, "Disponibilidad horaria declarada.")
    bullet(doc, "Expectativa salarial.")
    bullet(doc, "Historial de cambios de estado.")

    screenshot_placeholder(doc, "Modal de Perfil del Aplicante con las 3 pestanas")

    # --- 32. Reviews Obligatorias ---
    h2(doc, "32. Reviews Obligatorias al Contratar y Finalizar", anchor="sec_recl_reviews")
    para(doc, "El sistema de reclutamiento integra un mecanismo de reviews obligatorias "
         "en dos momentos clave del proceso:")

    h3(doc, "Review al Contratar")
    para(doc, "Cuando el administrador selecciona (acepta) un candidato, el sistema genera "
         "automaticamente una review pendiente obligatoria. El administrador debe calificar "
         "al nuevo empleado con un rating de 1 a 5 estrellas y un comentario.")

    h3(doc, "Review al Finalizar")
    para(doc, "Si un empleado contratado via reclutamiento es dado de baja o finaliza su relacion "
         "laboral, se genera otra review obligatoria para documentar la experiencia.")

    para(doc, "Las reviews obligatorias aparecen como un modal (MandatoryReviewModal) que se "
         "muestra automaticamente al administrador. No se puede cerrar sin completar la review.")

    callout(doc, "Hook: useMandatoryReviews",
            "Las reviews obligatorias pendientes se consultan con el hook useMandatoryReviews. "
            "Si hay reviews pendientes, el modal se muestra automaticamente al entrar al modulo.",
            color=PURPLE)

    # --- 33. Notificaciones ---
    h2(doc, "33. Notificaciones Automaticas de Reclutamiento", anchor="sec_recl_notificaciones")
    para(doc, "El modulo de reclutamiento genera notificaciones automaticas en los siguientes eventos:")

    simple_table(doc, ["Evento", "Destinatario", "Canal", "Contenido"],
                 [["Nueva aplicacion recibida", "Admins del negocio", "In-app + Email", "\"Juan Perez aplico a Estilista Senior\""],
                  ["Aplicacion movida a revision", "Candidato", "Email", "\"Tu aplicacion esta siendo revisada\""],
                  ["Candidato seleccionado", "Candidato aceptado", "In-app + Email", "\"Felicitaciones! Has sido seleccionado\""],
                  ["Candidato rechazado", "Candidato rechazado", "Email", "\"Gracias por tu interes. En esta oportunidad...\""],
                  ["Vacante cerrada", "Todos los candidatos pendientes", "Email", "\"La vacante ha sido cerrada\""]])

    para(doc, "Las notificaciones se envian a traves de las Edge Functions "
         "send-selection-notifications y send-employee-request-notification.")

    # --- 34. Permisos ---
    h2(doc, "34. Permisos del Modulo de Reclutamiento", anchor="sec_recl_permisos")

    simple_table(doc, ["Permiso", "Que protege", "Modo PermissionGate"],
                 [["recruitment.create_vacancy", "Boton \"Nueva Vacante\"", "hide"],
                  ["recruitment.edit_vacancy", "Editar y cerrar vacantes", "hide"],
                  ["recruitment.delete_vacancy", "Eliminar vacantes", "hide"],
                  ["recruitment.manage_applications", "Cambiar estado de aplicaciones, seleccionar, rechazar", "hide"]])

    para(doc, "El propietario del negocio (owner) tiene todos los permisos automaticamente. "
         "Los demas usuarios necesitan permisos asignados explicitamente o a traves de una plantilla.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # MODULO DE RECURSOS FISICOS
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "MODULO DE RECURSOS FISICOS", anchor="sec_recursos", color=DARK)

    # --- 35. Vista General ---
    h2(doc, "35. Modulo de Recursos Fisicos", anchor="sec_recursos")
    para(doc, "El modulo de Recursos Fisicos (ResourcesManager) permite a los administradores gestionar "
         "recursos fisicos del negocio como salas, canchas, mesas, equipos, vehiculos y mas. "
         "Es exclusivo del Plan Pro y esta disenado para negocios cuyo modelo no se basa "
         "unicamente en profesionales sino tambien en la reserva de espacios o equipos.")

    para(doc, "Casos de uso principales:")
    bullet(doc, "Hoteles: gestionar habitaciones (standard, suite, deluxe).", bold_prefix="Hoteles")
    bullet(doc, "Restaurantes: gestionar mesas (2, 4, 6, 8 personas).", bold_prefix="Restaurantes")
    bullet(doc, "Centros Deportivos: gestionar canchas (tenis, futbol, padel).", bold_prefix="Centros Deportivos")
    bullet(doc, "Gimnasios: gestionar equipos (caminadora, bicicleta, banco de pesas).", bold_prefix="Gimnasios")
    bullet(doc, "Co-working: gestionar espacios (escritorio, sala de reuniones).", bold_prefix="Co-working")
    bullet(doc, "Bowling: gestionar carriles de boliche.", bold_prefix="Bowling")
    bullet(doc, "Parqueaderos: gestionar espacios de estacionamiento.", bold_prefix="Parqueaderos")
    bullet(doc, "Hospitales: gestionar camas y consultorios.", bold_prefix="Hospitales")

    callout(doc, "Condicion de Visibilidad",
            "El modulo de Recursos solo es visible si el negocio tiene configurado un resource_model "
            "diferente de 'professional'. Si el negocio opera unicamente con empleados (salon de belleza, "
            "clinica), el modulo no aparece en la barra lateral.",
            color=PURPLE)

    screenshot_placeholder(doc, "Vista general del modulo de Recursos Fisicos")

    # --- 36. Vista General ---
    h2(doc, "36. Vista General de Recursos", anchor="sec_rec_dashboard")
    para(doc, "La vista principal del modulo muestra:")
    bullet(doc, "Cabecera con titulo \"Recursos\" y boton \"Agregar Recurso\".")
    bullet(doc, "Filtro por tipo de recurso (select con 15 opciones + \"Todos\").")
    bullet(doc, "Tabla con todos los recursos del negocio.")

    # --- 37. Los 15 Tipos ---
    h2(doc, "37. Los 15 Tipos de Recursos", anchor="sec_rec_tipos")
    para(doc, "El sistema soporta 15 tipos de recursos fisicos, cubriendo una amplia gama "
         "de verticales de negocio:")

    simple_table(doc, ["Tipo", "Clave", "Ejemplo de Uso"],
                 [["Habitacion", "room", "Hotel: Suite 101, Habitacion Estandar"],
                  ["Mesa", "table", "Restaurante: Mesa 1 (4 personas), Mesa VIP"],
                  ["Cancha", "court", "Centro deportivo: Cancha de Padel, Cancha de Tenis"],
                  ["Escritorio", "desk", "Co-working: Escritorio A1, Escritorio Premium"],
                  ["Equipo", "equipment", "Gimnasio: Caminadora 1, Bicicleta Estatica"],
                  ["Vehiculo", "vehicle", "Rent-a-car: Toyota Corolla, Van de 12 pasajeros"],
                  ["Espacio", "space", "Salon de eventos: Salon Principal (100 personas)"],
                  ["Carril", "lane", "Bowling: Carril 1, Carril 2"],
                  ["Campo/Cancha Exterior", "field", "Club: Cancha de Futbol, Campo de Golf (Hoyo 1-9)"],
                  ["Estacion", "station", "Fabrica: Estacion de Soldadura, Estacion de Corte"],
                  ["Parqueadero", "parking_spot", "Edificio: Espacio P-01, Espacio P-02"],
                  ["Cama", "bed", "Hospital: Cama UCI 1, Cama Recuperacion 3"],
                  ["Estudio", "studio", "Fotografia: Estudio A (luces incluidas)"],
                  ["Sala de Reuniones", "meeting_room", "Oficina: Sala Board, Sala de Videoconferencia"],
                  ["Otro", "other", "Cualquier recurso que no encaje en las categorias anteriores"]])

    # --- 38. Filtros ---
    h2(doc, "38. Filtros por Tipo de Recurso", anchor="sec_rec_filtros")
    para(doc, "En la parte superior de la tabla hay un select que permite filtrar los recursos "
         "por tipo. Las opciones incluyen los 15 tipos mas la opcion \"Todos\" para ver "
         "todos los recursos sin filtro.")

    para(doc, "El filtro se aplica en tiempo real — la tabla se actualiza inmediatamente al "
         "seleccionar un tipo diferente.")

    # --- 39. Tabla ---
    h2(doc, "39. Tabla de Recursos — Columnas", anchor="sec_rec_tabla")

    simple_table(doc, ["Columna", "Descripcion"],
                 [["Nombre", "Nombre identificador del recurso (ej: \"Cancha de Padel 1\")"],
                  ["Tipo", "Badge con el tipo de recurso (ej: court, room, table)"],
                  ["Sede", "Sede donde se ubica el recurso"],
                  ["Capacidad", "Numero de personas/unidades que soporta"],
                  ["Precio/hora", "Tarifa por hora en COP (formato con separadores de miles)"],
                  ["Estado", "Badge de estado: Disponible (verde), En Mantenimiento (amarillo), Fuera de Servicio (rojo)"],
                  ["Acciones", "Botones Editar y Eliminar"]])

    para(doc, "La tabla se ordena por nombre del recurso de forma alfabetica.")

    screenshot_placeholder(doc, "Tabla de Recursos con columnas y badges de estado")

    # --- 40. Crear/Editar Modal ---
    h2(doc, "40. Crear o Editar Recurso — Modal", anchor="sec_rec_crear")
    para(doc, "Al hacer clic en \"Agregar Recurso\" o en \"Editar\" de un recurso existente, "
         "se abre un modal con los siguientes campos:")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Nombre", "Texto", "Si", "Nombre identificador del recurso"],
                  ["Tipo", "Select (15 opciones)", "Si", "Tipo de recurso (room, table, court, etc.)"],
                  ["Sede", "Select (sedes del negocio)", "Si", "Sede donde se ubica el recurso"],
                  ["Capacidad", "Numero (min: 1)", "No", "Capacidad del recurso en personas/unidades"],
                  ["Precio por Hora", "Numero (COP)", "No", "Tarifa por hora de uso"],
                  ["Descripcion", "Textarea", "No", "Descripcion detallada del recurso"],
                  ["Amenidades", "Texto (separado por comas)", "No", "Lista de amenidades incluidas"],
                  ["Estado", "Select", "Si", "available, maintenance, out_of_service"]])

    h4(doc, "Validaciones:")
    bullet(doc, "El nombre es obligatorio — el formulario no se puede enviar sin nombre.")
    bullet(doc, "El tipo es obligatorio — debe seleccionarse uno de los 15 tipos.")
    bullet(doc, "La sede es obligatoria — el recurso debe estar asociado a una sede del negocio.")
    bullet(doc, "La capacidad minima es 1 si se proporciona.")
    bullet(doc, "El precio debe ser mayor a 0 si se proporciona.")
    bullet(doc, "Las amenidades se ingresan como texto libre separado por comas "
           "(ej: \"WiFi, Proyector, Aire acondicionado\") y se almacenan como JSONB.")

    h4(doc, "Flujo de guardado:")
    bullet(doc, "Crear: INSERT en business_resources con los datos del formulario.")
    bullet(doc, "Editar: UPDATE del recurso existente con los nuevos datos.")
    bullet(doc, "En ambos casos se invalida el cache de React Query y se muestra un toast de confirmacion.")

    screenshot_placeholder(doc, "Modal de Crear/Editar Recurso")

    # --- 41. Disponibilidad ---
    h2(doc, "41. Validacion de Disponibilidad de Recursos", anchor="sec_rec_disponibilidad")
    para(doc, "Cuando un cliente intenta reservar una cita que involucra un recurso fisico, "
         "el sistema valida automaticamente que el recurso este disponible en el horario solicitado.")

    para(doc, "La validacion utiliza la funcion RPC is_resource_available() que verifica:")
    bullet(doc, "Que no exista una cita ya agendada para ese recurso en el mismo horario (overlap detection).")
    bullet(doc, "Que el recurso tenga estado \"available\" (no en mantenimiento ni fuera de servicio).")
    bullet(doc, "Que el recurso pertenezca a la sede seleccionada.")

    para(doc, "El algoritmo de overlap es identico al de empleados: slotStart < aptEnd AND slotEnd > aptStart.")

    callout(doc, "Hook: useAssigneeAvailability",
            "La validacion de disponibilidad se realiza a traves del hook useAssigneeAvailability, "
            "que funciona tanto para empleados como para recursos fisicos. Automaticamente detecta "
            "si el assignee es un empleado o un recurso y aplica la logica correspondiente.",
            color=PURPLE)

    # --- 42. Modelo de Negocio ---
    h2(doc, "42. Condicion de Modelo de Negocio", anchor="sec_rec_modelo")
    para(doc, "El modulo de Recursos solo esta disponible para negocios cuyo campo "
         "resource_model en la tabla businesses no sea 'professional'. Los modelos soportados son:")

    simple_table(doc, ["Modelo", "Descripcion", "Ejemplo de Negocios"],
                 [["professional", "Solo empleados (modelo tradicional)", "Salon de belleza, Clinica dental, Consultorio medico"],
                  ["physical_resource", "Solo recursos fisicos", "Bowling, Parqueadero, Alquiler de vehiculos"],
                  ["hybrid", "Empleados + Recursos fisicos", "Hotel (recepcionista + habitaciones), Gimnasio (trainer + equipos)"],
                  ["group_class", "Clases grupales con recurso", "Yoga (instructor + sala), Crossfit (coach + box)"]])

    para(doc, "Los negocios existentes que fueron creados antes de esta funcionalidad tienen "
         "resource_model = 'professional' por defecto y no veran el modulo de Recursos. "
         "El administrador puede cambiar el modelo desde la configuracion del negocio.")

    # --- 43. Permisos ---
    h2(doc, "43. Permisos del Modulo de Recursos", anchor="sec_rec_permisos")

    simple_table(doc, ["Permiso", "Que protege", "Modo PermissionGate"],
                 [["resources.create", "Boton \"Agregar Recurso\"", "hide"],
                  ["resources.edit", "Boton \"Editar\" en cada recurso", "hide"],
                  ["resources.delete", "Boton \"Eliminar\" en cada recurso", "hide"],
                  ["resources.view", "Acceso al modulo de Recursos", "hide"]])

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # SISTEMA DE NOMINA Y SALARIOS
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "SISTEMA DE NOMINA Y SALARIOS", anchor="sec_nomina", color=DARK)

    # --- 44. Config Salario ---
    h2(doc, "44. Sistema de Nomina y Salarios", anchor="sec_nomina")
    para(doc, "El sistema de Nomina (Payroll) es un sub-modulo integrado dentro de la gestion "
         "de empleados que permite configurar el salario base, frecuencia de pago y dia de pago "
         "de cada empleado. Es exclusivo del Plan Pro.")

    para(doc, "Acceso: Modal de Perfil de Empleado → Pestana \"Salario\" (solo visible si el "
         "negocio tiene Plan Pro activo, verificado con planId === 'pro').")

    callout(doc, "Condicion de Visibilidad",
            "La pestana de Salario en el perfil del empleado solo es visible si el negocio "
            "tiene el Plan Pro activo. Si el plan es Gratuito o Basico, la pestana no aparece.",
            color=PURPLE)

    screenshot_placeholder(doc, "Pestana de Salario en el perfil de empleado")

    # --- 45. Config ---
    h2(doc, "45. Configuracion de Salario del Empleado", anchor="sec_nom_config")

    simple_table(doc, ["Campo", "Tipo", "Requerido", "Descripcion"],
                 [["Salario Base", "Numero (COP)", "No", "Monto del salario base del empleado"],
                  ["Frecuencia de Pago", "Select", "Si", "Con que periodicidad se paga al empleado"],
                  ["Dia de Pago", "Select", "Condicional", "Dia del mes/semana para el pago (solo si frecuencia es mensual)"],
                  ["Generar Egreso Recurrente", "Switch", "No", "Si el sistema debe crear automaticamente el egreso de nomina"]])

    h4(doc, "Logica de guardado:")
    bullet(doc, "Al guardar, se ejecuta un UPDATE en business_employees con los campos salary_base y salary_type.")
    bullet(doc, "Si el salario es mayor a $0 y el switch \"Generar Egreso Recurrente\" esta activo: "
           "se ejecuta un UPSERT en recurring_expenses creando un egreso de nomina con la categoria \"payroll\".")
    bullet(doc, "Si el salario se establece en $0 o se desactiva el switch: "
           "se desactiva el egreso recurrente correspondiente (no se elimina, solo se marca como inactivo).")

    # --- 46. Frecuencias ---
    h2(doc, "46. Frecuencias de Pago Disponibles", anchor="sec_nom_frecuencia")

    simple_table(doc, ["Frecuencia", "Clave", "Descripcion"],
                 [["Mensual", "monthly", "Pago una vez al mes en el dia configurado"],
                  ["Quincenal", "biweekly", "Pago dos veces al mes (1 y 15, o segun configuracion)"],
                  ["Semanal", "weekly", "Pago cada semana en el dia configurado"],
                  ["Diario", "daily", "Pago diario (tipico en jornaleros o trabajo temporal)"],
                  ["Por Hora", "hourly", "Pago basado en horas trabajadas"]])

    para(doc, "La frecuencia mas comun en Colombia es mensual, con pago el ultimo dia del mes "
         "o el dia 30. Las frecuencias quincenal y semanal son comunes en sectores como "
         "construccion y servicios temporales.")

    # --- 47. Dia de Pago ---
    h2(doc, "47. Dia de Pago Configurable", anchor="sec_nom_dia")
    para(doc, "El campo \"Dia de Pago\" solo se muestra cuando la frecuencia seleccionada "
         "es \"mensual\". Las opciones disponibles son:")

    bullet(doc, "Dias 1 al 28: dia fijo del mes.")
    bullet(doc, "Ultimo dia: se ajusta automaticamente al ultimo dia calendario del mes (28, 29, 30 o 31).")

    para(doc, "Para frecuencias quincenales, el sistema usa por defecto los dias 1 y 15 del mes. "
         "Para frecuencias semanales, el dia corresponde al dia de la semana.")

    callout(doc, "Nota: Dias 29-31",
            "El sistema permite como maximo el dia 28 para evitar inconsistencias con meses "
            "de distinta duracion (febrero tiene 28/29 dias). Para pagos al final del mes, "
            "use la opcion \"Ultimo dia\".",
            color=PURPLE)

    # --- 48. Egreso Recurrente ---
    h2(doc, "48. Generacion Automatica de Egreso Recurrente", anchor="sec_nom_recurrente")
    para(doc, "Cuando el switch \"Generar Egreso Recurrente\" esta activado y el salario es mayor "
         "a $0, el sistema crea automaticamente un registro en la tabla recurring_expenses con:")

    bullet(doc, "Nombre: \"Nomina — [Nombre del Empleado]\".")
    bullet(doc, "Categoria: payroll (nomina y salarios).")
    bullet(doc, "Monto: el salario base configurado.")
    bullet(doc, "Frecuencia: la frecuencia de pago seleccionada.")
    bullet(doc, "Dia: el dia de pago configurado.")
    bullet(doc, "Automatizado: true (el sistema genera la transaccion automaticamente).")
    bullet(doc, "Estado: activo.")

    para(doc, "Este egreso recurrente aparece en el modulo de Gastos, pestana Egresos Recurrentes, "
         "y genera transacciones automaticas de tipo expense en las fechas configuradas.")

    para(doc, "Si el administrador cambia el salario a $0 o desactiva el switch, el egreso "
         "recurrente se desactiva (is_active = false) pero no se elimina, preservando "
         "el historial de transacciones anteriores.")

    # --- 49. Permisos Nomina ---
    h2(doc, "49. Permisos del Sistema de Nomina", anchor="sec_nom_permisos")

    simple_table(doc, ["Permiso", "Que protege", "Modo PermissionGate"],
                 [["employees.edit_salary", "Pestana de Salario y todos sus campos", "disable"]])

    para(doc, "Si el usuario no tiene el permiso employees.edit_salary, la pestana de Salario "
         "aparece pero todos los campos estan deshabilitados (modo disable). El propietario "
         "del negocio siempre puede editar salarios.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # MAPA JERARQUICO
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "MAPA JERARQUICO DEL EQUIPO", anchor="sec_jerarquia", color=DARK)

    # --- 50. Mapa Jerarquico ---
    h2(doc, "50. Mapa Jerarquico del Equipo", anchor="sec_jerarquia")
    para(doc, "El Mapa Jerarquico (HierarchyMapView) es una visualizacion interactiva del "
         "organigrama del equipo en forma de arbol. Muestra la estructura de mando y las "
         "relaciones de reporte entre los empleados del negocio.")

    para(doc, "Acceso: Gestion de Empleados → Boton \"Mapa Jerarquico\" (vista alternativa a la lista).")

    screenshot_placeholder(doc, "Mapa Jerarquico del equipo con 5 niveles")

    # --- 51. 5 Niveles ---
    h2(doc, "51. Los 5 Niveles Jerarquicos", anchor="sec_jer_niveles")

    simple_table(doc, ["Nivel", "Rol", "Descripcion"],
                 [["1", "Owner (Propietario)", "Dueno del negocio — nodo raiz del arbol, siempre presente"],
                  ["2", "Admin (Administrador)", "Administrador con permisos elevados, reporta al owner"],
                  ["3", "Manager (Gerente)", "Gerente de sede o area, reporta al admin"],
                  ["4", "Team Lead (Lider de Equipo)", "Lider que coordina un grupo, reporta al gerente"],
                  ["5", "Staff (Personal)", "Personal operativo, reporta al lider de equipo"]])

    para(doc, "Cada empleado se muestra como un nodo con su avatar, nombre, rol y nivel. "
         "Las conexiones entre nodos representan relaciones de reporte.")

    # --- 52. Controles ---
    h2(doc, "52. Controles del Mapa: Zoom, Expandir, Pantalla Completa", anchor="sec_jer_zoom")

    simple_table(doc, ["Control", "Rango/Accion", "Descripcion"],
                 [["Zoom", "50% — 150%", "Slider para acercar o alejar el mapa"],
                  ["Expandir", "Toggle por nodo", "Expandir o colapsar los nodos hijos de un empleado"],
                  ["Pantalla Completa", "On/Off", "Expandir el mapa para ocupar toda la pantalla"],
                  ["Enfocar Empleado", "Click en nodo", "Centra la vista en el empleado seleccionado con auto-scroll"]])

    # --- 53. Editar Nivel ---
    h2(doc, "53. Editar Nivel Jerarquico de un Empleado", anchor="sec_jer_editar")
    para(doc, "Para cambiar el nivel jerarquico de un empleado:")

    numbered(doc, "Hacer clic derecho o seleccionar el nodo del empleado en el mapa.")
    numbered(doc, "Seleccionar \"Cambiar Nivel\" del menu contextual.")
    numbered(doc, "Elegir el nuevo nivel (1-5) del selector HierarchyLevelSelector.")
    numbered(doc, "Confirmar el cambio.")

    para(doc, "El cambio se guarda mediante el hook useUpdateEmployeeHierarchy que llama "
         "a la Edge Function update-hierarchy. El mapa se actualiza automaticamente.")

    callout(doc, "Nota: Solo Admins y Owners",
            "Solo los administradores y el propietario del negocio pueden cambiar niveles "
            "jerarquicos. Los empleados de nivel 3-5 solo pueden visualizar el mapa.",
            color=PURPLE)

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # FACTURACION Y PRICING
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "DASHBOARD DE FACTURACION Y PAGINA DE PRECIOS", anchor="sec_facturacion_pro", color=DARK)

    # --- 54. Dashboard Facturacion ---
    h2(doc, "54. Dashboard de Facturacion — Vista Pro", anchor="sec_facturacion_pro")
    para(doc, "El Dashboard de Facturacion (BillingDashboard) muestra diferente contenido "
         "segun el estado de la suscripcion del negocio. Para usuarios del Plan Pro, "
         "la vista activa incluye informacion completa del plan y opciones de gestion.")

    # --- 55. Estados ---
    h2(doc, "55. Estados del Dashboard de Facturacion", anchor="sec_fact_estados")

    simple_table(doc, ["Estado", "Vista Mostrada", "Descripcion"],
                 [["Sin plan (Gratuito)", "Invitacion a suscribirse", "Muestra beneficios del Plan Basico/Pro y CTA para activar"],
                  ["Plan cancelado", "Mensaje de expiracion", "Muestra fecha de expiracion y boton para reactivar"],
                  ["Plan activo", "Dashboard completo", "Informacion del plan, uso, historial de pagos, opcion de cancelar"],
                  ["Plan en prueba (trialing)", "Dashboard + banner", "Igual que activo pero con banner indicando dias restantes de prueba"]])

    # --- 56. Uso del Plan ---
    h2(doc, "56. Estadisticas de Uso del Plan", anchor="sec_fact_uso")
    para(doc, "Para negocios con plan activo, el dashboard muestra 3 tarjetas de uso:")

    simple_table(doc, ["Tarjeta", "Dato", "Limite Pro"],
                 [["Sedes Activas", "Numero de sedes creadas", "10"],
                  ["Empleados Activos", "Numero de empleados registrados", "15"],
                  ["Citas del Mes", "Citas agendadas en el mes actual", "Ilimitadas"]])

    para(doc, "Cada tarjeta muestra una barra de progreso visual que indica el porcentaje "
         "de uso respecto al limite del plan. Para citas (ilimitadas en Pro), la barra "
         "muestra el conteo sin limite superior.")

    # --- 57. Historial ---
    h2(doc, "57. Historial de Pagos", anchor="sec_fact_historial")
    para(doc, "El historial de pagos muestra una tabla con todas las transacciones de la suscripcion:")

    simple_table(doc, ["Columna", "Descripcion"],
                 [["Fecha", "Fecha del cargo"],
                  ["Concepto", "\"Plan Pro — Mensual\" o \"Plan Pro — Anual\""],
                  ["Monto", "Monto cobrado en COP"],
                  ["Estado", "pagado, pendiente, fallido"],
                  ["Factura", "Enlace para descargar la factura en PDF (si disponible)"]])

    # --- 58. Cancelacion ---
    h2(doc, "58. Flujo de Cancelacion", anchor="sec_fact_cancelar")
    para(doc, "Para cancelar el Plan Pro:")

    numbered(doc, "Navegar a Facturacion en la barra lateral.")
    numbered(doc, "Hacer clic en \"Cancelar Suscripcion\".")
    numbered(doc, "Confirmar la cancelacion en el dialogo de confirmacion.")
    numbered(doc, "El plan se mantiene activo hasta el final del periodo de facturacion actual.")
    numbered(doc, "Al expirar, el negocio vuelve al Plan Gratuito con los limites originales.")

    callout(doc, "Datos No Se Eliminan",
            "Al cancelar el Plan Pro, los datos del negocio (gastos, vacantes, recursos, salarios) "
            "no se eliminan. Solo se bloquea el acceso a los modulos exclusivos. Si el plan "
            "se reactiva posteriormente, toda la informacion estara disponible nuevamente.",
            color=ACCENT)

    page_break(doc)

    # --- 59. Pagina de Precios ---
    h2(doc, "59. Pagina de Precios — Vista Completa", anchor="sec_pricing")
    para(doc, "La pagina de Precios (PricingPage) es accesible desde Facturacion o desde "
         "cualquier PlanGate que invite al usuario a actualizar su plan. Muestra los 3 planes "
         "disponibles con sus caracteristicas, precios y botones de accion.")

    # --- 60. Grid de Planes ---
    h2(doc, "60. Grid de 3 Planes", anchor="sec_pricing_planes")

    simple_table(doc, ["Plan", "Precio Mensual", "Precio Anual", "Badge", "Boton"],
                 [["Gratuito", "$0 COP", "$0 COP", "—", "Deshabilitado (plan actual si aplica)"],
                  ["Basico", "$89,900 COP", "$899,000 COP", "Mas Popular", "Activo — seleccionar plan"],
                  ["Pro", "$159,900 COP", "$1,599,000 COP", "Proximamente", "Deshabilitado (opacity 60%)"]])

    para(doc, "Cada plan muestra:")
    bullet(doc, "Icono representativo.")
    bullet(doc, "Nombre del plan y subtitulo descriptivo.")
    bullet(doc, "Precio mensual y anual (con tachado si hay descuento activo).")
    bullet(doc, "Lista de funcionalidades incluidas (check verde) y no incluidas (X gris).")
    bullet(doc, "Boton de accion (segun estado del plan).")

    para(doc, "El Plan Basico tiene estilo destacado con borde primario, sombra y escala 105% "
         "para resaltar como la opcion mas popular.")

    para(doc, "El Plan Pro aparece con opacidad reducida (60%) y un badge \"Proximamente\" "
         "que indica que aun no esta disponible para compra.")

    screenshot_placeholder(doc, "Grid de 3 planes con el Plan Pro marcado como Proximamente")

    # --- 61. Codigos ---
    h2(doc, "61. Codigos de Descuento", anchor="sec_pricing_descuento")
    para(doc, "La pagina de precios incluye un campo de texto para ingresar codigos de descuento "
         "promocionales. El flujo es:")

    numbered(doc, "El usuario escribe el codigo en el campo (se convierte a mayusculas automaticamente).")
    numbered(doc, "Hace clic en \"Aplicar\".")
    numbered(doc, "El sistema valida el codigo — si es valido, muestra un toast de exito y "
             "actualiza los precios con el descuento aplicado (tachado + nuevo precio).")
    numbered(doc, "Si el codigo no es valido o ha expirado, se muestra un toast de error.")

    # --- 62. FAQ ---
    h2(doc, "62. Preguntas Frecuentes en Pricing", anchor="sec_pricing_faq")
    para(doc, "La pagina de precios incluye 4 preguntas frecuentes:")

    h3(doc, "Puedo cambiar de plan?")
    para(doc, "Si. Los upgrades se prorratean (se cobra la diferencia proporcional). "
         "Los downgrades se aplican al final del periodo de facturacion actual.")

    h3(doc, "Hay periodo de prueba?")
    para(doc, "El Plan Basico ofrece 14 dias de prueba gratuita sin requerir datos de pago. "
         "El Plan Pro ofrecera condiciones similares cuando se habilite.")

    h3(doc, "Que metodos de pago aceptan?")
    para(doc, "Visa, Mastercard, American Express via Stripe. PayU para metodos locales colombianos. "
         "MercadoPago para otros paises de Latinoamerica.")

    h3(doc, "Puedo cancelar en cualquier momento?")
    para(doc, "Si, sin penalidades. El plan se mantiene activo hasta el final del periodo pagado.")

    para(doc, "Al final de la seccion hay un CTA: \"Necesitas ayuda para elegir?\" con enlace "
         "a soporte@gestabiz.com.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # LIMITES AMPLIADOS Y FUNCIONALIDADES ADICIONALES
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "LIMITES AMPLIADOS Y FUNCIONALIDADES ADICIONALES", anchor="sec_limites_pro", color=DARK)

    # --- 63. Limites ---
    h2(doc, "63. Limites Ampliados del Plan Pro", anchor="sec_limites_pro")
    para(doc, "El Plan Pro expande significativamente los limites operativos del negocio, "
         "permitiendo gestionar operaciones mas grandes y complejas:")

    simple_table(doc, ["Recurso", "Plan Basico", "Plan Pro", "Incremento"],
                 [["Sedes", "3", "10", "+233%"],
                  ["Empleados", "6", "15", "+150%"],
                  ["Citas/mes", "Ilimitadas", "Ilimitadas", "Sin cambio"],
                  ["Clientes", "Ilimitados", "Ilimitados", "Sin cambio"],
                  ["Servicios", "Ilimitados", "Ilimitados", "Sin cambio"]])

    # --- 64. 10 Sedes ---
    h2(doc, "64. Gestion de hasta 10 Sedes", anchor="sec_sedes_10")
    para(doc, "Con el Plan Pro, el negocio puede tener hasta 10 sedes activas simultaneamente. "
         "Cada sede funciona de forma independiente con sus propios:")
    bullet(doc, "Horarios de apertura y cierre.")
    bullet(doc, "Empleados asignados.")
    bullet(doc, "Servicios disponibles.")
    bullet(doc, "Recursos fisicos.")
    bullet(doc, "Direccion y coordenadas de geolocalizacion.")

    para(doc, "Esto permite a cadenas de negocios, franquicias o negocios con multiples puntos "
         "de atencion gestionar toda su operacion desde una sola cuenta de Gestabiz.")

    callout(doc, "Limite de Sedes",
            "Si el negocio intenta crear una sede mas alla del limite de 10, el sistema muestra "
            "un PlanLimitBanner informando que se ha alcanzado el maximo del plan. No se ofrece "
            "un plan superior — 10 es el maximo actual.",
            color=PURPLE)

    # --- 65. 15 Empleados ---
    h2(doc, "65. Hasta 15 Empleados", anchor="sec_empleados_15")
    para(doc, "El Plan Pro permite registrar hasta 15 empleados activos en el negocio "
         "(vs. 6 del Plan Basico). Todos los empleados tienen acceso a:")
    bullet(doc, "Horarios configurables por dia de la semana.")
    bullet(doc, "Hora de almuerzo personalizada.")
    bullet(doc, "Asignacion de servicios.")
    bullet(doc, "Configuracion de salario (exclusivo Plan Pro).")
    bullet(doc, "Posicion en el mapa jerarquico.")
    bullet(doc, "Permisos personalizados.")

    para(doc, "Si el negocio ya tiene 15 empleados activos y necesita mas, "
         "el sistema muestra un PlanLimitBanner. Actualmente no existe un plan superior al Pro.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # FUNCIONALIDADES PROXIMAMENTE
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "FUNCIONALIDADES MARCADAS COMO PROXIMAMENTE", anchor="sec_proximamente", color=DARK)

    h2(doc, "66. Funcionalidades Marcadas como Proximamente", anchor="sec_proximamente")
    para(doc, "Las siguientes funcionalidades estan parcialmente implementadas (backend listo) "
         "pero su interfaz de usuario aun esta en desarrollo. Se habilitaran en futuras "
         "actualizaciones sin costo adicional para suscriptores del Plan Pro:")

    simple_table(doc, ["Funcionalidad", "Estado Backend", "Estado UI", "Descripcion"],
                 [["Compra del Plan Pro", "Listo (Stripe/PayU/MP)", "Boton deshabilitado", "El flujo de checkout esta listo pero el boton muestra \"Proximamente\""],
                  ["Editor de Permisos Individuales", "Listo (79 permisos)", "Proximamente", "Asignar/revocar los 79 tipos de permisos individualmente a cada usuario"],
                  ["Plantillas de Permisos UI", "Listo (9 templates)", "Proximamente", "Crear, editar y aplicar plantillas de permisos personalizadas desde la UI"],
                  ["Historial de Auditoria", "Listo (permission_audit_log)", "Proximamente", "Ver registro completo de cambios de permisos con fecha, usuario y accion"],
                  ["Emojis en Chat", "—", "Proximamente", "Selector de emojis en la ventana de chat"],
                  ["Llamada de Voz", "—", "Proximamente", "Llamadas de voz entre usuarios dentro de la plataforma"],
                  ["Videollamada", "—", "Proximamente", "Videollamadas entre usuarios dentro de la plataforma"],
                  ["Agregar a Google Calendar", "Listo", "Proximamente", "Boton para agregar la cita confirmada directamente al Google Calendar del cliente"]])

    # --- 67. Permisos Individuales ---
    h2(doc, "67. Editor de Permisos Individuales (Proximamente)", anchor="sec_permisos_individuales")
    para(doc, "Actualmente en el Plan Basico, los permisos se asignan mediante 3 plantillas "
         "predefinidas (Recepcionista, Profesional, Contador). El Plan Pro habilitara un editor "
         "completo que permite:")
    bullet(doc, "Ver los 79 tipos de permisos organizados por categoria (services.*, employees.*, appointments.*, etc.).")
    bullet(doc, "Asignar o revocar cada permiso individualmente a cada usuario.")
    bullet(doc, "Crear plantillas personalizadas combinando permisos a medida.")
    bullet(doc, "Aplicar plantillas en bulk a multiples usuarios.")

    para(doc, "El backend esta completamente funcional (PermissionRPCService con 5 metodos: "
         "revoke, assign, applyTemplate, bulkRevoke, bulkAssign). Solo falta la interfaz de usuario.")

    # --- 68. Auditoria ---
    h2(doc, "68. Historial de Auditoria (Proximamente)", anchor="sec_audit_trail")
    para(doc, "El Historial de Auditoria permitira al administrador ver un registro completo "
         "de todos los cambios de permisos realizados en el negocio:")
    bullet(doc, "Quien asigno o revoco el permiso.")
    bullet(doc, "A quien se le asigno o revoco.")
    bullet(doc, "Que permiso fue afectado.")
    bullet(doc, "Fecha y hora del cambio.")
    bullet(doc, "Metodo utilizado (individual, plantilla, bulk).")

    para(doc, "La tabla permission_audit_log ya registra estos datos automaticamente mediante "
         "triggers. Solo falta la interfaz de consulta.")

    page_break(doc)

    # ═══════════════════════════════════════════════════════════════════
    # UPGRADE Y SOPORTE
    # ═══════════════════════════════════════════════════════════════════

    h1(doc, "UPGRADE, SOPORTE Y GLOSARIO", anchor="sec_upgrade_basico_pro", color=DARK)

    # --- 69. Como Actualizar ---
    h2(doc, "69. Como Actualizar de Basico a Pro", anchor="sec_upgrade_basico_pro")
    para(doc, "Cuando el Plan Pro se habilite para compra, los negocios con Plan Basico "
         "podran actualizarse facilmente:")

    numbered(doc, "Navegar a Facturacion → Mejorar Plan.")
    numbered(doc, "Seleccionar el Plan Pro.")
    numbered(doc, "El sistema calcula la diferencia prorrateada por los dias restantes del ciclo actual.")
    numbered(doc, "Completar el pago adicional.")
    numbered(doc, "Los modulos Pro se desbloquean inmediatamente.")

    para(doc, "Ejemplo: Si quedan 15 dias del mes y el negocio paga $89,900/mes (Basico), "
         "la diferencia prorrateada seria aproximadamente $35,000 COP por los 15 dias restantes. "
         "A partir del siguiente ciclo se cobrara el monto completo del Plan Pro.")

    callout(doc, "Sin Perdida de Datos",
            "Al actualizar de Basico a Pro, no se pierden datos ni configuraciones. "
            "Todos los empleados, ausencias, ventas, reportes y permisos existentes "
            "se mantienen intactos. Solo se agregan nuevas capacidades.",
            color=ACCENT)

    # --- 70. Soporte ---
    h2(doc, "70. Soporte y Ayuda del Plan Pro", anchor="sec_soporte_pro")
    para(doc, "Los suscriptores del Plan Pro cuentan con:")
    bullet(doc, "Soporte prioritario por email con respuesta en 12-24 horas habiles.", bold_prefix="Email")
    bullet(doc, "Boton de reporte de bugs integrado en la aplicacion (FloatingBugReportButton).", bold_prefix="Bug Reports")
    bullet(doc, "Documentacion completa en este manual y en la Propuesta de Valor.", bold_prefix="Documentacion")
    bullet(doc, "Actualizaciones automaticas — las nuevas funcionalidades se despliegan sin intervencion del usuario.", bold_prefix="Updates")

    para(doc, "Contacto: soporte@gestabiz.com | WhatsApp: +57 XXX XXX XXXX | www.gestabiz.com")

    # --- 71. Glosario ---
    h2(doc, "71. Glosario de Terminos del Plan Pro", anchor="sec_glosario_pro")

    simple_table(doc, ["Termino", "Definicion"],
                 [["PlanGate", "Pantalla de bloqueo que aparece al intentar acceder a un modulo no incluido en el plan activo"],
                  ["PlanLimitBanner", "Banner que informa cuando se ha alcanzado el limite del plan (sedes, empleados)"],
                  ["resource_model", "Campo de la tabla businesses que define el modelo operativo (professional, physical_resource, hybrid, group_class)"],
                  ["recurring_expenses", "Tabla que almacena egresos que se repiten periodicamente"],
                  ["Payroll", "Sistema de nomina — configuracion de salarios y generacion automatica de egresos"],
                  ["UPSERT", "Operacion que inserta un registro nuevo o actualiza el existente si ya existe"],
                  ["JSONB", "Tipo de dato PostgreSQL para almacenar datos JSON binarios (usado en amenidades)"],
                  ["Overlap detection", "Algoritmo que detecta conflictos de horario entre citas/reservas"],
                  ["RPC", "Remote Procedure Call — funciones SQL ejecutadas desde el cliente via supabase.rpc()"],
                  ["SECURITY DEFINER", "Modo de ejecucion de funciones SQL con privilegios del creador (no del usuario)"],
                  ["Prorrateado", "Calculo proporcional del costo al cambiar de plan a mitad de ciclo"],
                  ["Edge Function", "Funcion serverless ejecutada en Supabase (Deno) para operaciones privilegiadas"],
                  ["MandatoryReviewModal", "Modal que se muestra automaticamente para reviews obligatorias en reclutamiento"],
                  ["Soft delete", "Desactivacion logica (is_active = false) en vez de eliminacion fisica"],
                  ["Materialized view", "Vista precalculada en PostgreSQL que se refresca periodicamente"],
                  ["COP", "Peso colombiano — moneda utilizada en toda la plataforma"],
                  ["Owner bypass", "El propietario del negocio tiene todos los permisos sin verificacion"],
                  ["PlanId", "Identificador del plan activo del negocio (free, basico, pro)"]])

    # ── Footer ────────────────────────────────────────────────────────────
    page_break(doc)
    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing — www.tituring.com")
    style_run(r, size=11, bold=True, color=PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gestabiz v1.0.3 — Manual de Usuario — Parte 5 de 5")
    style_run(r, size=10, italic=True, color=GREY)

    return doc


# ============================================================================
# PROPUESTA DE VALOR — PARTE 5 (ADMIN PLAN PRO)
# ============================================================================

def build_proposal_part5() -> Document:
    doc = setup_document(
        "Propuesta de Valor — Gestabiz",
        "Por que Gestabiz es la mejor opcion para tu negocio de servicios",
    )

    # ── Table of Contents ─────────────────────────────────────────────────
    h1(doc, "INDICE DE CONTENIDOS", anchor="toc_pv5")
    para(doc, "Haga clic en cualquier titulo para navegar directamente a la seccion.",
         italic=True, color=GREY)

    toc_items = [
        ("pv_intro", "1. Tu Negocio Necesita Herramientas de Crecimiento"),
        ("pv_pro_que", "2. Que Desbloqueas con el Plan Pro"),
        ("pv_gastos", "3. Control Total de Gastos — Cada Peso Cuenta"),
        ("pv_reclutamiento", "4. Reclutamiento Profesional — El Talento Correcto"),
        ("pv_recursos", "5. Recursos Fisicos — Mas Alla de los Profesionales"),
        ("pv_nomina", "6. Nomina Automatizada — Sin Hojas de Calculo"),
        ("pv_jerarquia", "7. Mapa Jerarquico — Visualiza tu Organizacion"),
        ("pv_sin_limites", "8. Limites Ampliados — Opera sin Restricciones"),
        ("pv_todo_incluido", "9. Todo lo del Basico + Cuatro Modulos Nuevos"),
        ("pv_roi", "10. Retorno de Inversion — Los Numeros Hablan"),
        ("pv_comparativa_planes", "11. Comparativa: Gratuito vs. Basico vs. Pro"),
        ("pv_vs_competencia", "12. Gestabiz Pro vs. La Competencia"),
        ("pv_escenarios", "13. Escenarios Reales de Negocios"),
        ("pv_verticales", "14. Verticales Atendidos por el Plan Pro"),
        ("pv_proximamente", "15. Roadmap — Lo que Viene"),
        ("pv_faq", "16. Preguntas Frecuentes"),
        ("pv_cta", "17. Prepara tu Negocio para el Plan Pro"),
    ]
    for anchor_id, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        add_internal_hyperlink(p, anchor_id, label)

    page_break(doc)

    # ══════════════════════════════════════════════════════════════════════
    # PARTE 1 — POR QUE ACTUALIZAR AL PLAN PRO
    # ══════════════════════════════════════════════════════════════════════

    h1(doc, "PARTE 1 — POR QUE TU NEGOCIO NECESITA EL PLAN PRO", anchor="parte1_pv5")

    # --- 1. Intro ---
    h2(doc, "1. Tu Negocio Necesita Herramientas de Crecimiento", anchor="pv_intro")
    para(doc, "El Plan Basico de Gestabiz te ayudo a profesionalizar la gestion de tu equipo, "
         "tus ventas y tus reportes. Pero crecer trae nuevos desafios: necesitas controlar gastos, "
         "contratar al talento correcto, gestionar recursos fisicos y automatizar la nomina. "
         "Sin estas herramientas, el crecimiento se convierte en caos.")

    para(doc, "El Plan Pro de Gestabiz es la respuesta. Por $159,900 COP al mes — menos de lo que "
         "cuesta un almuerzo diario para tu equipo — obtienes las herramientas que usan las "
         "grandes empresas, adaptadas al tamano y presupuesto de tu negocio.")

    callout(doc, "EL PROBLEMA",
            "Tu negocio esta creciendo: mas empleados, mas sedes, mas gastos, mas complejidad. "
            "Pero sigues gestionando gastos en hojas de Excel, contratando por WhatsApp, y pagando "
            "nomina de memoria. Cada mes que pasa sin control, pierdes dinero sin darte cuenta.",
            color=DANGER)

    callout(doc, "LA SOLUCION",
            "Gestabiz Plan Pro centraliza gastos, reclutamiento, recursos y nomina en una sola "
            "plataforma — la misma donde ya gestionas tus citas, empleados y clientes. "
            "Sin apps adicionales. Sin integraciones complicadas. Todo en un solo lugar.",
            color=ACCENT)

    # --- 2. Que desbloqueas ---
    h2(doc, "2. Que Desbloqueas con el Plan Pro", anchor="pv_pro_que")

    simple_table(doc, ["Modulo", "Problema que Resuelve", "Impacto"],
                 [["Gastos y Egresos", "No sabes a donde se va tu dinero", "Cada peso categorizado y rastreable"],
                  ["Reclutamiento", "Contratar es improvisado y lento", "Proceso profesional con 6 estados"],
                  ["Recursos Fisicos", "Reservas manuales causan conflictos", "Disponibilidad en tiempo real"],
                  ["Nomina / Payroll", "La nomina se calcula en Excel", "Egresos automaticos sin errores"]])

    para(doc, "Ademas, los limites del plan se amplian a 10 sedes y 15 empleados, "
         "permitiendo gestionar operaciones mas grandes desde la misma plataforma.")

    screenshot_placeholder(doc, "Dashboard admin con los 4 modulos Pro desbloqueados")

    page_break(doc)

    # --- 3. Gastos ---
    h2(doc, "3. Control Total de Gastos — Cada Peso Cuenta", anchor="pv_gastos")
    para(doc, "Imagina saber exactamente en que se gasta cada peso de tu negocio. No en una hoja "
         "de Excel que nadie actualiza — en un sistema en tiempo real que categoriza, suma y "
         "presenta tus gastos automaticamente.")

    h3(doc, "Lo que obtienes:")
    bullet(doc, "48 categorias de gastos que cubren TODO: nomina, arriendo, marketing, impuestos, equipos, transporte y mas.")
    bullet(doc, "Egresos unicos para gastos puntuales y egresos recurrentes para gastos fijos (arriendo, servicios publicos, nomina).")
    bullet(doc, "3 tarjetas de resumen: gastos de hoy, de la semana y del mes — siempre visibles.")
    bullet(doc, "Resumen por categoria para identificar donde se concentran tus mayores gastos.")
    bullet(doc, "9 metodos de pago soportados: efectivo, tarjeta, transferencia, PSE, Nequi, Daviplata y mas.")

    callout(doc, "CASO REAL",
            "Un salon de belleza en Medellin descubrio que gastaba $2,400,000 COP/mes en productos "
            "de limpieza — el triple de lo necesario. Con el modulo de gastos, identifico el exceso "
            "en la primera semana y redujo el gasto a $800,000 COP/mes. Ahorro: $1,600,000 COP/mes. "
            "El Plan Pro se pago solo en 3 dias.",
            color=ACCENT)

    # --- 4. Reclutamiento ---
    h2(doc, "4. Reclutamiento Profesional — El Talento Correcto", anchor="pv_reclutamiento")
    para(doc, "Contratar a la persona equivocada cuesta entre 3 y 6 meses de salario. "
         "Con el modulo de Reclutamiento de Gestabiz, el proceso de contratacion es estructurado, "
         "transparente y profesional — desde la publicacion de la vacante hasta la seleccion final.")

    h3(doc, "Lo que obtienes:")
    bullet(doc, "Publicacion de vacantes con 4 secciones: informacion basica, detalles, compensacion y estado.")
    bullet(doc, "Recepcion de aplicaciones organizada por 6 estados: pendiente, en revision, en proceso, aceptada, rechazada, retirada.")
    bullet(doc, "7 tarjetas de estadisticas para evaluar el pipeline de candidatos en tiempo real.")
    bullet(doc, "Perfil detallado de cada candidato: experiencia, habilidades, CV, disponibilidad.")
    bullet(doc, "Reviews obligatorias al contratar y al finalizar — documentacion del talento.")
    bullet(doc, "Notificaciones automaticas a candidatos en cada cambio de estado.")

    callout(doc, "DATO CLAVE",
            "Los negocios que usan un sistema de reclutamiento estructurado reducen el tiempo "
            "de contratacion en un 40% y aumentan la retencion de empleados en un 25%. "
            "Gestabiz te da ese sistema sin necesitar otra herramienta.",
            color=PURPLE)

    # --- 5. Recursos ---
    h2(doc, "5. Recursos Fisicos — Mas Alla de los Profesionales", anchor="pv_recursos")
    para(doc, "No todos los negocios se basan en profesionales. Si tienes canchas, mesas, salas, "
         "equipos o vehiculos, necesitas un sistema que gestione reservas de recursos — no solo "
         "de personas. Gestabiz es la unica plataforma que maneja ambos en un solo sistema.")

    h3(doc, "Lo que obtienes:")
    bullet(doc, "15 tipos de recursos: habitaciones, mesas, canchas, escritorios, equipos, vehiculos, "
           "espacios, carriles, campos, estaciones, parqueaderos, camas, estudios, salas de reuniones y otros.")
    bullet(doc, "Cada recurso con capacidad, precio por hora, amenidades y estado.")
    bullet(doc, "Validacion de disponibilidad en tiempo real — cero reservas duplicadas.")
    bullet(doc, "Asignacion de servicios a recursos: una cancha de tenis puede ofrecer \"Clase grupal\" "
           "y \"Alquiler libre\" con precios diferentes.")

    simple_table(doc, ["Vertical", "Recursos Tipicos", "Beneficio"],
                 [["Hotel", "Habitaciones (standard, suite, deluxe)", "Revenue management por tipo de habitacion"],
                  ["Restaurante", "Mesas (2, 4, 6, 8 personas)", "Optimizar rotacion y reservas"],
                  ["Centro Deportivo", "Canchas (tenis, padel, futbol)", "Alquiler por hora sin conflictos"],
                  ["Gimnasio", "Equipos (caminadora, bicicleta)", "Control de uso por horario"],
                  ["Co-working", "Escritorios y salas de reunion", "Disponibilidad en tiempo real"],
                  ["Bowling", "Carriles de boliche", "Reservas por grupo y horario"],
                  ["Hospital/Clinica", "Camas, consultorios", "Gestion de ocupacion"]])

    callout(doc, "VENTAJA COMPETITIVA",
            "Ninguna otra plataforma de gestion de citas en Colombia ofrece gestion de recursos "
            "fisicos integrada. Calendly no lo tiene. Booksy no lo tiene. Fresha no lo tiene. "
            "Gestabiz es la unica opcion para negocios que necesitan reservar espacios y profesionales.",
            color=ACCENT)

    # --- 6. Nomina ---
    h2(doc, "6. Nomina Automatizada — Sin Hojas de Calculo", anchor="pv_nomina")
    para(doc, "Calcular y registrar la nomina manualmente es tedioso, propenso a errores "
         "y consume tiempo que podrias invertir en tu negocio. Con el sistema de Nomina de Gestabiz, "
         "configuras una vez y el sistema se encarga del resto.")

    h3(doc, "Lo que obtienes:")
    bullet(doc, "Configuracion de salario base por empleado en COP.")
    bullet(doc, "5 frecuencias de pago: mensual, quincenal, semanal, diario, por hora.")
    bullet(doc, "Dia de pago configurable (1-28 o ultimo dia del mes).")
    bullet(doc, "Generacion automatica de egresos recurrentes de nomina.")
    bullet(doc, "Integracion con el modulo de Gastos — la nomina aparece categorizada como \"payroll\".")

    callout(doc, "CASO REAL",
            "Una clinica estetica con 12 empleados dedicaba 4 horas mensuales a calcular y registrar "
            "la nomina en Excel. Con Gestabiz, la configuracion tomo 15 minutos (una sola vez) "
            "y ahora los egresos se generan automaticamente cada mes. Ahorro: 47 horas al ano.",
            color=ACCENT)

    page_break(doc)

    # --- 7. Jerarquia ---
    h2(doc, "7. Mapa Jerarquico — Visualiza tu Organizacion", anchor="pv_jerarquia")
    para(doc, "A medida que tu equipo crece, necesitas claridad sobre quien reporta a quien. "
         "El Mapa Jerarquico de Gestabiz muestra tu organigrama de forma visual e interactiva "
         "con 5 niveles: Owner → Admin → Gerente → Lider → Personal.")

    para(doc, "Con zoom, expansion de nodos, pantalla completa y enfoque automatico, "
         "el mapa es la herramienta perfecta para onboarding de nuevos empleados y para "
         "mantener claridad organizacional.")

    # --- 8. Sin limites ---
    h2(doc, "8. Limites Ampliados — Opera sin Restricciones", anchor="pv_sin_limites")

    simple_table(doc, ["Lo que cambia", "Plan Basico", "Plan Pro"],
                 [["Sedes", "3", "10 (+233%)"],
                  ["Empleados", "6", "15 (+150%)"],
                  ["Gastos registrables", "No disponible", "Ilimitados"],
                  ["Vacantes activas", "No disponible", "Ilimitadas"],
                  ["Recursos fisicos", "No disponible", "Ilimitados"]])

    para(doc, "Con el Plan Pro, tu negocio puede escalar a multiples sedes y un equipo grande "
         "sin preocuparse por limites artificiales. Las citas, clientes y servicios ya eran "
         "ilimitados en el Plan Basico y lo siguen siendo en el Pro.")

    # --- 9. Todo incluido ---
    h2(doc, "9. Todo lo del Basico + Cuatro Modulos Nuevos", anchor="pv_todo_incluido")
    para(doc, "El Plan Pro no reemplaza al Plan Basico — lo extiende. Todo lo que ya funcionaba "
         "sigue funcionando, mas 4 modulos nuevos y limites ampliados:")

    h3(doc, "Del Plan Gratuito (Parte 3):")
    bullet(doc, "Dashboard Admin, Servicios, Sedes, Calendario, CRM de Clientes, Facturacion, "
           "Configuraciones, Perfil Publico, QR, Notificaciones, Chat.")

    h3(doc, "Del Plan Basico (Parte 4):")
    bullet(doc, "Empleados (6→15), Ausencias/Vacaciones, Historial de Ventas, Ventas Rapidas, "
           "Reportes Financieros, Permisos (3 plantillas), WhatsApp, Google Calendar.")

    h3(doc, "Nuevo en Plan Pro (esta Parte 5):")
    bullet(doc, "Gastos y Egresos (48 categorias), Reclutamiento (6 estados), "
           "Recursos Fisicos (15 tipos), Nomina/Payroll, Mapa Jerarquico.")

    page_break(doc)

    # --- 10. ROI ---
    h2(doc, "10. Retorno de Inversion — Los Numeros Hablan", anchor="pv_roi")
    para(doc, "El Plan Pro cuesta $159,900 COP/mes ($5,330 COP/dia). Veamos como se paga solo:")

    simple_table(doc, ["Fuente de Ahorro", "Ahorro Estimado/Mes", "Como"],
                 [["Control de gastos innecesarios", "$500,000 — $2,000,000", "Identificar y eliminar gastos excesivos con categorias detalladas"],
                  ["Reduccion de tiempo en contratacion", "$300,000 — $800,000", "Proceso estructurado reduce tiempo de 30 a 18 dias"],
                  ["Eliminacion de reservas duplicadas", "$200,000 — $500,000", "Validacion automatica de disponibilidad de recursos"],
                  ["Automatizacion de nomina", "$100,000 — $300,000", "4+ horas mensuales ahorradas en calculo manual"],
                  ["Mejor retencion de empleados", "$500,000 — $1,500,000", "Procesos profesionales aumentan satisfaccion y retencion"]])

    callout(doc, "AHORRO TOTAL ESTIMADO: $1,600,000 — $5,100,000 COP/MES",
            "El Plan Pro de $159,900 COP se paga entre 1 y 3 dias de uso. "
            "El retorno de inversion es de 10x a 32x al mes. "
            "No es un gasto — es la mejor inversion que puedes hacer en tu negocio.",
            color=ACCENT)

    para(doc, "Con el plan anual ($1,599,000 COP/ano — 17% de ahorro), la inversion equivale a "
         "$4,381 COP por dia. Menos que un cafe.")

    # --- 11. Comparativa Planes ---
    h2(doc, "11. Comparativa: Gratuito vs. Basico vs. Pro", anchor="pv_comparativa_planes")

    simple_table(doc, ["Caracteristica", "Gratuito", "Basico", "Pro"],
                 [["Precio/mes", "$0", "$89,900", "$159,900"],
                  ["Precio/ano", "$0", "$899,000", "$1,599,000"],
                  ["Ahorro anual", "—", "2 meses gratis", "2 meses gratis (17%)"],
                  ["Sedes", "1", "3", "10"],
                  ["Empleados", "1", "6", "15"],
                  ["Modulos exclusivos", "0", "6", "10 (6+4)"],
                  ["Gastos", "No", "No", "Si"],
                  ["Reclutamiento", "No", "No", "Si"],
                  ["Recursos", "No", "No", "Si"],
                  ["Nomina", "No", "No", "Si"],
                  ["Soporte", "Basico", "Email 24-48h", "Email 12-24h"]])

    # --- 12. vs. Competencia ---
    h2(doc, "12. Gestabiz Pro vs. La Competencia", anchor="pv_vs_competencia")
    para(doc, "Comparemos Gestabiz Plan Pro con las alternativas mas populares en el mercado:")

    simple_table(doc, ["Funcionalidad", "Gestabiz Pro", "Calendly", "Booksy", "Fresha"],
                 [["Precio/mes", "$159,900 COP (~$38 USD)", "$12-16 USD", "$29-49 USD", "$0-49 USD"],
                  ["Gestion de citas", "Incluido", "Incluido", "Incluido", "Incluido"],
                  ["Multi-sede (10+)", "10 sedes", "No", "1-3 sedes", "1-5 sedes"],
                  ["Gestion de empleados", "15 empleados", "No", "Basico", "Basico"],
                  ["CRM de clientes", "Incluido", "No", "Basico", "Basico"],
                  ["Gastos y egresos", "48 categorias", "No", "No", "No"],
                  ["Reclutamiento", "6 estados + reviews", "No", "No", "No"],
                  ["Recursos fisicos", "15 tipos", "No", "No", "No"],
                  ["Nomina/Payroll", "Automatizado", "No", "No", "No"],
                  ["Reportes financieros", "Graficos interactivos", "Basico", "Basico", "Basico"],
                  ["Chat en tiempo real", "Incluido", "No", "No", "No"],
                  ["Permisos granulares", "79 tipos + plantillas", "No", "No", "No"],
                  ["Mapa jerarquico", "5 niveles", "No", "No", "No"],
                  ["Idiomas", "ES + EN", "Multi", "Multi", "Multi"],
                  ["Pagos Colombia", "Stripe + PayU + MP", "Solo Stripe", "Solo Stripe", "Propio"],
                  ["Perfil publico SEO", "Incluido", "No", "Basico", "Basico"]])

    callout(doc, "CONCLUSION",
            "Gestabiz Plan Pro ofrece MAS funcionalidades que cualquier competidor a un precio "
            "competitivo para el mercado colombiano. Las plataformas globales como Calendly o Booksy "
            "cobran en dolares, no ofrecen gestion de gastos ni reclutamiento, y tienen soporte limitado "
            "en espanol. Gestabiz fue disenado desde cero para el mercado latinoamericano.",
            color=ACCENT)

    page_break(doc)

    # --- 13. Escenarios ---
    h2(doc, "13. Escenarios Reales de Negocios", anchor="pv_escenarios")

    h3(doc, "Escenario 1: Cadena de Salones de Belleza (8 sedes, 14 empleados)")
    para(doc, "Carolina tiene una cadena de salones de belleza en Bogota con 8 sedes y 14 estilistas. "
         "Antes de Gestabiz, usaba 5 herramientas diferentes: Google Calendar (citas), Excel (gastos y nomina), "
         "WhatsApp (reclutamiento), un cuaderno (control de sedes) y Computrabajo (vacantes).")
    para(doc, "Con Gestabiz Plan Pro:")
    bullet(doc, "Gestiona las 8 sedes desde un solo dashboard.")
    bullet(doc, "Publico 3 vacantes y recibio 47 aplicaciones en la primera semana.")
    bullet(doc, "Configuro la nomina de 14 empleados en 20 minutos (una sola vez).")
    bullet(doc, "Identifico $1,800,000 COP en gastos innecesarios de productos el primer mes.")
    bullet(doc, "Ahorro: 12 horas semanales en gestion administrativa.")

    h3(doc, "Escenario 2: Centro Deportivo (canchas + equipos + instructores)")
    para(doc, "Carlos tiene un centro deportivo en Cali con 3 canchas de padel, 2 de tenis, "
         "un gimnasio con 20 equipos y 8 instructores. Su mayor problema: reservas duplicadas "
         "de canchas y equipos.")
    para(doc, "Con Gestabiz Plan Pro:")
    bullet(doc, "Registro las 5 canchas y 20 equipos como recursos fisicos con precio por hora.")
    bullet(doc, "Los clientes reservan en linea y el sistema valida disponibilidad en tiempo real.")
    bullet(doc, "Elimino el 100% de reservas duplicadas (antes: 3-5 conflictos por semana).")
    bullet(doc, "Ingreso adicional estimado: $2,000,000 COP/mes por mejor ocupacion.")

    h3(doc, "Escenario 3: Clinica Estetica en Expansion (contratando)")
    para(doc, "Ana tiene una clinica estetica en Medellin con 4 sedes y 10 empleados. "
         "Esta contratando 3 nuevos profesionales y necesita un proceso estructurado.")
    para(doc, "Con Gestabiz Plan Pro:")
    bullet(doc, "Publico 3 vacantes con requisitos detallados, salario y beneficios.")
    bullet(doc, "Recibio 62 aplicaciones y las filtro por experiencia y disponibilidad.")
    bullet(doc, "Selecciono 3 candidatos con review obligatoria de cada uno.")
    bullet(doc, "Tiempo de contratacion: 12 dias (vs. 35 dias sin el sistema).")
    bullet(doc, "Ahorro en costo de contratacion: ~$3,000,000 COP.")

    # --- 14. Verticales ---
    h2(doc, "14. Verticales Atendidos por el Plan Pro", anchor="pv_verticales")
    para(doc, "El Plan Pro esta disenado para negocios de servicios en crecimiento. "
         "Los verticales que mas se benefician son:")

    simple_table(doc, ["Vertical", "Modulos Pro Clave", "Tamano Tipico"],
                 [["Salones de Belleza / Barrerias", "Gastos, Nomina, Reclutamiento", "3-10 sedes, 5-15 empleados"],
                  ["Clinicas Esteticas / Dentales", "Gastos, Nomina, Reclutamiento", "2-5 sedes, 8-15 empleados"],
                  ["Centros Deportivos / Gimnasios", "Recursos, Gastos, Nomina", "1-3 sedes, 5-15 empleados"],
                  ["Hoteles / Hostales", "Recursos (habitaciones), Gastos", "1-5 sedes, 10-15 empleados"],
                  ["Restaurantes / Cafes", "Recursos (mesas), Gastos", "2-8 sedes, 8-15 empleados"],
                  ["Co-working / Oficinas", "Recursos (escritorios, salas)", "1-3 sedes, 3-10 empleados"],
                  ["Spas / Centros de Bienestar", "Reclutamiento, Nomina, Gastos", "1-3 sedes, 5-12 empleados"],
                  ["Talleres Automotrices", "Recursos (estaciones), Gastos", "1-2 sedes, 5-10 empleados"],
                  ["Consultorios Medicos", "Recursos (consultorios), Nomina", "2-5 sedes, 8-15 empleados"],
                  ["Estudios de Fotografia / Arte", "Recursos (estudios), Reclutamiento", "1-2 sedes, 3-8 empleados"]])

    page_break(doc)

    # --- 15. Roadmap ---
    h2(doc, "15. Roadmap — Lo que Viene", anchor="pv_proximamente")
    para(doc, "El Plan Pro de Gestabiz esta en constante evolucion. Las siguientes funcionalidades "
         "se habilitaran en futuras actualizaciones sin costo adicional:")

    simple_table(doc, ["Funcionalidad", "Estado", "Impacto Esperado"],
                 [["Compra del Plan Pro", "Backend listo, UI pendiente", "Activacion directa desde la app"],
                  ["Editor de 79 permisos individuales", "Backend listo", "Control granular total sobre cada accion"],
                  ["Plantillas de permisos personalizadas", "Backend listo", "Crear roles a medida en segundos"],
                  ["Historial de auditoria de permisos", "Backend listo", "Trazabilidad completa de cambios"],
                  ["Emojis en chat", "Pendiente", "Comunicacion mas expresiva"],
                  ["Llamadas de voz y videollamadas", "Pendiente", "Comunicacion interna sin salir de la app"],
                  ["Agregar cita a Google Calendar", "Backend listo", "Sincronizacion directa con un clic"]])

    callout(doc, "COMPROMISO",
            "Todas las funcionalidades del roadmap se habilitaran SIN COSTO ADICIONAL para los "
            "suscriptores del Plan Pro. El precio del plan no aumentara al agregar nuevas funcionalidades.",
            color=ACCENT)

    # --- FAQ ---
    h2(doc, "16. Preguntas Frecuentes", anchor="pv_faq")

    h3(doc, "Cuando estara disponible el Plan Pro para compra?")
    para(doc, "Estamos preparando el lanzamiento. Puedes registrar tu interes enviando un email "
         "a soporte@gestabiz.com y te notificaremos tan pronto este disponible. "
         "Los primeros suscriptores recibiran un descuento exclusivo de lanzamiento.")

    h3(doc, "Puedo probar el Plan Pro gratis?")
    para(doc, "Si. Se ofrecera un periodo de prueba gratuita (similar al Plan Basico) "
         "cuando se habilite la compra. Sin datos de pago durante la prueba.")

    h3(doc, "Si ya tengo el Plan Basico, como actualizo?")
    para(doc, "Desde Facturacion → Mejorar Plan → Seleccionar Pro. "
         "El cobro sera prorrateado: solo pagas la diferencia por los dias restantes del ciclo actual.")

    h3(doc, "Que pasa con mis datos si bajo del Plan Pro al Basico?")
    para(doc, "Tus datos no se eliminan. Los modulos exclusivos del Pro se bloquean "
         "pero la informacion (gastos, vacantes, recursos, salarios) se mantiene intacta. "
         "Si reactivas el Pro, todo estara disponible inmediatamente.")

    h3(doc, "El Plan Pro funciona para negocios fuera de Colombia?")
    para(doc, "Si. Gestabiz soporta multiples paises de Latinoamerica. Los pagos se procesan "
         "con Stripe (global), PayU (Colombia/Latam) o MercadoPago (Argentina/Brasil/Mexico/Chile). "
         "Las categorias de gastos y la moneda son configurables.")

    h3(doc, "Cuantos empleados puedo tener?")
    para(doc, "Hasta 15 empleados activos con el Plan Pro. Si necesitas mas, contactanos "
         "para un plan Empresarial personalizado.")

    h3(doc, "El precio va a aumentar?")
    para(doc, "No para los suscriptores actuales. Si el precio cambia en el futuro, "
         "los suscriptores existentes mantienen su tarifa original (grandfather pricing).")

    h3(doc, "Hay descuento por pago anual?")
    para(doc, "Si. El Plan Pro anual cuesta $1,599,000 COP — equivalente a 10 meses, "
         "obteniendo 2 meses gratis (ahorro de $319,800 COP, 17%).")

    page_break(doc)

    # --- CTA ---
    h2(doc, "17. Prepara tu Negocio para el Plan Pro", anchor="pv_cta")

    para(doc, "El Plan Pro de Gestabiz es la herramienta definitiva para negocios de servicios "
         "en crecimiento. Mientras se habilita la compra, puedes preparar tu negocio:")

    numbered(doc, "Registrate gratis y configura tu negocio con el Plan Gratuito.")
    numbered(doc, "Activa el Plan Basico para gestionar empleados, ventas y reportes.")
    numbered(doc, "Cuando el Plan Pro este disponible, actualiza para desbloquear gastos, "
             "reclutamiento, recursos y nomina.")

    callout(doc, "SE EL PRIMERO EN SABERLO",
            "Envia un email a soporte@gestabiz.com con el asunto \"Interes Plan Pro\" y te "
            "notificaremos el dia del lanzamiento con un descuento exclusivo de early adopter.\n\n"
            "www.gestabiz.com | soporte@gestabiz.com | WhatsApp: +57 XXX XXX XXXX",
            color=ACCENT)

    para(doc, "Gestabiz no es solo una herramienta de citas — es el sistema nervioso central de tu negocio. "
         "Y el Plan Pro es su version mas poderosa.", italic=True, color=PURPLE)

    # ── Footer ────────────────────────────────────────────────────────────
    page_break(doc)
    if LOGO_TITURING.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(str(LOGO_TITURING), width=Cm(3))
        except Exception:
            pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Desarrollado por Ti Turing — www.tituring.com")
    style_run(r, size=11, bold=True, color=PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Gestabiz v1.0.3 — Propuesta de Valor — Parte 5 de 5")
    style_run(r, size=10, italic=True, color=GREY)

    return doc


# ============================================================================
# MAIN
# ============================================================================

def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    # --- Manual de Usuario ---
    print("Generando Manual de Usuario — Parte 5: Admin Plan Pro...")
    manual = build_manual_part5()
    manual_path = DOCS_DIR / "Manual_Usuario_Gestabiz - copilot - parte5.docx"
    manual.save(str(manual_path))
    print(f"  -> {manual_path}")

    # --- Propuesta de Valor ---
    print("Generando Propuesta de Valor — Parte 5: Admin Plan Pro...")
    propuesta = build_proposal_part5()
    propuesta_path = DOCS_DIR / "Propuesta_Valor_Gestabiz - copilot - parte5.docx"
    propuesta.save(str(propuesta_path))
    print(f"  -> {propuesta_path}")

    print("\nListo. Ambos documentos de la Parte 5 generados exitosamente.")
    print("=" * 60)
    print("SERIE COMPLETA (5 de 5 partes):")
    print("  Parte 1: Rol Cliente")
    print("  Parte 2: Rol Empleado")
    print("  Parte 3: Administrador — Plan Gratuito")
    print("  Parte 4: Administrador — Plan Basico")
    print("  Parte 5: Administrador — Plan Pro  <-- ESTA")
    print("=" * 60)


if __name__ == "__main__":
    main()
